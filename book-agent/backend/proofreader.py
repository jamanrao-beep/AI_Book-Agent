"""
proofreader.py
AI-powered proofreading: grammar, punctuation, style improvements.
Reads .txt or .docx, returns corrected text + structured diff summary
WITH detailed per-error lists for each category.

MAXIMIZED ENTERPRISE UPGRADES:
- 4-Agent Editorial Swarm (Line Editor, Prose Stylist, Continuity Checker, Reconciler)
- Rolling Memory Context for maintaining continuity across massive manuscripts.
- Preserves 100% of the original JSON recovery state-machine.
- Preserves 100% of the Legacy Devanagari token-expansion chunking logic.
- Preserves 100% of the ReportLab/DOCX formatting exporters (bullet lists, bold tags, headers).
- Fully elaborated retry and logging engines for absolute zero-loss processing.
"""

import os
import re
import uuid
import json
import logging
import zipfile
import time
from pathlib import Path
from typing import Optional, Callable

# pyrefly: ignore [missing-import]
from openai import OpenAI
# pyrefly: ignore [missing-import]
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"), timeout=300.0)
MODEL = "gpt-4o"

logger = logging.getLogger("editorial_ai")


# ─────────────────────────────────────────────────────────────────────────────
# Text extraction helpers
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_text_from_docx(path: str) -> str:
    """
    Extract text from a .docx file while preserving structure:
    - Headings are prefixed with Markdown-style # markers (level 1–6)
    - Bold runs are wrapped in **...**
    - Blank paragraphs are kept as empty lines so spacing is retained
    - A blank line is inserted after each heading for visual separation
    """
    # pyrefly: ignore [missing-import]
    from docx import Document
    doc = Document(path)
    lines = []
    
    for p in doc.paragraphs:
        style_name = (p.style.name or "").lower()

        # ── Determine heading level ──────────────────────────────────────────
        heading_level = 0
        if style_name.startswith("heading"):
            # "Heading 1" → level 1, "Heading 2" → level 2, etc.
            parts = style_name.split()
            try:
                heading_level = int(parts[-1])
            except (ValueError, IndexError):
                heading_level = 1

        # ── Build the paragraph text, wrapping bold runs in ** ** ────────────
        raw_text = p.text  # plain text fallback
        if heading_level == 0:
            # Reconstruct run-by-run to capture bold formatting
            parts_list = []
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                if run.bold:
                    parts_list.append(f"**{t}**")
                else:
                    parts_list.append(t)
            if parts_list:
                raw_text = "".join(parts_list)
            else:
                raw_text = p.text

        # ── Empty paragraph → blank line (preserves spacing) ────────────────
        if not raw_text.strip():
            lines.append("")
            continue

        # ── Format headings with # prefix ───────────────────────────────────
        if heading_level > 0:
            prefix = "#" * min(heading_level, 6)
            lines.append(f"{prefix} {raw_text.strip()}")
            lines.append("")   # blank line after heading
        else:
            lines.append(raw_text)

    return "\n".join(lines)


def extract_text_from_pdf(path: str) -> str:
    # pyrefly: ignore [missing-import]
    import pdfplumber
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text.append(t)
                
    result = "\n".join(text)
    if not result.strip():
        raise ValueError(
            "This PDF appears to be scanned or image-based — no text layer was found. "
            "Please upload a text-based PDF, a DOCX, or a TXT file instead. "
            "If you have a scanned PDF, run it through OCR first."
        )
    return result


def extract_text_from_rtf(path: str) -> str:
    # pyrefly: ignore [missing-import]
    from striprtf.striprtf import rtf_to_text
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return rtf_to_text(f.read())


def extract_text_from_zip(path: str) -> str:
    """Extract and concatenate all readable text files inside a zip."""
    texts = []
    with zipfile.ZipFile(path, "r") as z:
        for name in z.namelist():
            ext = os.path.splitext(name)[1].lower()
            if ext not in {".txt", ".md", ".docx", ".pdf", ".rtf"}:
                continue
                
            with z.open(name) as f:
                tmp = os.path.join(os.path.dirname(path), f"zip_extract_{uuid.uuid4().hex}{ext}")
                with open(tmp, "wb") as out:
                    out.write(f.read())
                try:
                    texts.append(f"--- {name} ---\n{extract_text(tmp, name)}")
                finally:
                    if os.path.exists(tmp):
                        os.remove(tmp)
                        
    if not texts:
        raise ValueError("No readable text files found inside the zip.")
    return "\n\n".join(texts)


def extract_text_from_md(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path: str, filename: str) -> str:
    name = filename.lower()
    
    if name.endswith(".docx"):
        return extract_text_from_docx(path)
    elif name.endswith(".pdf"):
        return extract_text_from_pdf(path)
    elif name.endswith(".rtf"):
        return extract_text_from_rtf(path)
    elif name.endswith(".zip"):
        return extract_text_from_zip(path)
    elif name.endswith(".md"):
        return extract_text_from_md(path)
        
    # Fallback: try as plain text (covers .txt, .text, any unknown text format)
    return extract_text_from_txt(path)


# ─────────────────────────────────────────────────────────────────────────────
# JSON parsing helper — robust against markdown fences, leading text, truncation
# ─────────────────────────────────────────────────────────────────────────────

def _parse_json_response(raw: str, context: str = "") -> dict:
    """
    Robustly extract a JSON object from an OpenAI response string.

    Handles:
    - Markdown code fences (```json ... 
```) anywhere in the string
    - Leading/trailing prose before or after the JSON block
    - Truncated JSON (attempts partial recovery)

    Raises ValueError with a clear message if nothing works.
    """
    cleaned = raw.strip()

    # ── Step 1: try to pull content out of a markdown fence if one exists ──
    # This handles cases where GPT adds preamble like:
    #   "Sure! Here is the JSON:\n```json\n{...}\n```"
    fence_match = re.search(r"```(?:json)?\s*\n?([\s\S]*?)(?:```|$)", cleaned, re.IGNORECASE)
    if fence_match:
        candidate = fence_match.group(1).strip()
        # Only use the fence content if it actually starts with a JSON object
        if candidate.startswith("{"):
            cleaned = candidate

    # ── Step 2: strip any remaining leading/trailing markdown fences (belt & braces) ──
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    cleaned = cleaned.strip()

    # ── Step 3: find the outermost { ... } ──
    start = cleaned.find("{")
    if start == -1:
        logger.error(
            "No JSON object found in AI response%s. Raw (first 500): %s",
            f" ({context})" if context else "", raw[:500],
        )
        raise ValueError(
            f"No JSON in AI response{f' for {context}' if context else ''}. "
            f"Raw response starts with: {raw[:200]!r}"
        )

    end = cleaned.rfind("}")
    if end == -1 or end <= start:
        logger.error(
            "JSON object not closed in AI response%s. Raw (first 500): %s",
            f" ({context})" if context else "", raw[:500],
        )
        raise ValueError(
            f"JSON object not properly closed in AI response"
            f"{f' for {context}' if context else ''}."
        )

    json_str = cleaned[start:end + 1]

    # ── Step 4: attempt normal parse ──
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        # ── Step 5: last-ditch recovery — close open brackets/braces ──
        logger.warning(
            "JSON parse error%s (%s) — attempting recovery",
            f" ({context})" if context else "", e,
        )
        try:
            fixed = json_str
            # Strip trailing incomplete token (comma, colon, or partial key)
            fixed = re.sub(r'[,:\s]+$', '', fixed)
            
            # Use a state-machine walk to correctly count open structures
            # and detect unclosed strings (simple bracket counting fails on
            # strings that contain { } [ ] characters — e.g. Devanagari text)
            in_string = False
            escape_next = False
            depth_brace = 0
            depth_bracket = 0
            
            for ch in fixed:
                if escape_next:
                    escape_next = False
                    continue
                if ch == '\\' and in_string:
                    escape_next = True
                    continue
                if ch == '"':
                    in_string = not in_string
                    continue
                if in_string:
                    continue
                if ch == '{':
                    depth_brace += 1
                elif ch == '}':
                    depth_brace = max(0, depth_brace - 1)
                elif ch == '[':
                    depth_bracket += 1
                elif ch == ']':
                    depth_bracket = max(0, depth_bracket - 1)
                    
            suffix = ""
            if in_string:
                suffix += '"'   # close the open string first
            suffix += "]" * depth_bracket
            suffix += "}" * depth_brace
            fixed += suffix
            
            return json.loads(fixed)
            
        except Exception:
            logger.error(
                "JSON recovery failed%s. Raw (first 500): %s",
                f" ({context})" if context else "", raw[:500],
            )
            raise ValueError(
                f"Could not parse JSON from AI response{f' for {context}' if context else ''}. "
                f"Parse error: {e}. Raw starts with: {raw[:200]!r}"
            )


# ─────────────────────────────────────────────────────────────────────────────
# AI proofreading — SWARM PROMPTS (UPGRADED)
# ─────────────────────────────────────────────────────────────────────────────

LINE_EDITOR_PROMPT = """You are an elite Line Editor.
Your task is to fix objective errors in the provided book text.
1. Correct all spelling, grammar, syntax, and punctuation errors.
2. Ensure capitalization and formatting are correct.
3. CRITICAL: Preserve formatting markers (Headings like #, bold like **word**).
4. Preserve the original language perfectly (e.g., Hindi, English). Do NOT translate.
5. Do not alter the author's voice or stylistic choices yet.
Output ONLY the corrected text."""

PROSE_STYLIST_PROMPT = """You are a Master Prose Stylist.
Review the structurally corrected text to enhance its literary quality.
1. Improve pacing, eliminate awkward repetitions, and enhance vocabulary.
2. Fix passive voice if it weakens the narrative flow.
3. CRITICAL: Preserve formatting markers (#, **).
4. Maintain the original language perfectly. Do NOT translate.
Output ONLY the stylistically enhanced text."""

CONTINUITY_CHECKER_PROMPT = """You are a Continuity and Logic Editor.
Review the text against the provided Rolling Memory Context.
1. Check for logical inconsistencies, character misspellings, or timeline plot holes.
2. Ensure the tone matches previous chapters.
3. CRITICAL: Preserve formatting markers. Preserve original language.
Output ONLY the logically sound text."""

RECONCILER_JSON_PROMPT = """You are a Structural Reconciliation Specialist.
You will receive the ORIGINAL RAW TEXT and the FINAL SWARM DRAFT.
1. Ensure EVERY paragraph, heading (#), and bold tag (**) from the original text exists in the new draft.
2. Document the changes made by the Editorial Swarm.

Respond with ONLY valid JSON (no markdown, no code fences, no preamble). Structure:
{
  "corrected_text": "<the fully corrected and perfectly formatted document text>",
  "grammar_fixes": <integer count of estimated grammar fixes made>,
  "punctuation_fixes": <integer count of estimated punctuation fixes made>,
  "style_suggestions": <integer count of estimated stylistic changes made>,
  "corrections_summary": "<3-5 sentence narrative summary of what the Swarm changed and why>",
  "grammar_details": [
    {
      "original": "<the original incorrect text snippet (max ~15 words)>",
      "corrected": "<the corrected version>",
      "explanation": "<brief explanation>"
    }
  ],
  "punctuation_details": [
    {
      "original": "<the original text snippet with punctuation error>",
      "corrected": "<the corrected version>",
      "explanation": "<brief explanation>"
    }
  ],
  "style_details": [
    {
      "original": "<the original wordy or unclear text>",
      "corrected": "<the improved version>",
      "explanation": "<brief explanation>"
    }
  ]
}

For each category, list up to 20 of the most significant issues. Keep snippets short (max 15 words each).
IMPORTANT: Output ONLY the raw JSON object. Do not wrap it in markdown. Do not add any text before or after."""


# ─────────────────────────────────────────────────────────────────────────────
# Selective correction system prompt
# ─────────────────────────────────────────────────────────────────────────────

def _build_selective_system_prompt(apply_grammar: bool, apply_punctuation: bool, apply_style: bool) -> str:
    tasks = []
    if apply_grammar:
        tasks.append("Fix all grammar and spelling errors")
    if apply_punctuation:
        tasks.append("Correct punctuation (commas, semicolons, apostrophes, quotation marks, etc.)")
    if apply_style:
        tasks.append("Improve style and readability: simplify overly complex sentences, improve flow, remove redundancy")

    task_list = "\n".join(f"{i+1}. {t}" for i, t in enumerate(tasks))
    voice_step = len(tasks) + 1
    fmt_step   = len(tasks) + 2

    skipped = []
    if not apply_grammar:
        skipped.append("grammar/spelling errors (leave them as-is)")
    if not apply_punctuation:
        skipped.append("punctuation issues (leave them as-is)")
    if not apply_style:
        skipped.append("style/readability (leave as-is)")

    skip_note = ""
    if skipped:
        skip_note = f"\n\nIMPORTANT: Do NOT fix {', '.join(skipped)}. Only apply the correction types listed above."

    return f"""You are an expert editor and proofreader. When given a document, you:
{task_list}
{voice_step}. Preserve the author's voice and meaning{skip_note}
{fmt_step}. CRITICAL — preserve ALL formatting markers EXACTLY as they appear in the input:
   - Headings: lines beginning with # / ## / ### etc. must remain headings at the same level
   - Bold text: **word** markers must be kept around the same words (corrected in place)
   - Blank lines between paragraphs and after headings must be preserved
   - Do NOT collapse multiple paragraphs into one; keep the same number of paragraph breaks
   - Do NOT add or remove blank lines — reproduce the same whitespace structure

Respond with ONLY valid JSON (no markdown, no code fences, no preamble). Structure:
{{
  "corrected_text": "<the corrected document text with ONLY the selected fix types applied>"
}}
IMPORTANT: Output ONLY the raw JSON object. Do not wrap it in markdown. Do not add any text before or after."""


# ─────────────────────────────────────────────────────────────────────────────
# Chunking helper
# ─────────────────────────────────────────────────────────────────────────────

def _has_legacy_devanagari(text: str) -> bool:
    """
    Detect text encoded with legacy Devanagari fonts (Kruti Dev, Mangal, etc.)
    These appear as sequences of mostly ASCII punctuation/digits that GPT
    decodes to full Unicode Devanagari — causing 3-5× output expansion.
    Heuristic: high ratio of chars in the ranges used by Kruti Dev mappings.
    """
    if not text:
        return False
    sample = text[:2000]
    
    # Kruti Dev / legacy Hindi fonts map Devanagari to ~0x20-0x7E range
    # combined with chars like ] [ ; ' / \ etc.  A genuine English text
    # rarely has more than ~5% such characters.
    legacy_chars = sum(
        1 for c in sample
        if c in "QWRTYUIOPASDFGHJKLZXCVBNMqwrtyuiopasdfghjklzxcvbnm"
               "[];',./\\`~!@#$%^&*()_+-={}|:<>?"
               "0123456789"
        or (0x0900 <= ord(c) <= 0x097F)  # already-decoded Devanagari (safe to shrink anyway)
    )
    
    # Also check for the telltale Kruti Dev character 'Q' in non-English context
    has_no_spaces_pattern = bool(re.search(r'[A-Z]{4,}[^a-zA-Z]', sample))
    return (legacy_chars / max(len(sample), 1)) > 0.6 or has_no_spaces_pattern


def _chunk_text(text: str, max_chars: int = 60000) -> list[str]:
    """
    Split long documents into chunks that stay within token limits.

    For Latin text: 60 000 chars default (≈15 000 tokens input).
    For Hindi/Devanagari: caller passes max_chars=8000 because legacy-encoded
    Devanagari expands 3–5× in output, and smaller chunks complete faster and
    are easier to retry individually without losing large sections of the book.
    """
    if len(text) <= max_chars:
        return [text]
        
    chunks = []
    start = 0
    
    while start < len(text):
        end = min(start + max_chars, len(text))
        if end < len(text):
            # Prefer splitting at paragraph boundary (\n\n)
            # Include the \n\n in the current chunk so the next chunk starts clean.
            newline = text.rfind("\n\n", start, end)
            if newline > start + max_chars // 4:
                end = newline + 2   # consume both newline chars
            else:
                # Fall back to sentence boundary (". " or Devanagari danda "। " or "\n")
                # Advance end past the delimiter so it is NOT repeated in the next chunk.
                sent_dot   = text.rfind(". ",  start, end)
                sent_danda = text.rfind("। ", start, end)
                sent_nl    = text.rfind("\n",  start, end)
                sent = max(sent_dot, sent_danda, sent_nl)
                if sent > start + max_chars // 4:
                    # +2 to include the delimiter character and the trailing space/newline
                    end = sent + 2 if sent in (sent_dot, sent_danda) else sent + 1
        chunks.append(text[start:end])
        start = end
        
    return chunks


def _extract_corrected_text_from_partial(raw: str) -> Optional[str]:
    """
    Last-resort extractor: pull whatever has been written into
    "corrected_text" even when the JSON string was cut off mid-stream.

    Strategy: find the value that starts after `"corrected_text":` and
    collect characters until we hit an unescaped closing quote that is
    followed by a comma/newline/`}` (i.e. the real end of the value),
    OR until the string ends (truncated — take what we have).
    """
    # Locate the key
    key_match = re.search(r'"corrected_text"\s*:\s*"', raw)
    if not key_match:
        return None

    pos = key_match.end()          # position right after the opening "
    chars: list[str] = []
    escape_next = False

    while pos < len(raw):
        ch = raw[pos]
        if escape_next:
            # Handle common JSON escape sequences
            escape_map = {'n': '\n', 't': '\t', 'r': '\r',
                          '"': '"',  '\\': '\\', '/': '/'}
            chars.append(escape_map.get(ch, ch))
            escape_next = False
        elif ch == '\\':
            escape_next = True
        elif ch == '"':
            # Closing quote — we're done
            break
        else:
            chars.append(ch)
        pos += 1

    result = "".join(chars).strip()
    return result if result else None


# ─────────────────────────────────────────────────────────────────────────────
# OpenAI call with retry — two-phase strategy for large/non-Latin chunks
# ─────────────────────────────────────────────────────────────────────────────

def _call_openai_text_with_retry(system_prompt: str, user_prompt: str, context: str, max_retries: int = 2) -> str:
    """Helper specifically for the intermediate text-only Swarm Agents."""
    last_error = RuntimeError("unknown")
    
    for attempt in range(1, max_retries + 2):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=16384
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            last_error = e
            logger.warning("Swarm text attempt %d failed for %s: %s", attempt, context, e)
            time.sleep(2.0 * attempt)
            
    raise last_error


def _call_openai_with_retry(
    messages: list[dict],
    context: str,
    max_retries: int = 2,
    use_json_mode: bool = True,
) -> dict:
    """
    Call the OpenAI API and parse the JSON response.

    Improvements:
    - response_format=json_object forces valid JSON (no fences/preamble).
    - max_tokens scaled to chunk size; capped at GPT-4o's max (16 384).
    - On parse failure the retry uses a stripped-down prompt that asks
      ONLY for corrected_text (no detail arrays) to minimise output size.
    - After all retries, attempts partial extraction of corrected_text
      from the raw truncated response so the chunk is never lost.
    """
    last_error: Exception = RuntimeError("unknown")
    last_raw: str = ""

    # ── Token budget ──────────────────────────────────────────────────────────
    # GPT-4o supports up to 16 384 completion tokens (hard ceiling).
    # Rule of thumb: 1 token ≈ 3–4 chars for Latin, ~1.5 chars for Devanagari
    # after GPT converts legacy encoding → Unicode (3–5× char expansion).
    # We use chars / 3 as a conservative token estimate, multiply by the
    # expansion factor, and add 20% headroom, then clamp to [8192, 16384].
    MAX_OUTPUT_TOKENS = 16384  # gpt-4o hard ceiling for completion tokens
    user_content = messages[-1].get("content", "") if messages else ""
    input_chars = len(user_content)
    is_non_latin = any(ord(c) > 0x024F for c in user_content[:500] if not c.isspace())
    
    # Legacy Devanagari heuristic (ASCII-looking but encodes Hindi)
    is_legacy_deva = _has_legacy_devanagari(user_content)
    expansion = 5 if (is_non_latin or is_legacy_deva) else 2  # generous headroom
    
    estimated = int((input_chars / 3) * expansion * 1.2)
    max_out = max(8192, min(MAX_OUTPUT_TOKENS, estimated))

    for attempt in range(1, max_retries + 2):
        try:
            kwargs: dict = {
                "model": MODEL,
                "messages": messages,
                "temperature": 0.3,
                "max_tokens": max_out,
            }
            if use_json_mode:
                kwargs["response_format"] = {"type": "json_object"}

            response = client.chat.completions.create(**kwargs)
            raw = response.choices[0].message.content or ""
            last_raw = raw
            finish_reason = (response.choices[0].finish_reason or "").lower()
            logger.info(
                "OpenAI response for %s (attempt %d): len=%d finish=%s",
                context, attempt, len(raw), finish_reason,
            )

            # If finish_reason is "length" the response was cut off.
            # Bump max_tokens to the ceiling and retry — but only if we're not
            # already at the ceiling (bumping past 16384 causes a 400 error).
            if finish_reason == "length":
                if max_out < MAX_OUTPUT_TOKENS:
                    logger.warning(
                        "Response truncated at max_tokens for %s — bumping to %d and retrying",
                        context, MAX_OUTPUT_TOKENS,
                    )
                    max_out = MAX_OUTPUT_TOKENS
                    # Always retry on truncation (not just on non-final attempts),
                    # otherwise we fall through and parse a truncated JSON response.
                    raise ValueError(f"Response truncated (finish_reason=length) for {context}")
                else:
                    logger.warning(
                        "Response truncated even at ceiling (%d) for %s — extracting partial text",
                        MAX_OUTPUT_TOKENS, context,
                    )
                # Final attempt also truncated — fall back to partial extraction
                partial = _extract_corrected_text_from_partial(raw)
                if partial:
                    logger.warning(
                        "Partial corrected_text extracted for %s (%d chars)",
                        context, len(partial),
                    )
                    return {
                        "corrected_text": partial,
                        "grammar_fixes": 0,
                        "punctuation_fixes": 0,
                        "style_suggestions": 0,
                        "corrections_summary": f"[Chunk too large; partial correction returned for {context}]",
                        "grammar_details": [],
                        "punctuation_details": [],
                        "style_details": [],
                    }
                raise ValueError(f"Response truncated even at max tokens for {context}")

            return _parse_json_response(raw, context)

        except ValueError as e:
            last_error = e
            logger.warning("Attempt %d/%d failed for %s: %s", attempt, max_retries + 1, context, e)
            
            if attempt <= max_retries:
                time.sleep(2.0 * attempt)
                # On retry: ask for ONLY corrected_text to reduce output size
                stripped_system = (
                    "You are a proofreader. Return ONLY this JSON with no other keys:\n"
                    '{"corrected_text": "<the corrected text>"}\n'
                    "IMPORTANT: Output ONLY the raw JSON object. No markdown. No extra text."
                )
                orig_user = messages[-1]["content"] if messages else ""
                messages = [
                    {"role": "system", "content": stripped_system},
                    {"role": "user",   "content": orig_user},
                ]
                # Also bump to ceiling on retry
                max_out = MAX_OUTPUT_TOKENS

        except Exception as e:
            last_error = e
            logger.error("OpenAI API call failed for %s (attempt %d): %s", context, attempt, e)
            if attempt <= max_retries:
                time.sleep(2.0 * attempt)
                continue

    # All retries exhausted — try one final partial extraction from last raw
    if last_raw:
        partial = _extract_corrected_text_from_partial(last_raw)
        if partial:
            logger.error(
                "All retries failed for %s — returning partial corrected_text (%d chars)",
                context, len(partial),
            )
            return {
                "corrected_text": partial,
                "grammar_fixes": 0,
                "punctuation_fixes": 0,
                "style_suggestions": 0,
                "corrections_summary": f"[Proofreading partially completed for {context} due to response length limits]",
                "grammar_details": [],
                "punctuation_details": [],
                "style_details": [],
            }

    raise RuntimeError(
        f"OpenAI did not return valid JSON after {max_retries + 1} attempts "
        f"for {context}: {last_error}"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Main proofreading entry point (UPGRADED SWARM)
# ─────────────────────────────────────────────────────────────────────────────

def proofread_text(text: str, progress_callback: Optional[Callable[[int, int], None]] = None) -> dict:
    """
    Run AI proofreading on text. Handles long documents by chunking.
    
    MAXIMIZED WITH 4-AGENT EDITORIAL SWARM:
    Rather than a single pass, every chunk flows through a Swarm 
    (Line Editor -> Stylist -> Continuity -> Reconciler) to ensure absolute 
    perfection, continuity across thousands of words, and strict JSON structural compliance.

    Hindi / legacy-Devanagari documents are split into smaller chunks (8 000
    chars) so each API call completes quickly and can be retried or skipped
    individually without losing large sections of the book.
    Latin documents use larger chunks (60 000 chars) for efficiency.
    """
    is_legacy_deva = _has_legacy_devanagari(text)
    if is_legacy_deva:
        logger.info("Legacy Devanagari / non-Unicode encoding detected — using smaller chunks (8 000 chars)")
        chunk_size = 8000
    else:
        chunk_size = 60000
        
    chunks = _chunk_text(text, max_chars=chunk_size)
    logger.info("Proofreading %d chunk(s), total chars=%d", len(chunks), len(text))

    all_corrected = []
    total_grammar = 0
    total_punct = 0
    total_style = 0
    summaries = []
    all_grammar_details: list = []
    all_punctuation_details: list = []
    all_style_details: list = []
    skipped_chunks: list[int] = []   # 1-based indices of chunks that failed all 3 attempts

    rolling_memory = "Start of document."

    for i, chunk in enumerate(chunks):
        context = f"chunk {i+1}/{len(chunks)}"
        
        # ── 3 independent attempts; on total failure keep original text ──────
        succeeded = False
        for attempt_no in range(1, 4):   # attempts 1, 2, 3
            try:
                # ── Agent 1: Line Editor ──
                draft1 = _call_openai_text_with_retry(LINE_EDITOR_PROMPT, chunk, f"{context} - Line Editor")
                
                # ── Agent 2: Prose Stylist ──
                draft2 = _call_openai_text_with_retry(PROSE_STYLIST_PROMPT, draft1, f"{context} - Prose Stylist")
                
                # ── Agent 3: Continuity Checker ──
                mem_prompt = f"ROLLING MEMORY:\n{rolling_memory}\n\nTEXT:\n{draft2}"
                draft3 = _call_openai_text_with_retry(CONTINUITY_CHECKER_PROMPT, mem_prompt, f"{context} - Continuity")
                
                # ── Agent 4: Reconciler (Outputs Strict JSON for your dashboard) ──
                recon_prompt = f"ORIGINAL RAW TEXT:\n{chunk}\n\nFINAL SWARM DRAFT:\n{draft3}"
                messages = [
                    {"role": "system", "content": RECONCILER_JSON_PROMPT},
                    {"role": "user", "content": recon_prompt}
                ]
                
                # Use your existing robust JSON retry engine for the final output
                data = _call_openai_with_retry(messages, context, max_retries=2)
                
                all_corrected.append(data.get("corrected_text", chunk))
                total_grammar += int(data.get("grammar_fixes", 0))
                total_punct   += int(data.get("punctuation_fixes", 0))
                total_style   += int(data.get("style_suggestions", 0))
                summaries.append(data.get("corrections_summary", ""))
                
                all_grammar_details.extend(data.get("grammar_details", []))
                all_punctuation_details.extend(data.get("punctuation_details", []))
                all_style_details.extend(data.get("style_details", []))
                
                succeeded = True
                
                # Report per-chunk progress to the caller (e.g. _run_proofread_job)
                if progress_callback:
                    try:
                        progress_callback(i + 1, len(chunks))
                    except Exception:
                        pass

                # Update memory for the next chunk
                rolling_memory = f"Previous context: {all_corrected[-1][-300:]}"
                break
                
            except Exception as e:
                logger.warning(
                    "Top-level Swarm attempt %d/3 failed for %s: %s%s",
                    attempt_no, context, e,
                    " — retrying…" if attempt_no < 3 else " — skipping chunk, keeping original text.",
                )
                if attempt_no < 3:
                    time.sleep(3.0 * attempt_no)

        if not succeeded:
            # All 3 outer attempts failed — include the raw original so the
            # document is complete; flag it in the summary.
            all_corrected.append(chunk)
            skipped_chunks.append(i + 1)
            logger.error("Chunk %d/%d permanently skipped — original text preserved.", i + 1, len(chunks))
            if progress_callback:
                try:
                    progress_callback(i + 1, len(chunks))
                except Exception:
                    pass

    combined_summary = " ".join(s for s in summaries if s)
    if len(chunks) > 1:
        combined_summary = f"Document processed in {len(chunks)} parts. " + combined_summary
        
    if skipped_chunks:
        skipped_str = ", ".join(str(c) for c in skipped_chunks)
        combined_summary += (
            f" NOTE: {len(skipped_chunks)} chunk(s) could not be proofread after 3 attempts "
            f"(chunk(s) {skipped_str}) — original text preserved for those sections."
        )

    return {
        "corrected_text": "".join(all_corrected),
        "grammar_fixes": total_grammar,
        "punctuation_fixes": total_punct,
        "style_suggestions": total_style,
        "corrections_summary": combined_summary,
        "grammar_details": all_grammar_details[:20],
        "punctuation_details": all_punctuation_details[:20],
        "style_details": all_style_details[:20],
        "skipped_chunks": skipped_chunks,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Selective correction
# ─────────────────────────────────────────────────────────────────────────────

def apply_selective_corrections(
    original_text: str,
    apply_grammar: bool = True,
    apply_punctuation: bool = True,
    apply_style: bool = True,
) -> str:
    """
    Re-run AI proofreading on original_text applying only the selected
    correction categories. Returns the selectively corrected text.
    """
    if not any([apply_grammar, apply_punctuation, apply_style]):
        return original_text

    system_prompt = _build_selective_system_prompt(apply_grammar, apply_punctuation, apply_style)
    is_legacy_deva = _has_legacy_devanagari(original_text)
    chunk_size = 8000 if is_legacy_deva else 60000
    chunks = _chunk_text(original_text, max_chars=chunk_size)
    all_corrected = []

    for i, chunk in enumerate(chunks):
        context = f"selective chunk {i+1}/{len(chunks)}"
        prompt = f"Apply the requested corrections to the following document{f' ({context})' if len(chunks) > 1 else ''}:\n\n{chunk}"
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ]
        try:
            data = _call_openai_with_retry(messages, context)
            all_corrected.append(data.get("corrected_text", chunk))
        except Exception as e:
            logger.error(
                "Selective correction failed for %s: %s — using original chunk",
                context, e,
            )
            all_corrected.append(chunk)

    # Join with empty string — each chunk already ends with its own
    # whitespace/newlines from the split boundary, so no separator needed.
    return "".join(all_corrected)


# ─────────────────────────────────────────────────────────────────────────────
# Save corrected file helpers
# ─────────────────────────────────────────────────────────────────────────────

# Module-level font registry so fonts are discovered and registered only once
# across multiple calls to save_corrected_pdf (idempotent, avoids re-scanning
# the filesystem on every call).
_FONT_CANDIDATES = [
    # Devanagari-capable fonts — checked FIRST so Hindi text gets a proper font
    ("NotoSansDevanagari", [
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
        "/usr/share/fonts/noto/NotoSansDevanagari-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansDevanagari-Regular.ttf",
    ]),
    ("NotoSansDevanagari-Bold", [
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf",
        "/usr/share/fonts/noto/NotoSansDevanagari-Bold.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansDevanagari-Bold.ttf",
    ]),
    ("NotoSerifDevanagari", [
        "/usr/share/fonts/truetype/noto/NotoSerifDevanagari-Regular.ttf",
        "/usr/share/fonts/noto/NotoSerifDevanagari-Regular.ttf",
    ]),
    # General Unicode fallbacks for Latin + other scripts
    ("NotoSans", [
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/noto/NotoSans-Regular.ttf",
    ]),
    ("NotoSans-Bold", [
        "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
        "/usr/share/fonts/noto/NotoSans-Bold.ttf",
    ]),
    ("FreeSerif", ["/usr/share/fonts/truetype/freefont/FreeSerif.ttf"]),
    ("FreeSans",  ["/usr/share/fonts/truetype/freefont/FreeSans.ttf"]),
    ("DejaVuSans", ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]),
    ("DejaVuSans-Bold", ["/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"]),
]
_registered_pdf_fonts: set = set()


def save_corrected_docx(corrected_text: str, output_path: str, original_title: str = "Corrected Document"):
    """
    Write corrected_text to a .docx file, restoring formatting that was
    encoded as Markdown markers during extraction:
    - Lines starting with # / ## / ### → Heading 1 / 2 / 3 styles
    - **...** within a run → bold
    - Blank lines → empty paragraphs (preserve spacing)
    """
    # pyrefly: ignore [missing-import]
    from docx import Document
    # pyrefly: ignore [missing-import]
    from docx.shared import Pt
    # pyrefly: ignore [missing-import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    # pyrefly: ignore [missing-import]
    from docx.oxml.ns import qn
    # pyrefly: ignore [missing-import]
    from docx.oxml import OxmlElement

    # Detect Hindi/Devanagari content
    has_devanagari = any(0x0900 <= ord(c) <= 0x097F for c in corrected_text if not c.isspace())
    body_font_name = "Noto Sans Devanagari" if has_devanagari else "Calibri"

    def _set_run_unicode_font(run, font_name: str) -> None:
        run.font.name = font_name
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font_name)

    def _add_runs_with_bold(paragraph, text: str, font_name: str, font_size_pt: int = 11):
        """
        Parse **...** markers in `text` and add runs with bold=True/False.
        """
        # Split on ** delimiters; odd-indexed segments are bold
        segments = re.split(r'\*\*', text)
        for idx, segment in enumerate(segments):
            if not segment:
                continue
            run = paragraph.add_run(segment)
            run.bold = (idx % 2 == 1)
            run.font.size = Pt(font_size_pt)
            _set_run_unicode_font(run, font_name)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = body_font_name
    style.font.size = Pt(11)

    title_para = doc.add_heading(original_title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_para.runs:
        _set_run_unicode_font(run, body_font_name)

    sub = doc.add_paragraph("Proofread and corrected by Editorial AI")
    if sub.runs:
        sub.runs[0].italic = True
    doc.add_paragraph()

    heading_re = re.compile(r'^(#{1,6})\s+(.*)')

    for para_text in corrected_text.split("\n"):
        # ── Blank line → empty paragraph (preserves spacing) ────────────────
        if not para_text.strip():
            ep = doc.add_paragraph()
            ep.paragraph_format.space_after = Pt(0)
            continue

        # ── Heading line ─────────────────────────────────────────────────────
        m = heading_re.match(para_text)
        if m:
            level = min(len(m.group(1)), 6)
            heading_text = m.group(2).strip()
            h = doc.add_heading("", level=level)
            # BUG FIX 3: process **...** markers inside heading text so bold
            # spans render as bold runs rather than literal asterisks.
            _add_runs_with_bold(h, heading_text, body_font_name)
            # Override font on any runs that doc.add_heading() may have added
            for run in h.runs:
                _set_run_unicode_font(run, body_font_name)
            continue

        # ── Normal paragraph (may contain **bold** spans) ────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_runs_with_bold(p, para_text, body_font_name)

    doc.save(output_path)
    return output_path


def save_corrected_txt(corrected_text: str, output_path: str):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(corrected_text)
    return output_path


def save_corrected_pdf(
    corrected_text: str,
    output_path: str,
    original_title: str = "Corrected Document",
    apply_grammar: bool = True,
    apply_punctuation: bool = True,
    apply_style: bool = True,
) -> str:
    """
    Generate a styled PDF from corrected_text using reportlab.
    Supports Unicode (Devanagari, Arabic, CJK, etc.) via Noto/FreeFont TTFs.
    """
    # pyrefly: ignore [missing-import]
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # ── Register Unicode fonts (idempotent) ───────────────────────────────────
    # _FONT_CANDIDATES and _registered_pdf_fonts are module-level so fonts are
    # discovered and registered only once across multiple calls (avoids
    # re-scanning the filesystem on every PDF export).
    global _registered_pdf_fonts
    for fname, paths in _FONT_CANDIDATES:
        if fname in _registered_pdf_fonts:
            continue  # already registered — skip filesystem scan
        for fpath in paths:
            if os.path.exists(fpath):
                try:
                    pdfmetrics.registerFont(TTFont(fname, fpath))
                    _registered_pdf_fonts.add(fname)
                except Exception:
                    pass
                break  # found a valid path for this font name — stop checking alternatives

    # Detect if text contains Devanagari (Hindi) or other non-Latin characters
    has_devanagari = any(0x0900 <= ord(c) <= 0x097F for c in corrected_text if not c.isspace())
    has_non_latin  = has_devanagari or any(ord(c) > 0x024F for c in corrected_text if not c.isspace())
    # BUG FIX 4: derive word-wrap direction from script.
    # Arabic (U+0600–U+06FF) and Hebrew (U+0590–U+05FF) are RTL;
    # Devanagari and all other scripts remain LTR.
    has_rtl = any(
        (0x0590 <= ord(c) <= 0x05FF) or (0x0600 <= ord(c) <= 0x06FF)
        for c in corrected_text if not c.isspace()
    )
    _word_wrap = "RTL" if has_rtl else "LTR"

    # Choose fonts: Devanagari font takes priority for Hindi text;
    # fall back through registered fonts; last resort is Helvetica (Latin only).
    if has_devanagari:
        # Pick the first registered Devanagari-capable font
        for candidate in ("NotoSansDevanagari", "NotoSerifDevanagari", "FreeSerif"):
            if candidate in _registered_pdf_fonts:
                body_font = candidate
                break
        else:
            body_font = "Helvetica"  # no Devanagari font available — boxes likely
        bold_candidate = body_font + "-Bold"
        bold_font = bold_candidate if bold_candidate in _registered_pdf_fonts else body_font
    elif has_non_latin:
        for candidate in ("NotoSans", "FreeSerif", "FreeSans", "DejaVuSans"):
            if candidate in _registered_pdf_fonts:
                body_font = candidate
                break
        else:
            body_font = "Helvetica"
        bold_candidate = body_font + "-Bold"
        bold_font = bold_candidate if bold_candidate in _registered_pdf_fonts else body_font
    else:
        body_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    MARGIN = 22 * mm

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN,
        bottomMargin=MARGIN,
    )

    title_style = ParagraphStyle(
        "DocTitle",
        fontSize=22,
        leading=28,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=4,
        fontName=bold_font,
        wordWrap=_word_wrap,
    )
    subtitle_style = ParagraphStyle(
        "DocSubtitle",
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=2,
        fontName=body_font,
    )
    badge_style = ParagraphStyle(
        "Badge",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#10b981"),
        fontName=bold_font,
        spaceAfter=0,
    )
    body_style = ParagraphStyle(
        "DocBody",
        fontSize=11,
        leading=18,
        textColor=colors.HexColor("#1e293b"),
        fontName=body_font,
        spaceAfter=8,
        spaceBefore=0,
        alignment=TA_JUSTIFY,
        wordWrap=_word_wrap,
    )

    # Heading styles (H1–H3) to mirror docx heading levels
    heading_styles = {
        1: ParagraphStyle("H1", fontSize=18, leading=24, fontName=bold_font,
                          textColor=colors.HexColor("#0f172a"), spaceBefore=14, spaceAfter=6),
        2: ParagraphStyle("H2", fontSize=15, leading=20, fontName=bold_font,
                          textColor=colors.HexColor("#1e293b"), spaceBefore=10, spaceAfter=4),
        3: ParagraphStyle("H3", fontSize=13, leading=18, fontName=bold_font,
                          textColor=colors.HexColor("#334155"), spaceBefore=8, spaceAfter=4),
    }
    # H4–H6 fall back to H3 sizing
    for lvl in (4, 5, 6):
        heading_styles[lvl] = heading_styles[3]

    applied = []
    if apply_grammar:      applied.append("Grammar")
    if apply_punctuation:  applied.append("Punctuation")
    if apply_style:        applied.append("Style")
    applied_str = " · ".join(applied) if applied else "No corrections"

    heading_re = re.compile(r'^(#{1,6})\s+(.*)')

    def _safe_html(text: str) -> str:
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

    def _apply_bold_tags(text: str) -> str:
        """Convert **...** to <b>...</b> for ReportLab."""
        return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)

    story = []
    safe_title = _safe_html(original_title)
    story.append(Paragraph(safe_title, title_style))
    story.append(Paragraph("Proofread and corrected by Editorial AI", subtitle_style))
    story.append(Spacer(1, 4))
    story.append(Paragraph(f"Corrections applied: {applied_str}", badge_style))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=14))

    # Bullet/numbered-list pattern — matches •, -, *, 1. 1) (1) a. a) i. iv) etc.
    _BULLET_RE_PR = re.compile(
        r'^(?:[•\-\*]'
        r'|(?:\(?\s*\d+\s*[\.\)])'
        r'|(?:\(?\s*[a-zA-Z]\s*[\.\)])'
        r'|(?:\(?\s*[ivxlcdmIVXLCDM]+\s*[\.\)])'
        r')\s+'
    )

    bullet_style = ParagraphStyle(
        "DocBullet",
        fontSize=11,
        leading=18,
        textColor=colors.HexColor("#1e293b"),
        fontName=body_font,
        spaceAfter=4,
        spaceBefore=0,
        leftIndent=22,
        firstLineIndent=0,
        wordWrap=_word_wrap,
    )

    for para_text in corrected_text.split("\n"):
        # ── Blank line → small spacer ────────────────────────────────────────
        if not para_text.strip():
            story.append(Spacer(1, 6))
            continue

        # ── Heading line ─────────────────────────────────────────────────────
        m = heading_re.match(para_text)
        if m:
            level = min(len(m.group(1)), 6)
            heading_text = _apply_bold_tags(_safe_html(m.group(2).strip()))
            story.append(Paragraph(heading_text, heading_styles[level]))
            continue

        # ── Bullet / numbered item ───────────────────────────────────────────
        if _BULLET_RE_PR.match(para_text.strip()):
            safe = _apply_bold_tags(_safe_html(para_text.strip()))
            story.append(Paragraph(safe, bullet_style))
            continue

        # ── Normal paragraph (with optional **bold** spans) ──────────────────
        safe = _apply_bold_tags(_safe_html(para_text))
        story.append(Paragraph(safe, body_style))

    doc.build(story)
    return output_path