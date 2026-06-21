"""
cover_designer.py  ·  v8.0 (Nano Banana 6-Tier Edition)
AI-powered book cover designer — fully personalised, image-aware covers.

Pipeline:
  1. Extract title from filename if not provided
  2. Extract a representative page image from the book (PDF page → PNG, or
     first image found in DOCX) to use as a visual source
  3. GPT-4o Vision analyses the book image AND the title to produce a
     deeply personalised cover concept with unique visual DNA per book
  4. Generate Cover Illustration via Nano Banana's 6-Tier Image Cluster
  5. render_cover_pdf  — full-bleed A4 cover with:
       • Nano Banana / Stability / SVG / Procedural background
       • Premium multi-pass drop shadow typography
       • Style-specific illustration overlay & motif texture layers
       • Genre badge & author lines
  6. Prepend cover to original PDF or DOCX
  7. Return output path + concept metadata

NANO BANANA MIGRATION (cover designer ONLY — all other features still use
OpenAI for text exactly as before):
  - Image generation no longer calls any OpenAI image endpoint
    (dall-e-2 / dall-e-3 removed entirely from this pipeline).
  - All cover artwork now flows through nano_banana.run_image_cluster(),
    which implements its own internal 6-Tier Failover Cluster:
        Tier 1: Nano Banana Pro (Gemini 2.5 Flash image preview)
        Tier 2: Nano Banana 2  (Gemini 2.0 Flash, sanitised prompt)
        Tier 3: Stability AI REST fallback
        Tier 4: SVG template fallback
        Tier 5: Procedural gradient fallback
        Tier 6: Complex-shape mosaic — guaranteed even fully offline
  - GPT-4o is still used for the text-only concept + prompt-crafting calls.
"""

from __future__ import annotations

import base64
import io
import math
import os
import json
import uuid
import shutil
import random
import logging
import tempfile
import urllib.request
import traceback
import time
from pathlib import Path

# pyrefly: ignore [missing-import]
from openai import OpenAI           
# pyrefly: ignore [missing-import]
from dotenv import load_dotenv      

import nano_banana

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
MODEL  = "gpt-4o"

logger = logging.getLogger("editorial_ai")


# ─────────────────────────────────────────────────────────────────────────────
# Book image extraction (PRESERVED EXACTLY)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_book_image(file_path: str, ext: str) -> bytes | None:
    """
    Extract one representative page/image from the book as JPEG bytes.
    - PDF  → rasterise a content-rich page (not page 0 which may be blank)
    - DOCX → extract first embedded image
    Returns None if extraction fails (cover still works without it).
    """
    try:
        if ext == ".pdf":
            return _extract_pdf_page_image(file_path)
        elif ext == ".docx":
            return _extract_docx_image(file_path)
    except Exception as ex:
        logger.error(f"  ⚠️  Book image extraction failed: {ex}\n{traceback.format_exc()}")
    return None


def _extract_pdf_page_image(pdf_path: str) -> bytes | None:
    """Rasterise a mid-book page to JPEG bytes using pymupdf (fitz)."""
    try:
        import fitz  # PyMuPDF   # pyrefly: ignore [missing-import]
    except ImportError:
        logger.warning("  ℹ️  pymupdf not installed; skipping book image extraction")
        return None

    try:
        doc = fitz.open(pdf_path)
        n   = doc.page_count
        
        # Pick a content-rich page: ~20% into the book, skip blank cover pages
        target = max(1, min(int(n * 0.20), n - 1))
        
        for attempt in range(min(5, n)):
            page_idx = (target + attempt) % n
            page = doc[page_idx]
            
            # 1.2× zoom → reasonable resolution for concept extraction
            mat  = fitz.Matrix(1.2, 1.2)   
            pix  = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_bytes = pix.tobytes("jpeg")
            
            if len(img_bytes) > 8_000:      # skip near-blank pages
                doc.close()
                return img_bytes
                
        doc.close()
    except Exception as e:
        logger.error(f"  ⚠️  Error in _extract_pdf_page_image: {e}\n{traceback.format_exc()}")
    return None


def _extract_docx_image(docx_path: str) -> bytes | None:
    """Extract the first embedded image from a DOCX as JPEG bytes."""
    try:
        import zipfile
        from PIL import Image   # pyrefly: ignore [missing-import]

        with zipfile.ZipFile(docx_path, "r") as z:
            media = [n for n in z.namelist()
                     if n.startswith("word/media/") and
                     any(n.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif"))]
            if not media:
                return None
                
            raw = z.read(media[0])
            img = Image.open(io.BytesIO(raw)).convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=80)
            return buf.getvalue()
            
    except Exception as e:
        logger.error(f"  ⚠️  Error in _extract_docx_image: {e}\n{traceback.format_exc()}")
    return None


def _image_to_b64(img_bytes: bytes) -> str:
    """Encode raw bytes to base64 string for the API."""
    return base64.b64encode(img_bytes).decode()


# ─────────────────────────────────────────────────────────────────────────────
# Book TEXT extraction (PRESERVED EXACTLY)
# ─────────────────────────────────────────────────────────────────────────────

def extract_book_text(file_path: str, ext: str, max_chars: int = 8000) -> str:
    """
    Extract plain text from the book (PDF or DOCX) to inform AI cover design.
    Returns up to max_chars characters sampled from across the document so
    the AI sees an early chapter, a mid-book section, and a later section.
    """
    text_parts: list[str] = []

    try:
        if ext == ".pdf":
            try:
                # pyrefly: ignore [missing-import]
                import fitz  # PyMuPDF
            except ImportError:
                # Fallback: pypdf
                # pyrefly: ignore [missing-import]
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                pages = reader.pages
                n = len(pages)
                sample_indices = sorted(set([
                    0, 1, max(0, n // 5), max(0, n // 2), max(0, n * 3 // 4)
                ]))
                for i in sample_indices:
                    if i < n:
                        chunk = pages[i].extract_text() or ""
                        text_parts.append(chunk)
                return "\n\n".join(text_parts)[:max_chars]

            doc = fitz.open(file_path)
            n = doc.page_count
            
            # Skip cover/blank pages; sample intro, early, mid, late
            sample_indices = sorted(set([
                min(1, n-1),
                min(2, n-1),
                max(0, n // 5),
                max(0, n // 2),
                max(0, n * 3 // 4),
            ]))
            
            for i in sample_indices:
                page = doc[i]
                chunk = page.get_text("text").strip()
                if len(chunk) > 200:
                    text_parts.append(chunk)
            doc.close()

        elif ext == ".docx":
            # pyrefly: ignore [missing-import]
            from docx import Document
            doc = Document(file_path)
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            
            # Sample beginning, middle, and later thirds
            n = len(all_text)
            parts = [
                all_text[:n // 3],
                all_text[n // 3: 2 * n // 3],
                all_text[2 * n // 3:],
            ]
            for part in parts:
                text_parts.append(part[:2500])

    except Exception as ex:
        logger.error(f"  ⚠️  Book text extraction failed: {ex}\n{traceback.format_exc()}")

    full = "\n\n---\n\n".join(text_parts)
    return full[:max_chars]


# ─────────────────────────────────────────────────────────────────────────────
# 5-TIER AI IMAGING CLUSTER
# ─────────────────────────────────────────────────────────────────────────────

DALLE_PROMPT_SYSTEM = """You are a world-class book cover art director and illustrator.
Your job: read the actual book content provided and craft one precise, vivid
DALL-E 3 image-generation prompt that will produce a stunning, professional
illustrated book cover background image that is UNIQUE to this specific book.

Rules:
- Read the book excerpt carefully. Identify: main characters, key settings,
  central conflict, emotional tone, historical/cultural context, iconic symbols.
- Build the prompt around SPECIFIC elements from the book — not generic themes.
- Describe: the main visual scene or character, setting, lighting, mood,
  color palette, and art style in vivid concrete terms.
- Art style should match genre:
    biography/history → "detailed oil painting", "dramatic realist illustration"
    fantasy → "epic digital painting", "dark fantasy art"
    sci-fi → "cinematic sci-fi concept art", "futuristic render"
    romance → "painterly watercolor", "soft impressionist"
    thriller → "noir illustration", "high-contrast graphic novel art"
- NEVER include generic placeholders. Every detail must be specific to THIS book.
- IMPORTANT: Ensure your prompt strictly follows OpenAI Safety Guidelines. Do not include excessive gore, real world politicians, or copyrighted names.
- 4–6 sentences. No text or words anywhere in the image.
- End EVERY prompt with: "No text, letters, titles, or words anywhere in the image. Portrait orientation, photorealistic painting."
"""

SANITIZER_SYSTEM_PROMPT = """You are an OpenAI Safety Policy expert.
The following DALL-E 3 image prompt was REJECTED by the safety filters.
Your job is to rewrite the prompt to be 100% compliant and brand-safe while maintaining the artistic vibe.
1. Remove all violence, gore, weapons, or self-harm references.
2. Remove all names of real public figures, politicians, or copyrighted characters.
3. Remove any sexually explicit or adult themes.
4. Replace them with abstract, safe, beautiful, or metaphorical imagery that conveys the same genre.
Output ONLY the rewritten, perfectly safe DALL-E prompt."""


def generate_dalle_prompt(concept: dict, book_title: str, book_text: str = "") -> str:
    """Ask GPT-4o to craft a scene-specific DALL-E 3 prompt derived from actual content."""
    genre   = concept.get("genre_label", "")
    style   = concept.get("style", "premium")
    tagline = concept.get("tagline", "")
    palette = concept.get("palette", {})
    bg1     = palette.get("bg_primary", "#1a1a2e")
    acc     = palette.get("accent", "#f59e0b")

    book_excerpt = book_text[:4000] if book_text else "(no excerpt available)"

    user_msg = (
        f"Book title: {book_title}\n"
        f"Genre: {genre}\n"
        f"Style: {style}\n"
        f"Tagline: {tagline}\n"
        f"Primary palette color: {bg1}, accent: {acc}\n"
        f"Design rationale: {concept.get('design_rationale', '')}\n\n"
        f"=== BOOK CONTENT EXCERPT ===\n{book_excerpt}\n"
        f"=== END EXCERPT ===\n\n"
        "Based on the book excerpt above, identify the most visually dramatic "
        "scene, character, or symbol from this specific book and write a "
        "DALL-E 3 prompt for the cover illustration. Be very specific."
    )
    
    resp = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": DALLE_PROMPT_SYSTEM},
            {"role": "user",   "content": user_msg},
        ],
        temperature=0.85,
        max_tokens=450,
    )
    return (resp.choices[0].message.content or "").strip()


def _generate_via_stability_ai(prompt: str) -> bytes | None:
    """
    Tier 3 Backup Engine: Calls the Stability AI REST API directly using standard urllib.
    Avoids heavy external dependencies and executes flawlessly if Stability API key is active.
    """
    api_key = os.getenv("STABILITY_API_KEY")
    if not api_key:
        logger.warning("  ℹ️  Stability AI API Key missing from environment; skipping fallback engine.")
        return None
        
    logger.info("  🚀 Stability AI failover cluster engaged. Generating via SD3/Ultra...")
    
    url = "https://api.stability.ai/v2beta/stable-image/generate/ultra"
    boundary = "----WebKitFormBoundaryEnterpriseAICluster"
    
    # Assemble multipart body payload parameters manually to maintain zero external lock-in
    parts = []
    parts.append(f"--{boundary}")
    parts.append('Content-Disposition: form-data; name="prompt"')
    parts.append("")
    parts.append(prompt[:2000]) # Protect length boundaries
    parts.append(f"--{boundary}")
    parts.append('Content-Disposition: form-data; name="output_format"')
    parts.append("")
    parts.append("jpeg")
    parts.append(f"--{boundary}")
    parts.append('Content-Disposition: form-data; name="aspect_ratio"')
    parts.append("")
    parts.append("2:3") # Clean book portrait dimension bounds
    parts.append(f"--{boundary}--")
    parts.append("")
    
    body = "\r\n".join(parts).encode("utf-8")
    
    try:
        req = urllib.request.Request(url, data=body)
        req.add_header("Authorization", f"Bearer {api_key}")
        req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
        req.add_header("accept", "image/*")
        
        with urllib.request.urlopen(req, timeout=45) as response:
            if response.status == 200:
                img_data = response.read()
                logger.info(f"  ✅ Stability AI image successfully generated and downloaded ({len(img_data)//1024} KB)")
                return img_data
    except Exception as e:
        logger.warning(f"  ⚠️ Stability AI generation pass failed: {e}")
        
    return None


def _generate_procedural_fallback_image(concept: dict, width: int = 1024, height: int = 1792) -> bytes | None:
    """
    Tier 5 Backup Engine: Local Procedural Generation.
    If all external generative APIs fail or the network drops completely, this algorithm
    generates a mathematically beautiful abstract composition using the exact book palette.
    """
    try:
        # pyrefly: ignore [missing-import]
        from PIL import Image, ImageDraw, ImageFilter
        
        palette = concept.get("palette", {})
        title = concept.get("title", "Cover")
        
        def _hex_to_rgb(h_str):
            h_str = h_str.lstrip("#")
            if len(h_str) == 3: 
                h_str = "".join(c*2 for c in h_str)
            return (int(h_str[0:2], 16), int(h_str[2:4], 16), int(h_str[4:6], 16))
            
        bg_rgb = _hex_to_rgb(palette.get("bg_primary", "#0f172a"))
        bg2_rgb = _hex_to_rgb(palette.get("bg_secondary", "#1e3a5f"))
        acc_rgb = _hex_to_rgb(palette.get("accent", "#f59e0b"))
        
        # Create base canvas
        img = Image.new("RGBA", (width, height))
        draw = ImageDraw.Draw(img)
        
        # 1. Base Gradient Interpolation
        for y in range(height):
            t = y / height
            r = int(bg_rgb[0] * (1 - t) + bg2_rgb[0] * t)
            g = int(bg_rgb[1] * (1 - t) + bg2_rgb[1] * t)
            b = int(bg_rgb[2] * (1 - t) + bg2_rgb[2] * t)
            draw.line([(0, y), (width, y)], fill=(r, g, b, 255))
            
        # 2. Procedural Soft Orbs
        rng = random.Random(hash(title))
        for _ in range(15):
            rad = rng.randint(100, 600)
            cx = rng.randint(-200, width + 200)
            cy = rng.randint(-200, height + 200)
            alpha = rng.randint(10, 40)
            draw.ellipse([cx-rad, cy-rad, cx+rad, cy+rad], fill=acc_rgb + (alpha,))
            
        # 3. Procedural Sine Waves (adds motion/flow)
        for i in range(5):
            amp = rng.randint(50, 300)
            freq = rng.uniform(0.001, 0.005)
            y_offset = rng.randint(200, height - 200)
            pts = []
            for x in range(0, width, 10):
                y = y_offset + math.sin(x * freq + i) * amp
                pts.append((x, y))
            if len(pts) > 1:
                draw.line(pts, fill=acc_rgb + (rng.randint(20, 80),), width=rng.randint(2, 8))
                
        # Blur the composition into an atmospheric backdrop
        img = img.filter(ImageFilter.GaussianBlur(15))
        
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=90)
        
        logger.info(f"  ✅ Tier 5 Procedural Fallback engaged successfully ({len(buf.getvalue())//1024} KB)")
        return buf.getvalue()
        
    except ImportError:
        logger.warning("  ⚠️ PIL not installed, cannot generate Tier 5 procedural fallback.")
        return None
    except Exception as e:
        logger.warning(f"  ⚠️ Tier 5 procedural generation failed: {e}")
        return None


def generate_cover_from_svg_template(
    title: str,
    concept: dict | None = None,
    author: str = "",
) -> bytes | None:
    """
    Genre-aware SVG template fallback.

    Template naming convention (put your SVGs in the templates/ folder):
        thriller_01.svg, thriller_02.svg ...
        scifi_01.svg, scifi_dark_02.svg ...
        romance_floral_01.svg ...
        fantasy_01.svg ...
        academic_01.svg, academic_blue_02.svg ...
        business_01.svg ...
        history_01.svg ...
        horror_01.svg ...
        general_01.svg   <- catch-all for unmatched genres

    Scoring: each SVG filename + parent folder is matched against the book's
    genre and style keywords. Highest-scoring group is collected, then one
    is chosen at random -- so 80 thriller templates each get equal chance.

    PLACEHOLDER TABLE (use these strings inside your .svg files):
        {{TITLE}}           -> book title  (<=30 chars)
        {{AUTHOR}}          -> author line (<=28 chars)
        {{GENRE}}           -> genre label (e.g. THRILLER)
        {{TAGLINE}}         -> one-line tagline
        {{SUBTITLE}}        -> subtitle
        {{BG_PRIMARY}}      -> hex bg colour    (e.g. #0f172a)
        {{BG_SECONDARY}}    -> hex gradient colour
        {{ACCENT}}          -> hex accent colour
        {{ACCENT2}}         -> hex secondary accent
        {{TITLE_COLOR}}     -> hex title text colour
        {{SUBTITLE_COLOR}}  -> hex subtitle text colour

        Legacy placeholders still honoured:
        "Your Title Here", "Your Title", "Author Name"
    """
    import glob

    concept   = concept or {}
    palette   = concept.get("palette", {})
    genre_raw = concept.get("genre_label", "").lower().strip()
    style_raw = concept.get("style",       "").lower().strip()

    GENRE_ALIASES: dict[str, list[str]] = {
        "thriller":  ["thriller", "suspense", "crime", "mystery", "noir"],
        "scifi":     ["scifi", "sci-fi", "science fiction", "sf", "space", "futur"],
        "fantasy":   ["fantasy", "magic", "epic", "sword", "dragon", "mythic"],
        "romance":   ["romance", "love", "romantic", "contemporary romance"],
        "horror":    ["horror", "dark", "gothic", "paranormal", "occult"],
        "academic":  ["academic", "textbook", "education", "science", "research",
                      "nonfiction", "non-fiction", "reference"],
        "business":  ["business", "finance", "economics", "entrepreneur", "leadership",
                      "self-help", "productivity", "management"],
        "history":   ["history", "historical", "biography", "memoir", "war",
                      "political", "true crime"],
        "children":  ["children", "kids", "juvenile", "middle grade", "picture book"],
        "poetry":    ["poetry", "poems", "verse", "literary"],
    }

    search_keywords: list[str] = []
    for canonical, aliases in GENRE_ALIASES.items():
        if any(a in genre_raw or a in style_raw for a in aliases):
            search_keywords.append(canonical)
    search_keywords += [genre_raw, style_raw]

    templates_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
    svgs = glob.glob(os.path.join(templates_dir, "**", "*.svg"), recursive=True)
    if not svgs:
        logger.warning("  No SVG templates found in templates/ folder.")
        return None

    def _score(svg_path: str) -> int:
        name     = os.path.basename(svg_path).lower()
        folder   = os.path.basename(os.path.dirname(svg_path)).lower()
        combined = folder + "_" + name
        score    = 0
        for rank, kw in enumerate(reversed(search_keywords)):
            if kw and kw in combined:
                score += (rank + 1) * 10
        if "general" in combined or "default" in combined:
            score -= 5
        return score

    scored    = sorted(svgs, key=_score, reverse=True)
    top_score = _score(scored[0])
    top_pool  = [s for s in scored if _score(s) == top_score]
    svg_path  = random.choice(top_pool)

    logger.info(
        f"  SVG template selected: {os.path.relpath(svg_path, templates_dir)!r}"
        f"  (score={top_score}, pool={len(top_pool)} of {len(svgs)},"
        f"  genre={genre_raw!r}, style={style_raw!r})"
    )

    try:
        with open(svg_path, "r", encoding="utf-8") as f:
            svg = f.read()

        short_title    = title[:30]  if len(title)  > 30 else title
        short_author   = (concept.get("author_line") or author or "")[:28]
        short_genre    = concept.get("genre_label", "").upper()[:20]
        short_tagline  = concept.get("tagline",  "")[:60]
        short_subtitle = concept.get("subtitle", "")[:60]

        bg1  = palette.get("bg_primary",    "#0f172a")
        bg2  = palette.get("bg_secondary",  "#1e3a5f")
        acc  = palette.get("accent",        "#f59e0b")
        acc2 = palette.get("accent2",       "#fbbf24")
        tcol = palette.get("title_color",   "#ffffff")
        scol = palette.get("subtitle_color","#e2e8f0")

        substitutions = {
            "{{TITLE}}":          short_title,
            "{{AUTHOR}}":         short_author,
            "{{GENRE}}":          short_genre,
            "{{TAGLINE}}":        short_tagline,
            "{{SUBTITLE}}":       short_subtitle,
            "{{BG_PRIMARY}}":     bg1,
            "{{BG_SECONDARY}}":   bg2,
            "{{ACCENT}}":         acc,
            "{{ACCENT2}}":        acc2,
            "{{TITLE_COLOR}}":    tcol,
            "{{SUBTITLE_COLOR}}": scol,
        }
        for placeholder, value in substitutions.items():
            svg = svg.replace(placeholder, value)

        # Legacy backward-compatibility placeholders
        svg = svg.replace("Your Title Here", short_title)
        svg = svg.replace("Your Title",      short_title)
        svg = svg.replace("Author Name",     short_author)

        return svg.encode("utf-8")

    except Exception as e:
        logger.warning(f"  SVG template fallback failed: {e}\n{traceback.format_exc()}")
        return None


def generate_cover_image(concept: dict, book_title: str, book_text: str = "") -> bytes | None:
    """
    Cover illustration generation — fully migrated to Nano Banana.

    GPT-4o still crafts the scene-specific image prompt (text-only call,
    unchanged), but ALL actual image generation now runs through
    nano_banana.run_image_cluster(), which implements its own internal
    6-tier failover:
        Tier 1: Nano Banana Pro (Gemini 2.5 Flash image preview)
        Tier 2: Nano Banana 2  (Gemini 2.0 Flash, sanitised prompt)
        Tier 3: Stability AI REST fallback
        Tier 4: SVG template fallback
        Tier 5: Procedural gradient fallback
        Tier 6: Complex-shape mosaic (absolute last-resort guarantee)

    No OpenAI image endpoint (dall-e-2 / dall-e-3) is called anywhere in
    this pipeline anymore — this keeps cover design fully on Nano Banana
    while every other feature (editor, scanner, etc.) keeps using OpenAI
    for text exactly as before.
    """
    image_prompt = generate_dalle_prompt(concept, book_title, book_text)
    logger.info(f"  🎨 Base Image Swarm prompt ({len(image_prompt)} chars): {image_prompt[:200]}…")

    result = nano_banana.run_image_cluster(image_prompt, book_title, concept)
    if result:
        return result

    logger.error("  ❌ All 6 Nano Banana tiers failed. Returning empty asset buffer.")
    return None


# ─────────────────────────────────────────────────────────────────────────────
# AI: Generate cover concept  (vision-aware)
# ─────────────────────────────────────────────────────────────────────────────

COVER_SYSTEM_PROMPT = """You are a world-class book cover designer and creative director with deep expertise in typography, colour theory, and visual storytelling.

Your task: produce a HIGHLY PERSONALISED cover design brief for the specific book provided.
Every cover must feel UNIQUE to its title and content — never generic.

━━ DESIGN STYLE RULES ━━
If a `design_style` hint is given, honour it strictly:
• "normal"      → clean, readable, balanced; neutral accessible tones; sans-serif feel
• "premium"     → rich dark backgrounds, gold/silver/copper accents, luxury serif weight
• "scifi"       → deep space blacks/navy, neon cyan/electric purple, sharp geometric motifs
• "minimalist"  → maximum whitespace, monochrome + ONE accent, ultra-sparse, no ornament
• "fantasy"     → jewel tones (emerald, burgundy, midnight blue), ornate flourishes, mystical
• "thriller"    → extreme contrast, blood reds or cold greys, shattered/diagonal motifs, urgency
• "romance"     → blush/rose/gold, soft watercolour feel, floral or ribbon motif, warmth
• "academic"    → muted navy/slate, grid or rule motifs, structured, no decorative excess
• "vibrant"     → bold saturated primary/secondary colours, energetic wave or dot motifs, loud
• "retro"       → sepia/mustard/burnt orange, halftone dots or diagonal stripe, vintage press
If no style is given, infer the best style from the title and book content.

━━ VISUAL PERSONALISATION RULES ━━
1. Derive palette DIRECTLY from the book's subject and emotional tone.
   Examples: a book about the ocean → deep teals + seafoam; a book about fire → deep crimson + amber
2. Choose motif and illustration_shape to MATCH the subject. A maths book → hexagons + cross_lines.
   A nature book → wave_curves + arch. A tech book → grid_lines + sunburst. NEVER default to circles.
3. The layout_template must vary: split books by feel into one of 6 templates (see below).
4. image_treatment tells the renderer how to blend the extracted book page into the background.
5. accent_elements is a list of up to 4 specific decorative elements to draw (e.g. "constellation dots", "circuit trace lines", "watercolour wash", "ink splatter").

━━ LAYOUT TEMPLATES ━━
• "split_horizon"   → horizontal colour split at 40% height; image in lower half, text top-left
• "full_bleed"      → image fills entire background (heavy tint overlay), text floats centred
• "left_panel"      → solid colour left 45%, image right 55%, text on left panel
• "top_image"       → image occupies top 50%, solid colour bottom 50%, text bottom
• "diagonal_cut"    → diagonal colour divide from bottom-left to top-right; text upper-left
• "magazine"        → clean white/light background, bold oversized title, image small top-right

━━ OUTPUT FORMAT ━━
Respond ONLY with valid JSON. No markdown, no code fences, no explanation.
{
  "title": "<display title — use \\n for line breaks to improve layout>",
  "subtitle": "<compelling subtitle that adds context, or empty string>",
  "tagline": "<one punchy sentence capturing the book's essence — specific to THIS book>",
  "author_line": "<descriptive line or author name, e.g. 'A definitive guide to quantum mechanics'>",
  "palette": {
    "bg_primary":    "<hex — dominant background colour>",
    "bg_secondary":  "<hex — secondary/gradient colour>",
    "panel_color":   "<hex — text panel background>",
    "accent":        "<hex — accent: rules, ornaments, badge — must pop against bg>",
    "accent2":       "<hex — secondary accent for depth, e.g. a lighter tint of accent>",
    "title_color":   "<hex — title text — must have WCAG AA contrast on bg_primary>",
    "subtitle_color":"<hex — subtitle text>",
    "tagline_color": "<hex — tagline text, may be softer>"
  },
  "style": "<one of: normal|premium|scifi|minimalist|fantasy|thriller|romance|academic|vibrant|retro>",
  "layout_template": "<one of: split_horizon|full_bleed|left_panel|top_image|diagonal_cut|magazine>",
  "motif": "<one of: concentric_circles|diagonal_stripes|scattered_dots|grid_lines|wave_curves|hexagons|triangles|stars|arcs|circuit_traces|halftone|ink_drops|none>",
  "illustration_shape": "<one of: large_circle|diamond|arch|triangle|cross_lines|sunburst|polygon|none>",
  "image_treatment": "<one of: tinted_overlay|grayscale_fade|duotone|vignette|blur_bg|color_burn>",
  "accent_elements": ["<element 1>", "<element 2>"],
  "genre_label": "<UPPERCASE genre, max 20 chars, e.g. BUSINESS|SCIENCE|FICTION|HISTORY|SELF-HELP>",
  "design_rationale": "<1-2 sentences explaining why these specific choices suit THIS book>"
}"""


def generate_cover_concept(
    book_title   : str,
    description  : str  = "",
    design_style : str  = "",
    book_image   : bytes | None = None,
    book_text    : str  = "",
) -> dict:
    """
    Call GPT-4o (with optional vision input) to generate a personalised cover concept.
    If book_image is provided, GPT-4o sees an actual page from the book.
    If book_text is provided, GPT-4o reads the actual prose for deeper personalisation.
    """
    user_text = f"Book title: {book_title}"
    if description:
        user_text += f"\nDescription / book summary: {description}"
    user_text += f"\nDesign style: {design_style or 'auto — infer from title and content'}"

    if book_text:
        user_text += (
            f"\n\n=== BOOK CONTENT EXCERPT ===\n{book_text[:5000]}\n=== END ===\n"
            "\nUse the book content above to personalise EVERY design decision: "
            "the palette should reflect the emotional tone, the tagline should "
            "echo a theme from the text, the genre_label should be precise, "
            "and the design_rationale should reference specific content from the book."
        )

    if book_image:
        user_text += ("\n\nI have also attached an image of a page from this book. "
                      "Analyse its visual density, subject matter, and mood to "
                      "inform the image_treatment and layout_template choices.")

    # Build message content
    if book_image:
        b64 = _image_to_b64(book_image)
        content = [
            {"type": "text",       "text": user_text},
            {"type": "image_url",  "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "low"}},
        ]
    else:
        content = user_text

    concept: dict = {}
    last_exc: Exception | None = None

    for attempt in range(3):
        try:
            # Use json_object mode so GPT-4o is forced to return valid JSON.
            # Vision (image_url) content is incompatible with response_format,
            # so we only enable it for text-only requests.
            extra_kwargs: dict = {}
            if not book_image:
                extra_kwargs["response_format"] = {"type": "json_object"}

            response = client.chat.completions.create(
                model    = MODEL,
                messages = [
                    {"role": "system", "content": COVER_SYSTEM_PROMPT},
                    {"role": "user",   "content": content},
                ],
                temperature = 0.85,
                max_tokens  = 1800,   # raised: large JSON schema needs room
                **extra_kwargs,
            )
            raw = (response.choices[0].message.content or "").strip()
            raw = raw.replace("```json", "").replace("```", "").strip()

            # Extract outermost JSON object (handles any surrounding text)
            s = raw.find("{")
            e = raw.rfind("}") + 1
            if s == -1 or e <= s:
                raise ValueError(f"No JSON object in cover concept response. Raw: {raw[:300]!r}")

            concept = json.loads(raw[s:e])
            break   # success

        except (json.JSONDecodeError, ValueError) as parse_err:
            last_exc = parse_err
            logger.warning(
                "  ⚠️  Cover concept attempt %d/3 — parse error: %s",
                attempt + 1, parse_err,
            )
            if attempt < 2:
                time.sleep(3 * (attempt + 1))

        except Exception as api_err:
            last_exc = api_err
            logger.warning(
                "  ⚠️  Cover concept attempt %d/3 — API error: %s",
                attempt + 1, api_err,
            )
            if attempt < 2:
                time.sleep(3 * (attempt + 1))

    if not concept:
        logger.error(
            "  ❌ generate_cover_concept failed after 3 attempts. Last error: %s\n%s",
            last_exc, traceback.format_exc(),
        )
        raise ValueError(
            f"Could not get a valid JSON cover concept from GPT-4o after 3 attempts. "
            f"Last error: {last_exc}"
        )

    # Defaults for any missing fields
    concept.setdefault("illustration_shape", "large_circle")
    concept.setdefault("layout_template",    "split_horizon")
    concept.setdefault("image_treatment",    "tinted_overlay")
    concept.setdefault("accent_elements",    [])
    concept.setdefault("motif",              "none")
    p = concept.setdefault("palette", {})
    p.setdefault("bg_primary",    p.get("bg_top",    "#0f172a"))
    p.setdefault("bg_secondary",  p.get("bg_bottom", "#1e3a5f"))
    p.setdefault("panel_color",   p.get("bg_secondary", "#1e293b"))
    p.setdefault("accent2",       p.get("accent", "#f59e0b"))
    # Ensure frontend-expected keys are always present (bg_top / bg_bottom)
    p.setdefault("bg_top",    p.get("bg_primary",   "#0f172a"))
    p.setdefault("bg_bottom", p.get("bg_secondary", "#1e3a5f"))

    rationale = concept.get("design_rationale", "")
    if rationale:
        logger.info(f"  🎨 Design rationale: {rationale}")

    return concept


# ─────────────────────────────────────────────────────────────────────────────
# Colour helpers
# ─────────────────────────────────────────────────────────────────────────────

def _hex(h: str) -> tuple[float, float, float]:
    h = h.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return (int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255)


def _lerp(a: float, b: float, t: float) -> float:
    return a + (b - a) * t


def _blend(c1: tuple, c2: tuple, t: float) -> tuple:
    return (_lerp(c1[0], c2[0], t), _lerp(c1[1], c2[1], t), _lerp(c1[2], c2[2], t))


def _luminance(rgb: tuple) -> float:
    return 0.2126 * rgb[0] + 0.7152 * rgb[1] + 0.0722 * rgb[2]


# ─────────────────────────────────────────────────────────────────────────────
# Font helpers
# ─────────────────────────────────────────────────────────────────────────────

def _get_cover_font(text: str, bold: bool = False) -> str:
    """
    Return a registered ReportLab font for the given text.
    Falls back to Helvetica for Latin so we don't need to download for English-only titles.
    """
    try:
        from pdf_generator import get_font_for_text, detect_script  # pyrefly: ignore [missing-import]
        script = detect_script(text)
        if script == "Latin":
            return "Helvetica-Bold" if bold else "Helvetica"
        return get_font_for_text(text)
    except Exception:
        return "Helvetica-Bold" if bold else "Helvetica"


# ─────────────────────────────────────────────────────────────────────────────
# Image compositing helpers
# ─────────────────────────────────────────────────────────────────────────────

def _prepare_book_image(img_bytes: bytes, treatment: str,
                        bg_primary: tuple, accent: tuple,
                        width_px: int = 600, height_px: int = 800) -> bytes | None:
    """
    Apply image treatment (tint, grayscale, duotone, vignette, etc.) using Pillow.
    Returns processed JPEG bytes, or None if Pillow is unavailable.
    """
    try:
        from PIL import Image, ImageFilter, ImageEnhance, ImageOps, ImageDraw  # pyrefly: ignore [missing-import]
        import numpy as np  # pyrefly: ignore [missing-import]
    except ImportError as e:
        logger.warning(f"  ⚠️  Pillow/Numpy unavailable for image compositing. Error details: {e}")
        return img_bytes   

    try:
        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        # Resize to fill target canvas (cover crop)
        ratio = max(width_px / img.width, height_px / img.height)
        new_w = int(img.width  * ratio)
        new_h = int(img.height * ratio)
        img   = img.resize((new_w, new_h), Image.LANCZOS)
        # Centre crop
        left  = (new_w - width_px) // 2
        top   = (new_h - height_px) // 2
        img   = img.crop((left, top, left + width_px, top + height_px))

        if treatment == "grayscale_fade":
            img = ImageOps.grayscale(img).convert("RGB")
            img = ImageEnhance.Brightness(img).enhance(0.55)

        elif treatment == "duotone":
            gray  = ImageOps.grayscale(img)
            dark  = tuple(int(x * 255) for x in bg_primary)
            light = tuple(int(x * 255) for x in accent)
            arr   = np.array(gray, dtype=np.float32) / 255.0
            r = (dark[0] + (light[0] - dark[0]) * arr).clip(0, 255).astype(np.uint8)
            g = (dark[1] + (light[1] - dark[1]) * arr).clip(0, 255).astype(np.uint8)
            b = (dark[2] + (light[2] - dark[2]) * arr).clip(0, 255).astype(np.uint8)
            img = Image.fromarray(np.stack([r, g, b], axis=2))

        elif treatment == "blur_bg":
            img = img.filter(ImageFilter.GaussianBlur(radius=12))
            img = ImageEnhance.Brightness(img).enhance(0.60)

        elif treatment == "vignette":
            img = ImageEnhance.Brightness(img).enhance(0.65)
            vignette = Image.new("RGB", (width_px, height_px), (0, 0, 0))
            mask     = Image.new("L",   (width_px, height_px), 0)
            draw     = ImageDraw.Draw(mask)
            for i in range(60):
                alpha = int(200 * (i / 60))
                draw.rectangle([i, i, width_px - i, height_px - i], outline=alpha)
            img = Image.composite(img, vignette, ImageOps.invert(mask))

        elif treatment == "color_burn":
            tint_col = tuple(int(x * 220) for x in bg_primary)
            tint     = Image.new("RGB", (width_px, height_px), tint_col)
            img      = Image.blend(img, tint, alpha=0.72)

        else:  
            img = ImageEnhance.Brightness(img).enhance(0.50)
            img = ImageEnhance.Saturation(img).enhance(0.45)

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"  ⚠️  Error during _prepare_book_image compositing: {e}\n{traceback.format_exc()}")
        return img_bytes


def _draw_image_on_canvas(c, img_bytes: bytes, x: float, y: float, w: float, h: float) -> None:
    """Draw JPEG bytes onto a ReportLab canvas at given position/size."""
    try:
        from reportlab.lib.utils import ImageReader   # pyrefly: ignore [missing-import]
        reader = ImageReader(io.BytesIO(img_bytes))
        c.drawImage(reader, x, y, width=w, height=h, preserveAspectRatio=False, mask="auto")
    except Exception as ex:
        logger.error(f"  ⚠️  Could not draw book image on cover: {ex}\n{traceback.format_exc()}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF cover renderer
# ─────────────────────────────────────────────────────────────────────────────

def render_cover_pdf(concept: dict, output_path: str,
                     book_image_bytes: bytes | None = None,
                     dalle_image_bytes: bytes | None = None,
                     cover_image_bytes: bytes | None = None) -> str:
    """
    Render a professional A4 book cover PDF.
    Features a soft drop shadow text multi-pass pipeline to maintain accessibility over vivid background imagery.
    """
    from reportlab.pdfgen import canvas as rl_canvas   # pyrefly: ignore [missing-import]
    from reportlab.lib.pagesizes import A4              # pyrefly: ignore [missing-import]
    from reportlab.lib.utils import ImageReader         # pyrefly: ignore [missing-import]

    W_pt, H_pt = A4   

    DPI  = 150
    W_px = int(W_pt / 72 * DPI)   
    H_px = int(H_pt / 72 * DPI)   

    palette = concept.get("palette", {})

    def _hex_to_rgb(h: str) -> tuple:
        h = h.lstrip("#")
        if len(h) == 3: 
            h = "".join(c * 2 for c in h)
        return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    bg1_rgb   = _hex_to_rgb(palette.get("bg_primary",    "#0f172a"))
    bg2_rgb   = _hex_to_rgb(palette.get("bg_secondary",  "#1e3a5f"))
    acc_rgb   = _hex_to_rgb(palette.get("accent",        "#f59e0b"))
    acc2_rgb  = _hex_to_rgb(palette.get("accent2",       "#fbbf24"))
    tcol_rgb  = _hex_to_rgb(palette.get("title_color",   "#ffffff"))
    scol_rgb  = _hex_to_rgb(palette.get("subtitle_color","#e2e8f0"))
    tgcol_rgb = _hex_to_rgb(palette.get("tagline_color", "#94a3b8"))

    try:
        # pyrefly: ignore [missing-import]
        from PIL import Image, ImageDraw, ImageFilter, ImageFont, ImageEnhance, ImageOps
    except ImportError as e:
        logger.warning(f"  ⚠️  Pillow not available; using legacy ReportLab renderer. Error details: {e}")
        return _render_cover_pdf_legacy(concept, output_path, book_image_bytes)

    # ── Build background canvas ───────────────────────────────────────────────
    try:
        def _fit_to_canvas(img_bytes: bytes) -> "Image.Image":
            img   = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
            ratio = max(W_px / img.width, H_px / img.height)
            nw    = int(img.width  * ratio)
            nh    = int(img.height * ratio)
            img   = img.resize((nw, nh), Image.LANCZOS)
            left  = (nw - W_px) // 2
            top_  = (nh - H_px) // 2
            return img.crop((left, top_, left + W_px, top_ + H_px))

        if cover_image_bytes:
            bg = _fit_to_canvas(cover_image_bytes)
            logger.info("  ✅ User-supplied cover image composited as full-bleed background")
        elif dalle_image_bytes:
            bg = _fit_to_canvas(dalle_image_bytes)
            logger.info("  ✅ Generative imagery swarm composited as cover background")
        elif book_image_bytes:
            bk_img = Image.open(io.BytesIO(book_image_bytes)).convert("RGBA")
            ratio  = max(W_px / bk_img.width, H_px / bk_img.height)
            nw, nh = int(bk_img.width * ratio), int(bk_img.height * ratio)
            
            bk_img = bk_img.resize((nw, nh), Image.LANCZOS).crop(
                ((nw - W_px)//2, (nh - H_px)//2,
                 (nw - W_px)//2 + W_px, (nh - H_px)//2 + H_px)
            ).convert("RGBA")
            
            bk_img = bk_img.filter(ImageFilter.GaussianBlur(radius=18))
            bk_img = ImageEnhance.Brightness(bk_img).enhance(0.40)
            
            grad = Image.new("RGBA", (W_px, H_px))
            gd   = ImageDraw.Draw(grad)
            for yi in range(H_px):
                t = yi / H_px
                r = int(bg1_rgb[0]*(1-t) + bg2_rgb[0]*t)
                g = int(bg1_rgb[1]*(1-t) + bg2_rgb[1]*t)
                b = int(bg1_rgb[2]*(1-t) + bg2_rgb[2]*t)
                gd.line([(0, yi), (W_px, yi)], fill=(r, g, b, 170))
                
            bg = Image.alpha_composite(bk_img, grad)
        else:
            bg = Image.new("RGBA", (W_px, H_px))
            gd = ImageDraw.Draw(bg)
            for yi in range(H_px):
                t = yi / H_px
                r = int(bg1_rgb[0]*(1-t) + bg2_rgb[0]*t)
                g = int(bg1_rgb[1]*(1-t) + bg2_rgb[1]*t)
                b = int(bg1_rgb[2]*(1-t) + bg2_rgb[2]*t)
                gd.line([(0, yi), (W_px, yi)], fill=(r, g, b, 255))

        draw = ImageDraw.Draw(bg, "RGBA")

        # ── Overlay strategy ──────────────────────────────────────────────────────
        has_image = (cover_image_bytes is not None) or (dalle_image_bytes is not None)
        overlay = Image.new("RGBA", (W_px, H_px), (0, 0, 0, 0))
        od = ImageDraw.Draw(overlay)

        if has_image:
            for yi in range(H_px):
                t = yi / H_px
                if t < 0.70:
                    alpha = int(8 * t)
                else:
                    tt = (t - 0.70) / 0.30
                    alpha = int(8 + 220 * (tt ** 0.55))
                od.line([(0, yi), (W_px, yi)], fill=(0, 0, 0, alpha))
        else:
            grad_start = int(H_px * 0.40)
            for yi in range(grad_start, H_px):
                t = (yi - grad_start) / (H_px - grad_start)
                alpha = int(215 * (t ** 0.65))
                od.line([(0, yi), (W_px, yi)], fill=(0, 0, 0, alpha))

        bg = Image.alpha_composite(bg, overlay)
        draw = ImageDraw.Draw(bg, "RGBA")

        # ── Left accent bar ───────────────────────────────────────────────────────
        bar_w = max(8, int(W_px * 0.012))
        draw.rectangle([0, 0, bar_w, H_px], fill=acc_rgb + (255,))
        
        draw.rectangle(
            [bar_w, 0, bar_w + max(3, int(W_px * 0.004)), H_px],
            fill=acc2_rgb + (200,)
        )

        # ── Font loading ──────────────────────────────────────────────────────────
        def _load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
            candidates = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf" if bold else
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf" if bold else
                "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
                "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
                "/System/Library/Fonts/Helvetica.ttc",
            ]
            for path in candidates:
                if path and os.path.exists(path):
                    try:
                        return ImageFont.truetype(path, size)
                    except Exception:
                        continue
            return ImageFont.load_default()

        # ── Text helpers ──────────────────────────────────────────────────────────
        def _wrap_text(text: str, font: ImageFont.ImageFont, max_width: int) -> list:
            words, lines, current = text.split(), [], ""
            for w in words:
                test = (current + " " + w).strip()
                bb   = draw.textbbox((0, 0), test, font=font)
                if bb[2] - bb[0] <= max_width:
                    current = test
                else:
                    if current: 
                        lines.append(current)
                    current = w
            if current: 
                lines.append(current)
            return lines or [text]

        def _premium_shadow_text(xy: tuple, text: str, font, fill: tuple):
            base_x, base_y = xy
            draw.text((base_x + 6, base_y + 6), text, font=font, fill=(0, 0, 0, 40))
            draw.text((base_x + 3, base_y + 3), text, font=font, fill=(0, 0, 0, 90))
            draw.text((base_x + 1, base_y + 1), text, font=font, fill=(0, 0, 0, 180))
            draw.text(xy, text, font=font, fill=fill + (255,))

        # ── Genre badge ───────────────────────────────────────────────────────────
        genre = concept.get("genre_label", "").upper()[:20].strip()
        if genre:
            bf = _load_font(int(H_px * 0.014), bold=True)
            bx, by = int(W_px * 0.065), int(H_px * 0.032)
            
            bb = draw.textbbox((0, 0), genre, font=bf)
            bw = bb[2] - bb[0] + int(W_px * 0.035)
            bh = bb[3] - bb[1] + int(H_px * 0.010)
            
            draw.rounded_rectangle([bx, by, bx + bw, by + bh],
                                    radius=int(bh * 0.3), fill=(0, 0, 0, 160))
            draw.rounded_rectangle([bx, by, bx + int(W_px * 0.008), by + bh],
                                    radius=int(bh * 0.3), fill=acc_rgb + (255,))
            
            draw.text((bx + int(W_px * 0.012), by + int(H_px * 0.004)),
                      genre, font=bf, fill=acc_rgb + (255,))

        # ── Title ─────────────────────────────────────────────────────────────────
        title_lines_raw = [ln.strip() for ln in concept.get("title", "Book Title").split("\n") if ln.strip()]
        max_chars = max(len(ln) for ln in title_lines_raw)
        
        title_size = (
            int(H_px * 0.082) if max_chars <= 10 else
            int(H_px * 0.068) if max_chars <= 16 else
            int(H_px * 0.055) if max_chars <= 22 else
            int(H_px * 0.044)
        )
        
        title_font = _load_font(title_size, bold=True)
        text_left  = int(W_px * 0.065)
        text_width = W_px - text_left - int(W_px * 0.065)

        title_lines = []
        for raw_ln in title_lines_raw:
            title_lines.extend(_wrap_text(raw_ln, title_font, text_width))

        BOTTOM_BAND_H = int(H_px * 0.085)
        line_gap      = int(title_size * 1.15)
        total_title_h = len(title_lines) * line_gap

        if has_image:
            ty_start  = int(H_px * 0.705)
            ty_bottom = ty_start + total_title_h
        else:
            ty_bottom = H_px - BOTTOM_BAND_H - int(H_px * 0.025)
            ty_start  = ty_bottom - total_title_h

        for i, ln in enumerate(title_lines):
            _premium_shadow_text((text_left, ty_start + i * line_gap), ln, title_font, tcol_rgb)

        rule_y = ty_bottom + int(H_px * 0.008)
        
        draw.rectangle([text_left, rule_y, text_left + int(text_width * 0.72), rule_y + int(H_px * 0.004)],
                       fill=acc_rgb + (230,))
        draw.rectangle([text_left, rule_y + int(H_px * 0.007), text_left + int(text_width * 0.42), rule_y + int(H_px * 0.009)],
                       fill=acc2_rgb + (180,))

        # ── Subtitle ──────────────────────────────────────────────────────────────
        subtitle = concept.get("subtitle", "").strip()
        sub_y    = rule_y + int(H_px * 0.016)
        
        if subtitle:
            sub_font = _load_font(int(H_px * 0.026))
            for ln in _wrap_text(subtitle, sub_font, text_width):
                _premium_shadow_text((text_left, sub_y), ln, sub_font, scol_rgb)
                sub_y += int(H_px * 0.030)

        # ── Tagline ───────────────────────────────────────────────────────────────
        tagline = concept.get("tagline", "").strip()
        if tagline and has_image:
            tag_font = _load_font(int(H_px * 0.012))
            tag_y    = sub_y + int(H_px * 0.005)
            
            if tag_y < (H_px - BOTTOM_BAND_H - int(H_px * 0.04)):
                for ln in _wrap_text(tagline, tag_font, text_width):
                    draw.text((text_left, tag_y), ln, font=tag_font, fill=tgcol_rgb + (180,))
                    tag_y += int(H_px * 0.015)

        # ── Bottom author band ────────────────────────────────────────────────────
        band_top = H_px - BOTTOM_BAND_H
        band_col = tuple(max(0, int(c * 0.25)) for c in bg1_rgb) + (230,)
        
        draw.rectangle([0, band_top, W_px, H_px], fill=band_col)
        draw.rectangle([0, band_top, W_px, band_top + int(H_px * 0.003)], fill=acc_rgb + (255,))

        author_line = concept.get("author_line", "").strip()
        if author_line:
            auth_font = _load_font(int(H_px * 0.020))
            draw.text((text_left, band_top + int(BOTTOM_BAND_H * 0.32)),
                      author_line, font=auth_font, fill=scol_rgb + (230,))

        pub_font = _load_font(int(H_px * 0.016), bold=True)
        pub_text = "ENTERPRISE AI"
        pb = draw.textbbox((0, 0), pub_text, font=pub_font)
        draw.text((W_px - text_left - (pb[2] - pb[0]), band_top + int(BOTTOM_BAND_H * 0.32)),
                  pub_text, font=pub_font, fill=acc_rgb + (230,))

        # ── Top highlight bar ─────────────────────────────────────────────────────
        draw.rectangle([0, 0, W_px, int(H_px * 0.006)], fill=acc_rgb + (80,))

        # ── Embed into ReportLab PDF ──────────────────────────────────────────────
        final_buf = io.BytesIO()
        bg.convert("RGB").save(final_buf, format="JPEG", quality=92)
        final_buf.seek(0)

        c = rl_canvas.Canvas(output_path, pagesize=A4)
        c.drawImage(ImageReader(final_buf), 0, 0, width=W_pt, height=H_pt, preserveAspectRatio=False)
        c.showPage()
        c.save()
        
        return output_path
        
    except Exception as e:
        logger.error(f"  ⚠️  Error in render_cover_pdf: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# PDF cover renderer (legacy fallback — used when Pillow unavailable)
# ─────────────────────────────────────────────────────────────────────────────

def _render_cover_pdf_legacy(concept: dict, output_path: str,
                              book_image_bytes: bytes | None = None) -> str:
    try:
        from reportlab.pdfgen import canvas as rl_canvas   # pyrefly: ignore [missing-import]
        from reportlab.lib.pagesizes import A4              # pyrefly: ignore [missing-import]
        from reportlab.lib.units import mm                  # pyrefly: ignore [missing-import]

        W, H = A4   
        c = rl_canvas.Canvas(output_path, pagesize=A4)

        palette   = concept.get("palette", {})
        motif     = concept.get("motif", "none").lower()
        illus     = concept.get("illustration_shape", "large_circle").lower()
        style     = concept.get("style", "premium").lower()
        layout    = concept.get("layout_template", "split_horizon").lower()
        treatment = concept.get("image_treatment", "tinted_overlay").lower()
        acc_els   = concept.get("accent_elements", [])

        bg1   = _hex(palette.get("bg_primary",    "#0f172a"))
        bg2   = _hex(palette.get("bg_secondary",  "#1e3a5f"))
        panel = _hex(palette.get("panel_color",   "#1e293b"))
        acc   = _hex(palette.get("accent",        "#f59e0b"))
        acc2  = _hex(palette.get("accent2",       "#fbbf24"))
        tcol  = _hex(palette.get("title_color",   "#ffffff"))
        scol  = _hex(palette.get("subtitle_color","#e2e8f0"))
        tgcol = _hex(palette.get("tagline_color", "#94a3b8"))

        TEXT_LEFT = 20 * mm + 8

        processed_img = None
        if book_image_bytes:
            processed_img = _prepare_book_image(
                book_image_bytes, treatment, bg1, acc,
                width_px=int(W * 2), height_px=int(H * 2)
            )

        BANDS = 140
        for i in range(BANDS):
            t  = i / BANDS
            rc = _blend(bg1, bg2, t)
            c.setFillColorRGB(*rc)
            bh = H / BANDS + 1
            c.rect(0, H - (i + 1) * bh, W, bh + 1, fill=1, stroke=0)

        text_panel_x = 0
        text_panel_y = H * 0.22
        text_panel_w = W
        text_panel_h = H * 0.50

        if processed_img:
            if layout == "full_bleed":
                _draw_image_on_canvas(c, processed_img, 0, 0, W, H)
                text_panel_x = W * 0.08
                text_panel_w = W * 0.84
                text_panel_y = H * 0.20
                text_panel_h = H * 0.55

            elif layout == "left_panel":
                _draw_image_on_canvas(c, processed_img, W * 0.45, 0, W * 0.55, H)
                c.setFillColorRGB(*panel)
                c.rect(0, 0, W * 0.47, H, fill=1, stroke=0)
                
                text_panel_x = 16 * mm
                text_panel_y = H * 0.15
                text_panel_w = W * 0.42
                text_panel_h = H * 0.70

            elif layout == "top_image":
                _draw_image_on_canvas(c, processed_img, 0, H * 0.48, W, H * 0.52)
                c.setFillColorRGB(*_blend(bg1, (0,0,0), 0.2))
                c.rect(0, 0, W, H * 0.50, fill=1, stroke=0)
                
                text_panel_x = TEXT_LEFT
                text_panel_y = H * 0.05
                text_panel_w = W - text_panel_x - 16 * mm
                text_panel_h = H * 0.44

            elif layout == "diagonal_cut":
                c.saveState()
                path = c.beginPath()
                path.moveTo(0, 0)
                path.lineTo(W, 0)
                path.lineTo(W, H)
                path.lineTo(W * 0.30, H)
                path.close()
                c.clipPath(path, stroke=0)
                _draw_image_on_canvas(c, processed_img, 0, 0, W, H)
                c.restoreState()
                
                c.saveState()
                c.setFillColorRGB(*bg1)
                path2 = c.beginPath()
                path2.moveTo(0, 0)
                path2.lineTo(W * 0.60, 0)
                path2.lineTo(0, H)
                path2.close()
                c.drawPath(path2, fill=1, stroke=0)
                c.restoreState()
                
                text_panel_x = TEXT_LEFT
                text_panel_y = H * 0.22
                text_panel_w = W * 0.52
                text_panel_h = H * 0.55

            elif layout == "magazine":
                img_w = W * 0.38
                img_h = H * 0.28
                _draw_image_on_canvas(c, processed_img, W - img_w - 14*mm, H - img_h - 14*mm, img_w, img_h)
                
                c.setFillColorRGB(*_blend(bg1, (1,1,1), 0.06))
                c.rect(0, 0, W * 0.58, H, fill=1, stroke=0)
                
                text_panel_x = TEXT_LEFT
                text_panel_y = H * 0.18
                text_panel_w = W * 0.56
                text_panel_h = H * 0.60

            else:  
                _draw_image_on_canvas(c, processed_img, 0, 0, W, H * 0.55)
                text_panel_x = 0
                text_panel_y = H * 0.50
                text_panel_w = W
                text_panel_h = H * 0.46

        rng = random.Random(hash(concept.get("title", "")) & 0xFFFFFF)
        _draw_motif(c, motif, acc, W, H, rng)
        _draw_accent_elements(c, acc_els, acc, acc2, bg1, W, H, rng)

        if layout not in ("left_panel", "magazine"):
            shape_x = W * 0.75 if layout in ("split_horizon", "full_bleed") else W * 0.80
            shape_y = H * 0.60 if layout == "split_horizon" else H * 0.50
            _draw_illustration(c, illus, acc, shape_x, shape_y)

        c.setFillColorRGB(*acc)
        c.rect(0, 0, 6.5, H, fill=1, stroke=0)
        c.setFillColorRGB(*acc2)
        c.rect(6.5, 0, 2.5, H, fill=1, stroke=0)

        if layout not in ("left_panel",):
            c.saveState()
            c.setFillAlpha(0.78)
            c.setFillColorRGB(*panel)
            c.rect(text_panel_x, text_panel_y, text_panel_w, text_panel_h, fill=1, stroke=0)
            c.restoreState()
            
            c.saveState()
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(2.0)
            c.line(text_panel_x, text_panel_y + text_panel_h, text_panel_x + text_panel_w, text_panel_y + text_panel_h)
            c.line(text_panel_x, text_panel_y, text_panel_x + text_panel_w, text_panel_y)
            c.restoreState()

        genre = concept.get("genre_label", "").upper()[:20].strip()
        if genre:
            bx = 20 * mm
            by = H - 24 * mm
            bw = len(genre) * 7.0 + 28
            bh = 21
            c.setFillColorRGB(*_blend(bg1, (0,0,0), 0.40))
            c.roundRect(bx, by, bw, bh, 3, fill=1, stroke=0)
            c.setFillColorRGB(*acc)
            c.roundRect(bx, by, 5, bh, 2, fill=1, stroke=0)
            c.setFillColorRGB(*acc)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(bx + 10, by + 6, genre)

        title_lines = [ln.strip() for ln in concept.get("title", "Book Title").split("\n") if ln.strip()]
        max_line    = max(len(ln) for ln in title_lines)
        font_size   = 50 if max_line <= 12 else 42 if max_line <= 18 else 34 if max_line <= 26 else 26
        title_font  = _get_cover_font(title_lines[0], bold=True)

        ty = text_panel_y + text_panel_h - 16
        tx = text_panel_x + (16 * mm if text_panel_x < 20 else 8 * mm)

        c.setFillColorRGB(*tcol)
        c.setFont(title_font, font_size)
        for line in title_lines:
            ty -= (font_size + 5)
            c.drawString(tx, ty, line)
        ty -= 8

        rule_w = min(text_panel_w - 30, W - tx - 16 * mm)
        c.setFillColorRGB(*acc)
        c.rect(tx, ty, rule_w, 3, fill=1, stroke=0)
        c.setFillColorRGB(*acc2)
        c.rect(tx, ty - 4, rule_w * 0.55, 1.5, fill=1, stroke=0)
        ty -= 16

        subtitle = concept.get("subtitle", "").strip()
        if subtitle:
            sub_font = _get_cover_font(subtitle)
            c.setFillColorRGB(*scol)
            c.setFont(sub_font, 15)
            for ln in _wrap(subtitle, max_chars=int(rule_w / 8.5)):
                ty -= 21
                c.drawString(tx, ty, ln)
            ty -= 8

        tagline = concept.get("tagline", "").strip()
        if tagline:
            tag_font = _get_cover_font(tagline)
            c.setFillColorRGB(*tgcol)
            c.setFont(tag_font, 10.5)
            for ln in _wrap(tagline, max_chars=int(rule_w / 6.2)):
                ty -= 15
                c.drawString(tx, ty, ln)

        BAND_H = 20 * mm
        bot_band = _blend(bg1, (0,0,0), 0.60)
        c.setFillColorRGB(*bot_band)
        c.rect(0, 0, W, BAND_H, fill=1, stroke=0)
        c.setFillColorRGB(*acc)
        c.rect(0, BAND_H, W, 2, fill=1, stroke=0)

        author_line = concept.get("author_line", "").strip()
        if author_line:
            af = _get_cover_font(author_line)
            c.setFillColorRGB(*scol)
            c.setFont(af, 9.5)
            c.drawString(20 * mm, BAND_H * 0.38, author_line)
            
        c.setFillColorRGB(*acc)
        c.setFont("Helvetica-Bold", 8.5)
        c.drawRightString(W - 20 * mm, BAND_H * 0.38, "EDITORIAL AI")

        c.saveState()
        c.setFillAlpha(0.28)
        c.setFillColorRGB(*acc)
        c.rect(0, H - 5, W, 5, fill=1, stroke=0)
        c.restoreState()

        c.showPage()
        c.save()
        return output_path
    except Exception as e:
        logger.error(f"  ⚠️  Error in _render_cover_pdf_legacy: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# Motif renderer (Fully Expanded PEP-8)
# ─────────────────────────────────────────────────────────────────────────────

def _draw_motif(c, motif: str, acc: tuple, W: float, H: float, rng: random.Random):
    if motif == "concentric_circles":
        for radius in range(20, int(W * 0.9), 30):
            c.saveState()
            c.setStrokeAlpha(max(0.02, 0.09 - radius * 0.00015))
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(0.7)
            c.circle(W * 0.5, H * 0.5, radius, fill=0, stroke=1)
            c.restoreState()

    elif motif == "diagonal_stripes":
        for x in range(-int(H), int(W) + int(H), 22):
            c.saveState()
            c.setStrokeAlpha(0.07)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(1.1)
            c.line(x, 0, x + H, H)
            c.restoreState()

    elif motif == "scattered_dots":
        for _ in range(180):
            c.saveState()
            c.setFillAlpha(rng.uniform(0.03, 0.15))
            c.setFillColorRGB(*acc)
            c.circle(rng.uniform(0, W), rng.uniform(0, H), rng.uniform(1.5, 5), fill=1, stroke=0)
            c.restoreState()

    elif motif == "grid_lines":
        for x in range(0, int(W), 26):
            c.saveState()
            c.setStrokeAlpha(0.055)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(0.5)
            c.line(x, 0, x, H)
            c.restoreState()
            
        for y in range(0, int(H), 26):
            c.saveState()
            c.setStrokeAlpha(0.055)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(0.5)
            c.line(0, y, W, y)
            c.restoreState()

    elif motif == "wave_curves":
        for offset in range(-100, 340, 28):
            c.saveState()
            c.setStrokeAlpha(0.09)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(1.4)
            path = c.beginPath()
            path.moveTo(0, H * 0.45 + offset)
            for px in range(0, int(W) + 10, 5):
                py = H * 0.45 + offset + math.sin(px * 0.018) * 52
                path.lineTo(px, py)
            c.drawPath(path, stroke=1, fill=0)
            c.restoreState()

    elif motif == "hexagons":
        hr = 20
        for row in range(int(H / (hr * 1.52)) + 2):
            for col in range(int(W / (hr * 1.75)) + 2):
                hx = col * hr * 1.75 + (hr * 0.9 if row % 2 else 0)
                hy = row * hr * 1.52
                pts = [(hx + hr * math.cos(math.radians(60*i+30)),
                        hy + hr * math.sin(math.radians(60*i+30))) for i in range(6)]
                        
                c.saveState()
                c.setStrokeAlpha(0.055)
                c.setStrokeColorRGB(*acc)
                c.setLineWidth(0.6)
                
                path = c.beginPath()
                path.moveTo(*pts[0])
                for pt in pts[1:]: 
                    path.lineTo(*pt)
                path.close()
                
                c.drawPath(path, stroke=1, fill=0)
                c.restoreState()

    elif motif == "triangles":
        ts = 46
        for row in range(0, int(H)+ts, ts):
            for col in range(0, int(W)+ts, ts):
                c.saveState()
                c.setStrokeAlpha(0.055)
                c.setStrokeColorRGB(*acc)
                c.setLineWidth(0.6)
                
                path = c.beginPath()
                if (row//ts + col//ts) % 2 == 0:
                    path.moveTo(col, row)
                    path.lineTo(col+ts, row)
                    path.lineTo(col+ts/2, row+ts)
                else:
                    path.moveTo(col+ts/2, row)
                    path.lineTo(col, row+ts)
                    path.lineTo(col+ts, row+ts)
                path.close()
                
                c.drawPath(path, stroke=1, fill=0)
                c.restoreState()

    elif motif == "stars":
        for _ in range(90):
            c.saveState()
            c.setFillAlpha(rng.uniform(0.05, 0.25))
            c.setFillColorRGB(*acc)
            c.circle(rng.uniform(0, W), rng.uniform(0, H), rng.uniform(0.8, 2.8), fill=1, stroke=0)
            c.restoreState()
            
        for _ in range(14):
            sx, sy = rng.uniform(0, W), rng.uniform(0, H)
            c.saveState()
            c.setStrokeAlpha(0.13)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(0.5)
            
            for a in range(0, 360, 45):
                r = math.radians(a)
                c.line(sx+math.cos(r)*3, sy+math.sin(r)*3, sx+math.cos(r)*10, sy+math.sin(r)*10)
            c.restoreState()

    elif motif == "arcs":
        for ox, oy in [(0,0),(W,0),(0,H),(W,H)]:
            for r in range(40, 320, 36):
                c.saveState()
                c.setStrokeAlpha(max(0.03, 0.13-r*0.0003))
                c.setStrokeColorRGB(*acc)
                c.setLineWidth(0.9)
                c.arc(ox-r, oy-r, ox+r, oy+r, startAng=0, extent=90)
                c.restoreState()

    elif motif == "circuit_traces":
        for _ in range(25):
            x0 = rng.uniform(0, W)
            y0 = rng.uniform(0, H)
            
            c.saveState()
            c.setStrokeAlpha(0.10)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(0.8)
            
            path = c.beginPath()
            path.moveTo(x0, y0)
            for seg in range(rng.randint(2, 5)):
                if rng.random() > 0.5:
                    x0 += rng.choice([-1, 1]) * rng.uniform(20, 90)
                else:
                    y0 += rng.choice([-1, 1]) * rng.uniform(20, 90)
                path.lineTo(x0, y0)
                
            c.drawPath(path, stroke=1, fill=0)
            c.restoreState()
            
            c.saveState()
            c.setFillAlpha(0.20)
            c.setFillColorRGB(*acc)
            c.circle(x0, y0, 2.5, fill=1, stroke=0)
            c.restoreState()

    elif motif == "halftone":
        spacing = 14
        for row in range(0, int(H)+spacing, spacing):
            for col in range(0, int(W)+spacing, spacing):
                dist = math.sqrt((col - W/2)**2 + (row - H/2)**2)
                r = max(0.5, 4.5 - dist * 0.006)
                
                c.saveState()
                c.setFillAlpha(0.07)
                c.setFillColorRGB(*acc)
                c.circle(col, row, r, fill=1, stroke=0)
                c.restoreState()

    elif motif == "ink_drops":
        for _ in range(60):
            c.saveState()
            c.setFillAlpha(rng.uniform(0.04, 0.14))
            c.setFillColorRGB(*acc)
            rx = rng.uniform(3, 14)
            ry = rng.uniform(3, 14)
            bx = rng.uniform(0, W)
            by = rng.uniform(0, H)
            c.ellipse(bx - rx, by - ry, bx + rx, by + ry, fill=1, stroke=0)
            c.restoreState()


# ─────────────────────────────────────────────────────────────────────────────
# Illustration shape renderer (Fully Expanded PEP-8)
# ─────────────────────────────────────────────────────────────────────────────

def _draw_illustration(c, illus: str, acc: tuple, cx: float, cy: float):
    if illus == "large_circle":
        for radius, alpha in [(190, 0.14), (140, 0.18), (90, 0.26), (48, 0.38)]:
            c.saveState()
            c.setFillAlpha(alpha)
            c.setFillColorRGB(*acc)
            c.circle(cx, cy, radius, fill=1, stroke=0)
            c.restoreState()
            
        c.saveState()
        c.setFillAlpha(0.52)
        c.setFillColorRGB(*acc)
        c.circle(cx, cy, 26, fill=1, stroke=0)
        c.restoreState()

    elif illus == "diamond":
        for size, alpha in [(255, 0.11), (175, 0.17), (105, 0.25), (52, 0.42)]:
            c.saveState()
            c.setFillAlpha(alpha)
            c.setFillColorRGB(*acc)
            c.translate(cx, cy)
            c.rotate(45)
            c.rect(-size/2, -size/2, size, size, fill=1, stroke=0)
            c.restoreState()

    elif illus == "arch":
        for ow, oh, alpha in [(170, 230, 0.18), (220, 290, 0.10)]:
            c.saveState()
            c.setFillAlpha(alpha)
            c.setFillColorRGB(*acc)
            c.roundRect(cx-ow/2, cy-oh/2, ow, oh, ow/2, fill=1, stroke=0)
            c.restoreState()

    elif illus == "triangle":
        for size, alpha in [(300, 0.09), (210, 0.15), (130, 0.23), (68, 0.38)]:
            c.saveState()
            c.setFillAlpha(alpha)
            c.setFillColorRGB(*acc)
            
            path = c.beginPath()
            path.moveTo(cx, cy+size*0.6)
            path.lineTo(cx-size*0.52, cy-size*0.4)
            path.lineTo(cx+size*0.52, cy-size*0.4)
            path.close()
            
            c.drawPath(path, fill=1, stroke=0)
            c.restoreState()

    elif illus == "sunburst":
        for angle in range(0, 360, 14):
            r = math.radians(angle)
            c.saveState()
            c.setStrokeAlpha(0.16)
            c.setStrokeColorRGB(*acc)
            c.setLineWidth(3.0)
            c.line(cx+math.cos(r)*34, cy+math.sin(r)*34,
                   cx+math.cos(r)*205, cy+math.sin(r)*205)
            c.restoreState()
            
        c.saveState()
        c.setFillAlpha(0.38)
        c.setFillColorRGB(*acc)
        c.circle(cx, cy, 36, fill=1, stroke=0)
        c.restoreState()

    elif illus == "cross_lines":
        c.saveState()
        c.setFillAlpha(0.16)
        c.setFillColorRGB(*acc)
        c.rect(cx-9, cy-195, 18, 390, fill=1, stroke=0)
        c.rect(cx-195, cy-9, 390, 18, fill=1, stroke=0)
        c.restoreState()
        
        c.saveState()
        c.setFillAlpha(0.32)
        c.setFillColorRGB(*acc)
        c.circle(cx, cy, 28, fill=1, stroke=0)
        c.restoreState()

    elif illus == "polygon":
        for scale, alpha in [(160, 0.12), (110, 0.18), (65, 0.30)]:
            pts = [(cx + scale * math.cos(math.radians(60*i-30)),
                    cy + scale * math.sin(math.radians(60*i-30))) for i in range(6)]
                    
            c.saveState()
            c.setFillAlpha(alpha)
            c.setFillColorRGB(*acc)
            
            path = c.beginPath()
            path.moveTo(*pts[0])
            for pt in pts[1:]: 
                path.lineTo(*pt)
            path.close()
            
            c.drawPath(path, fill=1, stroke=0)
            c.restoreState()


# ─────────────────────────────────────────────────────────────────────────────
# Accent elements renderer
# ─────────────────────────────────────────────────────────────────────────────

def _draw_accent_elements(c, elements: list, acc: tuple, acc2: tuple,
                           bg: tuple, W: float, H: float, rng: random.Random):
    """Render up to 4 named accent elements for per-book personalisation."""
    for el in elements[:4]:
        el = el.lower()
        if "constellation" in el or "star" in el:
            stars = [(rng.uniform(W*0.4, W*0.95), rng.uniform(H*0.4, H*0.95)) for _ in range(12)]
            for sx, sy in stars:
                c.saveState()
                c.setFillAlpha(0.30)
                c.setFillColorRGB(*acc2)
                c.circle(sx, sy, rng.uniform(1.5, 3.5), fill=1, stroke=0)
                c.restoreState()
                
            for i in range(len(stars)-1):
                c.saveState()
                c.setStrokeAlpha(0.12)
                c.setStrokeColorRGB(*acc)
                c.setLineWidth(0.6)
                c.line(stars[i][0], stars[i][1], stars[i+1][0], stars[i+1][1])
                c.restoreState()

        elif "circuit" in el or "trace" in el:
            for _ in range(8):
                x0, y0 = rng.uniform(0, W), rng.uniform(0, H)
                c.saveState()
                c.setStrokeAlpha(0.14)
                c.setStrokeColorRGB(*acc)
                c.setLineWidth(1.0)
                
                path = c.beginPath()
                path.moveTo(x0, y0)
                
                x0 += rng.choice([-1,1]) * rng.uniform(30, 80)
                path.lineTo(x0, y0)
                y0 += rng.choice([-1,1]) * rng.uniform(30, 80)
                path.lineTo(x0, y0)
                
                c.drawPath(path, stroke=1, fill=0)
                c.restoreState()

        elif "watercolour" in el or "watercolor" in el or "wash" in el:
            for _ in range(6):
                c.saveState()
                c.setFillAlpha(rng.uniform(0.04, 0.10))
                c.setFillColorRGB(*acc)
                rx = rng.uniform(60, 160)
                ry = rng.uniform(40, 120)
                bx = rng.uniform(0, W)
                by = rng.uniform(0, H)
                c.ellipse(bx-rx, by-ry, bx+rx, by+ry, fill=1, stroke=0)
                c.restoreState()

        elif "ink" in el or "splatter" in el:
            for _ in range(30):
                c.saveState()
                c.setFillAlpha(rng.uniform(0.06, 0.18))
                c.setFillColorRGB(*acc)
                r = rng.uniform(1.5, 8)
                c.circle(rng.uniform(0,W), rng.uniform(0,H), r, fill=1, stroke=0)
                c.restoreState()

        elif "rule" in el or "line" in el:
            for i in range(3):
                ly = H * (0.30 + i * 0.15)
                c.saveState()
                c.setStrokeAlpha(0.22)
                c.setStrokeColorRGB(*acc2)
                c.setLineWidth(0.8)
                c.line(20*6, ly, W - 20*6, ly)
                c.restoreState()

        elif "dot" in el or "grid" in el:
            for gx in range(0, int(W), 18):
                for gy in range(0, int(H), 18):
                    c.saveState()
                    c.setFillAlpha(0.055)
                    c.setFillColorRGB(*acc)
                    c.circle(gx, gy, 0.9, fill=1, stroke=0)
                    c.restoreState()


# ─────────────────────────────────────────────────────────────────────────────
# Word-wrap helper
# ─────────────────────────────────────────────────────────────────────────────

def _wrap(text: str, max_chars: int = 50) -> list[str]:
    words = text.split()
    lines, cur = [], []
    for w in words:
        if len(" ".join(cur + [w])) > max_chars:
            if cur: 
                lines.append(" ".join(cur))
            cur = [w]
        else:
            cur.append(w)
            
    if cur: 
        lines.append(" ".join(cur))
        
    return lines or [""]


# ─────────────────────────────────────────────────────────────────────────────
# PDF prepend
# ─────────────────────────────────────────────────────────────────────────────

def prepend_cover_to_pdf(cover_pdf: str, original_pdf: str, output_pdf: str) -> str:
    """Add cover as page 1, then ALL pages of original (never replaces them)."""
    from pypdf import PdfWriter, PdfReader   # pyrefly: ignore [missing-import]
    writer = PdfWriter()
    
    for page in PdfReader(cover_pdf).pages:
        writer.add_page(page)
        
    for page in PdfReader(original_pdf).pages:
        writer.add_page(page)
        
    with open(output_pdf, "wb") as f:
        writer.write(f)
        
    return output_pdf


def replace_first_page_of_pdf(cover_pdf: str, original_pdf: str, output_pdf: str) -> str:
    """Replace the first page of original_pdf with the cover page from cover_pdf."""
    from pypdf import PdfWriter, PdfReader   # pyrefly: ignore [missing-import]
    writer  = PdfWriter()
    cover   = PdfReader(cover_pdf)
    content = PdfReader(original_pdf)

    writer.add_page(cover.pages[0])

    for page in content.pages[1:]:
        writer.add_page(page)

    with open(output_pdf, "wb") as f:
        writer.write(f)
        
    return output_pdf


# ─────────────────────────────────────────────────────────────────────────────
# DOCX cover renderer + prepend
# ─────────────────────────────────────────────────────────────────────────────

def prepend_cover_to_docx(concept: dict, original_docx: str, output_docx: str,
                          cover_image_bytes: bytes | None = None) -> str:
    try:
        from docx import Document                           # pyrefly: ignore [missing-import]
        from docx.shared import Pt, RGBColor, Cm           # pyrefly: ignore [missing-import]
        from docx.enum.text import WD_ALIGN_PARAGRAPH      # pyrefly: ignore [missing-import]
        from docx.oxml.ns import qn                        # pyrefly: ignore [missing-import]
        from docx.oxml import OxmlElement                  # pyrefly: ignore [missing-import]
        import copy

        palette = concept.get("palette", {})

        def rgb(key: str, fallback: str = "#1a1a1a") -> RGBColor:
            h = palette.get(key, fallback).lstrip("#")
            if len(h) == 3: 
                h = "".join(c*2 for c in h)
            return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

        from docx.shared import Inches, Emu              # pyrefly: ignore [missing-import]
        from docx.oxml.ns import nsmap                      # pyrefly: ignore [missing-import]
        from lxml import etree                              # pyrefly: ignore [missing-import]

        # ── Rasterise SVG cover image to JPEG if needed ──────────────────────
        # DOCX cannot embed SVG directly; convert to JPEG via Pillow+cairosvg
        # or fall back to a gradient JPEG generated from the palette.
        def _cover_img_bytes_to_jpeg(raw: bytes) -> bytes | None:
            """Convert any cover bytes (JPEG, PNG, or SVG) to JPEG for DOCX embedding."""
            try:
                if raw[:5] in (b"<?xml", b"<svg ") or b"<svg" in raw[:100]:
                    # SVG -> PNG via cairosvg, then JPEG via Pillow
                    try:
                        import cairosvg                     # pyrefly: ignore [missing-import]
                        from PIL import Image               # pyrefly: ignore [missing-import]
                        png_data = cairosvg.svg2png(bytestring=raw, output_width=595, output_height=842)
                        img = Image.open(io.BytesIO(png_data)).convert("RGB")
                    except ImportError:
                        logger.warning("  cairosvg/PIL not available; SVG->JPEG conversion skipped.")
                        return None
                else:
                    from PIL import Image                   # pyrefly: ignore [missing-import]
                    img = Image.open(io.BytesIO(raw)).convert("RGB")
                # Resize to A4 at 150 DPI: 595 x 842 pt = ~1240 x 1754 px
                img = img.resize((1240, 1754), Image.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=88)
                return buf.getvalue()
            except Exception as ex:
                logger.warning(f"  Cover image conversion failed: {ex}")
                return None

        def _make_gradient_jpeg(palette: dict) -> bytes:
            """Generate a palette-based gradient JPEG when no image is available."""
            try:
                from PIL import Image, ImageDraw            # pyrefly: ignore [missing-import]
                def _hx(h):
                    h = h.lstrip("#")
                    if len(h)==3: h="".join(c*2 for c in h)
                    return (int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))
                w, h = 1240, 1754
                bg1 = _hx(palette.get("bg_primary",   "#0f172a"))
                bg2 = _hx(palette.get("bg_secondary", "#1e3a5f"))
                img = Image.new("RGB", (w, h))
                draw = ImageDraw.Draw(img)
                for yi in range(h):
                    t = yi / h
                    r = int(bg1[0]*(1-t) + bg2[0]*t)
                    g = int(bg1[1]*(1-t) + bg2[1]*t)
                    b = int(bg1[2]*(1-t) + bg2[2]*t)
                    draw.line([(0,yi),(w,yi)], fill=(r,g,b))
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=85)
                return buf.getvalue()
            except Exception:
                return b""  # absolute last resort: empty

        # Convert/prepare the cover image for embedding
        jpeg_cover: bytes | None = None
        if cover_image_bytes:
            jpeg_cover = _cover_img_bytes_to_jpeg(cover_image_bytes)
            if not jpeg_cover:
                jpeg_cover = _make_gradient_jpeg(palette)
        else:
            jpeg_cover = _make_gradient_jpeg(palette)

        # ── Helper: set paragraph page background via XML shading ────────────
        def _set_para_shading(para, fill_hex: str):
            """Apply solid background fill to a paragraph (simulates bg colour)."""
            h = fill_hex.lstrip("#")
            if len(h)==3: h="".join(c*2 for c in h)
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  h.upper())
            pPr.append(shd)

        cover_doc = Document()
        sec = cover_doc.sections[0]
        sec.page_height = Cm(29.7)
        sec.page_width  = Cm(21.0)
        sec.left_margin = Cm(0.0)
        sec.right_margin= Cm(0.0)
        sec.top_margin  = Cm(0.0)
        sec.bottom_margin= Cm(0.0)

        # ── Embed cover image as full-page inline picture ─────────────────────
        if jpeg_cover:
            try:
                img_para = cover_doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.paragraph_format.space_before = Pt(0)
                img_para.paragraph_format.space_after  = Pt(0)
                run = img_para.add_run()
                run.add_picture(io.BytesIO(jpeg_cover), width=Cm(21.0), height=Cm(29.7))
                cover_doc.add_page_break()
                logger.info("  DOCX: full-bleed cover image embedded successfully.")
            except Exception as img_err:
                logger.warning(f"  DOCX image embed failed: {img_err}; falling back to text cover.")
                # Reset: start a fresh document for text-only fallback
                cover_doc = Document()
                sec = cover_doc.sections[0]
                sec.page_height = Cm(29.7)
                sec.page_width  = Cm(21.0)
        else:
            # No image — set page background colour via document XML
            pass

        # Restore margins for the text content pages
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
        sec.top_margin    = Cm(3.0)
        sec.bottom_margin = Cm(2.0)

        def add_para(text, size, bold=False, italic=False,
                     color_key="title_color", fallback="#ffffff",
                     align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
            p = cover_doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after  = Pt(sa)
            
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = rgb(color_key, fallback)

        def add_rule(color_key="accent", fallback="#f59e0b"):
            h = palette.get(color_key, fallback).lstrip("#")
            if len(h)==3: 
                h="".join(x*2 for x in h)
                
            p = cover_doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bt = OxmlElement("w:bottom")
            
            bt.set(qn("w:val"),"single")
            bt.set(qn("w:sz"),"16")
            bt.set(qn("w:space"),"1")
            bt.set(qn("w:color"),h)
            
            pBdr.append(bt)
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        genre = concept.get("genre_label","").upper().strip()
        if genre:
            add_para(f"— {genre} —", 9, bold=True, color_key="accent", fallback="#f59e0b", sa=2)
            add_para(concept.get("style","").upper(), 8, color_key="accent2", fallback=palette.get("accent","#f59e0b"), sa=2)

        for _ in range(3): 
            cover_doc.add_paragraph()

        for line in [ln.strip() for ln in concept.get("title","").split("\n") if ln.strip()]:
            add_para(line, 38, bold=True, color_key="title_color", fallback="#ffffff", sa=4)
            
        add_rule()

        subtitle = concept.get("subtitle","").strip()
        if subtitle:
            add_para(subtitle, 15, color_key="subtitle_color", fallback="#e2e8f0", sb=4, sa=8)

        tagline = concept.get("tagline","").strip()
        if tagline:
            add_para(tagline, 10.5, italic=True, color_key="tagline_color", fallback="#94a3b8", sa=6)

        rationale = concept.get("design_rationale","").strip()
        if rationale:
            add_para(rationale, 8, italic=True, color_key="tagline_color", fallback="#94a3b8", sa=4)

        for _ in range(5): 
            cover_doc.add_paragraph()
            
        add_rule()
        
        author_line = concept.get("author_line","").strip()
        if author_line:
            add_para(author_line, 11, color_key="subtitle_color", fallback="#e2e8f0", sb=6)

        cover_doc.add_page_break()
        tmp = output_docx + ".covertmp.docx"
        cover_doc.save(tmp)

        orig = Document(original_docx)
        out  = Document(tmp)
        
        for el in orig.element.body:
            if el.tag == qn("w:sectPr"): 
                continue
            out.element.body.append(copy.deepcopy(el))
            
        out.save(output_docx)
        if os.path.exists(tmp): 
            os.remove(tmp)
            
        return output_docx
        
    except Exception as e:
        logger.error(f"  ⚠️  Error in prepend_cover_to_docx: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# Main orchestrator
# ─────────────────────────────────────────────────────────────────────────────

def design_cover(
    file_path         : str,
    filename          : str,
    output_dir        : str,
    book_title        : str = "",
    description       : str = "",
    design_style      : str = "",
    cover_image_bytes : bytes | None = None,
) -> dict:
    """
    Full pipeline:
      1. Infer title from filename if not provided
      2. Extract TEXT from the book (fed into imaging cluster for scene-specific art)
      3. Extract a page image from the book (visual input for GPT-4o)
      4. Generate deeply personalised AI cover concept (text + image aware)
      5. Generate background via 5-Tier Multi-Engine failover cluster
      6. Render cover PDF with premium soft-shadow typography
      7. Prepend cover to original PDF or DOCX file
    """
    try:
        os.makedirs(output_dir, exist_ok=True)
        ext = os.path.splitext(filename)[1].lower()

        if not book_title:
            book_title = Path(filename).stem.replace("_"," ").replace("-"," ").title()

        logger.info("  📖 Extracting book text for content-aware cover design…")
        book_text = extract_book_text(file_path, ext)
        if book_text:
            logger.info(f"  ✅ Book text extracted ({len(book_text)} chars)")
        else:
            logger.info("  ℹ️  No book text extracted; using title only")

        logger.info("  📄 Extracting book page image for visual personalisation…")
        book_image = _extract_book_image(file_path, ext)
        if book_image:
            logger.info(f"  ✅ Book image extracted ({len(book_image)//1024} KB)")
        else:
            logger.info("  ℹ️  No book image extracted")

        concept = generate_cover_concept(
            book_title, description, design_style,
            book_image=book_image,
            book_text=book_text,
        )

        job_id   = uuid.uuid4().hex
        out_path = os.path.join(output_dir, f"cover_{job_id}{ext}")

        if ext == ".pdf":
            cover_pdf = os.path.join(output_dir, f"coverpage_{job_id}.pdf")

            generated_illustration: bytes | None = None
            if cover_image_bytes:
                logger.info("  🖼️  User-supplied cover image provided — skipping generation cluster step.")
            else:
                logger.info("  🖼️  No cover image supplied — engaging Nano Banana 6-Tier image generation cluster…")
                generated_illustration = generate_cover_image(concept, book_title, book_text)
                if generated_illustration:
                    logger.info(f"  ✅ Image pipeline resolved successfully ({len(generated_illustration)//1024} KB)")
                else:
                    logger.warning("  ⚠️ All 6 Nano Banana tiers offline — cover using gradient background.")
                    concept["_dalle_failed"] = True
                    concept["_dalle_note"] = "Image generation unavailable — cover uses gradient background. Check your Nano Banana / Gemini API key."

            render_cover_pdf(concept, cover_pdf,
                             book_image_bytes=book_image,
                             dalle_image_bytes=generated_illustration,
                             cover_image_bytes=cover_image_bytes)

            logger.info(f"  📎 Prepending cover to original PDF ({os.path.basename(file_path)})…")
            prepend_cover_to_pdf(cover_pdf, file_path, out_path)
            logger.info(f"  ✅ Output written: {out_path}")

            if os.path.exists(cover_pdf):
                os.remove(cover_pdf)

        elif ext == ".docx":
            # Generate cover image for DOCX the same way as PDF
            docx_illustration: bytes | None = None
            if cover_image_bytes:
                docx_illustration = cover_image_bytes
                logger.info("  🖼️  User-supplied cover image will be embedded in DOCX cover.")
            else:
                logger.info("  🖼️  Generating cover image for DOCX via Nano Banana 6-Tier cluster…")
                docx_illustration = generate_cover_image(concept, book_title, book_text)
                if docx_illustration:
                    logger.info(f"  ✅ Cover image ready for DOCX ({len(docx_illustration)//1024} KB)")
                else:
                    logger.warning("  ⚠️  All 6 Nano Banana tiers offline for DOCX — text-only layout.")
            prepend_cover_to_docx(concept, file_path, out_path,
                                  cover_image_bytes=docx_illustration)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Upload a .pdf or .docx")

        return {
            "output_path": out_path,
            "concept":     concept,
            "ext":         ext,
            "job_id":      job_id,
        }
    except Exception as e:
        logger.error(f"  🚨 CRITICAL ERROR in design_cover: {e}\n{traceback.format_exc()}")
        raise