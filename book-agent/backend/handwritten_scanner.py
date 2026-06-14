"""
handwritten_scanner.py
AI-powered handwritten book scanner.
- Accepts: image files (jpg/png/webp/heic/bmp), PDF, DOCX, or ZIP of any above
- Uses GPT-4o vision to transcribe each page/image
- Assembles transcription into a clean, structured book
- Exports PDF + DOCX
Supports any language — GPT-4o handles multilingual handwriting (English, Hindi, etc).

MAXIMIZED ENTERPRISE UPGRADES:
  - Layout & Reading Order Analysis (Multi-column & sidebar support)
  - Anti-Hallucination & Garbage Detection Guard
  - Adaptive Binarization & Lighting Normalization Pre-Processing
  - Context Healer Agent to stitch broken sentences across page boundaries
  - Dynamic Unicode/Devanagari Font Mapping for flawless PDF/DOCX generation
"""

import os
import io
import json
import re
import uuid
import time
import zipfile
import shutil
import base64
import tempfile
import traceback
import threading
import unicodedata
import hashlib
import urllib.request
from collections import Counter
from pathlib import Path
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

# pyrefly: ignore [missing-import]
from openai import OpenAI
# pyrefly: ignore [missing-import]
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
MODEL = "gpt-4o"

SUPPORTED_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff", ".tif"}
SUPPORTED_UPLOAD_EXTS = SUPPORTED_IMAGE_EXTS | {".pdf", ".docx", ".zip"}

# ─────────────────────────────────────────────────────────────────────────────
# Tuning knobs
# ─────────────────────────────────────────────────────────────────────────────

# How many pages to send per vision API call.
# MUST stay at 1 — sending multiple pages relies on GPT-4o splitting on
# ---PAGE_BREAK--- which it does inconsistently, causing pages to be silently
# merged or lost. 1 page per call is slower but reliable.
PAGES_PER_BATCH = 1

# Max output tokens per transcription call.
# gpt-4o's hard limit is 16384 output tokens — anything higher causes an
# immediate 400 Bad Request on EVERY call (no retry, since 400s aren't
# treated as transient). 16384 is plenty for one page of dense handwriting.
TRANSCRIPTION_MAX_TOKENS = 16384

# Number of parallel workers for transcription.
# Keep at 4 — 8 workers hammers the OpenAI rate limit on standard tiers
# (gpt-4o Vision is ~100 RPM on Tier 1). 4 workers with retries is safer.
TRANSCRIPTION_WORKERS = 4

# Global semaphore: hard cap on simultaneous OpenAI calls across ALL workers.
# Prevents a single large job from exhausting rate limits and causing all
# workers to 429-backoff simultaneously, stalling the entire pipeline.
_OPENAI_SEMAPHORE = threading.Semaphore(4)

# Retries on transient API errors (429, 500, 502, 503).
MAX_RETRIES = 3
RETRY_BASE_DELAY = 2.0  

# Structuring: max chars per chunk fed to GPT-4o.
# Larger = fewer chunks = fewer chapter-boundary stitching errors.
STRUCTURE_CHUNK_SIZE = 40_000

# Max tokens for the structuring / chapter-assembly call.
STRUCTURE_MAX_TOKENS = 16384

# PDF render zoom. 3.0 ≈ 225 DPI (better for dense writing).
PDF_RENDER_ZOOM = 3.0


# ─────────────────────────────────────────────────────────────────────────────
# Unicode / Devanagari font system
# ─────────────────────────────────────────────────────────────────────────────

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_FONTS_DIR   = os.path.join(_SCRIPT_DIR, "fonts")

_FONT_URLS: dict[str, str] = {
    "NotoSerifDevanagari-Regular.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSerifDevanagari/NotoSerifDevanagari-Regular.ttf"
    ),
    "NotoSerifDevanagari-Bold.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSerifDevanagari/NotoSerifDevanagari-Bold.ttf"
    ),
    "NotoSansDevanagari-Regular.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSansDevanagari/NotoSansDevanagari-Regular.ttf"
    ),
    "NotoSansDevanagari-Bold.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSansDevanagari/NotoSansDevanagari-Bold.ttf"
    ),
}

_NOTO_FONT_FILES: dict[str, str] = {
    "NotoSerifDevanagari":      "NotoSerifDevanagari-Regular.ttf",
    "NotoSerifDevanagari-Bold": "NotoSerifDevanagari-Bold.ttf",
    "NotoSansDevanagari":       "NotoSansDevanagari-Regular.ttf",
    "NotoSansDevanagari-Bold":  "NotoSansDevanagari-Bold.ttf",
}

_SYSTEM_FONT_DIRS = [
    "/usr/share/fonts/truetype/noto",
    "/usr/share/fonts/noto",
    "/usr/share/fonts/opentype/noto",
    "/usr/share/fonts/truetype/freefont",
    "/usr/share/fonts/freefont",
    "/usr/share/fonts",
    "/usr/share/fonts/noto-cjk",
    "/Library/Fonts",
    "/System/Library/Fonts",
    "C:/Windows/Fonts",
]

_REGISTERED_FONTS: set[str] = set()
_FONTS_REGISTERED  = False
_FONT_LOCK         = threading.Lock()


def _find_font_on_system(filename: str) -> Optional[str]:
    """Search known system font directories for a TTF file."""
    local = os.path.join(_FONTS_DIR, filename)
    if os.path.isfile(local):
        return local
        
    for d in _SYSTEM_FONT_DIRS:
        p = os.path.join(d, filename)
        if os.path.isfile(p):
            return p
            
    base_search = "/usr/share/fonts"
    if os.path.isdir(base_search):
        for root, _dirs, files in os.walk(base_search):
            if filename in files:
                found = os.path.join(root, filename)
                print(f"  🔍  Found font via recursive search: {found}")
                return found
                
    return None


def _download_font(filename: str) -> Optional[str]:
    """Download a font from GitHub/Google Fonts into _FONTS_DIR."""
    url = _FONT_URLS.get(filename)
    if not url:
        return None
        
    dest = os.path.join(_FONTS_DIR, filename)
    if os.path.isfile(dest):
        return dest
        
    try:
        os.makedirs(_FONTS_DIR, exist_ok=True)
        print(f"  ⬇️   Downloading font {filename} …")
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp, \
             open(dest, "wb") as f:
            f.write(resp.read())
        print(f"  ✅  Downloaded {filename} → {dest}")
        return dest
    except Exception as e:
        # Surface this prominently — silent font failure = boxes in output
        print(
            f"\n  ❌  FONT DOWNLOAD FAILED: {filename}\n"
            f"      Error: {e}\n"
            f"      Non-Latin text (Hindi, Arabic, etc.) will render as □□□ in PDF output.\n"
            f"      Fix: run  apt-get install -y fonts-noto-core  OR\n"
            f"               place {filename} in ./fonts/\n"
        )
        if os.path.isfile(dest):
            try:
                os.remove(dest)
            except Exception:
                pass
        return None


def _ensure_unicode_fonts() -> None:
    """Register Noto Devanagari TTFs with ReportLab (idempotent, thread-safe)."""
    global _FONTS_REGISTERED, _REGISTERED_FONTS
    if _FONTS_REGISTERED:
        return
        
    with _FONT_LOCK:
        if _FONTS_REGISTERED:
            return
        try:
            from reportlab.pdfbase import pdfmetrics      
            from reportlab.pdfbase.ttfonts import TTFont  

            for rl_name, filename in _NOTO_FONT_FILES.items():
                path = _find_font_on_system(filename)
                if not path:
                    path = _download_font(filename)
                if not path:
                    print(f"  ⚠️   Font unavailable: {rl_name} ({filename}) — skipping")
                    continue
                try:
                    pdfmetrics.registerFont(TTFont(rl_name, path))
                    _REGISTERED_FONTS.add(rl_name)
                    print(f"  ✅  Registered: {rl_name} from {path}")
                except Exception as e:
                    print(f"  ⚠️   registerFont failed for {rl_name}: {e}")

            if _REGISTERED_FONTS:
                print(f"  ✅  Unicode fonts ready: {sorted(_REGISTERED_FONTS)}")
                # FIX: Set fonts strictly to True here when registration loop finishes safely.
                _FONTS_REGISTERED = True
            else:
                print(
                    "  ⚠️  WARN: No Unicode/Devanagari fonts could be registered. "
                    "Hindi text may render as square boxes. "
                    "Fix: run `apt-get install -y fonts-noto-core` or place "
                    "NotoSansDevanagari-Regular.ttf in a ./fonts/ folder."
                )
        except Exception as e:
            print(f"  ⚠️   _ensure_unicode_fonts failed: {e}\n{traceback.format_exc()}")

# Initialize fonts in the background immediately
threading.Thread(target=_ensure_unicode_fonts, daemon=True).start()

def _has_non_latin(text: str) -> bool:
    """True if text contains Devanagari or other Indic scripts (U+0900+)."""
    return any(ord(c) >= 0x0900 for c in text
               if not unicodedata.category(c).startswith("Z"))


def _unicode_body_font(rl_name: str, has_unicode: bool) -> str:
    """Return the best available Unicode-capable font for the requested style."""
    if not has_unicode:
        return rl_name

    _PREF: dict[str, list[str]] = {
        "Helvetica":         ["NotoSansDevanagari",  "NotoSerifDevanagari"],
        "Helvetica-Bold":    ["NotoSansDevanagari-Bold", "NotoSansDevanagari"],
        "Helvetica-Oblique": ["NotoSansDevanagari",  "NotoSerifDevanagari"],
        "Times-Roman":       ["NotoSerifDevanagari", "NotoSansDevanagari"],
        "Times-Bold":        ["NotoSerifDevanagari-Bold", "NotoSerifDevanagari"],
        "Courier":           ["NotoSansDevanagari",  "NotoSerifDevanagari"],
    }
    
    for candidate in _PREF.get(rl_name, ["NotoSansDevanagari", "NotoSerifDevanagari"]):
        if candidate in _REGISTERED_FONTS:
            return candidate

    print(f"  🚨  No Unicode font for '{rl_name}' — text may render as boxes.")
    return rl_name


_LATIN_FALLBACK: dict[str, str] = {
    "NotoSerifDevanagari":      "Times-Roman",
    "NotoSerifDevanagari-Bold": "Times-Roman",
    "NotoSansDevanagari":       "Helvetica",
    "NotoSansDevanagari-Bold":  "Helvetica",
}

_DOCX_LATIN_FALLBACK: dict[str, str] = {
    "NotoSerifDevanagari":      "Times New Roman",
    "NotoSerifDevanagari-Bold": "Times New Roman",
    "NotoSansDevanagari":       "Arial",
    "NotoSansDevanagari-Bold":  "Arial",
}


def _is_latin_char(ch: str) -> bool:
    """Return True for Latin / punctuation characters, False for Devanagari / whitespace."""
    cp = ord(ch)
    if ch.isspace():
        return False
    if 0x0900 <= cp <= 0x097F:
        return False   # Devanagari
    if 0x0020 <= cp <= 0x024F:
        return True    # Basic Latin + Latin Extended
    if 0x2000 <= cp <= 0x206F:
        return True    # General Punctuation
    return False


def _mixed_font_html(safe_escaped_text: str, deva_font: str) -> str:
    """Return ReportLab XML markup with dual-font tags for Latin runs."""
    latin_font = _LATIN_FALLBACK.get(deva_font)
    if not latin_font:
        return safe_escaped_text

    fragments = re.split(r"(<br\s*/>)", safe_escaped_text)
    result_parts: list[str] = []

    for frag in fragments:
        if re.fullmatch(r"<br\s*/>", frag):
            result_parts.append(frag)
            continue
            
        if not frag:
            continue

        frag_plain = (
            frag.replace("&amp;", "&")
                .replace("&lt;",  "<")
                .replace("&gt;",  ">")
                .replace("&quot;", '"')
                .replace("&#39;",  "'")
        )

        runs: list[tuple[bool, str]] = []
        if frag_plain:
            cur_latin = _is_latin_char(frag_plain[0])
            cur_buf   = frag_plain[0]
            for ch in frag_plain[1:]:
                ch_latin = _is_latin_char(ch)
                if ch_latin == cur_latin or ch.isspace():
                    cur_buf += ch
                else:
                    runs.append((cur_latin, cur_buf))
                    cur_latin = ch_latin
                    cur_buf   = ch
            if cur_buf:
                runs.append((cur_latin, cur_buf))

        frag_html = ""
        for is_latin, chunk in runs:
            esc = (
                chunk.replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;")
            )
            if is_latin and chunk.strip():
                frag_html += f'<font name="{latin_font}">{esc}</font>'
            else:
                frag_html += esc

        result_parts.append(frag_html)

    return "".join(result_parts)


# ─────────────────────────────────────────────────────────────────────────────
# MAXIMIZED UPGRADE: Adaptive Binarization & Image Pre-Processing Engine
# ─────────────────────────────────────────────────────────────────────────────
def _enhance_image_for_ocr(path: str, intensity: str = "normal") -> str:
    """
    Agency-Grade Pre-processing: Automatically normalizes lighting and shadows
    (common in phone photos of notebooks). 
    'intensity' allows the anti-hallucination engine to request a harsher scrub.
    """
    try:
        # pyrefly: ignore [missing-import]
        from PIL import Image, ImageEnhance, ImageOps
        img = Image.open(path)
        
        # 1. Convert to grayscale to remove distracting background noise/stains
        img = ImageOps.grayscale(img)
        
        # 2. Lighting Normalization (AutoContrast to equalize shadows)
        img = ImageOps.autocontrast(img, cutoff=1)
        
        # 3. Dynamic Contrast & Sharpness Tuning
        contrast_multiplier = 2.0 if intensity == "high" else 1.5
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(contrast_multiplier)
        
        sharpness_multiplier = 3.0 if intensity == "high" else 2.0
        sharpness = ImageEnhance.Sharpness(img)
        img = sharpness.enhance(sharpness_multiplier)
        
        # Save to a temporary optimized file — use Path so dots in parent
        # directories (e.g. /home/user.name/page.png) are never corrupted.
        p = Path(path)
        enhanced_path = str(p.with_name(f"{p.stem}_{intensity}_enhanced.png"))
        img.save(enhanced_path, format="PNG")
        
        return enhanced_path
        
    except ImportError:
        print("  ⚠️ PIL not installed, skipping image enhancement. (Run: pip install Pillow)")
        return path
    except Exception as e:
        print(f"  ⚠️ Image enhancement failed, using original: {e}")
        return path

# ─────────────────────────────────────────────────────────────────────────────
# Image → base64
# ─────────────────────────────────────────────────────────────────────────────

def _image_to_b64(path: str, intensity: str = "normal") -> tuple[str, str, str]:
    """Passes images through the enhancement pipeline first. Returns (b64, mime, enhanced_path)."""
    enhanced_path = _enhance_image_for_ocr(path, intensity)
    
    ext = Path(enhanced_path).suffix.lower()
    
    if ext in {".bmp", ".tiff", ".tif"}:
        try:
            # pyrefly: ignore [missing-import]
            from PIL import Image
            buf = io.BytesIO()
            Image.open(enhanced_path).save(buf, format="PNG")
            b64_data = base64.b64encode(buf.getvalue()).decode()
            media_type = "image/png"
        except Exception as e:
            print(f"  ⚠️  Unexpected error during image conversion. {e}")
            with open(enhanced_path, "rb") as f:
                b64_data = base64.b64encode(f.read()).decode()
            # FIX: Fallback uses actual formatted type of what we just generated
            media_type = "image/png"
    else:
        if ext == ".png":
            media_type = "image/png"
        elif ext == ".webp":
            media_type = "image/webp"
        elif ext == ".gif":
            media_type = "image/gif"
        else:
            media_type = "image/jpeg"
            
        with open(enhanced_path, "rb") as f:
            b64_data = base64.b64encode(f.read()).decode()
            
    return b64_data, media_type, enhanced_path


# ─────────────────────────────────────────────────────────────────────────────
# PDF → images (one per page)
# ─────────────────────────────────────────────────────────────────────────────

def _pdf_to_images(pdf_path: str, out_dir: str) -> list[str]:
    """Render each PDF page to a PNG at PDF_RENDER_ZOOM."""
    try:
        # pyrefly: ignore [missing-import]
        import fitz  
        doc = fitz.open(pdf_path)
        paths = []
        mat = fitz.Matrix(PDF_RENDER_ZOOM, PDF_RENDER_ZOOM)
        
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            img_path = os.path.join(out_dir, f"page_{i:04d}.png")
            pix.save(img_path)
            paths.append(img_path)
            
        print(f"  📄 Rendered {len(paths)} pages at {PDF_RENDER_ZOOM}x zoom ({int(72 * PDF_RENDER_ZOOM)} DPI)")
        return paths
        
    except ImportError:
        # FIX: Provide descriptive error tracking if fallback is missing
        print("  ⚠️  fitz (PyMuPDF) missing. Falling back to pdf2image.")
        try:
            # pyrefly: ignore [missing-import]
            from pdf2image import convert_from_path
            dpi = int(72 * PDF_RENDER_ZOOM)
            images = convert_from_path(pdf_path, dpi=dpi)
            paths = []
            
            for i, img in enumerate(images):
                img_path = os.path.join(out_dir, f"page_{i:04d}.png")
                img.save(img_path, "PNG")
                paths.append(img_path)
                
            return paths
        except ImportError:
            raise RuntimeError(
                "Neither PyMuPDF (fitz) nor pdf2image is installed. "
                "Run: pip install pymupdf  OR  pip install pdf2image"
            )
        
    except Exception as e:
        print(f"  ⚠️  PDF to images conversion failed completely. Error details: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# DOCX → images
# ─────────────────────────────────────────────────────────────────────────────

def _docx_to_images(docx_path: str, out_dir: str) -> list[str]:
    """Extract embedded images from a DOCX."""
    # pyrefly: ignore [missing-import]
    from docx import Document
    # pyrefly: ignore [missing-import]
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = Document(docx_path)
    paths = []
    idx = 0
    
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                ext = Path(rel.target_partname).suffix.lower() or ".png"
                img_path = os.path.join(out_dir, f"image_{idx:04d}{ext}")
                
                with open(img_path, "wb") as f:
                    f.write(img_data)
                    
                paths.append(img_path)
                idx += 1
                
    except Exception as e:
        print(f"  ⚠️  Error extracting embedded images from DOCX. Error details: {e}")

    if not paths:
        try:
            import subprocess
            pdf_path = os.path.join(out_dir, "converted.pdf")
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                capture_output=True, 
                timeout=30
            )
            base = Path(docx_path).stem
            pdf_candidate = os.path.join(out_dir, f"{base}.pdf")
            
            if os.path.exists(pdf_candidate):
                return _pdf_to_images(pdf_candidate, out_dir)
                
        except Exception as e:
            print(f"  ⚠️  DOCX to PDF fallback conversion failed. Error details: {e}")

    return sorted(paths)


# ─────────────────────────────────────────────────────────────────────────────
# Collect all images from any supported input
# ─────────────────────────────────────────────────────────────────────────────

def collect_images(file_path: str, filename: str, scratch_dir: str) -> list[str]:
    """Given any supported input file, return an ordered list of image paths."""
    ext = Path(filename).suffix.lower()
    os.makedirs(scratch_dir, exist_ok=True)

    if ext in SUPPORTED_IMAGE_EXTS:
        dest = os.path.join(scratch_dir, f"img_0000{ext}")
        shutil.copy2(file_path, dest)
        return [dest]

    if ext == ".pdf":
        return _pdf_to_images(file_path, scratch_dir)

    if ext == ".docx":
        return _docx_to_images(file_path, scratch_dir)

    if ext == ".zip":
        image_paths = []
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                # Sort ALL members together so the final page order matches the
                # original ZIP ordering (images and embedded PDFs/DOCXs interleaved).
                all_members = sorted([
                    m for m in zf.namelist()
                    if not m.startswith("__MACOSX")
                    and not os.path.basename(m).startswith(".")
                    and Path(m).suffix.lower() in (SUPPORTED_IMAGE_EXTS | {".pdf", ".docx"})
                ])

                img_counter = 0
                for member in all_members:
                    member_ext = Path(member).suffix.lower()

                    if member_ext in SUPPORTED_IMAGE_EXTS:
                        dest = os.path.join(scratch_dir, f"img_{img_counter:04d}{member_ext}")
                        img_counter += 1
                        with zf.open(member) as src, open(dest, "wb") as dst:
                            shutil.copyfileobj(src, dst)
                        image_paths.append(dest)

                    elif member_ext in {".pdf", ".docx"}:
                        tmp = os.path.join(scratch_dir, f"inner_{uuid.uuid4().hex}{member_ext}")
                        with zf.open(member) as src, open(tmp, "wb") as dst:
                            shutil.copyfileobj(src, dst)

                        sub_dir = os.path.join(scratch_dir, f"sub_{uuid.uuid4().hex}")
                        os.makedirs(sub_dir, exist_ok=True)

                        if member_ext == ".pdf":
                            image_paths.extend(_pdf_to_images(tmp, sub_dir))
                        else:
                            image_paths.extend(_docx_to_images(tmp, sub_dir))

                        os.remove(tmp)

            # FIX: Hash-based deduplication
            seen_hashes: set[str] = set()
            ordered: list[str] = []
            for p in image_paths:
                with open(p, 'rb') as f:
                    file_hash = hashlib.md5(f.read()).hexdigest()
                if file_hash not in seen_hashes:
                    seen_hashes.add(file_hash)
                    ordered.append(p)
            return ordered

        except Exception as e:
            print(f"  ⚠️  ZIP file processing failed. Error details: {e}")
            raise

    raise ValueError(f"Unsupported file type: {ext}")


# ─────────────────────────────────────────────────────────────────────────────
# Retry helper
# ─────────────────────────────────────────────────────────────────────────────

def _api_call_with_retry(fn, *args, **kwargs):
    """Call fn(*args, **kwargs) up to MAX_RETRIES times with exponential backoff.
    Acquires the global semaphore before each attempt so concurrent workers
    can't collectively exceed the OpenAI rate limit.
    """
    delay = RETRY_BASE_DELAY
    last_exc = None
    
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            with _OPENAI_SEMAPHORE:
                return fn(*args, **kwargs)
        except Exception as e:
            last_exc = e
            err_str = str(e).lower()
            
            if any(k in err_str for k in ("rate limit", "429", "500", "502", "503", "timeout")):
                if attempt < MAX_RETRIES:
                    wait = delay * (2 ** (attempt - 1))
                    print(f"  ⏳  API error (attempt {attempt}/{MAX_RETRIES}), retrying in {wait:.1f}s: {e}")
                    time.sleep(wait)
                    continue
            else:
                # Non-transient error (e.g. 400 Bad Request from a bad request
                # body/param) — surface it immediately, it will NOT fix itself
                # on retry.
                print(f"  ❌  Non-retryable API error (attempt {attempt}/{MAX_RETRIES}): {e}")
            raise
            
    raise last_exc


# ─────────────────────────────────────────────────────────────────────────────
# MAXIMIZED UPGRADE: GPT-4o Vision with Reading Order & Anti-Hallucination
# ─────────────────────────────────────────────────────────────────────────────

# Upgraded Prompt to handle complex physical layouts and multi-lingual text
TRANSCRIPTION_SYSTEM = """You are an Enterprise-Grade Handwriting & Physical Document Layout AI.
Your job is to read handwritten text from page images and produce a faithful, COMPLETE transcription.

CRITICAL RULES:
1. COMPLETENESS: Transcribe ALL text you can see on the page — do NOT stop early or summarize. If the page is dense, keep going until every word is captured.
2. LAYOUT PREDICTION (READING ORDER): If the page has multiple columns, sidebars, or text wrapped around diagrams, you MUST read the text in logical human reading order (generally Top to Bottom, Left to Right within columns). Do not mix column text horizontally.
3. MULTI-LINGUAL PRESERVATION: The text may contain mixed languages (e.g. English and Hindi). Transcribe each exactly as written in its native script. Do not translate.
4. ILLEGIBILITY: If a word is truly illegible, write [illegible]. Do NOT hallucinate or guess wildly.
5. FORMATTING: Produce plain text paragraphs. Maintain natural paragraph breaks. Do NOT add markdown fences, headers, or commentary.
6. BLANK PAGES: If the page has no text, output strictly: [PAGE: no text]
7. NO TRUNCATION: If you are running out of space, continue transcribing — never add "..." or stop mid-sentence.
"""

def _is_hallucinating(text: str) -> bool:
    """Detects if the OCR fell into an infinite repeating loop (common failure on faded text).
    
    Conservative threshold: only flags clear machine-loop failures, NOT legitimate
    repetitive text like poetry, prayers, lists, or rhetorical repetition.
    """
    if len(text) < 100:
        return False
        
    words = text.split()
    # Require at least 30 words before checking — short responses can't be loops
    if len(words) > 30:
        most_common_word, most_common_count = Counter(words).most_common(1)[0]
        # Only flag if >60% of words are the same AND it's a common function word pattern
        # (real hallucinations loop on "the the the" or "और और और", not content words)
        if most_common_count > len(words) * 0.60:
            return True
    
    # Secondary check: detect pure phrase-level looping (e.g., "hello world hello world hello world")
    # Split text into 4-word ngrams and check if any repeats excessively
    if len(words) >= 20:
        ngram_size = 4
        ngrams = [" ".join(words[i:i+ngram_size]) for i in range(len(words) - ngram_size + 1)]
        if ngrams:
            top_ngram, top_count = Counter(ngrams).most_common(1)[0]
            # If a 4-word phrase appears more than 25% of possible positions, it's looping
            if top_count > len(ngrams) * 0.25:
                return True
            
    return False


def _transcribe_single_batch(batch: list[str], batch_start: int) -> list[dict]:
    """Transcribes one batch of images, featuring a dynamic hallucination-retry loop."""
    
    def attempt_transcription(intensity: str = "normal") -> list[dict]:
        successful_indices: list[int] = []
        failed_indices: list[int] = []
        api_content: list[dict] = [{"type": "text", "text": ""}] 
        temp_paths: list[str] = []

        for slot, img_path in enumerate(batch):
            try:
                b64, mime, enhanced_path = _image_to_b64(img_path, intensity)
                temp_paths.append(enhanced_path)
                api_content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"},
                })
                successful_indices.append(slot)
            except Exception as e:
                failed_indices.append(slot)

        batch_results: list[dict] = [{} for _ in range(len(batch))]
        
        for slot in failed_indices:
            batch_results[slot] = {
                "page_num": batch_start + slot + 1,
                "text": "[illegible - could not process image]",
                "has_content": False,
                "hallucination_flag": False
            }

        # FIX: Avoid wasted empty calls
        if not successful_indices:
            return batch_results

        n_sent = len(successful_indices)
        if n_sent == 1:
            api_content[0]["text"] = "Transcribe the handwritten text from this page image. Obey layout reading order."
        else:
            api_content[0]["text"] = f"Transcribe the handwritten text from the following {n_sent} page image(s). Separate each page's content with the marker ---PAGE_BREAK--- on its own line."

        if successful_indices:
            try:
                def _call():
                    return client.chat.completions.create(
                        model=MODEL,
                        messages=[
                            {"role": "system", "content": TRANSCRIPTION_SYSTEM},
                            {"role": "user", "content": api_content},
                        ],
                        max_tokens=TRANSCRIPTION_MAX_TOKENS,
                        temperature=0.1 if intensity == "normal" else 0.4 
                    )

                response = _api_call_with_retry(_call)
                raw = (response.choices[0].message.content or "").strip()

                if n_sent == 1:
                    gpt_pages = [raw]
                else:
                    gpt_pages = raw.split("---PAGE_BREAK---")

                for i, slot in enumerate(successful_indices):
                    page_text = gpt_pages[i].strip() if i < len(gpt_pages) else ""
                    has_content = bool(page_text) and "[PAGE: no text]" not in page_text
                    
                    # Anti-Hallucination Guard
                    is_garbage = _is_hallucinating(page_text)
                    
                    batch_results[slot] = {
                        "page_num": batch_start + slot + 1,
                        "text": page_text if has_content and not is_garbage else "",
                        "has_content": has_content and not is_garbage,
                        "hallucination_flag": is_garbage
                    }

                for slot in successful_indices[len(gpt_pages):]:
                    # The model returned fewer splits than pages sent — retry this page solo
                    print(f"  ⚠️  Page {batch_start + slot + 1} not returned in batch split — queuing solo retry")
                    batch_results[slot] = {
                        "page_num": batch_start + slot + 1,
                        "text": "[transcription incomplete - will retry]",
                        "has_content": False,
                        "hallucination_flag": False,
                        "_needs_solo_retry": True,
                        "_img_path": batch[slot],
                    }

            except Exception as e:
                print(f"  ❌  Transcription API call failed for page(s) {[batch_start + s + 1 for s in successful_indices]}: {e}")
                for slot in successful_indices:
                    batch_results[slot] = {
                        "page_num": batch_start + slot + 1,
                        "text": "[transcription failed for this page]",
                        "has_content": False,
                        "hallucination_flag": False
                    }
                    
        # Cleanup temp enhanced files — only delete files that are NOT originals
        # (enhanced_path == path when PIL is missing, so we skip those).
        original_paths = set(batch)
        for p in temp_paths:
            if os.path.exists(p) and p not in original_paths:
                try:
                    os.remove(p)
                except Exception:
                    pass

        return batch_results

    # First Pass (Normal Binarization)
    results = attempt_transcription("normal")
    
    # Check if any page in this batch hallucinated. If so, trigger the High-Intensity scrub.
    if any(r.get("hallucination_flag") for r in results if r):
        print(f"  🚨 Garbage loop detected on page {batch_start + 1}. Triggering High-Intensity Binarization Re-read...")
        results = attempt_transcription("high")

    return [r for r in results if r]


def transcribe_images(image_paths: list[str], book_title: str = "") -> list[dict]:
    """Transcribe a list of images using GPT-4o vision in parallel with ThreadPoolExecutor."""
    total = len(image_paths)
    print(f"  🖊️  Transcribing {total} pages | batch_size={PAGES_PER_BATCH} | workers={TRANSCRIPTION_WORKERS}")

    batches: list[tuple[list[str], int]] = []
    for batch_start in range(0, total, PAGES_PER_BATCH):
        batch = image_paths[batch_start: batch_start + PAGES_PER_BATCH]
        batches.append((batch, batch_start))

    results_by_start: dict[int, list[dict]] = {}
    completed = 0

    with ThreadPoolExecutor(max_workers=TRANSCRIPTION_WORKERS) as executor:
        future_to_start = {
            executor.submit(_transcribe_single_batch, batch, batch_start): batch_start
            for batch, batch_start in batches
        }
        for future in as_completed(future_to_start):
            batch_start = future_to_start[future]
            try:
                batch_result = future.result()
                results_by_start[batch_start] = batch_result
            except Exception as e:
                batch_size = min(PAGES_PER_BATCH, total - batch_start)
                results_by_start[batch_start] = [
                    {
                        "page_num": batch_start + i + 1, 
                        "text": "[transcription failed]", 
                        "has_content": False
                    }
                    for i in range(batch_size)
                ]
                
            completed += 1
            if completed % 20 == 0 or completed == len(batches):
                # FIX: Bound pages log count properly
                print(f"  ✅  {completed}/{len(batches)} batches done ({min(completed * PAGES_PER_BATCH, total)}/{total} pages)")

    all_results: list[dict] = []
    for batch_start in sorted(results_by_start.keys()):
        all_results.extend(results_by_start[batch_start])

    # ── Solo retry pass: re-transcribe any pages that were missed in batch splits ──
    needs_retry = [r for r in all_results if r.get("_needs_solo_retry")]
    if needs_retry:
        print(f"  🔁 Solo retry pass: {len(needs_retry)} page(s) missed in batch splits")
        for r in needs_retry:
            img_path = r.get("_img_path", "")
            page_num = r["page_num"]
            if not img_path or not os.path.exists(img_path):
                r["text"] = "[transcription failed — image unavailable for retry]"
                r["has_content"] = False
                continue
            try:
                solo_result = _transcribe_single_batch([img_path], page_num - 1)
                if solo_result and solo_result[0].get("has_content"):
                    r.update(solo_result[0])
                    r.pop("_needs_solo_retry", None)
                    r.pop("_img_path", None)
                    print(f"    ✅ Solo retry succeeded for page {page_num}")
                else:
                    r["text"] = "[transcription failed after solo retry]"
                    r["has_content"] = False
            except Exception as e:
                print(f"    ⚠️  Solo retry failed for page {page_num}: {e}")
                r["text"] = "[transcription failed after solo retry]"
                r["has_content"] = False

    content_count = sum(1 for r in all_results if r["has_content"])
    print(f"  📝  Transcription complete: {content_count}/{total} pages have content")
    return all_results

# ─────────────────────────────────────────────────────────────────────────────
# UPGRADE: The "Context Healer" Agent (Page Boundary Stitching)
# ─────────────────────────────────────────────────────────────────────────────

HEALER_SYSTEM_PROMPT = """You are a post-OCR correction agent. 
You are receiving transcribed text from handwritten pages that were processed in parallel.
Your strict mission:
1. Fix hyphenation and broken sentences that occur across page boundaries.
2. If there is an [illegible] marker, use the context of the surrounding sentences to deduce what the word likely was. If you cannot confidently guess, leave it as [illegible].
3. DO NOT paraphrase, summarize, or change the style. Only fix OCR artifacts.
4. Preserve the original language perfectly (especially Hindi/Devanagari if present).
5. Keep ALL the content — do not skip, summarize, or drop any pages.
6. Return the healed text separated by ---PAGE_BREAK--- between pages, in the same order.
"""

# How many content pages to heal per API call.
# Smaller = more API calls but more reliable healing of boundaries.
HEALER_CHUNK_PAGES = 20

def heal_transcription_context(pages: list[dict]) -> list[dict]:
    """
    Chunk-aware Context Healer: heals page boundaries in groups of HEALER_CHUNK_PAGES.
    This avoids the single-call char limit that silently truncated large books.
    Each chunk has a 2-page overlap with the previous chunk so cross-boundary
    sentence breaks at chunk edges are also repaired.
    """
    print("  🩹 Initiating Context Healer Agent (chunked mode) to fix page boundaries…")
    
    content_indices = [i for i, p in enumerate(pages) if p["has_content"]]
    if not content_indices:
        return pages

    # Work on a copy so we don't mutate in-place until we're sure healing succeeded
    healed_pages = [dict(p) for p in pages]

    # Process in overlapping chunks of HEALER_CHUNK_PAGES
    chunk_size = HEALER_CHUNK_PAGES
    overlap = 2  # pages of overlap between chunks to catch cross-chunk boundaries

    i = 0
    chunk_num = 0
    total_chunks = max(1, (len(content_indices) + chunk_size - 1) // chunk_size)

    while i < len(content_indices):
        chunk_end = min(i + chunk_size, len(content_indices))
        chunk_content_indices = content_indices[i:chunk_end]
        chunk_num += 1
        
        # Build the text block for this chunk
        chunk_blocks = []
        for idx in chunk_content_indices:
            p = healed_pages[idx]
            chunk_blocks.append(f"[PAGE {p['page_num']}]\n{p['text']}")
        
        compiled_text = "\n\n---PAGE_BREAK---\n\n".join(chunk_blocks)
        
        try:
            response = _api_call_with_retry(
                client.chat.completions.create,
                model=MODEL,
                messages=[
                    {"role": "system", "content": HEALER_SYSTEM_PROMPT},
                    {"role": "user", "content": (
                        f"Heal this raw OCR text (chunk {chunk_num}/{total_chunks}).\n"
                        f"There are {len(chunk_content_indices)} pages. Return ALL of them separated by ---PAGE_BREAK---.\n\n"
                        f"{compiled_text}"
                    )},
                ],
                max_tokens=16384,
                temperature=0.05,
            )
            healed_raw = (response.choices[0].message.content or "").strip()

            # Split on the page break marker
            raw_splits = healed_raw.split("---PAGE_BREAK---")

            # FIX: Do not overwrite the overlapping chunk data we already healed
            write_start = overlap if chunk_num > 1 else 0
            
            for j, split_text in enumerate(raw_splits):
                if j < write_start or j >= len(chunk_content_indices):
                    continue
                original_idx = chunk_content_indices[j]
                # Strip [PAGE N] header that the healer may have echoed back
                clean_text = re.sub(r'^\s*\[PAGE\s+\d+\]\s*\n?', '', split_text.strip()).strip()
                if clean_text:  # only update if we got content back
                    healed_pages[original_idx]["text"] = clean_text

            print(f"  ✅ Healer chunk {chunk_num}/{total_chunks} done ({len(chunk_content_indices)} pages)")

        except Exception as e:
            print(f"  ⚠️  Healer chunk {chunk_num} failed, keeping raw transcription for those pages: {e}")
            # Don't update — keep original transcription for this chunk

        # Advance, but keep `overlap` pages in the next chunk to heal cross-chunk boundaries
        i = chunk_end - overlap if chunk_end < len(content_indices) else chunk_end

    print(f"  ✅ Context Healer complete ({total_chunks} chunk(s) processed).")
    return healed_pages

# ─────────────────────────────────────────────────────────────────────────────
# Post-processing: structure the transcribed text into a clean book
# ─────────────────────────────────────────────────────────────────────────────

STRUCTURE_SYSTEM = """You are a professional book editor and formatter.
You receive raw transcribed text (possibly from handwritten pages) and must:
1. Clean up obvious OCR/transcription artifacts (repeated words, stray characters)
2. Identify and mark chapters/sections if they exist (look for numbered headings, "Chapter X", etc.)
3. Fix clear run-on sentences caused by page breaks mid-sentence
4. Preserve the author's original voice, language, and style — do NOT paraphrase or rewrite
5. Preserve the original language (do not translate)
6. Include ALL the text — do not summarize, skip, or drop any content. Every word of the input must appear in the output.
7. Return structured JSON only

Return ONLY valid JSON with no preamble, no markdown fences:
{
  "title": "<inferred or provided title>",
  "language": "<detected language, e.g. English, Hindi, Tamil, etc.>",
  "chapters": [
    {
      "chapter_number": 1,
      "title": "<chapter title or 'Chapter 1' if untitled>",
      "content": "<cleaned chapter text with paragraph breaks as \\n\\n>"
    }
  ],
  "total_words": <integer word count>
}
"""

def structure_transcription(pages: list[dict], book_title: str = "") -> dict:
    content_pages = [p for p in pages if p["has_content"]]
    if not content_pages:
        return {
            "title": book_title or "Untitled Manuscript",
            "language": "Unknown",
            "chapters": [{"chapter_number": 1, "title": "Content", "content": "[No readable text found in the uploaded pages]"}],
            "total_words": 0,
        }

    sample_text = "\n\n".join(p["text"] for p in content_pages[:5])
    detected_language = "Unknown"
    
    try:
        lang_resp = _api_call_with_retry(
            client.chat.completions.create,
            model=MODEL,
            messages=[{"role": "user", "content": f"What language is this text written in? Reply with just the language name.\n\n{sample_text[:500]}"}],
            max_tokens=20,
        )
        detected_language = (lang_resp.choices[0].message.content or "").strip()
    except Exception:
        pass

    text_chunks: list[str] = []
    current_chunk: list[str] = []
    current_size = 0

    for p in content_pages:
        page_block = f"[Page {p['page_num']}]\n{p['text']}"
        block_size = len(page_block) + 2

        if current_size + block_size > STRUCTURE_CHUNK_SIZE and current_chunk:
            text_chunks.append("\n\n".join(current_chunk))
            current_chunk = []
            current_size = 0

        current_chunk.append(page_block)
        current_size += block_size

    if current_chunk:
        text_chunks.append("\n\n".join(current_chunk))

    print(f"  📖 Structuring transcription: {len(content_pages)} pages → {len(text_chunks)} chunk(s)")

    all_chapters: list[dict] = []
    chapter_counter = 0

    for chunk_idx, chunk_text in enumerate(text_chunks):
        # Calculate input word count so we can validate the model didn't drop content
        input_word_count = len(chunk_text.split())

        prompt = (
            f"Book title (if known): {book_title or 'Unknown — infer from content if possible'}\n"
            f"Language: {detected_language}\n"
            + (f"(This is part {chunk_idx + 1} of {len(text_chunks)} — continue chapter numbering from {chapter_counter + 1})\n" if len(text_chunks) > 1 else "")
            + f"\nRaw transcribed pages:\n{chunk_text}"
        )
        try:
            def _structure_call(p=prompt):
                return client.chat.completions.create(
                    model=MODEL,
                    messages=[
                        {"role": "system", "content": STRUCTURE_SYSTEM},
                        {"role": "user", "content": p},
                    ],
                    max_tokens=STRUCTURE_MAX_TOKENS,
                )

            response = _api_call_with_retry(_structure_call)

            # Check finish_reason BEFORE parsing — if truncated, go straight to fallback
            finish_reason = response.choices[0].finish_reason
            if finish_reason == "length":
                raise ValueError(
                    f"Structure response was truncated (finish_reason=length) on chunk {chunk_idx + 1}. "
                    f"Input had ~{input_word_count} words. Falling back to flat text."
                )

            raw = (response.choices[0].message.content or "").strip()
            raw = raw.replace("```json", "").replace("```", "").strip()
            
            s = raw.find("{")
            e = raw.rfind("}") + 1
            
            if s == -1 or e == 0:
                raise ValueError("No JSON object found in structure response")
                
            chunk_result = json.loads(raw[s:e])

            # Validate output word count is at least 60% of input — guards against
            # the model summarising or dropping content silently
            output_words = sum(len(ch.get("content", "").split()) for ch in chunk_result.get("chapters", []))
            if output_words < input_word_count * 0.60:
                raise ValueError(
                    f"Structure output suspiciously short: {output_words} words out vs "
                    f"{input_word_count} words in ({output_words/max(input_word_count,1)*100:.0f}%). "
                    f"Falling back to flat text to avoid data loss."
                )

            for ch in chunk_result.get("chapters", []):
                chapter_counter += 1
                all_chapters.append({
                    "chapter_number": chapter_counter,
                    "title": ch.get("title", f"Chapter {chapter_counter}"),
                    "content": ch.get("content", ""),
                })

            if chunk_idx == 0:
                if chunk_result.get("title") and not book_title:
                    book_title = chunk_result["title"]
                if chunk_result.get("language") and detected_language == "Unknown":
                    detected_language = chunk_result["language"]

        except Exception as exc:
            print(f"  ⚠️  Structuring chunk {chunk_idx + 1} failed ({exc}). Using flat fallback to preserve all content.")
            chapter_counter += 1
            # Strip [Page N] markers that would appear in the final PDF/DOCX
            clean_fallback = re.sub(r'\[Page\s+\d+\]\s*\n?', '', chunk_text).strip()
            all_chapters.append({
                "chapter_number": chapter_counter,
                "title": f"Part {chapter_counter}",
                "content": clean_fallback,
            })

    if not all_chapters:
        full_text = "\n\n".join(f"[Page {p['page_num']}]\n{p['text']}" for p in content_pages)
        all_chapters = [{"chapter_number": 1, "title": "Full Content", "content": full_text}]

    total_words = sum(len(ch["content"].split()) for ch in all_chapters)

    return {
        "title": book_title or "Transcribed Manuscript",
        "language": detected_language,
        "chapters": all_chapters,
        "total_words": total_words,
    }


# ─────────────────────────────────────────────────────────────────────────────
# PDF generator for scanned book
# ─────────────────────────────────────────────────────────────────────────────

def generate_scanned_pdf(structure: dict, output_path: str) -> str:
    try:
        from reportlab.lib.pagesizes import A4                                      
        from reportlab.lib.units import mm                                          
        from reportlab.lib.styles import ParagraphStyle                             
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT             
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable  
        from reportlab.lib.colors import HexColor, white                            
        import datetime

        _ensure_unicode_fonts()

        # FIX: proper concatenation spacing
        all_text = structure.get("title", "") + " " + " ".join(
            ch.get("title", "") + " " + ch.get("content", "")
            for ch in structure.get("chapters", [])
        )
        has_unicode = _has_non_latin(all_text)

        body_font_base = _unicode_body_font("Helvetica", has_unicode)
        ch_title_font  = _unicode_body_font("Helvetica", has_unicode)
        header_font    = "Helvetica"

        DARK   = HexColor("#1E293B")
        ACCENT = HexColor("#7C3AED")
        MARGIN = 22 * mm

        def on_cover(canvas, doc):
            w, h = A4
            canvas.setFillColor(HexColor("#0f0a1e"))
            canvas.rect(0, 0, w, h, fill=1, stroke=0)
            
            canvas.setFillColor(HexColor("#7C3AED"))
            canvas.rect(0, h * 0.38, w, 3, fill=1, stroke=0)
            
            canvas.setFillColor(HexColor("#0a0614"))
            canvas.rect(0, 0, w, 18*mm, fill=1, stroke=0)
            
            canvas.setFillColor(HexColor("#475569"))
            canvas.setFont(header_font, 8)
            canvas.drawCentredString(w/2, 7*mm, f"Transcribed by Enterprise AI Scanner  ·  {datetime.date.today()}")

        def on_page(canvas, doc):
            w, h = A4
            canvas.saveState()
            
            canvas.setFillColor(DARK)
            canvas.rect(0, 0, w, 9*mm, fill=1, stroke=0)
            
            canvas.setFillColor(white)
            # Use the unicode-capable body font if available so Devanagari/other
            # scripts in the title aren't silently dropped by Helvetica.
            footer_font = body_font_base if has_unicode and body_font_base != "Helvetica" else header_font
            canvas.setFont(footer_font, 7.5)
            title_str = structure.get("title", "Manuscript")
            canvas.drawString(MARGIN, 3*mm, title_str)
            canvas.setFont(header_font, 7.5)
            canvas.drawRightString(w - MARGIN, 3*mm, f"Page {doc.page}")
            
            canvas.restoreState()

        cover_title_font = _unicode_body_font("Helvetica-Bold", has_unicode)

        S = lambda name, **kw: ParagraphStyle(name, **kw)
        
        cover_title = S("ct", fontName=cover_title_font, fontSize=30, textColor=white,
                        leading=38, alignment=TA_CENTER, spaceAfter=8, wordWrap="LTR")
                        
        cover_sub   = S("cs", fontName="Helvetica",      fontSize=12, textColor=HexColor("#94a3b8"),
                        leading=16, alignment=TA_CENTER, wordWrap="LTR")
                        
        ch_label    = S("cl", fontName="Helvetica", fontSize=10, textColor=ACCENT,
                        leading=14, spaceBefore=0, spaceAfter=3, wordWrap="LTR")
                        
        ch_title    = S("cht", fontName=ch_title_font, fontSize=20, textColor=DARK,
                        leading=26, spaceBefore=2, spaceAfter=5, wordWrap="LTR")
                        
        body        = S("bs", fontName=body_font_base, fontSize=10.5,
                        textColor=HexColor("#334155"), leading=17, spaceAfter=8,
                        alignment=TA_JUSTIFY, wordWrap="LTR")

        doc = SimpleDocTemplate(
            output_path, 
            pagesize=A4,
            leftMargin=MARGIN, 
            rightMargin=MARGIN,
            topMargin=MARGIN, 
            bottomMargin=18*mm,
            title=structure.get("title", "Manuscript")
        )

        story = []
        story.append(Spacer(1, 50*mm))
        
        safe_cov_title = (structure.get("title", "Manuscript")
                          .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
                          
        if has_unicode:
            safe_cov_title = _mixed_font_html(safe_cov_title, body_font_base)
            
        story.append(Paragraph(safe_cov_title, cover_title))
        story.append(Spacer(1, 5*mm))
        
        lang = structure.get("language", "")
        if lang:
            story.append(Paragraph(f"Language: {lang}", cover_sub))
            
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(f"Transcribed on {datetime.date.today().strftime('%B %d, %Y')}", cover_sub))
        story.append(PageBreak())

        for ch in structure.get("chapters", []):
            story.append(Spacer(1, 8*mm))
            story.append(Paragraph(f"Chapter {ch['chapter_number']}", ch_label))
            
            safe_ch_title = (ch["title"]
                             .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
                             
            if has_unicode:
                safe_ch_title = _mixed_font_html(safe_ch_title, ch_title_font)
                
            story.append(Paragraph(safe_ch_title, ch_title))
            story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=10))
            
            for para in ch["content"].split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                    
                safe_para = (para
                             .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
                             
                if has_unicode:
                    safe_para = _mixed_font_html(safe_para, body_font_base)
                    
                story.append(Paragraph(safe_para, body))
                
            story.append(PageBreak())

        doc.build(story, onFirstPage=on_cover, onLaterPages=on_page)
        return output_path
        
    except Exception as e:
        print(f"  ⚠️  generate_scanned_pdf failed. Error details: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# DOCX generator for scanned book
# ─────────────────────────────────────────────────────────────────────────────

def generate_scanned_docx(structure: dict, output_path: str) -> str:
    try:
        # pyrefly: ignore [missing-import]
        from docx import Document
        # pyrefly: ignore [missing-import]
        from docx.shared import Pt, RGBColor, Cm
        # pyrefly: ignore [missing-import]
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        # pyrefly: ignore [missing-import]
        from docx.oxml.ns import qn
        # pyrefly: ignore [missing-import]
        from docx.oxml import OxmlElement
        import datetime

        # FIX: proper concatenation spacing
        all_text_docx = structure.get("title", "") + " " + " ".join(
            ch.get("title", "") + " " + ch.get("content", "")
            for ch in structure.get("chapters", [])
        )
        has_unicode_docx = _has_non_latin(all_text_docx)

        _RL_TO_WORD: dict[str, str] = {
            "NotoSansDevanagari":       "Noto Sans Devanagari",
            "NotoSansDevanagari-Bold":  "Noto Sans Devanagari",
            "NotoSerifDevanagari":      "Noto Serif Devanagari",
            "NotoSerifDevanagari-Bold": "Noto Serif Devanagari",
        }
        
        rl_body_font = _unicode_body_font("Helvetica", has_unicode_docx)
        word_body_font = _RL_TO_WORD.get(rl_body_font, "Calibri")
        word_latin_fallback = _DOCX_LATIN_FALLBACK.get(rl_body_font, "Calibri")

        def _set_run_font_cs(run, font_name: str) -> None:
            rPr = run._r.get_or_add_rPr()
            existing = rPr.find(qn("w:rFonts"))
            
            if existing is None:
                existing = OxmlElement("w:rFonts")
                rPr.insert(0, existing)
                
            existing.set(qn("w:ascii"),    font_name)
            existing.set(qn("w:hAnsi"),    font_name)
            existing.set(qn("w:cs"),       font_name)   
            existing.set(qn("w:eastAsia"), font_name)

        def _is_latin_char_docx(ch: str) -> bool:
            cp = ord(ch)
            if ch.isspace():
                return False
            if 0x0900 <= cp <= 0x097F:
                return False   
            if 0x0020 <= cp <= 0x024F:
                return True    
            if 0x2000 <= cp <= 0x206F:
                return True    
            return False

        def add_para_unicode(p, text: str, size_pt: float, bold: bool = False,
                             italic: bool = False, color_rgb: RGBColor = None,
                             align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
            p.alignment = align
            
            if not has_unicode_docx or not text.strip():
                run = p.add_run(text)
                run.font.name   = word_body_font
                run.font.size   = Pt(size_pt)
                run.font.bold   = bold
                run.font.italic = italic
                if color_rgb:
                    run.font.color.rgb = color_rgb
                _set_run_font_cs(run, word_body_font)
                return

            runs_list: list[tuple[bool, str]] = []
            
            if text:
                cur_latin = _is_latin_char_docx(text[0])
                cur_buf   = text[0]
                for ch in text[1:]:
                    ch_lat = _is_latin_char_docx(ch)
                    if ch.isspace() or ch_lat == cur_latin:
                        cur_buf += ch
                    else:
                        runs_list.append((cur_latin, cur_buf))
                        cur_latin = ch_lat
                        cur_buf   = ch
                if cur_buf:
                    runs_list.append((cur_latin, cur_buf))

            for is_latin, chunk in runs_list:
                chosen_font = word_latin_fallback if (is_latin and chunk.strip()) else word_body_font
                
                run = p.add_run(chunk)
                run.font.name   = chosen_font
                run.font.size   = Pt(size_pt)
                run.font.bold   = bold
                run.font.italic = italic
                
                if color_rgb:
                    run.font.color.rgb = color_rgb
                    
                _set_run_font_cs(run, chosen_font)

        doc = Document()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = Cm(2.5)
        section.bottom_margin = Cm(2.0)

        style_normal = doc.styles['Normal']
        style_normal.font.name = word_body_font
        style_normal.font.size = Pt(11)

        for _ in range(4): 
            doc.add_paragraph()
            
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        add_para_unicode(
            t, 
            structure.get("title", "Manuscript"), 
            26, 
            bold=True,
            color_rgb=RGBColor(0x1E, 0x29, 0x3B), 
            align=WD_ALIGN_PARAGRAPH.CENTER
        )

        lang = structure.get("language", "")
        if lang:
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            lr = lp.add_run(f"Language: {lang}")
            lr.font.size = Pt(11)
            lr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        dr = dp.add_run(f"Transcribed on {datetime.date.today().strftime('%B %d, %Y')}")
        dr.font.size = Pt(10)
        dr.font.italic = True
        dr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        
        doc.add_page_break()

        for ch in structure.get("chapters", []):
            lbl = doc.add_paragraph()
            lr = lbl.add_run(f"Chapter {ch['chapter_number']}")
            lr.font.size = Pt(10)
            lr.font.bold = True
            lr.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

            heading = doc.add_heading("", level=1)
            add_para_unicode(
                heading, 
                ch["title"], 
                16, 
                bold=True,
                color_rgb=RGBColor(0x1E, 0x29, 0x3B),
                align=WD_ALIGN_PARAGRAPH.LEFT
            )
            
            doc.add_paragraph()

            for para in ch["content"].split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                    
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = Pt(16)
                add_para_unicode(p, para, 11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

            doc.add_page_break()

        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"  ⚠️  generate_scanned_docx failed. Error details: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# UPGRADED MAIN ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────

def scan_handwritten_book(
    file_path: str,
    filename: str,
    output_dir: str,
    book_title: str = "",
    progress_callback=None,
) -> dict:
    """
    Maximized Enterprise Pipeline:
    1. Extract/collect images
    2. Image Pre-Processing & Adaptive Contrast Enhancement
    3. Transcribe via GPT-4o vision (Parallel) with Anti-Hallucination Guard
    4. Context Healer Agent fixes page breaks & illegible tags
    5. Structure into chapters
    6. Generate Premium Multi-lingual PDF + DOCX
    """
    os.makedirs(output_dir, exist_ok=True)
    job_id = uuid.uuid4().hex
    scratch_dir = os.path.join(output_dir, f"scratch_{job_id}")

    try:
        if progress_callback: 
            progress_callback("collecting", 0, "Extracting pages…")
            
        images = collect_images(file_path, filename, scratch_dir)
        total_pages = len(images)
        
        if total_pages == 0:
            raise ValueError("No readable images found in the uploaded file.")

        if progress_callback: 
            progress_callback("transcribing", 5, f"Found {total_pages} pages. Enhancing and transcribing…")

        transcribed = transcribe_images(images, book_title)

        if progress_callback: 
            progress_callback("healing", 60, "AI Healer stitching broken sentences and fixing illegible words…")
            
        healed_transcription = heal_transcription_context(transcribed)

        content_pages = sum(1 for p in healed_transcription if p["has_content"])
        
        if progress_callback: 
            progress_callback("structuring", 78, f"Structuring {content_pages}/{total_pages} pages into chapters…")

        structure = structure_transcription(healed_transcription, book_title)
        
        if book_title:
            structure["title"] = book_title

        if progress_callback: 
            progress_callback("assembling", 90, "Generating PDF and DOCX (Applying Unicode layout)…")

        safe_title = "".join(c for c in structure["title"] if c.isalnum() or c in (" ", "-", "_")).strip() or "manuscript"
        pdf_path  = os.path.join(output_dir, f"{safe_title}_{job_id}.pdf")
        docx_path = os.path.join(output_dir, f"{safe_title}_{job_id}.docx")

        generate_scanned_pdf(structure, pdf_path)
        generate_scanned_docx(structure, docx_path)

        if progress_callback: 
            progress_callback("done", 100, "Complete!")

        return {
            "job_id": job_id,
            "title": structure["title"],
            "language": structure.get("language", "Unknown"),
            "total_pages": total_pages,
            "content_pages": content_pages,
            "total_words": structure.get("total_words", 0),
            "chapters": len(structure.get("chapters", [])),
            "chapter_titles": [c["title"] for c in structure.get("chapters", [])],
            "pdf_path": pdf_path,
            "docx_path": docx_path,
        }
        
    except Exception as e:
        print(f"  🚨 CRITICAL ERROR in scan_handwritten_book: {e}\n{traceback.format_exc()}")
        if progress_callback: 
            progress_callback("error", -1, f"Failed: {e}")
        raise

    finally:
        shutil.rmtree(scratch_dir, ignore_errors=True)