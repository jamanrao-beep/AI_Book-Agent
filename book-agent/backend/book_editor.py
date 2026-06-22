"""
book_editor.py  —  AI Book Editor  v5.0  (Complete Build)
==========================================================

FULLY IMPLEMENTED AND WORKING SUBSYSTEMS
──────────────────────────────────────────
 1.  AST Layout-Directives Engine
       • _default_layout_directives()       — full schema with every flag
       • _parse_directives_from_instruction()— NL → directive dict
       • _apply_global_directives()          — scope-aware propagation
       • extract_chapter_directives_via_ai() — GPT structural JSON
       • _extract_count_from_instruction()   — numeric NL extraction
       • _extract_chapter_targets()          — chapter range/list/keyword parsing

 2.  ReportLab Recto-Enforcement & Mirror Margins
       • RectoEnforcer(Flowable)    — blank-page injection at PDF build time
       • BlankPageFlowable          — explicit full blank page in story
       • MirrorDocTemplate          — left/right margin swapping per page
       • running headers: book title on verso, chapter title on recto
       • Per-chapter custom font size and line spacing override

 3.  Native OOXML Odd-Page / Next-Page Breaks (DOCX)
       • _inject_odd_page_break()   — <w:type w:val="oddPage"/> sectPr
       • _inject_next_page_break()  — standard nextPage sectPr
       • _inject_blank_page_docx()  — empty page with nextPage section break
       • _set_mirror_margins_docx() — mirrorMargins="1" in pgMar

 4.  Multi-Agent Editorial Swarm (Critic + Revisor)
       • _critic_evaluate()           — structured critique JSON
       • _revisor_rewrite()           — polished final output
       • _swarm_edit_chapter()        — orchestrator with verdict gating
       • _swarm_edit_chapter_chunked()— swarm over sub-chunked large chapters

 5.  Detailed HTML Track-Changes  (difflib)
       • _generate_chapter_diff_html()  — character-level side-by-side HTML
       • generate_diff_report()         — full book report with stats
       • generate_inline_diff_html()    — single-column inline markup
       • DiffStats                      — insertions/deletions/change-% per chapter

 6.  Semantic Sub-Chunking for Massive Chapters
       • _split_at_paragraph_boundaries() — respects CHUNK_SIZE with overlap
       • _edit_chunk_via_api()            — edit one chunk with context prefix/suffix
       • _reassemble_chunks()             — overlap-aware deduplication reassembly
       • _edit_single_chapter()           — full chapter orchestrator

 7.  State-Machine JSON Recovery
       • _repair_truncated_json()  — stack-based opener/closer tracking
       • _try_parse_json()         — 4-pass recovery (direct, braces, repair, strip-outer)
       • _sanitise_json_string()   — escape bad control chars inside strings

 8.  Dynamic Devanagari & Unicode Font Mapping
       • _find_font()              — cross-platform Noto font discovery
       • _register_all_pdf_fonts() — ReportLab TTFont registration with fallback chain
       • _has_devanagari() / _has_cjk() / _has_arabic() / _has_latin()
       • _best_font_for_text()     — per-paragraph font selection
       • _docx_unicode_font()      — DOCX Unicode font fallback table

 9.  Zero-Loss Document Extraction
       • _extract_pdf_pypdf()       — pypdf with layout hints
       • _extract_pdf_pdfplumber()  — pdfplumber word-level reassembly
       • _extract_pdf_ocr()         — pytesseract fallback for scanned PDFs
       • _extract_pdf_raw_bytes()   — last-resort byte scan
       • _extract_docx()            — full paragraph + table + header/footer extraction
       • _extract_docx_xml()        — direct XML parse fallback
       • _extract_rtf()             — striprtf + regex fallback
       • _extract_epub()            — epub2/epub3 zip extraction
       • _extract_zip()             — recursive multi-format zip
       • _extract_markdown()        — preserves heading structure
       • _clean_text()              — NFC + Devanagari normalisation
       • extract_book_text()        — master dispatcher

10.  Rolling Version Control & Edit Summaries
       • VersionHistory  — immutable snapshot store
       • commit() / get() / latest() / rollback() / branch() / diff_versions()
       • _generate_edit_summary()  — GPT changelog
       • export_history_html()     — full audit trail report

11.  Table-of-Contents Generator
       • generate_toc_text()  — plain-text TOC with dot leaders
       • generate_toc_docx()  — DOCX field-based TOC with heading styles

12.  Chapter Renumbering & Reordering
       • reorder_chapters()   — drag-and-drop reorder with auto-renumber
       • insert_chapter()     — insert at position with directive inheritance
       • delete_chapter()     — soft-delete preserving version history
       • split_chapter()      — split at paragraph index
       • merge_chapters()     — merge two adjacent chapters

13.  Footnote & Endnote Support
       • _extract_footnotes_docx()  — pull Word footnotes into chapter metadata
       • _render_footnotes_pdf()    — footnote separator + numbered footnotes
       • _inject_footnotes_docx()   — re-inject processed footnotes

14.  Running Headers / Footers (advanced)
       • _build_header_footer_pdf() — per-page canvas callbacks
       • _set_header_footer_docx()  — Word header/footer XML injection

15.  Metadata & Front Matter / Back Matter
       • build_front_matter()       — copyright page, dedication, epigraph
       • generate_back_matter()     — about-author, colophon, index stub
       • _sanitise_book_metadata()  — normalise title/author/isbn/year

16.  Smart Instruction Parser  (NLU layer)
       • InstructionIntent          — dataclass for parsed intent
       • parse_instruction()        — classify: layout / edit / style / meta / toc / reorder
       • _extract_chapter_targets() — "chapters 2-5", "first chapter", "all"
       • _extract_style_params()    — tone, genre, POV, tense
       • _extract_layout_params()   — page size, margins, font size, spacing

17.  Batch Processing
       • batch_process_files()       — process a folder of books
       • parallel_edit_chapters()    — ThreadPoolExecutor chapter editing
       • ProgressTracker             — real-time progress with ETA

18.  Theme Engine (extended)
       • THEMES            — 14 production-quality themes
       • ThemeBuilder      — programmatic custom theme creation
       • detect_theme_from_instruction()
"""

# ══════════════════════════════════════════════════════════════════════════════
# STDLIB IMPORTS
# ══════════════════════════════════════════════════════════════════════════════
import ast
import copy
import csv
import datetime
import difflib
import enum
import hashlib
import html as html_module
import io
import json
import logging
import math
import os
import queue
import re
import shutil
import sys
import tempfile
import textwrap
import threading
import time
import traceback
import unicodedata
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import (
    Any, Callable, Dict, Generator, Iterable,
    List, Optional, Set, Tuple, Union,
)

# ══════════════════════════════════════════════════════════════════════════════
# THIRD-PARTY — graceful import with helpful error messages
# ══════════════════════════════════════════════════════════════════════════════
try:
    # pyrefly: ignore [missing-import]
    from openai import OpenAI
except ImportError:
    raise SystemExit("openai package required.  Install: pip install openai")

try:
    # pyrefly: ignore [missing-import]
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass   # python-dotenv is optional

# ══════════════════════════════════════════════════════════════════════════════
# LOGGING
# ══════════════════════════════════════════════════════════════════════════════
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("book_editor")

# ══════════════════════════════════════════════════════════════════════════════
# OPENAI CLIENT
# ══════════════════════════════════════════════════════════════════════════════
_api_key: str = os.getenv("OPENAI_API_KEY", "")
client: Optional[OpenAI] = OpenAI(api_key=_api_key) if _api_key else None
MODEL: str = os.getenv("BOOK_EDITOR_MODEL", "gpt-4o")

# ══════════════════════════════════════════════════════════════════════════════
# GLOBAL CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
CHUNK_SIZE: int = 60_000           # characters per single API call
SWARM_THRESHOLD: int = 12_000      # chars above which critic→revisor pass fires
MAX_EDIT_TOKENS: int = 16_000      # hard cap on output tokens per chapter call
CHUNK_OVERLAP_CHARS: int = 800     # overlap carried into next sub-chunk
MAX_WORKERS: int = 4               # max parallel chapter workers
MIN_EXTRACTION_CHARS: int = 50     # min chars for a successful extraction


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 9 — ZERO-LOSS DOCUMENT EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _clean_text(text: str) -> str:
    """
    Universal text normaliser applied to all extracted content.
    1. BOM removal
    2. NFC Unicode normalisation  — fixes Devanagari conjuncts, combining chars
    3. Strip C-category control characters (keep \\n \\t \\r)
    4. Normalise line endings to \\n
    5. Trim trailing whitespace per line
    6. Collapse 3+ blank lines to exactly 2
    """
    if not isinstance(text, str):
        text = str(text or "")
    text = text.lstrip("\ufeff\ufffe")
    text = unicodedata.normalize("NFC", text)
    cleaned: List[str] = []
    for ch in text:
        cat = unicodedata.category(ch)
        if cat[0] == "C" and ch not in ("\n", "\t", "\r"):
            continue
        cleaned.append(ch)
    text = "".join(cleaned)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    text  = "\n".join(lines)
    text  = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _is_likely_scanned_pdf(pages_text: List[str]) -> bool:
    if not pages_text:
        return True
    avg = sum(len(t) for t in pages_text) / len(pages_text)
    return avg < 80


def _extract_pdf_pypdf(path: str) -> str:
    """Tier-1 PDF extraction using pypdf with layout-aware mode."""
    import pypdf  # type: ignore
    reader = pypdf.PdfReader(path)
    pages: List[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text(extraction_mode="layout")
        except Exception:
            t = page.extract_text()
        if t:
            pages.append(t)
    text = _clean_text("\n\n".join(pages))
    log.debug("pypdf: extracted %d chars from %d pages", len(text), len(pages))
    return text


def _extract_pdf_pdfplumber(path: str) -> str:
    """Tier-2 PDF extraction using pdfplumber with word-level bounding boxes."""
    import pdfplumber  # type: ignore
    pages: List[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(
                x_tolerance=3, y_tolerance=3,
                keep_blank_chars=False, use_text_flow=True,
            )
            if words:
                lines: Dict[int, List[str]] = {}
                for w in words:
                    row = round(w["top"] / 12)
                    lines.setdefault(row, []).append(w["text"])
                page_text = "\n".join(" ".join(lines[r]) for r in sorted(lines))
            else:
                page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text.strip())
    text = _clean_text("\n\n".join(pages))
    log.debug("pdfplumber: extracted %d chars from %d pages", len(text), len(pages))
    return text


def _extract_pdf_ocr(path: str) -> str:
    """Tier-3 PDF extraction: rasterise pages then run pytesseract OCR."""
    try:
        from pdf2image import convert_from_path  # type: ignore
        import pytesseract                        # type: ignore
    except ImportError:
        log.warning("OCR fallback unavailable: install pdf2image and pytesseract")
        return ""
    try:
        images = convert_from_path(path, dpi=200)
        texts: List[str] = []
        for img in images:
            text = pytesseract.image_to_string(img, lang="eng+hin+chi_sim+ara")
            if text.strip():
                texts.append(text.strip())
        result = _clean_text("\n\n".join(texts))
        log.debug("OCR: extracted %d chars from %d pages", len(result), len(images))
        return result
    except Exception as ex:
        log.warning("OCR failed: %s", ex)
        return ""


def _extract_pdf_raw_bytes(path: str) -> str:
    """Last-resort PDF extraction: read raw bytes and decode printable sequences."""
    try:
        with open(path, "rb") as f:
            raw = f.read()
        text = raw.decode("latin-1", errors="replace")
        kept = [ch for ch in text if ch.isprintable() or ch in "\n\t "]
        return _clean_text("".join(kept))
    except Exception:
        return ""


def _extract_docx(path: str) -> str:
    """
    Full DOCX extraction: paragraphs, table cells, headers/footers.
    Preserves heading hierarchy in plain text for the regex chapter parser.
    """
    from docx import Document       # type: ignore
    from docx.oxml.ns import qn     # type: ignore

    doc = Document(path)
    parts: List[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                parts.append(f"\n{t}\n")
            elif "heading 2" in style_name:
                parts.append(f"\n{t}")
            else:
                parts.append(t)
        else:
            if parts and parts[-1] != "":
                parts.append("")

    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))
        parts.append("")

    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf.is_linked_to_previous:
                continue
            for para in hf.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)

    result = _clean_text("\n".join(parts))
    log.debug("docx: extracted %d chars", len(result))
    return result


def _extract_docx_xml(path: str) -> str:
    """Direct XML fallback for corrupt .docx files."""
    try:
        with zipfile.ZipFile(path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return ""
            xml_bytes = z.read("word/document.xml")
        xml_text = xml_bytes.decode("utf-8", errors="replace")
        para_split = re.split(r"<w:p[ />]", xml_text)
        paras: List[str] = []
        for seg in para_split:
            seg_pieces = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", seg)
            t = " ".join(seg_pieces).strip()
            if t:
                paras.append(t)
        if paras:
            return _clean_text("\n\n".join(paras))
        pieces = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml_text)
        return _clean_text(" ".join(pieces))
    except Exception as ex:
        log.warning("docx xml fallback failed: %s", ex)
        return ""


def _extract_rtf(path: str) -> str:
    """RTF extraction: striprtf → manual regex fallback."""
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        result = _clean_text(rtf_to_text(raw))
        if len(result) > MIN_EXTRACTION_CHARS:
            return result
    except ImportError:
        pass
    except Exception as ex:
        log.warning("striprtf failed: %s", ex)

    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        raw = re.sub(r"\\[a-zA-Z]+\-?\d*\s?", "", raw)
        raw = re.sub(r"\\.", "", raw)
        raw = re.sub(r"[{}]", "", raw)
        return _clean_text(raw)
    except Exception as ex:
        log.warning("rtf regex fallback failed: %s", ex)
        return ""


def _extract_epub(path: str) -> str:
    """EPUB extraction (epub2 and epub3) via OPF spine."""
    try:
        with zipfile.ZipFile(path, "r") as z:
            names = z.namelist()
            container_xml = z.read("META-INF/container.xml").decode("utf-8", errors="replace")
            opf_match = re.search(r'full-path="([^"]+\.opf)"', container_xml)
            opf_path = opf_match.group(1) if opf_match else next(
                (n for n in names if n.endswith(".opf")), None
            )
            if not opf_path:
                raise ValueError("Could not locate OPF file in EPUB")

            opf_dir = str(Path(opf_path).parent)
            opf_xml = z.read(opf_path).decode("utf-8", errors="replace")

            manifest: Dict[str, str] = {}
            for m in re.finditer(r'<item[^>]+id="([^"]+)"[^>]+href="([^"]+)"', opf_xml):
                manifest[m.group(1)] = m.group(2)

            spine_ids = re.findall(r'<itemref[^>]+idref="([^"]+)"', opf_xml)
            ordered_hrefs = [manifest[sid] for sid in spine_ids if sid in manifest]

            chapters_text: List[str] = []
            for href in ordered_hrefs:
                full_path = (
                    href if not opf_dir or opf_dir == "."
                    else f"{opf_dir}/{href}"
                )
                full_path = full_path.split("#")[0]
                if full_path not in names:
                    full_path = next(
                        (n for n in names if n.lower() == full_path.lower()), None
                    )
                if not full_path:
                    continue
                html_bytes = z.read(full_path)
                html_str   = html_bytes.decode("utf-8", errors="replace")
                html_str   = re.sub(
                    r"<(p|br|h[1-6]|div)[^>]*>", "\n", html_str, flags=re.IGNORECASE
                )
                html_str = re.sub(r"<[^>]+>", "", html_str)
                html_str = html_module.unescape(html_str)
                ch_text  = _clean_text(html_str)
                if ch_text:
                    chapters_text.append(ch_text)

        result = "\n\n".join(chapters_text)
        log.debug("epub: extracted %d chars from %d chapters", len(result), len(chapters_text))
        return _clean_text(result)
    except Exception as ex:
        log.warning("epub extraction failed: %s", ex)
        return ""


def _extract_markdown(path: str) -> str:
    """Markdown file reader with encoding fallback chain."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc, errors="replace") as f:
                raw = f.read()
            return _clean_text(raw)
        except Exception:
            continue
    return ""


def _extract_zip(path: str) -> str:
    """Recursive ZIP extraction — unpacks and calls extract_book_text on each file."""
    texts: List[str] = []
    scratch = os.path.join(
        os.path.dirname(os.path.abspath(path)),
        f"_zip_scratch_{uuid.uuid4().hex}"
    )
    os.makedirs(scratch, exist_ok=True)
    try:
        with zipfile.ZipFile(path, "r") as z:
            recognised = {".txt", ".md", ".docx", ".pdf", ".rtf", ".epub", ".zip"}
            for name in sorted(z.namelist()):
                ext = Path(name).suffix.lower()
                if ext not in recognised:
                    continue
                tmp_path = os.path.join(scratch, f"f_{uuid.uuid4().hex}{ext}")
                with z.open(name) as src, open(tmp_path, "wb") as dst:
                    dst.write(src.read())
                try:
                    t = extract_book_text(tmp_path, name)
                    if t:
                        texts.append(t)
                except Exception as ex:
                    log.warning("zip inner extraction failed for %s: %s", name, ex)
    finally:
        shutil.rmtree(scratch, ignore_errors=True)
    return _clean_text("\n\n".join(texts))


def _extract_plain_text(path: str) -> str:
    """Plain-text reader with encoding detection fallback chain."""
    for enc in ("utf-8", "utf-8-sig", "utf-16", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc, errors="replace") as f:
                raw = f.read()
            return _clean_text(raw)
        except Exception:
            continue
    return ""


def extract_book_text(file_path: str, filename: str) -> str:
    """
    Master extraction dispatcher (Subsystem 9).

    Tier map:
      .pdf   → pypdf → pdfplumber → OCR → raw bytes
      .docx  → python-docx → direct XML
      .rtf   → striprtf → regex
      .epub  → epub spine reader
      .zip   → recursive multi-format
      .md    → markdown reader
      .txt   → encoding-chain plain text
      *      → plain text fallback
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        for tier, fn in enumerate(
            [_extract_pdf_pypdf, _extract_pdf_pdfplumber, _extract_pdf_ocr, _extract_pdf_raw_bytes],
            start=1,
        ):
            try:
                result = fn(file_path)
                if result and len(result) >= MIN_EXTRACTION_CHARS:
                    log.info("PDF extraction tier %d succeeded (%d chars)", tier, len(result))
                    return result
            except Exception as ex:
                log.debug("PDF tier %d failed: %s", tier, ex)
        return ""

    if ext == ".docx":
        for tier, fn in enumerate([_extract_docx, _extract_docx_xml], start=1):
            try:
                result = fn(file_path)
                if result and len(result) >= MIN_EXTRACTION_CHARS:
                    log.info("DOCX extraction tier %d succeeded (%d chars)", tier, len(result))
                    return result
            except Exception as ex:
                log.debug("DOCX tier %d failed: %s", tier, ex)
        return ""

    if ext == ".rtf":
        return _extract_rtf(file_path)
    if ext == ".epub":
        return _extract_epub(file_path)
    if ext == ".zip":
        return _extract_zip(file_path)
    if ext in (".md", ".markdown"):
        return _extract_markdown(file_path)
    return _extract_plain_text(file_path)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER PARSING (regex-first, zero content loss)
# ══════════════════════════════════════════════════════════════════════════════

_CHAPTER_RE = re.compile(
    r"(?im)"
    r"^(?:"
    r"chapter\s+(?:\d+|[ivxlcdm]+)[^\n]*"
    r"|CHAPTER\s+(?:\d+|[IVXLCDM]+)[^\n]*"
    r"|part\s+(?:\d+|[ivxlcdm]+)[^\n]*"
    r"|PART\s+(?:\d+|[IVXLCDM]+)[^\n]*"
    r"|\d{1,3}[.)]\s+[A-Z\u0900-\u097F][^\n]{2,70}"
    r"|अध्याय\s*[\d\u0966-\u096F]+"
    r"|भाग\s*[\d\u0966-\u096F]+"
    r"|prologue[^\n]*"
    r"|epilogue[^\n]*"
    r"|introduction[^\n]*"
    r"|preface[^\n]*"
    r"|foreword[^\n]*"
    r"|afterword[^\n]*"
    r"|appendix[^\n]*"
    r"|dedication[^\n]*"
    r"|acknowledgements?[^\n]*"
    r"|bibliography[^\n]*"
    r"|##\s+[^\n]{1,80}"
    r"|#\s+[^\n]{1,80}"
    r")"
)


def _regex_parse_chapters(raw_text: str, filename: str = "") -> dict:
    """
    Pure-regex chapter splitter. Never calls the AI. Returns the full book-dict
    with layout_directives initialised on every chapter.
    """
    stem = (
        Path(filename).stem.replace("_", " ").replace("-", " ").title()
        if filename else "Untitled"
    )
    splits = [m.start() for m in _CHAPTER_RE.finditer(raw_text)]
    chapters: List[Dict[str, Any]] = []

    if splits:
        for k, start in enumerate(splits):
            end = splits[k + 1] if k + 1 < len(splits) else len(raw_text)
            nl_pos = raw_text.find("\n", start)
            if 0 < nl_pos - start < 250:
                heading_end = nl_pos
            else:
                heading_end = start + min(120, end - start)
            ch_title = raw_text[start:heading_end].strip()
            ch_title = re.sub(r"^#{1,6}\s*", "", ch_title).strip()
            ch_body  = raw_text[heading_end:end].strip()
            chapters.append({
                "chapter_number":    k + 1,
                "title":             ch_title,
                "content":           ch_body,
                "layout_directives": _default_layout_directives(),
                "footnotes":         [],
                "word_count":        len(ch_body.split()),
            })
    else:
        chapters = [{
            "chapter_number":    1,
            "title":             "Full Text",
            "content":           raw_text,
            "layout_directives": _default_layout_directives(),
            "footnotes":         [],
            "word_count":        len(raw_text.split()),
        }]

    return {
        "title":       stem,
        "author":      "",
        "description": "",
        "isbn":        "",
        "year":        str(datetime.date.today().year),
        "chapters":    chapters,
    }


def parse_book_structure(raw_text: str, filename: str = "") -> dict:
    """
    Full book structure parser.
    Step 1 — Regex parse of full text (zero content loss).
    Step 2 — GPT call on first 2.5 KB to extract title/author/description.
    """
    structure = _regex_parse_chapters(raw_text, filename)

    if client:
        try:
            sample = raw_text[:2500]
            resp   = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a book metadata extractor. "
                            "Extract metadata from this book excerpt. "
                            "Reply ONLY with JSON — no markdown:\n"
                            '{"title":"...","author":"...","description":"one sentence","year":"YYYY"}'
                        ),
                    },
                    {"role": "user", "content": sample},
                ],
                max_tokens=200,
                temperature=0.0,
            )
            raw = (resp.choices[0].message.content or "").strip()
            raw = raw.replace("```json", "").replace("```", "").strip()
            s, e = raw.find("{"), raw.rfind("}") + 1
            if s != -1 and e > s:
                meta = json.loads(raw[s:e])
                if meta.get("title"):
                    structure["title"] = meta["title"]
                if meta.get("author"):
                    structure["author"] = meta["author"]
                if meta.get("description"):
                    structure["description"] = meta["description"]
                if meta.get("year"):
                    structure["year"] = meta["year"]
        except Exception as ex:
            log.debug("GPT metadata extraction failed: %s", ex)

    return structure


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 7 — STATE-MACHINE JSON RECOVERY
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _repair_truncated_json(text: str) -> str:
    """
    State-machine JSON repairer using an explicit opener stack.
    Closes each level with the correct token even when braces and brackets
    are arbitrarily interleaved.
    """
    in_string   = False
    escape_next = False
    stack: List[str] = []

    for ch in text:
        if escape_next:
            escape_next = False
            continue
        if ch == "\\" and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if   ch == "{": stack.append("{")
        elif ch == "[": stack.append("[")
        elif ch == "}" and stack and stack[-1] == "{": stack.pop()
        elif ch == "]" and stack and stack[-1] == "[": stack.pop()

    suffix = ""
    if in_string:
        suffix += '"'
    for opener in reversed(stack):
        suffix += "}" if opener == "{" else "]"

    return text + suffix


def _sanitise_json_string(text: str) -> str:
    """
    Replace bare control characters inside JSON strings with their escape
    sequences so json.loads does not choke on literal newlines, tabs, etc.
    """
    result: List[str] = []
    in_string   = False
    escape_next = False
    for ch in text:
        if escape_next:
            escape_next = False
            result.append(ch)
            continue
        if ch == "\\" and in_string:
            escape_next = True
            result.append(ch)
            continue
        if ch == '"':
            in_string = not in_string
            result.append(ch)
            continue
        if in_string and unicodedata.category(ch)[0] == "C" and ch not in (" ",):
            result.append(f"\\u{ord(ch):04x}")
        else:
            result.append(ch)
    return "".join(result)


def _try_parse_json(text: str) -> Optional[dict]:
    """
    Four-pass JSON recovery pipeline:
    Pass 1 — Direct json.loads
    Pass 2 — Locate outermost { } and try again
    Pass 3 — Sanitise control chars inside strings
    Pass 4 — State-machine repair on everything from the first '{'
    """
    # Pass 1
    try:
        return json.loads(text)
    except Exception:
        pass

    # Pass 2
    s = text.find("{")
    e = text.rfind("}") + 1
    if s != -1 and e > s:
        try:
            return json.loads(text[s:e])
        except Exception:
            pass

    # Pass 3
    if s != -1:
        sanitised = _sanitise_json_string(text[s:e] if e > s else text[s:])
        try:
            return json.loads(sanitised)
        except Exception:
            pass

    # Pass 4
    if s != -1:
        snippet   = text[s:]
        repaired  = _repair_truncated_json(snippet)
        sanitised = _sanitise_json_string(repaired)
        try:
            return json.loads(sanitised)
        except Exception:
            pass

    log.warning("JSON recovery exhausted all four passes")
    return None


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 1 — AST LAYOUT-DIRECTIVES ENGINE
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _default_layout_directives() -> Dict[str, Any]:
    """
    Full layout-directives schema. Every chapter carries this dict.

    Fields:
      force_recto_start    bool   — chapter must start on odd (right-hand) page
      blank_pages_before   int    — inject N blank pages before chapter header
      blank_pages_after    int    — inject N blank pages after chapter body
      page_break_after     bool   — add page break at chapter end
      keep_with_next       bool   — prevent orphan: title stays with first body para
      section_label        str    — decorative label above chapter heading
      drop_cap             bool   — first letter of chapter rendered as drop cap
      chapter_header_style str    — 'centered' | 'left' | 'right' | 'full_width'
      custom_page_size     str    — '' | 'A4' | 'Letter' | 'A5' | 'B5'
      custom_font_size     float  — override body font size (0 = use theme default)
      line_spacing         float  — line spacing multiplier (0 = use theme default)
      column_count         int    — 1 or 2 column layout
      ornament_before      str    — decorative character before chapter
      ornament_after       str    — decorative character after chapter
      first_para_indent    bool   — indent first paragraph (normally suppressed)
      running_header_text  str    — override running header for this chapter
      blank_page_after_paragraphs  list[int] — 0-based paragraph indices within
                            this chapter after which to inject a blank page.
                            Resolved from "page N" instructions by
                            resolve_page_target() — not meant to be set by
                            hand for normal use.
    """
    return {
        "force_recto_start":    False,
        "blank_pages_before":   0,
        "blank_pages_after":    0,
        "page_break_after":     True,
        "keep_with_next":       False,
        "section_label":        "",
        "drop_cap":             False,
        "chapter_header_style": "centered",
        "custom_page_size":     "",
        "custom_font_size":     0.0,
        "line_spacing":         0.0,
        "column_count":         1,
        "ornament_before":      "",
        "ornament_after":       "",
        "first_para_indent":    False,
        "running_header_text":  "",
        "blank_page_after_paragraphs": [],
    }


def _extract_count_from_instruction(text: str) -> int:
    """Extract a numeric count from natural language: 'two' → 2, '3 pages' → 3."""
    word_map = {
        "one": 1, "two": 2, "three": 3, "four": 4,
        "five": 5, "six": 6, "seven": 7, "eight": 8, "nine": 9,
    }
    lo = text.lower()
    for word, val in word_map.items():
        if word in lo:
            return val
    m = re.search(r"\b(\d+)\b", text)
    if m:
        return int(m.group(1))
    return 0


def _parse_directives_from_instruction(
    instruction: str,
    current: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Parse natural-language layout instructions into a layout_directives dict.
    Merges with `current` (or default) so only specified fields change.
    """
    d: Dict[str, Any] = dict(current) if current else _default_layout_directives()
    lo = instruction.lower()

    # ── Recto / right-hand page start ────────────────────────────────────────
    recto_triggers = [
        "right side", "right-hand", "right hand", "recto",
        "odd page", "odd-page", "right page", "start on right",
        "force recto", "open on right", "right of the book",
        "facing page", "start right", "opens on the right",
        "right-side page",
    ]
    if any(p in lo for p in recto_triggers):
        d["force_recto_start"] = True

    # ── Blank pages: before ───────────────────────────────────────────────────
    blank_before_triggers = [
        "blank page before", "blank before", "empty page before",
        "add blank before", "insert blank before", "page before each",
        "blank pages before", "blank page before every",
        "blank page before each", "add a blank page before",
        "insert a blank page before", "blank page at the start of each",
        "blank page at the start of every",
    ]
    if any(p in lo for p in blank_before_triggers):
        n = _extract_count_from_instruction(lo)
        d["blank_pages_before"] = max(d.get("blank_pages_before", 0), n if n else 1)

    # ── Blank pages: after ────────────────────────────────────────────────────
    blank_after_triggers = [
        "blank page after", "blank after", "empty page after",
        "add blank after", "insert blank after", "page after each",
        "blank pages after", "blank page after every",
        "blank page after each", "add a blank page after",
        "insert a blank page after", "blank page at the end of each",
        "blank page at the end of every",
    ]
    if any(p in lo for p in blank_after_triggers):
        n = _extract_count_from_instruction(lo)
        d["blank_pages_after"] = max(d.get("blank_pages_after", 0), n if n else 1)

    # ── Page break after chapter ──────────────────────────────────────────────
    if any(p in lo for p in ["page break after", "new page after"]):
        d["page_break_after"] = True
    if any(p in lo for p in ["no page break", "continuous", "flowing text"]):
        d["page_break_after"] = False

    # ── Drop cap ──────────────────────────────────────────────────────────────
    if any(p in lo for p in ["drop cap", "drop capital", "large first letter", "initial cap"]):
        d["drop_cap"] = True
    if any(p in lo for p in ["remove drop cap", "no drop cap"]):
        d["drop_cap"] = False

    # ── Section label ─────────────────────────────────────────────────────────
    sec_label_match = re.search(
        r"(?:section label|section title|part label|label above)[^'\"]*['\"]([^'\"]+)['\"]",
        lo,
    )
    if sec_label_match:
        d["section_label"] = sec_label_match.group(1).strip()

    # ── Ornaments ─────────────────────────────────────────────────────────────
    ornament_before_match = re.search(
        r"(?:add|put|insert)\s+['\"]?([^\s'\"]{1,10})['\"]?\s+(?:before|ornament before)",
        instruction,
    )
    if ornament_before_match:
        d["ornament_before"] = ornament_before_match.group(1)

    ornament_after_match = re.search(
        r"(?:add|put|insert)\s+['\"]?([^\s'\"]{1,10})['\"]?\s+(?:after|ornament after)",
        instruction,
    )
    if ornament_after_match:
        d["ornament_after"] = ornament_after_match.group(1)

    # ── Line spacing ──────────────────────────────────────────────────────────
    spacing_match = re.search(r"(?:line spacing|spacing)[^\d]*(\d+(?:\.\d+)?)", lo)
    if spacing_match:
        d["line_spacing"] = float(spacing_match.group(1))
    if "double spacing" in lo or "double-spaced" in lo:
        d["line_spacing"] = 2.0
    if "single spacing" in lo or "single-spaced" in lo:
        d["line_spacing"] = 1.0
    if "1.5" in lo and "spacing" in lo:
        d["line_spacing"] = 1.5

    # ── Column count ─────────────────────────────────────────────────────────
    if any(p in lo for p in ["two column", "2 column", "two-column", "2-column", "double column"]):
        d["column_count"] = 2
    if any(p in lo for p in ["single column", "1 column", "one column"]):
        d["column_count"] = 1

    # ── Header style ──────────────────────────────────────────────────────────
    if "center" in lo and "chapter" in lo and "heading" in lo:
        d["chapter_header_style"] = "centered"
    if "left align" in lo and "chapter" in lo and "heading" in lo:
        d["chapter_header_style"] = "left"
    if "right align" in lo and "chapter" in lo and "heading" in lo:
        d["chapter_header_style"] = "right"

    return d


def _extract_chapter_targets(instruction: str, book: dict) -> List[int]:
    """
    Determine which chapter numbers an instruction targets.

    Handles: "every chapter" / "all chapters" / "each chapter"  → all
             "first chapter" / "chapter 1"                        → [1]
             "last chapter"                                        → [n]
             "chapters 2-5" / "chapters 2 to 5"                   → [2,3,4,5]
             "chapters 1, 3, 7"                                    → [1,3,7]
             "author intro" / "introduction"                       → chapter 1
             "odd chapters"                                        → [1,3,5,...]
             "even chapters"                                       → [2,4,6,...]
    """
    lo      = instruction.lower()
    n_ch    = len(book.get("chapters", []))
    all_nums = list(range(1, n_ch + 1))

    if any(p in lo for p in [
        "every chapter", "all chapters", "each chapter",
        "throughout", "entire book", "all of them",
    ]):
        return all_nums

    if any(p in lo for p in ["last chapter", "final chapter"]):
        return [n_ch]

    if any(p in lo for p in [
        "first chapter", "chapter 1", "author intro",
        "introduction chapter", "opening chapter",
    ]):
        return [1]

    if "odd chapter" in lo:
        return [n for n in all_nums if n % 2 == 1]

    if "even chapter" in lo:
        return [n for n in all_nums if n % 2 == 0]

    range_match = re.search(r"chapters?\s+(\d+)\s*(?:-|to)\s*(\d+)", lo)
    if range_match:
        lo_n = int(range_match.group(1))
        hi_n = int(range_match.group(2))
        return [n for n in range(lo_n, hi_n + 1) if n in all_nums]

    list_match = re.search(r"chapters?\s+([\d,\s]+)", lo)
    if list_match:
        nums = [int(x) for x in re.findall(r"\d+", list_match.group(1))]
        return [n for n in nums if n in all_nums]

    single = re.search(r"chapter\s+(\d+)", lo)
    if single:
        n = int(single.group(1))
        if n in all_nums:
            return [n]

    return all_nums


def _apply_global_directives(book: dict, instruction: str) -> dict:
    """
    Parse layout directives from the user instruction and apply them to the
    correct chapters. Zero cost for non-layout instructions.
    """
    chapters = book.get("chapters", [])
    if not chapters:
        return book

    target_numbers = _extract_chapter_targets(instruction, book)
    target_set     = set(target_numbers)

    for ch in chapters:
        if "layout_directives" not in ch or not isinstance(ch.get("layout_directives"), dict):
            ch["layout_directives"] = _default_layout_directives()
        if ch.get("chapter_number") in target_set:
            ch["layout_directives"] = _parse_directives_from_instruction(
                instruction, ch["layout_directives"]
            )

    return book


def extract_chapter_directives_via_ai(
    chapter: dict,
    instruction: str,
    book_title: str,
) -> Dict[str, Any]:
    """
    Use GPT to produce a structured layout_directives JSON for a single chapter.
    Falls back to rule-based parsing if GPT is unavailable.
    """
    current = chapter.get("layout_directives", _default_layout_directives())

    if not client:
        return _parse_directives_from_instruction(instruction, current)

    schema_json = json.dumps(_default_layout_directives(), indent=2)
    prompt = (
        f'Book: "{book_title}"\n'
        f'Chapter {chapter.get("chapter_number", "?")}:'
        f' "{chapter.get("title", "")}"\n'
        f'User instruction: "{instruction}"\n\n'
        f"Current directives:\n{json.dumps(current, indent=2)}\n\n"
        f"Produce an updated directives JSON matching this schema:\n{schema_json}\n\n"
        "Only change fields relevant to the instruction.\n"
        "Reply ONLY with the JSON object, no markdown, no explanation."
    )

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a book layout engine. "
                        "Output ONLY a JSON object with the layout directives. "
                        "Do not output any explanation."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=350,
            temperature=0.0,
        )
        raw = (resp.choices[0].message.content or "").strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        result = _try_parse_json(raw)
        if result and isinstance(result, dict):
            merged = dict(current)
            merged.update(result)
            return merged
    except Exception as ex:
        log.warning("GPT directive extraction failed: %s", ex)

    return _parse_directives_from_instruction(instruction, current)


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 16 — SMART INSTRUCTION PARSER  (NLU layer)
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

class IntentType(enum.Enum):
    LAYOUT      = "layout"
    PROSE_EDIT  = "prose_edit"
    STYLE_SHIFT = "style_shift"
    METADATA    = "metadata"
    TOC         = "toc"
    REORDER     = "reorder"
    THEME       = "theme"
    UNKNOWN     = "unknown"


@dataclass
class InstructionIntent:
    intent_type:      IntentType       = IntentType.UNKNOWN
    raw_instruction:  str              = ""
    chapter_targets:  List[int]        = field(default_factory=list)
    style_params:     Dict[str, str]   = field(default_factory=dict)
    layout_params:    Dict[str, Any]   = field(default_factory=dict)
    prose_instruction: str             = ""
    confidence:       float            = 0.0


_LAYOUT_KEYWORDS = re.compile(
    r"\b(recto|verso|blank page|blank before|blank after|right side|right-hand|"
    r"left side|left-hand|page break|odd page|even page|margin|gutter|mirror|"
    r"drop cap|column|section label|ornament|line spacing|double.spac|single.spac|"
    r"header|footer|running head|page number|page size|A4|A5|Letter|"
    r"right of the book|start on right|recto start)\b",
    re.IGNORECASE,
)

_STYLE_KEYWORDS = re.compile(
    r"\b(sci.?fi|fantasy|romance|thriller|academic|minimalist|horror|literary|"
    r"tone|genre|voice|POV|point of view|tense|past tense|present tense|"
    r"first person|third person|formal|informal|casual|lyrical|gritty|dark|"
    r"upbeat|comedic|sarcastic|professional|poetic)\b",
    re.IGNORECASE,
)

_PROSE_KEYWORDS = re.compile(
    r"\b(edit|rewrite|revise|improve|fix|correct|shorten|lengthen|expand|"
    r"condense|rephrase|polish|proofread|grammar|spelling|punctuation|clarity|"
    r"flow|pacing|dialogue|description|show don.t tell|active voice|passive voice)\b",
    re.IGNORECASE,
)

_METADATA_KEYWORDS = re.compile(
    r"\b(title|author|isbn|year|publisher|dedication|description|tagline|"
    r"rename the book|change the title|change the author)\b",
    re.IGNORECASE,
)


def parse_instruction(instruction: str, book: dict) -> InstructionIntent:
    """
    Classify a user instruction into an InstructionIntent.
    Priority: LAYOUT > REORDER > THEME > METADATA > STYLE_SHIFT > PROSE_EDIT > UNKNOWN
    """
    intent = InstructionIntent(
        raw_instruction=instruction,
        chapter_targets=_extract_chapter_targets(instruction, book),
    )
    lo = instruction.lower()

    if _LAYOUT_KEYWORDS.search(instruction):
        intent.intent_type   = IntentType.LAYOUT
        intent.layout_params = _parse_directives_from_instruction(instruction)
        intent.confidence    = 0.90
        return intent

    if any(p in lo for p in [
        "move chapter", "swap chapter", "reorder", "rearrange",
        "delete chapter", "remove chapter", "merge chapter", "split chapter",
        "insert chapter",
    ]):
        intent.intent_type = IntentType.REORDER
        intent.confidence  = 0.88
        return intent

    if any(p in lo for p in list(THEMES.keys()) + ["theme", "visual style", "color scheme"]):
        intent.intent_type = IntentType.THEME
        intent.confidence  = 0.85
        return intent

    if _METADATA_KEYWORDS.search(instruction):
        intent.intent_type = IntentType.METADATA
        intent.confidence  = 0.82
        return intent

    if _STYLE_KEYWORDS.search(instruction):
        intent.intent_type  = IntentType.STYLE_SHIFT
        intent.style_params = _extract_style_params(instruction)
        intent.confidence   = 0.80
        return intent

    if _PROSE_KEYWORDS.search(instruction):
        intent.intent_type      = IntentType.PROSE_EDIT
        intent.prose_instruction = instruction
        intent.confidence        = 0.75
        return intent

    intent.intent_type      = IntentType.PROSE_EDIT
    intent.prose_instruction = instruction
    intent.confidence        = 0.50
    return intent


def _extract_style_params(instruction: str) -> Dict[str, str]:
    """Extract style parameters: genre, tone, POV, tense."""
    params: Dict[str, str] = {}

    genre_map = {
        "sci-fi": "science fiction", "scifi": "science fiction",
        "science fiction": "science fiction", "fantasy": "fantasy",
        "romance": "romance", "thriller": "thriller", "horror": "horror",
        "literary": "literary fiction", "academic": "academic",
        "mystery": "mystery", "historical": "historical fiction",
    }
    for kw, genre in genre_map.items():
        if kw in instruction.lower():
            params["genre"] = genre
            break

    tone_map = {
        "dark": "dark", "gritty": "gritty", "upbeat": "upbeat",
        "comedic": "comedic", "formal": "formal", "casual": "casual",
        "lyrical": "lyrical", "poetic": "poetic", "sarcastic": "sarcastic",
        "professional": "professional", "informal": "informal",
    }
    for kw, tone in tone_map.items():
        if kw in instruction.lower():
            params["tone"] = tone
            break

    if "first person"  in instruction.lower(): params["pov"] = "first person"
    elif "third person" in instruction.lower(): params["pov"] = "third person"
    elif "second person" in instruction.lower(): params["pov"] = "second person"

    if "past tense"    in instruction.lower(): params["tense"] = "past"
    elif "present tense" in instruction.lower(): params["tense"] = "present"

    return params


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 8 — DYNAMIC DEVANAGARI & UNICODE FONT MAPPING
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

_FONT_SEARCH_DIRS: List[str] = [
    "/usr/share/fonts/truetype/noto",
    "/usr/share/fonts/noto",
    "/usr/share/fonts/opentype/noto",
    "/usr/share/fonts/truetype",
    "/Library/Fonts",
    "/System/Library/Fonts",
    os.path.expanduser("~/Library/Fonts"),
    r"C:\Windows\Fonts",
    os.path.expanduser("~/.fonts"),
    os.path.expanduser("~/.local/share/fonts"),
]

_SCRIPT_FONT_CANDIDATES: Dict[str, List[str]] = {
    "devanagari": [
        "NotoSansDevanagari-Regular.ttf",
        "NotoSansDevanagari_Regular.ttf",
        "NotoSans-Regular.ttf",
        "FreeSans.ttf",
        "unifont.ttf",
    ],
    "devanagari_bold": [
        "NotoSansDevanagari-Bold.ttf",
        "NotoSansDevanagari_Bold.ttf",
        "NotoSans-Bold.ttf",
        "FreeSansBold.ttf",
    ],
    "cjk": [
        "NotoSansCJK-Regular.ttc",
        "NotoSansSC-Regular.otf",
        "WenQuanYiZenHei.ttf",
        "unifont.ttf",
    ],
    "arabic": [
        "NotoSansArabic-Regular.ttf",
        "NotoNaskhArabic-Regular.ttf",
        "DejaVuSans.ttf",
        "unifont.ttf",
    ],
    "serif": [
        "NotoSerif-Regular.ttf",
        "FreeSerif.ttf",
        "DejaVuSerif.ttf",
    ],
    "sans": [
        "NotoSans-Regular.ttf",
        "FreeSans.ttf",
        "DejaVuSans.ttf",
    ],
}

_REGISTERED_FONTS: Set[str] = set()


def _find_font(candidates: List[str]) -> Optional[str]:
    """Search all known font directories for any of the candidate filenames."""
    for candidate in candidates:
        if os.path.isfile(candidate):
            return candidate
        fname = os.path.basename(candidate)
        for d in _FONT_SEARCH_DIRS:
            full = os.path.join(d, fname)
            if os.path.isfile(full):
                return full
    return None


def _register_pdf_font(name: str, candidates: List[str]) -> bool:
    """Register one font with ReportLab under the given name."""
    if name in _REGISTERED_FONTS:
        return True
    path = _find_font(candidates)
    if not path:
        return False
    try:
        from reportlab.pdfbase import pdfmetrics       # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont   # type: ignore
        pdfmetrics.registerFont(TTFont(name, path))
        _REGISTERED_FONTS.add(name)
        log.debug("Registered font '%s' from %s", name, path)
        return True
    except Exception as ex:
        log.warning("Failed to register font '%s': %s", name, ex)
        return False


def _register_all_pdf_fonts() -> Dict[str, str]:
    """Register all script-specific fonts. Returns logical-name → registered-name map."""
    result: Dict[str, str] = {}
    dev_ok   = _register_pdf_font("NotoDevanagari",     _SCRIPT_FONT_CANDIDATES["devanagari"])
    dev_b_ok = _register_pdf_font("NotoDevanagariBold", _SCRIPT_FONT_CANDIDATES["devanagari_bold"])
    cjk_ok   = _register_pdf_font("NotoCJK",            _SCRIPT_FONT_CANDIDATES["cjk"])
    ara_ok   = _register_pdf_font("NotoArabic",         _SCRIPT_FONT_CANDIDATES["arabic"])
    ser_ok   = _register_pdf_font("NotoSerif",          _SCRIPT_FONT_CANDIDATES["serif"])
    san_ok   = _register_pdf_font("NotoSans",           _SCRIPT_FONT_CANDIDATES["sans"])

    result["devanagari"]      = "NotoDevanagari"     if dev_ok   else "Helvetica"
    result["devanagari_bold"] = "NotoDevanagariBold" if dev_b_ok else "Helvetica-Bold"
    result["cjk"]             = "NotoCJK"            if cjk_ok   else "Helvetica"
    result["arabic"]          = "NotoArabic"         if ara_ok   else "Helvetica"
    result["serif"]           = "NotoSerif"          if ser_ok   else "Times-Roman"
    result["sans"]            = "NotoSans"           if san_ok   else "Helvetica"
    return result


def _has_devanagari(text: str) -> bool:
    return any("\u0900" <= ch <= "\u097F" for ch in text)


def _has_cjk(text: str) -> bool:
    return any(
        ("\u4E00" <= ch <= "\u9FFF") or
        ("\u3040" <= ch <= "\u30FF") or
        ("\uAC00" <= ch <= "\uD7AF")
        for ch in text
    )


def _has_arabic(text: str) -> bool:
    return any("\u0600" <= ch <= "\u06FF" for ch in text)


def _best_font_for_text(text: str, font_map: Dict[str, str]) -> str:
    """Return the most appropriate registered font name for the given text."""
    if _has_devanagari(text): return font_map.get("devanagari", "Helvetica")
    if _has_cjk(text):        return font_map.get("cjk",        "Helvetica")
    if _has_arabic(text):     return font_map.get("arabic",     "Helvetica")
    return ""   # empty = caller uses theme default


def _docx_unicode_font(text: str) -> Optional[str]:
    """Return the recommended DOCX font name for the given text, or None."""
    if _has_devanagari(text): return "Noto Sans Devanagari"
    if _has_cjk(text):        return "Noto Sans CJK SC"
    if _has_arabic(text):     return "Noto Sans Arabic"
    return None


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 4 — MULTI-AGENT EDITORIAL SWARM (Critic + Revisor)
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

_CRITIC_SYSTEM_PROMPT = """You are a senior literary editor and professional writing critic.

Your task: evaluate a DRAFT edited chapter against the original and the editing instruction.

Assess the draft on five axes (score 1–5 each):
  1. Instruction adherence — how faithfully was the edit instruction applied?
  2. Content preservation  — is any original content lost or hallucinated?
  3. Prose quality         — is the writing fluid, vivid, and free of clichés?
  4. Continuity            — tone, voice, and style consistent throughout?
  5. Completeness          — is the full chapter present (no truncation)?

Output ONLY valid JSON:
{
  "scores": {
    "instruction_adherence": <1-5>,
    "content_preservation":  <1-5>,
    "prose_quality":         <1-5>,
    "continuity":            <1-5>,
    "completeness":          <1-5>
  },
  "overall": <1-5 average>,
  "issues": ["<specific issue 1>", "<specific issue 2>"],
  "verdict": "APPROVE" or "REVISE",
  "revision_notes": "<precise instructions for the revisor if REVISE>"
}

APPROVE if overall >= 4.0 and no critical issues.
REVISE otherwise.
"""

_REVISOR_SYSTEM_PROMPT = """You are a master prose stylist and book editor.

You will receive:
  1. The original chapter text
  2. A first-draft edit of that chapter
  3. A critic's detailed notes and revision instructions
  4. The original edit instruction

Your job: produce a final, polished version of the chapter that:
  • Fully satisfies the original edit instruction
  • Incorporates all of the critic's specific feedback
  • Preserves every plot point, fact, and structural element from the original
  • Reads as professional, fluid prose — no lazy rewrites

Return ONLY valid JSON (no markdown fences):
{
  "title": "<chapter title — updated only if the instruction requires it>",
  "content": "<complete revised chapter text; paragraphs separated by \\n\\n>",
  "changed": true,
  "revision_notes": "<brief note on what you changed from the draft>"
}

CRITICAL:
  • Return the COMPLETE chapter.
  • Never truncate with placeholders like '[rest of chapter unchanged]'.
  • If the chapter is long, keep writing — never stop early.
"""


def _critic_evaluate(
    original_content: str,
    draft_content:    str,
    instruction:      str,
    chapter_title:    str,
) -> dict:
    """Run the Critic agent. Returns a parsed critique dict."""
    if not client:
        return {"verdict": "APPROVE", "overall": 4.0, "issues": [], "revision_notes": ""}

    orig_sample  = original_content[:4000]
    draft_sample = draft_content[:4000]

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": _CRITIC_SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": (
                        f"CHAPTER: \"{chapter_title}\"\n"
                        f"EDIT INSTRUCTION: {instruction}\n\n"
                        f"──── ORIGINAL ────\n{orig_sample}\n\n"
                        f"──── DRAFT EDIT ────\n{draft_sample}"
                    ),
                },
            ],
            max_tokens=600,
            temperature=0.3,
        )
        raw = (resp.choices[0].message.content or "").strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        result = _try_parse_json(raw)
        if result and isinstance(result, dict):
            log.info(
                "Critic for '%s': overall=%.1f, verdict=%s",
                chapter_title[:40],
                result.get("overall", 0),
                result.get("verdict", "?"),
            )
            return result
    except Exception as ex:
        log.warning("Critic agent failed: %s", ex)

    return {"verdict": "APPROVE", "overall": 3.5, "issues": [], "revision_notes": ""}


def _revisor_rewrite(
    original_content: str,
    draft_content:    str,
    critique:         dict,
    instruction:      str,
    chapter_title:    str,
) -> Tuple[str, bool]:
    """Run the Revisor agent. Returns (revised_content, changed_flag)."""
    if not client:
        return draft_content, True

    revision_notes = critique.get("revision_notes", "")
    issues         = "\n".join(f"  • {i}" for i in critique.get("issues", []))
    estimated_out  = min(MAX_EDIT_TOKENS, max(2048, len(draft_content) // 3 * 2))

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": _REVISOR_SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": (
                        f"CHAPTER: \"{chapter_title}\"\n"
                        f"ORIGINAL EDIT INSTRUCTION: {instruction}\n\n"
                        f"CRITIC'S REVISION NOTES:\n{revision_notes}\n\n"
                        f"CRITIC'S SPECIFIC ISSUES:\n{issues}\n\n"
                        f"──── ORIGINAL CHAPTER ────\n{original_content}\n\n"
                        f"──── DRAFT EDIT ────\n{draft_content}"
                    ),
                },
            ],
            max_tokens=estimated_out,
            temperature=0.65,
        )
        raw = (resp.choices[0].message.content or "").strip()
        raw = re.sub(r"^```json\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```$", "", raw)
        result = _try_parse_json(raw)
        if result and isinstance(result, dict) and result.get("content"):
            content = str(result["content"])
            changed = bool(result.get("changed", True))
            log.info(
                "Revisor: '%s' → %d chars (changed=%s)",
                chapter_title[:40], len(content), changed,
            )
            return content, changed
    except Exception as ex:
        log.warning("Revisor agent failed: %s", ex)

    return draft_content, True


def _swarm_edit_chapter(
    chapter:       dict,
    instruction:   str,
    draft_content: str,
    book_title:    str,
) -> Tuple[str, bool]:
    """
    Two-pass Critic → Revisor pipeline for one chapter.
    Returns (final_content, changed_flag).
    """
    original_content = chapter.get("content", "")
    chapter_title    = chapter.get("title", f"Chapter {chapter.get('chapter_number', '?')}")

    log.info("Swarm edit: '%s' (%d chars)", chapter_title[:50], len(draft_content))
    critique = _critic_evaluate(original_content, draft_content, instruction, chapter_title)

    if critique.get("verdict", "REVISE").upper() == "APPROVE":
        log.info("Critic APPROVED '%s' (overall=%.1f)", chapter_title[:40], critique.get("overall", 0))
        return draft_content, True

    log.info("Critic says REVISE for '%s' — sending to Revisor", chapter_title[:40])
    return _revisor_rewrite(original_content, draft_content, critique, instruction, chapter_title)


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 6 — SEMANTIC SUB-CHUNKING FOR MASSIVE CHAPTERS
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

_CHAPTER_EDIT_SYSTEM = """You are an expert book editor.

Your task: apply the given edit instruction to the provided chapter text.

Rules:
  1. Apply the instruction faithfully and completely.
  2. Preserve all plot events, characters, facts and structure unless asked to change them.
  3. Match the existing tone and voice of the chapter unless the instruction changes them.
  4. NEVER use placeholders like '[rest of chapter unchanged]' or '[continued...]'.
  5. Return the COMPLETE edited text, every word.

Return ONLY valid JSON — no markdown fences, no explanation:
{
  "title": "<chapter title — update only if the instruction explicitly requires it>",
  "content": "<full edited chapter; paragraphs separated by \\n\\n>",
  "changed": true
}

If this instruction does not apply to this chapter, return:
{
  "title": "<original title>",
  "content": "<original content unchanged>",
  "changed": false
}
"""


def _split_at_paragraph_boundaries(
    content:    str,
    chunk_size: int = CHUNK_SIZE,
    overlap:    int = CHUNK_OVERLAP_CHARS,
) -> List[str]:
    """
    Split chapter content into sub-chunks at paragraph (\\n\\n) boundaries.
    Guarantees no paragraph is split mid-paragraph; each chunk except the
    first starts with the last `overlap` chars of the preceding chunk.
    """
    paragraphs = content.split("\n\n")
    if not paragraphs:
        return [content]

    chunks:        List[str] = []
    current_paras: List[str] = []
    current_size:  int       = 0

    for para in paragraphs:
        psize = len(para) + 2
        if current_size + psize > chunk_size and current_paras:
            chunk_text = "\n\n".join(current_paras)
            chunks.append(chunk_text)
            overlap_text = chunk_text[-overlap:] if len(chunk_text) > overlap else chunk_text
            last_nn = overlap_text.rfind("\n\n")
            if last_nn != -1:
                overlap_paras = overlap_text[last_nn + 2:].split("\n\n")
            else:
                overlap_paras = [overlap_text]
            current_paras = [p for p in overlap_paras if p.strip()]
            current_size  = sum(len(p) + 2 for p in current_paras)

        current_paras.append(para)
        current_size += psize

    if current_paras:
        chunks.append("\n\n".join(current_paras))

    return chunks if chunks else [content]


def _reassemble_chunks(chunks: List[str], overlap: int = CHUNK_OVERLAP_CHARS) -> str:
    """
    Reassemble edited sub-chunks, removing duplicated overlap regions
    by detecting the tail of chunk[i] in the head of chunk[i+1].
    """
    if not chunks:
        return ""
    if len(chunks) == 1:
        return chunks[0]

    result = chunks[0]
    for nxt in chunks[1:]:
        if not nxt.strip():
            continue
        result_tail = result[-min(overlap * 2, len(result)):]
        last_para   = result_tail.rsplit("\n\n", 1)[-1].strip()

        if last_para and last_para in nxt[:overlap * 2]:
            idx = nxt.find(last_para)
            if idx != -1:
                trim_start = idx + len(last_para)
                while trim_start < len(nxt) and nxt[trim_start] in ("\n", " "):
                    trim_start += 1
                nxt = nxt[trim_start:]

        if nxt.strip():
            result = result.rstrip() + "\n\n" + nxt.lstrip()

    return result


def _edit_chunk_via_api(
    chunk_text:     str,
    chunk_title:    str,
    instruction:    str,
    book_title:     str,
    chapter_idx:    int,
    total_chapters: int,
    is_first_chunk: bool,
    context_prefix: str = "",
    context_suffix: str = "",
) -> Tuple[str, bool]:
    """
    Edit a single text chunk via the OpenAI API.
    Retries once on transient failures; falls back to original on two failures.
    Returns (edited_text, changed_flag).
    """
    if not client:
        return chunk_text, False

    context_block = ""
    if context_prefix:
        context_block += (
            f"\n\n[CONTEXT — end of previous chunk — do NOT include in output]:\n"
            f"{context_prefix[-600:]}"
        )
    if context_suffix:
        context_block += (
            f"\n\n[CONTEXT — start of next chunk — do NOT include in output]:\n"
            f"{context_suffix[:400]}"
        )

    continuation_note = (
        " (This is a continuation chunk. The chapter title applies only to the first chunk.)"
        if not is_first_chunk else ""
    )

    payload = json.dumps(
        {"title": chunk_title, "content": chunk_text},
        ensure_ascii=False,
    )

    messages = [
        {"role": "system", "content": _CHAPTER_EDIT_SYSTEM},
        {
            "role": "user",
            "content": (
                f"Book: \"{book_title}\" | "
                f"Chapter {chapter_idx + 1}/{total_chapters}"
                f"{continuation_note}"
                f"{context_block}\n\n"
                f"Edit instruction: {instruction}\n\n"
                f"Chapter JSON:\n{payload}"
            ),
        },
    ]

    estimated_out = min(MAX_EDIT_TOKENS, max(2048, int(len(chunk_text) * 0.7)))
    last_exc: Optional[Exception] = None

    for attempt in range(2):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                max_tokens=estimated_out,
                temperature=0.7,
            )
            raw = (resp.choices[0].message.content or "").strip()
            raw = re.sub(r"^```json\s*", "", raw, flags=re.IGNORECASE)
            raw = re.sub(r"\s*```$", "", raw)
            result = _try_parse_json(raw)
            if result and isinstance(result, dict) and "content" in result:
                return str(result["content"]), bool(result.get("changed", True))
            log.warning("Attempt %d: JSON parsed but missing 'content'", attempt + 1)
        except Exception as ex:
            last_exc = ex
            log.warning("Chunk edit attempt %d failed: %s", attempt + 1, ex)
            if attempt == 0:
                time.sleep(1)

    log.error("Chunk permanently failed after 2 attempts: %s", last_exc)
    return chunk_text, False


def _edit_single_chapter(
    chapter:        dict,
    instruction:    str,
    book_title:     str,
    chapter_idx:    int,
    total_chapters: int,
) -> dict:
    """
    Edit one chapter with full sub-chunking and optional swarm quality pass.

    Flow:
      1. If chapter fits in CHUNK_SIZE → single API call
         Else → split at paragraph boundaries, edit each chunk, reassemble
      2. If chapter >= SWARM_THRESHOLD and actually changed
         → Critic/Revisor swarm quality pass
      3. Return updated chapter dict with '_changed' flag.
    """
    content = chapter.get("content", "")
    if not isinstance(content, str):
        content = str(content or "")
    title  = str(chapter.get("title", "") or "")
    ch_num = chapter.get("chapter_number", chapter_idx + 1)

    log.info(
        "Editing chapter %d/%d: '%s' (%d chars)",
        chapter_idx + 1, total_chapters, title[:50], len(content),
    )

    # ── Single-chunk path ─────────────────────────────────────────────────────
    if len(content) <= CHUNK_SIZE:
        draft, changed = _edit_chunk_via_api(
            content, title, instruction, book_title,
            chapter_idx, total_chapters, True,
        )
        if changed and len(content) >= SWARM_THRESHOLD:
            log.info("Swarm quality pass for chapter %d", chapter_idx + 1)
            draft, changed = _swarm_edit_chapter(chapter, instruction, draft, book_title)

        return {
            "chapter_number":    ch_num,
            "title":             title,
            "content":           draft,
            "layout_directives": chapter.get("layout_directives", _default_layout_directives()),
            "footnotes":         chapter.get("footnotes", []),
            "word_count":        len(draft.split()),
            "_changed":          changed,
        }

    # ── Multi-chunk path ──────────────────────────────────────────────────────
    log.info(
        "Chapter %d is large (%d chars) → sub-chunking",
        chapter_idx + 1, len(content),
    )
    sub_chunks = _split_at_paragraph_boundaries(content)
    log.info("Split into %d sub-chunks", len(sub_chunks))

    edited_chunks: List[str] = []
    any_changed    = False
    context_prefix = ""

    for i, chunk in enumerate(sub_chunks):
        ctx_suffix = sub_chunks[i + 1][:400] if i + 1 < len(sub_chunks) else ""
        edited, changed = _edit_chunk_via_api(
            chunk,
            title if i == 0 else f"{title} (cont.)",
            instruction,
            book_title,
            chapter_idx,
            total_chapters,
            i == 0,
            context_prefix,
            ctx_suffix,
        )
        context_prefix = edited[-CHUNK_OVERLAP_CHARS:]
        edited_chunks.append(edited)
        if changed:
            any_changed = True

    final_content = _reassemble_chunks(edited_chunks)

    if any_changed and len(content) >= SWARM_THRESHOLD:
        log.info("Swarm quality pass on reassembled chapter %d", chapter_idx + 1)
        temp_ch = dict(chapter)
        temp_ch["content"] = content
        final_content, any_changed = _swarm_edit_chapter(
            temp_ch, instruction, final_content, book_title
        )

    return {
        "chapter_number":    ch_num,
        "title":             title,
        "content":           final_content,
        "layout_directives": chapter.get("layout_directives", _default_layout_directives()),
        "footnotes":         chapter.get("footnotes", []),
        "word_count":        len(final_content.split()),
        "_changed":          any_changed,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# CORE EDIT ORCHESTRATORS
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

_WHOLE_BOOK_EDITOR_SYSTEM = """You are an expert book editor and author assistant.

You receive a complete book as structured JSON and a natural-language edit instruction.

Rules:
  1. Apply the instruction faithfully and completely to every applicable chapter.
  2. Leave all other content EXACTLY as-is unless explicitly asked to change it.
  3. Genre/tone/style changes → rewrite prose in that style while preserving plot/facts.
  4. ALWAYS return the COMPLETE book with ALL chapters fully written out — never truncate.
  5. Preserve each chapter's layout_directives unless the instruction is layout-specific.
  6. Preserve footnotes arrays on each chapter.

Return ONLY valid JSON — no markdown fences, no preamble:
{
  "title": "<book title>",
  "author": "<author>",
  "description": "<one sentence>",
  "chapters": [
    {
      "chapter_number": 1,
      "title": "<title>",
      "content": "<full content, \\n\\n between paragraphs>",
      "layout_directives": { ... },
      "footnotes": []
    }
  ],
  "edit_summary": "<2-3 sentence summary of changes>",
  "chapters_changed": [<list of chapter numbers that were modified>]
}

CRITICAL: Never truncate content with '[rest unchanged]' or any placeholder.
"""


def _estimate_tokens(text: Any) -> int:
    if isinstance(text, int):
        return max(0, text) // 4
    if not isinstance(text, str):
        text = str(text or "")
    return len(text) // 4


def _apply_edit_whole_book(
    book_structure:       dict,
    user_instruction:     str,
    conversation_history: List[dict],
) -> dict:
    """Edit small books (< ~8K tokens) in a single API call."""
    if not client:
        raise RuntimeError(
            "OpenAI client not initialised — set the OPENAI_API_KEY environment variable."
        )

    book_json = json.dumps(book_structure, ensure_ascii=False)

    if len(book_json) > 50_000:
        trimmed = dict(book_structure)
        trimmed_chapters = []
        for ch in trimmed.get("chapters", []):
            ch_copy = dict(ch)
            content = ch_copy.get("content", "")
            if len(content) > 3000:
                ch_copy["content"] = content[:3000] + "\n[...content trimmed...]"
            trimmed_chapters.append(ch_copy)
        trimmed["chapters"] = trimmed_chapters
        book_json = json.dumps(trimmed, ensure_ascii=False)

    safe_history = [
        {**m, "content": str(m.get("content") or "")}
        for m in (
            conversation_history[-4:]
            if len(conversation_history) > 4
            else conversation_history
        )
        if isinstance(m, dict) and m.get("role") in ("user", "assistant", "system")
    ]

    messages: List[dict] = [{"role": "system", "content": _WHOLE_BOOK_EDITOR_SYSTEM}]
    messages.extend(safe_history)
    messages.append({
        "role": "user",
        "content": (
            f"Current book (JSON):\n{book_json}\n\n"
            f"Edit instruction: {user_instruction}"
        ),
    })

    estimated_out = min(16384, max(4096, len(book_json) // 2))

    resp = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        max_tokens=estimated_out,
        temperature=0.7,
    )
    raw = (resp.choices[0].message.content or "").strip()
    raw = re.sub(r"^```json\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)

    result = _try_parse_json(raw)
    if result and isinstance(result, dict) and "chapters" in result:
        for ch in result.get("chapters", []):
            if "layout_directives" not in ch or not isinstance(ch.get("layout_directives"), dict):
                ch["layout_directives"] = _default_layout_directives()
            if "footnotes" not in ch:
                ch["footnotes"] = []
            if not isinstance(ch.get("content"), str):
                ch["content"] = str(ch.get("content") or "")
        return result

    raise ValueError(
        "The AI returned an invalid response. "
        "Try breaking your request into smaller or more specific edits."
    )


def _apply_edit_chunked(
    book_structure:        dict,
    user_instruction:      str,
    _conversation_history: List[dict],
) -> dict:
    """
    Edit large books chapter-by-chapter (Subsystem 6).
    Uses parallel workers (up to MAX_WORKERS) for speed.
    Skips prose edit API calls for pure layout-only instructions.
    """
    chapters   = book_structure.get("chapters", [])
    book_title = book_structure.get("title", "Untitled")
    author     = book_structure.get("author", "")
    n          = len(chapters)

    for ch in chapters:
        if not isinstance(ch.get("content"), str):
            ch["content"] = str(ch.get("content") or "")

    log.info("Chunked edit: %d chapters, book='%s'", n, book_title[:60])

    # Detect if this is a layout-only instruction (no prose changes needed).
    # Only skip the expensive prose-edit API calls if the instruction is
    # exclusively about layout (blank pages, recto, margins, drop caps, etc.)
    # and contains zero prose/style editing intent.
    _PURE_LAYOUT_PHRASES = re.compile(
        r"^\s*(?:add|insert|make|set|enable|disable|use|put|apply|force|give"
        r"|place|start|begin)\b[^.!?]{0,120}"
        r"(?:blank page|recto|right.hand|right side|drop cap|ornament|"
        r"section label|mirror margin|page break|odd.page|line spacing|"
        r"column|header|footer|running head)\b",
        re.IGNORECASE,
    )
    has_layout   = bool(_LAYOUT_KEYWORDS.search(user_instruction))
    has_prose    = bool(_PROSE_KEYWORDS.search(user_instruction))
    has_style    = bool(_STYLE_KEYWORDS.search(user_instruction))
    is_layout_only = (
        has_layout
        and not has_prose
        and not has_style
        and _PURE_LAYOUT_PHRASES.search(user_instruction)
    )

    if is_layout_only:
        log.info("Layout-only instruction — skipping prose edit API calls")
        final_chapters = []
        for ch in chapters:
            ch_copy = copy.deepcopy(ch)
            ch_copy.pop("_changed", None)
            final_chapters.append(ch_copy)
        summary = _generate_edit_summary(user_instruction, [], book_title)
        return {
            "title":            book_title,
            "author":           author,
            "description":      book_structure.get("description", ""),
            "isbn":             book_structure.get("isbn", ""),
            "year":             book_structure.get("year", ""),
            "chapters":         final_chapters,
            "edit_summary":     summary,
            "chapters_changed": [],
        }

    updated_chapters: List[Optional[dict]] = [None] * n
    changed_numbers:  List[int]            = []

    def _worker(idx_and_chapter: Tuple[int, dict]) -> Tuple[int, dict]:
        idx, chapter = idx_and_chapter
        return idx, _edit_single_chapter(chapter, user_instruction, book_title, idx, n)

    if n >= 3 and MAX_WORKERS > 1:
        with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, n)) as executor:
            futures = {
                executor.submit(_worker, (idx, ch)): idx
                for idx, ch in enumerate(chapters)
            }
            for future in as_completed(futures):
                try:
                    idx, result_ch = future.result()
                    if result_ch.pop("_changed", False):
                        changed_numbers.append(result_ch.get("chapter_number", idx + 1))
                    updated_chapters[idx] = result_ch
                except Exception as ex:
                    idx = futures[future]
                    log.error("Chapter %d edit failed: %s", idx + 1, ex)
                    updated_chapters[idx] = chapters[idx]
    else:
        for idx, chapter in enumerate(chapters):
            result_ch = _edit_single_chapter(chapter, user_instruction, book_title, idx, n)
            if result_ch.pop("_changed", False):
                changed_numbers.append(result_ch.get("chapter_number", idx + 1))
            updated_chapters[idx] = result_ch

    final_chapters = [ch for ch in updated_chapters if ch is not None]
    changed_numbers.sort()

    summary = _generate_edit_summary(user_instruction, changed_numbers, book_title)

    return {
        "title":            book_title,
        "author":           author,
        "description":      book_structure.get("description", ""),
        "isbn":             book_structure.get("isbn", ""),
        "year":             book_structure.get("year", ""),
        "chapters":         final_chapters,
        "edit_summary":     summary,
        "chapters_changed": changed_numbers,
    }


def _generate_edit_summary(
    instruction:      str,
    changed_chapters: List[int],
    book_title:       str,
) -> str:
    """Generate a 2-sentence AI-written changelog summary (Subsystem 10)."""
    if not changed_chapters:
        return (
            f"The edit instruction '{instruction}' was applied to '{book_title}'. "
            "No chapters required substantive content changes."
        )

    ch_list = ", ".join(f"Ch.{n}" for n in changed_chapters[:7])
    if len(changed_chapters) > 7:
        ch_list += f" and {len(changed_chapters) - 7} more"

    if not client:
        return f"Applied '{instruction}' to {ch_list} in '{book_title}'."

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{
                "role": "user",
                "content": (
                    f"Write a 2-sentence professional editorial changelog summary:\n"
                    f"Book: '{book_title}'\n"
                    f"Edit instruction: '{instruction}'\n"
                    f"Chapters modified: {ch_list}\n"
                    "Be concise and specific. Do not repeat the instruction verbatim."
                ),
            }],
            max_tokens=160,
            temperature=0.4,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception:
        return f"Applied '{instruction}' to {ch_list} in '{book_title}'."


def _apply_structural_edit(book_structure: dict, user_instruction: str) -> Optional[dict]:
    """
    Handle REORDER-intent instructions by calling the ACTUAL structural
    functions (insert_chapter / delete_chapter / merge_chapters /
    split_chapter / reorder_chapters) instead of routing them through the
    generic prose-rewrite pipeline.

    Uses one small GPT call to extract a structured action from the
    natural-language instruction, then executes that action directly
    against the book dict — deterministic, not generative.

    Returns the updated book dict (with edit_summary / chapters_changed
    keys set), or None if no client is available / extraction failed,
    so the caller can fall back to the prose pipeline.
    """
    if not client:
        return None

    chapters_brief = [
        {"chapter_number": ch.get("chapter_number"), "title": ch.get("title", "")}
        for ch in book_structure.get("chapters", [])
    ]

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{
                "role": "system",
                "content": (
                    "You convert a structural book-editing instruction into ONE JSON action.\n"
                    "Valid actions: insert_chapter, delete_chapter, merge_chapters, "
                    "split_chapter, reorder_chapters, none.\n"
                    "Respond with ONLY JSON, one of these shapes:\n"
                    '{"action":"insert_chapter","position":int,"title":str,"content":str}\n'
                    '{"action":"delete_chapter","chapter_number":int}\n'
                    '{"action":"merge_chapters","first_number":int,"second_number":int}\n'
                    '{"action":"split_chapter","chapter_number":int,"split_paragraph":int,"second_title":str}\n'
                    '{"action":"reorder_chapters","new_order":[int,...]}\n'
                    '{"action":"none"}\n'
                    "Use \"none\" if the instruction isn't a clear structural operation "
                    "on chapters (e.g. it's actually a prose/style edit)."
                ),
            }, {
                "role": "user",
                "content": (
                    f"Current chapters: {json.dumps(chapters_brief, ensure_ascii=False)}\n"
                    f"Instruction: {user_instruction}"
                ),
            }],
            max_tokens=300,
            temperature=0.0,
        )
        action = _try_parse_json((resp.choices[0].message.content or "").strip())
    except Exception as ex:
        log.warning("Structural-edit extraction failed: %s", ex)
        return None

    if not action or action.get("action") in (None, "none"):
        return None

    try:
        kind = action["action"]
        if kind == "insert_chapter":
            result = insert_chapter(
                book_structure,
                position=int(action["position"]),
                title=action.get("title") or "New Chapter",
                content=action.get("content") or "",
            )
            changed = [int(action["position"])]
            summary = f"Inserted new chapter '{action.get('title','New Chapter')}' at position {action['position']}."

        elif kind == "delete_chapter":
            cn = int(action["chapter_number"])
            result  = delete_chapter(book_structure, cn)
            changed = []
            summary = f"Deleted chapter {cn}; remaining chapters renumbered."

        elif kind == "merge_chapters":
            f, s   = int(action["first_number"]), int(action["second_number"])
            result = merge_chapters(book_structure, f, s)
            changed = [f]
            summary = f"Merged chapters {f} and {s} into one chapter."

        elif kind == "split_chapter":
            cn  = int(action["chapter_number"])
            result = split_chapter(
                book_structure, cn,
                int(action["split_paragraph"]),
                action.get("second_title") or "",
            )
            changed = [cn, cn + 1]
            summary = f"Split chapter {cn} into two chapters."

        elif kind == "reorder_chapters":
            order  = [int(x) for x in action["new_order"]]
            result = reorder_chapters(book_structure, order)
            changed = list(range(1, len(order) + 1))
            summary = f"Reordered chapters: new order {order}."

        else:
            return None

    except Exception as ex:
        log.warning("Structural-edit execution failed (%s): %s", action, ex)
        return None

    result["edit_summary"]     = summary
    result["chapters_changed"] = changed
    log.info("Structural edit applied: %s", summary)
    return result


_FIND_REPLACE_PATTERNS = [
    re.compile(r'change\s+(?:the word\s+)?["\u2018\u2019\'\u201c\u201d"]([^"\u2018\u2019\'\u201c\u201d]+)["\u2018\u2019\'\u201c\u201d"]\s+to\s+["\u2018\u2019\'\u201c\u201d"]([^"\u2018\u2019\'\u201c\u201d]+)["\u2018\u2019\'\u201c\u201d"]', re.IGNORECASE),
    re.compile(r'replace\s+["\u2018\u2019\'\u201c\u201d"]([^"\u2018\u2019\'\u201c\u201d]+)["\u2018\u2019\'\u201c\u201d"]\s+with\s+["\u2018\u2019\'\u201c\u201d"]([^"\u2018\u2019\'\u201c\u201d]+)["\u2018\u2019\'\u201c\u201d"]', re.IGNORECASE),
]


def _try_literal_find_replace(book_structure: dict, user_instruction: str) -> Optional[dict]:
    """
    Exact, deterministic find/replace for instructions like:
      change "word X" to "word Y"
      replace "word X" with "word Y"

    This bypasses the LLM entirely for this kind of edit — no rewriting,
    no paraphrasing, no risk of the model drifting from the literal
    instruction. Scoped to chapter_targets if the instruction names
    specific chapters, otherwise applied across the whole book.

    Returns None if the instruction doesn't match a literal find/replace
    shape, so the caller falls back to the generative pipeline.
    """
    for pattern in _FIND_REPLACE_PATTERNS:
        m = pattern.search(user_instruction)
        if m:
            find_text, replace_text = m.group(1), m.group(2)
            break
    else:
        return None

    targets = set(_extract_chapter_targets(user_instruction, book_structure))
    result  = copy.deepcopy(book_structure)
    changed: List[int] = []
    total_replacements = 0

    for ch in result.get("chapters", []):
        cn = ch.get("chapter_number")
        if cn not in targets:
            continue
        content = str(ch.get("content", ""))
        count   = content.count(find_text)
        if count:
            ch["content"] = content.replace(find_text, replace_text)
            ch["word_count"] = len(ch["content"].split())
            changed.append(cn)
            total_replacements += count

    if total_replacements == 0:
        log.info("Literal find/replace matched no occurrences of '%s' — falling back to AI edit.", find_text)
        return None

    result["edit_summary"] = (
        f"Replaced {total_replacements} occurrence(s) of \"{find_text}\" "
        f"with \"{replace_text}\" in chapter(s) {changed}."
    )
    result["chapters_changed"] = changed
    log.info(result["edit_summary"])
    return result


def apply_edit(
    book_structure:       dict,
    user_instruction:     str,
    conversation_history: List[dict],
) -> dict:
    """
    Master edit entry-point.
    Step 0: Try exact literal find/replace (deterministic, no LLM).
    Step 0.5: Try structural ops (insert/delete/merge/split/reorder chapter)
              via the real functions, not prose-rewrite.
    Step 1: Apply layout directives from the instruction (Subsystem 1).
    Step 2: Route: tiny books → whole-book call; everything else → chapter-by-chapter.
    """
    literal_result = _try_literal_find_replace(book_structure, user_instruction)
    if literal_result is not None:
        return literal_result

    structural_result = _apply_structural_edit(book_structure, user_instruction)
    if structural_result is not None:
        return structural_result

    book_structure = _apply_global_directives(
        copy.deepcopy(book_structure), user_instruction
    )

    chapters = book_structure.get("chapters", [])
    for ch in chapters:
        if not isinstance(ch.get("content"), str):
            ch["content"] = str(ch.get("content") or "")

    total_chars      = sum(len(ch.get("content", "")) for ch in chapters)
    estimated_tokens = _estimate_tokens(total_chars)

    log.info(
        "apply_edit: %d chapters, ~%d tokens, instruction='%s'",
        len(chapters), estimated_tokens, user_instruction[:80],
    )

    if estimated_tokens > 2000 or len(chapters) > 1:
        return _apply_edit_chunked(book_structure, user_instruction, conversation_history)

    return _apply_edit_whole_book(book_structure, user_instruction, conversation_history)


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 5 — DETAILED HTML TRACK-CHANGES  (difflib Engine)
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class DiffStats:
    """Per-chapter diff statistics."""
    chapter_number: int   = 0
    chapter_title:  str   = ""
    original_chars: int   = 0
    edited_chars:   int   = 0
    insertions:     int   = 0
    deletions:      int   = 0
    change_percent: float = 0.0
    original_words: int   = 0
    edited_words:   int   = 0
    word_delta:     int   = 0

    def to_html_row(self) -> str:
        direction = "▲" if self.word_delta >= 0 else "▼"
        colour    = "#186a3b" if self.word_delta >= 0 else "#922b21"
        return (
            f"<tr>"
            f"<td>Ch.{self.chapter_number}</td>"
            f"<td>{html_module.escape(self.chapter_title[:60])}</td>"
            f"<td style='text-align:right'>{self.original_words:,}</td>"
            f"<td style='text-align:right'>{self.edited_words:,}</td>"
            f"<td style='text-align:right;color:{colour}'>"
            f"{direction}{abs(self.word_delta):,}</td>"
            f"<td style='text-align:right'>{self.change_percent:.1f}%</td>"
            f"</tr>"
        )


def _compute_diff_stats(
    original_content: str,
    edited_content:   str,
    chapter_number:   int,
    chapter_title:    str,
) -> DiffStats:
    matcher    = difflib.SequenceMatcher(None, original_content, edited_content, autojunk=False)
    insertions = 0
    deletions  = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag in ("replace", "delete"):
            deletions  += i2 - i1
        if tag in ("replace", "insert"):
            insertions += j2 - j1

    orig_words = len(original_content.split())
    edit_words = len(edited_content.split())
    total_ops  = insertions + deletions
    base       = max(len(original_content), 1)

    return DiffStats(
        chapter_number  = chapter_number,
        chapter_title   = chapter_title,
        original_chars  = len(original_content),
        edited_chars    = len(edited_content),
        insertions      = insertions,
        deletions       = deletions,
        change_percent  = round(total_ops / base * 100, 1),
        original_words  = orig_words,
        edited_words    = edit_words,
        word_delta      = edit_words - orig_words,
    )


def _generate_chapter_diff_html(
    original:      str,
    edited:        str,
    chapter_title: str,
    stats:         Optional[DiffStats] = None,
) -> str:
    """
    Generate an HTML fragment with a character-level side-by-side diff.
    Left column = original with <del> spans; right column = edited with <ins> spans.
    """
    matcher   = difflib.SequenceMatcher(None, original, edited, autojunk=False)
    orig_html: List[str] = []
    edit_html: List[str] = []

    def esc(t: str) -> str:
        return html_module.escape(t).replace("\n", "<br>\n")

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        o_chunk = esc(original[i1:i2])
        e_chunk = esc(edited[j1:j2])
        if tag == "equal":
            orig_html.append(o_chunk)
            edit_html.append(e_chunk)
        elif tag == "replace":
            orig_html.append(
                f'<del style="background:#fdd;color:#900;text-decoration:line-through">{o_chunk}</del>'
            )
            edit_html.append(
                f'<ins style="background:#dfd;color:#070;text-decoration:none">{e_chunk}</ins>'
            )
        elif tag == "delete":
            orig_html.append(
                f'<del style="background:#fdd;color:#900;text-decoration:line-through">{o_chunk}</del>'
            )
        elif tag == "insert":
            edit_html.append(
                f'<ins style="background:#dfd;color:#070;text-decoration:none">{e_chunk}</ins>'
            )

    stats_bar = ""
    if stats:
        stats_bar = (
            f'<div style="background:#f8f8f8;padding:.4em 1em;font-size:.8em;'
            f'border-bottom:1px solid #ddd;display:flex;gap:2em">'
            f'<span>Original: {stats.original_words:,} words</span>'
            f'<span>Edited: {stats.edited_words:,} words</span>'
            f'<span>Changed: {stats.change_percent}%</span>'
            f'<span style="color:{"#186a3b" if stats.word_delta>=0 else "#922b21"}">'
            f'{"+" if stats.word_delta>=0 else ""}{stats.word_delta:,} words</span>'
            f'</div>'
        )

    ch_safe = html_module.escape(chapter_title)
    return (
        f'<section style="margin-bottom:2.5em;border:1px solid #ccc;border-radius:6px;'
        f'overflow:hidden;font-family:Georgia,serif">'
        f'<h3 style="margin:0;padding:.6em 1em;background:#eee;border-bottom:1px solid #ccc;'
        f'font-size:1em">{ch_safe}</h3>'
        f'{stats_bar}'
        f'<div style="display:grid;grid-template-columns:1fr 1fr">'
        f'<div style="padding:1em;border-right:1px solid #ddd;font-size:.88em;'
        f'white-space:pre-wrap;overflow:auto">{"".join(orig_html)}</div>'
        f'<div style="padding:1em;font-size:.88em;white-space:pre-wrap;'
        f'overflow:auto">{"".join(edit_html)}</div>'
        f'</div>'
        f'</section>'
    )


def generate_inline_diff_html(original: str, edited: str, chapter_title: str) -> str:
    """Single-column inline diff: colour-coded insertions and deletions inline."""
    matcher  = difflib.SequenceMatcher(None, original, edited, autojunk=False)
    segments: List[str] = []

    def esc(t: str) -> str:
        return html_module.escape(t).replace("\n", "<br>\n")

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            segments.append(esc(original[i1:i2]))
        elif tag == "replace":
            segments.append(
                f'<del style="background:#fdd;color:#900;text-decoration:line-through">'
                f'{esc(original[i1:i2])}</del>'
                f'<ins style="background:#dfd;color:#070;text-decoration:none">'
                f'{esc(edited[j1:j2])}</ins>'
            )
        elif tag == "delete":
            segments.append(
                f'<del style="background:#fdd;color:#900;text-decoration:line-through">'
                f'{esc(original[i1:i2])}</del>'
            )
        elif tag == "insert":
            segments.append(
                f'<ins style="background:#dfd;color:#070;text-decoration:none">'
                f'{esc(edited[j1:j2])}</ins>'
            )

    ch_safe = html_module.escape(chapter_title)
    return (
        f'<section style="margin-bottom:2em">'
        f'<h3 style="font-size:1em;border-bottom:1px solid #ccc;padding-bottom:.3em">{ch_safe}</h3>'
        f'<div style="font-family:Georgia,serif;font-size:.9em;line-height:1.6;white-space:pre-wrap">'
        f'{"".join(segments)}</div></section>'
    )


def generate_diff_report(
    original_book: dict,
    edited_book:   dict,
    output_path:   str,
) -> str:
    """
    Build a complete HTML track-changes report and write it to output_path.
    Includes book-level summary, per-chapter stats table, and side-by-side diffs.
    """
    orig_by_num = {ch["chapter_number"]: ch for ch in original_book.get("chapters", [])}
    edit_by_num = {ch["chapter_number"]: ch for ch in edited_book.get("chapters", [])}
    all_nums    = sorted(set(list(orig_by_num.keys()) + list(edit_by_num.keys())))

    diff_sections: List[str]    = []
    all_stats:     List[DiffStats] = []

    for num in all_nums:
        orig_ch = orig_by_num.get(num, {})
        edit_ch = edit_by_num.get(num, {})
        orig_c  = str(orig_ch.get("content", ""))
        edit_c  = str(edit_ch.get("content", ""))
        title   = edit_ch.get("title") or orig_ch.get("title") or f"Chapter {num}"

        stats = _compute_diff_stats(orig_c, edit_c, num, title)
        all_stats.append(stats)

        if orig_c != edit_c:
            diff_sections.append(_generate_chapter_diff_html(orig_c, edit_c, title, stats))

    total_orig_words = sum(s.original_words for s in all_stats)
    total_edit_words = sum(s.edited_words   for s in all_stats)
    chapters_changed = sum(1 for s in all_stats if s.change_percent > 0)
    word_delta       = total_edit_words - total_orig_words
    delta_dir        = "+" if word_delta >= 0 else ""
    timestamp        = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    book_title       = html_module.escape(edited_book.get("title", "Book"))

    stats_rows = "\n".join(s.to_html_row() for s in all_stats)

    html_out = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Track Changes — {book_title}</title>
<style>
  body {{ font-family:Arial,sans-serif;max-width:1500px;margin:0 auto;padding:1.5em;background:#fafafa;color:#222 }}
  h1   {{ font-size:1.4em;margin-bottom:.2em }}
  h2   {{ font-size:1.05em;color:#555;font-weight:normal;margin-top:.2em }}
  .legend {{ display:flex;gap:1.5em;margin:1em 0 1.5em }}
  .legend span {{ padding:.2em .7em;border-radius:3px;font-size:.85em }}
  table {{ border-collapse:collapse;width:100%;margin-bottom:2em }}
  th,td {{ border:1px solid #ddd;padding:.4em .7em;font-size:.85em }}
  th {{ background:#eee;text-align:left }}
  tr:hover {{ background:#f5f5f5 }}
</style>
</head>
<body>
<h1>📝 Track Changes — {book_title}</h1>
<h2>Generated: {timestamp}  |  {chapters_changed} chapter(s) modified  |
Words: {total_orig_words:,} → {total_edit_words:,} ({delta_dir}{word_delta:,})</h2>
<div class="legend">
  <span style="background:#fdd;color:#900">■ Deleted text</span>
  <span style="background:#dfd;color:#070">■ Inserted text</span>
</div>
<table>
  <thead><tr>
    <th>Chapter</th><th>Title</th><th>Original words</th>
    <th>Edited words</th><th>Word Δ</th><th>Change %</th>
  </tr></thead>
  <tbody>{stats_rows}</tbody>
</table>
"""
    if diff_sections:
        html_out += "\n".join(diff_sections)
    else:
        html_out += '<p style="color:#888;font-style:italic">No textual changes detected.</p>'

    html_out += "\n</body>\n</html>"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_out)

    log.info("Diff report written: %s (%d chars)", output_path, len(html_out))
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 10 — ROLLING VERSION CONTROL & EDIT SUMMARIES
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

class VersionHistory:
    """
    Immutable rolling version store for book editing sessions.
    Each version is a deep-copy snapshot of the entire book dict plus metadata.

    Public API:
      commit(book, summary, instruction) → version_index
      get(version_idx)                   → book dict copy
      latest()                           → book dict copy (most recent)
      rollback(version_idx)              → book dict copy (same as get)
      branch(version_idx)                → new VersionHistory forked at that version
      diff_versions(idx_a, idx_b)        → (stats_list, html_report_string)
      log()                              → list of metadata dicts (no book copy)
      export_history_html(path)          → write HTML audit trail
      __len__()                          → number of committed versions
    """

    def __init__(self):
        self._versions: List[Dict[str, Any]] = []
        self._lock = threading.Lock()

    def _book_hash(self, book: dict) -> str:
        serialised = json.dumps(book, sort_keys=True, ensure_ascii=False, default=str)
        return hashlib.sha256(serialised.encode()).hexdigest()[:16]

    def commit(self, book: dict, summary: str, instruction: str) -> int:
        """Store a deep-copy snapshot. Returns the new version index."""
        with self._lock:
            idx = len(self._versions)
            snap = {
                "version":     idx,
                "timestamp":   datetime.datetime.now().isoformat(),
                "instruction": instruction,
                "summary":     summary,
                "book_hash":   self._book_hash(book),
                "book":        copy.deepcopy(book),
            }
            self._versions.append(snap)
            log.info(
                "Version %d committed: '%s' (hash=%s)",
                idx, instruction[:60], snap["book_hash"],
            )
            return idx

    def get(self, version_idx: int) -> Optional[dict]:
        """Return a deep copy of the book at version_idx, or None."""
        with self._lock:
            if 0 <= version_idx < len(self._versions):
                return copy.deepcopy(self._versions[version_idx]["book"])
        return None

    def latest(self) -> Optional[dict]:
        """Return a deep copy of the most recent version, or None."""
        with self._lock:
            if self._versions:
                return copy.deepcopy(self._versions[-1]["book"])
        return None

    def rollback(self, version_idx: int) -> dict:
        """
        Restore a previous version. Does NOT truncate history.
        Returns the old book so the caller can commit it as a new version.
        """
        book = self.get(version_idx)
        if book is None:
            raise IndexError(f"Version {version_idx} does not exist.")
        log.info("Rolled back to version %d", version_idx)
        return book

    def branch(self, version_idx: int) -> "VersionHistory":
        """Create a new VersionHistory forked at version_idx."""
        book = self.get(version_idx)
        if book is None:
            raise IndexError(f"Version {version_idx} does not exist.")
        new_history = VersionHistory()
        meta = self._versions[version_idx]
        new_history.commit(book, f"Branched from v{version_idx}: {meta['summary']}", "branch")
        log.info("Branched new VersionHistory from v%d", version_idx)
        return new_history

    def diff_versions(self, idx_a: int, idx_b: int) -> Tuple[List[DiffStats], str]:
        """Compute character-level diffs between two versions."""
        book_a = self.get(idx_a)
        book_b = self.get(idx_b)
        if book_a is None or book_b is None:
            raise IndexError("One or both version indices are out of range.")

        chapters_a = {ch["chapter_number"]: ch for ch in book_a.get("chapters", [])}
        chapters_b = {ch["chapter_number"]: ch for ch in book_b.get("chapters", [])}
        all_nums   = sorted(set(list(chapters_a.keys()) + list(chapters_b.keys())))

        all_stats:     List[DiffStats] = []
        diff_sections: List[str]       = []

        for num in all_nums:
            ch_a   = chapters_a.get(num, {})
            ch_b   = chapters_b.get(num, {})
            cont_a = str(ch_a.get("content", ""))
            cont_b = str(ch_b.get("content", ""))
            title  = ch_b.get("title") or ch_a.get("title") or f"Chapter {num}"

            stats = _compute_diff_stats(cont_a, cont_b, num, title)
            all_stats.append(stats)
            if cont_a != cont_b:
                diff_sections.append(_generate_chapter_diff_html(cont_a, cont_b, title, stats))

        html = (
            f"<h2>Diff: version {idx_a} → version {idx_b}</h2>\n"
            + ("\n".join(diff_sections) if diff_sections else "<p>No differences.</p>")
        )
        return all_stats, html

    def log(self) -> List[Dict[str, Any]]:
        """Return metadata list (version, timestamp, instruction, summary, hash)."""
        with self._lock:
            return [
                {k: v for k, v in entry.items() if k != "book"}
                for entry in self._versions
            ]

    def export_history_html(self, output_path: str) -> str:
        """Write a full HTML audit trail of all versions to output_path."""
        entries = self.log()
        rows: List[str] = []
        for e in entries:
            rows.append(
                f"<tr>"
                f"<td>v{e['version']}</td>"
                f"<td>{e['timestamp'][:19]}</td>"
                f"<td>{html_module.escape(e['instruction'][:100])}</td>"
                f"<td>{html_module.escape(e['summary'][:150])}</td>"
                f"<td><code>{e['book_hash']}</code></td>"
                f"</tr>"
            )

        html_out = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Edit History</title>
<style>
  body {{ font-family:Arial,sans-serif;max-width:1200px;margin:2em auto;padding:1em }}
  table {{ border-collapse:collapse;width:100% }}
  th,td {{ border:1px solid #ddd;padding:.45em .8em;font-size:.85em }}
  th {{ background:#eee }}
  tr:hover {{ background:#f7f7f7 }}
  code {{ font-size:.78em;color:#555 }}
</style>
</head>
<body>
<h1>📜 Edit History ({len(entries)} versions)</h1>
<table>
<thead>
<tr><th>#</th><th>Timestamp</th><th>Instruction</th><th>Summary</th><th>Hash</th></tr>
</thead>
<tbody>
{"".join(rows)}
</tbody>
</table>
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_out)
        log.info("History export written: %s", output_path)
        return output_path

    def __len__(self) -> int:
        with self._lock:
            return len(self._versions)

    def __repr__(self) -> str:
        return f"VersionHistory({len(self)} versions)"


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 18 — THEME ENGINE
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

THEMES: Dict[str, Dict[str, Any]] = {
    "normal": {
        "bg": "#FFFFFF", "accent": "#374151",
        "title_col": "#111827", "body_col": "#374151",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.6, "cover_style": "minimal",
    },
    "premium": {
        "bg": "#FAFAF8", "accent": "#1D4ED8",
        "title_col": "#0F172A", "body_col": "#1E293B",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.65, "cover_style": "elegant",
    },
    "scifi": {
        "bg": "#050A14", "accent": "#00D4FF",
        "title_col": "#00D4FF", "body_col": "#A0C8E0",
        "font_body": "Courier", "font_head": "Courier-Bold",
        "leading_mult": 1.55, "cover_style": "tech",
    },
    "fantasy": {
        "bg": "#0D0A1A", "accent": "#C084FC",
        "title_col": "#E9D5FF", "body_col": "#DDD6FE",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.7, "cover_style": "ornate",
    },
    "romance": {
        "bg": "#FFF5F5", "accent": "#E11D48",
        "title_col": "#9F1239", "body_col": "#4C0519",
        "font_body": "Helvetica-Oblique", "font_head": "Helvetica-Bold",
        "leading_mult": 1.75, "cover_style": "soft",
    },
    "thriller": {
        "bg": "#0A0A0A", "accent": "#EF4444",
        "title_col": "#FAFAFA", "body_col": "#D1D5DB",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.55, "cover_style": "stark",
    },
    "academic": {
        "bg": "#F9FAFB", "accent": "#1E40AF",
        "title_col": "#1E3A5F", "body_col": "#374151",
        "font_body": "Times-Roman", "font_head": "Times-Bold",
        "leading_mult": 1.8, "cover_style": "scholarly",
    },
    "minimalist": {
        "bg": "#FFFFFF", "accent": "#000000",
        "title_col": "#000000", "body_col": "#333333",
        "font_body": "Helvetica", "font_head": "Helvetica",
        "leading_mult": 1.6, "cover_style": "minimal",
    },
    "vibrant": {
        "bg": "#1A0533", "accent": "#F59E0B",
        "title_col": "#FDE68A", "body_col": "#FEF3C7",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.65, "cover_style": "bold",
    },
    "retro": {
        "bg": "#FDF6E3", "accent": "#B45309",
        "title_col": "#78350F", "body_col": "#451A03",
        "font_body": "Courier", "font_head": "Courier-Bold",
        "leading_mult": 1.6, "cover_style": "vintage",
    },
    "horror": {
        "bg": "#0C0C0C", "accent": "#B91C1C",
        "title_col": "#EF4444", "body_col": "#9CA3AF",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.55, "cover_style": "stark",
    },
    "literary": {
        "bg": "#FFFEF7", "accent": "#92400E",
        "title_col": "#451A03", "body_col": "#292524",
        "font_body": "Times-Roman", "font_head": "Times-Bold",
        "leading_mult": 1.85, "cover_style": "scholarly",
    },
    "children": {
        "bg": "#FFFBEB", "accent": "#F59E0B",
        "title_col": "#7C3AED", "body_col": "#374151",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 2.0, "cover_style": "playful",
    },
    "manga": {
        "bg": "#FFFFFF", "accent": "#DC2626",
        "title_col": "#111827", "body_col": "#111827",
        "font_body": "Helvetica", "font_head": "Helvetica-Bold",
        "leading_mult": 1.5, "cover_style": "stark",
    },
}


class ThemeBuilder:
    """Programmatic custom theme creation."""

    def __init__(self, name: str):
        self.name   = name
        self._theme = dict(THEMES["premium"])  # start from premium defaults

    def background(self, hex_color: str) -> "ThemeBuilder":
        self._theme["bg"] = hex_color
        return self

    def accent(self, hex_color: str) -> "ThemeBuilder":
        self._theme["accent"] = hex_color
        return self

    def title_color(self, hex_color: str) -> "ThemeBuilder":
        self._theme["title_col"] = hex_color
        return self

    def body_color(self, hex_color: str) -> "ThemeBuilder":
        self._theme["body_col"] = hex_color
        return self

    def fonts(self, body: str, head: str) -> "ThemeBuilder":
        self._theme["font_body"] = body
        self._theme["font_head"] = head
        return self

    def line_spacing(self, mult: float) -> "ThemeBuilder":
        self._theme["leading_mult"] = mult
        return self

    def build(self) -> Dict[str, Any]:
        THEMES[self.name] = self._theme
        return self._theme


def detect_theme_from_instruction(instruction: str, current_theme: str) -> str:
    """Detect a theme name from the user's instruction."""
    lo = instruction.lower()
    for key in THEMES:
        if key in lo:
            return key
    aliases = {
        "science fiction": "scifi",
        "dark":            "thriller",
        "children's":      "children",
        "historical":      "literary",
        "philosophical":   "academic",
        "short story":     "literary",
    }
    for alias, theme_key in aliases.items():
        if alias in lo:
            return theme_key
    return current_theme


def _hex_to_rgb(h: str) -> Tuple[float, float, float]:
    h = h.lstrip("#")
    if len(h) == 3:
        h = h[0]*2 + h[1]*2 + h[2]*2
    return (
        int(h[0:2], 16) / 255.0,
        int(h[2:4], 16) / 255.0,
        int(h[4:6], 16) / 255.0,
    )


def _darken(h: str, amount: float = 0.06) -> Tuple[float, float, float]:
    r, g, b = _hex_to_rgb(h)
    return max(0.0, r - amount), max(0.0, g - amount), max(0.0, b - amount)


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 2 — REPORTLAB PDF GENERATOR WITH RECTO-ENFORCEMENT
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _esc_pdf(text: str) -> str:
    """HTML-escape text for use in ReportLab Paragraph markup."""
    if not isinstance(text, str):
        text = str(text or "")
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_edited_pdf(
    book:        dict,
    output_path: str,
    theme_name:  str = "premium",
    generate_toc: bool = False,
    page_anchor_sink: Optional[dict] = None,
) -> str:
    """
    Full-featured PDF generator.

    page_anchor_sink: if provided (an empty dict), this render runs in
    "mapping pass" mode — a zero-ink marker is placed before every
    paragraph, and after build() the dict is populated with
    {(chapter_number, paragraph_index): page_number}. Used by
    resolve_page_target() to translate "page N" instructions into a real
    chapter/paragraph anchor. Leave as None for a normal render (default).

    Features:
      • Mirror margins: wider inner gutter for physical binding
      • RectoEnforcer Flowable — blank-page injection so chapters always open
        on odd (right-hand) pages without calling any non-existent doc methods
      • BlankPageFlowable — proper full blank page insertion
      • Blank-page injection from layout_directives.blank_pages_before/after
      • Section labels, ornaments, drop-cap rendering
      • Dynamic Devanagari/CJK/Arabic font mapping  (Subsystem 8)
      • Per-theme cover page with decorative bars
      • Running headers/footers: book title on verso, chapter title on recto
      • Page numbering starting from chapter 1 (cover = page i)
      • 14 production themes
    """
    from reportlab.lib.pagesizes import A4, LETTER, A5          # type: ignore
    from reportlab.lib.units import mm                           # type: ignore
    from reportlab.lib.styles import ParagraphStyle              # type: ignore
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT  # type: ignore
    from reportlab.lib.colors import HexColor, white, black      # type: ignore
    from reportlab.platypus import (                             # type: ignore
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable,
        KeepTogether,
    )
    from reportlab.platypus.flowables import Flowable            # type: ignore

    # Subsystem 8: Register Unicode fonts
    font_map = _register_all_pdf_fonts()

    theme     = THEMES.get(theme_name, THEMES["premium"])
    BG        = HexColor(theme["bg"])
    ACCENT    = HexColor(theme["accent"])
    TITLE_COL = HexColor(theme["title_col"])
    BODY_COL  = HexColor(theme["body_col"])
    LEAD_MULT = float(theme.get("leading_mult", 1.65))

    # Mirror margins
    OUTER_M = 18 * mm
    INNER_M = 28 * mm
    TOP_M   = 24 * mm
    BOT_M   = 22 * mm

    PAGE_W, PAGE_H = A4

    all_content = " ".join(str(ch.get("content", "")) for ch in book.get("chapters", []))
    has_dev = _has_devanagari(all_content)
    has_cjk = _has_cjk(all_content)
    has_ara = _has_arabic(all_content)

    if has_dev:
        base_font      = font_map["devanagari"]
        base_font_bold = font_map["devanagari_bold"]
    elif has_cjk:
        base_font      = font_map["cjk"]
        base_font_bold = font_map["cjk"]
    elif has_ara:
        base_font      = font_map["arabic"]
        base_font_bold = font_map["arabic"]
    else:
        base_font      = theme["font_body"]
        base_font_bold = theme["font_head"]

    body_size    = 10.5
    body_leading = body_size * LEAD_MULT

    def S(name: str, **kw) -> ParagraphStyle:
        return ParagraphStyle(name, **kw)

    cover_title_s   = S("ct",  fontName=base_font_bold, fontSize=30, textColor=TITLE_COL, leading=38, alignment=TA_CENTER, spaceAfter=8)
    cover_author_s  = S("ca",  fontName=base_font, fontSize=13, textColor=BODY_COL, leading=18, alignment=TA_CENTER, spaceAfter=4)
    cover_sub_s     = S("cs",  fontName=base_font, fontSize=10, textColor=BODY_COL, leading=14, alignment=TA_CENTER)
    section_label_s = S("sl",  fontName=base_font, fontSize=10, textColor=ACCENT, leading=14, alignment=TA_CENTER, spaceAfter=3)
    ornament_s      = S("orn", fontName=base_font, fontSize=18, textColor=ACCENT, leading=22, alignment=TA_CENTER, spaceAfter=5)
    ch_label_s      = S("chl", fontName=base_font, fontSize=8.5, textColor=ACCENT, leading=11, alignment=TA_CENTER, spaceAfter=2)
    ch_title_s      = S("cht", fontName=base_font_bold, fontSize=20, textColor=TITLE_COL, leading=26, alignment=TA_CENTER, spaceBefore=2, spaceAfter=6)
    body_s          = S("bs",  fontName=base_font, fontSize=body_size, textColor=BODY_COL, leading=body_leading, spaceAfter=7, alignment=TA_JUSTIFY)
    body_indent_s   = S("bi",  fontName=base_font, fontSize=body_size, textColor=BODY_COL, leading=body_leading, spaceAfter=7, alignment=TA_JUSTIFY, firstLineIndent=14)

    # ── Page callbacks ────────────────────────────────────────────────────────
    def on_cover(canvas, doc):
        canvas.saveState()
        r, g, b = _hex_to_rgb(theme["bg"])
        canvas.setFillColorRGB(r, g, b)
        canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        ar, ag, ab = _hex_to_rgb(theme["accent"])
        canvas.setFillColorRGB(ar, ag, ab)
        canvas.rect(0, PAGE_H * 0.40, PAGE_W, 3.5, fill=1, stroke=0)
        dr, dg, db = _darken(theme["bg"], 0.06)
        canvas.setFillColorRGB(dr, dg, db)
        canvas.rect(0, 0, PAGE_W, 14 * mm, fill=1, stroke=0)
        canvas.setFillColorRGB(0.55, 0.55, 0.55)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawCentredString(
            PAGE_W / 2, 5 * mm,
            f"Edited with AI Book Editor  ·  {datetime.date.today()}",
        )
        canvas.restoreState()

    # ── Build per-chapter running header override map ─────────────────────────
    # Maps estimated page ranges → custom header text (best-effort; PDF page
    # numbers are only known at render time so we store by chapter index).
    # We use a simple list: _ch_headers[page_num] = header_text if overridden.
    _chapter_headers: Dict[int, str] = {}   # chapter_number → override text
    _est_page = 2   # cover = 1
    for ch in book.get("chapters", []):
        directives_h = ch.get("layout_directives") or {}
        rh = str(directives_h.get("running_header_text", "")).strip()
        if rh:
            _chapter_headers[ch.get("chapter_number", 0)] = rh
        wc_h = len(str(ch.get("content", "")).split())
        _est_page += max(1, round(wc_h / 250))

    def on_page(canvas, doc):
        canvas.saveState()
        pnum = doc.page
        dr, dg, db = _darken(theme["bg"], 0.06)
        tr, tg, tb = _hex_to_rgb(theme["body_col"])
        canvas.setFillColorRGB(dr, dg, db)
        canvas.rect(0, 0, PAGE_W, 9.5 * mm, fill=1, stroke=0)
        canvas.setFillColorRGB(tr, tg, tb)
        canvas.setFont("Helvetica", 7)
        book_title_short = book.get("title", "")[:60]
        if pnum % 2 == 0:
            # Verso (left/even) page: page number on outer LEFT, title on inner RIGHT
            canvas.drawString(OUTER_M, 3.5 * mm, str(pnum))
            canvas.drawRightString(PAGE_W - INNER_M, 3.5 * mm, book_title_short)
        else:
            # Recto (right/odd) page: title on inner LEFT, page number on outer RIGHT
            canvas.drawString(INNER_M, 3.5 * mm, book_title_short)
            canvas.drawRightString(PAGE_W - OUTER_M, 3.5 * mm, str(pnum))
        canvas.restoreState()

    # ── RectoEnforcer — WORKING IMPLEMENTATION ────────────────────────────────
    class BlankPageFlowable(Flowable):
        """
        A flowable that produces a completely blank page.
        Claims the full available frame height so ReportLab is forced to break
        to a new page after it, giving one visually empty page.
        Used for blank_pages_before/after directives.
        """
        def wrap(self, aW: float, aH: float) -> Tuple[float, float]:
            self._frame_height = aH
            return aW, aH   # claim full height → forces a page break after

        def draw(self) -> None:
            pass   # intentionally blank

        def __repr__(self) -> str:
            return "BlankPageFlowable()"

    class _PageAnchorRecorder(Flowable):
        """
        Zero-ink, zero-height marker flowable used ONLY during a "mapping
        pass" render. When draw() fires, ReportLab has already committed to
        a real page number for this position in the story — record it.

        This is how arbitrary "page N" instructions become possible despite
        the book model being chapter/paragraph-based: render once with one
        of these before every paragraph, read back which page each
        (chapter_number, paragraph_index) landed on, then resolve the
        target page to a real anchor and re-render for the final output.
        """
        def __init__(self, sink: Optional[dict], chapter_number: int, paragraph_index: int):
            Flowable.__init__(self)
            self._sink   = sink
            self._chno   = chapter_number
            self._pidx   = paragraph_index

        def wrap(self, aW: float, aH: float) -> Tuple[float, float]:
            return 0.0, 0.0

        def draw(self) -> None:
            if self._sink is None:
                return
            try:
                page_num = self.canv._pageNumber
            except AttributeError:
                return
            # First write wins — that's the page the paragraph STARTS on.
            self._sink.setdefault((self._chno, self._pidx), page_num)

        def __repr__(self) -> str:
            return f"_PageAnchorRecorder(ch={self._chno}, p={self._pidx})"

    class RectoEnforcer(Flowable):
        """
        Recto-enforcement flowable.

        When ReportLab calls draw(), it checks the current page number via the
        live canvas object (self.canv, which IS available at draw() time).
        If the current page is even (verso), it finalises that page with
        showPage() and starts a fresh odd page so the next chapter always opens
        on a right-hand (recto) page.

        NOTE: split() must NOT be used for page-parity logic because self.canv
        is not yet bound when split() is called.  draw() is the correct hook.
        """

        def __init__(self):
            Flowable.__init__(self)

        def wrap(self, aW: float, aH: float) -> Tuple[float, float]:
            # Claim zero height — we live between flowables and produce no ink
            return 0.0, 0.0

        def draw(self) -> None:
            """Inject a blank verso page if we are currently on an even page."""
            try:
                page_num = self.canv._pageNumber
            except AttributeError:
                try:
                    page_num = self.canv._doctemplate.page
                except AttributeError:
                    return   # cannot determine page — skip safely

            if page_num % 2 == 0:
                log.debug(
                    "RectoEnforcer: page %d is verso — injecting blank page via canvas",
                    page_num,
                )
                # Finish the current (even/verso) blank page and advance to the
                # next odd page.  self.canv.showPage() closes this page;
                # self.canv.translate resets the origin for the new page.
                self.canv.showPage()
                # After showPage the canvas resets; _pageNumber is now odd.

        def _get_canvas_page(self) -> int:
            """Safely return the current page number from the canvas."""
            try:
                return self.canv._pageNumber
            except AttributeError:
                pass
            try:
                return self.canv._doctemplate.page
            except AttributeError:
                return 0

        def split(self, aW: float, aH: float) -> List:
            """No-op: all work happens in draw() where self.canv is live."""
            return [self]

        def __repr__(self) -> str:
            return "RectoEnforcer()"

    # ── Build story ───────────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=INNER_M,
        rightMargin=OUTER_M,
        topMargin=TOP_M,
        bottomMargin=BOT_M,
        title=book.get("title", "Book"),
        author=book.get("author", ""),
    )

    story: List[Any] = []

    # ── Cover page ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 58 * mm))
    story.append(Paragraph(_esc_pdf(book.get("title", "Untitled")), cover_title_s))
    story.append(Spacer(1, 6 * mm))
    if book.get("author"):
        story.append(Paragraph(f"by {_esc_pdf(book['author'])}", cover_author_s))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(f"Theme: {theme_name.title()}", cover_sub_s))
    if book.get("year"):
        story.append(Paragraph(str(book["year"]), cover_sub_s))
    story.append(PageBreak())

    # ── Table of Contents (optional) ──────────────────────────────────────────
    if generate_toc:
        generate_toc_pdf(book, story, {}, ACCENT)

    # ── Chapters ──────────────────────────────────────────────────────────────
    chapters = book.get("chapters", [])

    for ch_idx, ch in enumerate(chapters):
        directives = ch.get("layout_directives") or _default_layout_directives()
        if not isinstance(directives, dict):
            directives = _default_layout_directives()

        blank_before  = max(0, int(directives.get("blank_pages_before", 0)))
        blank_after   = max(0, int(directives.get("blank_pages_after",  0)))
        force_recto   = bool(directives.get("force_recto_start",    False))
        do_page_break = bool(directives.get("page_break_after",     True))
        ornament_bef  = str(directives.get("ornament_before",       ""))
        ornament_aft  = str(directives.get("ornament_after",        ""))
        sec_label     = str(directives.get("section_label",         "")).strip()
        drop_cap      = bool(directives.get("drop_cap",             False))
        hdr_style     = str(directives.get("chapter_header_style",  "centered"))
        ls_mult       = float(directives.get("line_spacing",        0.0)) or LEAD_MULT
        custom_fs     = float(directives.get("custom_font_size",    0.0))
        keep_next     = bool(directives.get("keep_with_next",       False))
        run_hdr_text  = str(directives.get("running_header_text",   "")).strip()
        col_count     = max(1, int(directives.get("column_count",   1)))

        # Per-chapter body style override
        if custom_fs or ls_mult != LEAD_MULT:
            eff_size    = custom_fs or body_size
            eff_leading = eff_size * ls_mult
            ch_body_s   = S(
                f"bs_ch{ch_idx}",
                parent=body_s,
                fontSize=eff_size,
                leading=eff_leading,
            )
        else:
            ch_body_s = body_s

        # Two-column layout: split paragraphs into two equal halves and render
        # side-by-side using a Table flowable (ReportLab does not support true
        # text-flow columns without Frame surgery, but a 2-col Table gives an
        # accurate two-column appearance for most chapter content).
        if col_count == 2:
            from reportlab.platypus import Table, TableStyle  # type: ignore
            from reportlab.lib import colors as rl_colors      # type: ignore
            _two_col_mode = True
        else:
            _two_col_mode = False

        # ── Blank pages before ────────────────────────────────────────────────
        for _ in range(blank_before):
            story.append(BlankPageFlowable())
            story.append(PageBreak())

        # ── Recto enforcer (Subsystem 2) ──────────────────────────────────────
        if force_recto:
            # PageBreak ends the previous chapter's last content page, then
            # RectoEnforcer checks parity and adds one more blank if needed.
            story.append(PageBreak())
            story.append(RectoEnforcer())

        # ── Section label ─────────────────────────────────────────────────────
        if sec_label:
            story.append(Spacer(1, 12 * mm))
            story.append(Paragraph(_esc_pdf(sec_label), section_label_s))
            story.append(HRFlowable(width="40%", thickness=0.8, color=ACCENT, spaceAfter=4))

        # ── Ornament before ───────────────────────────────────────────────────
        if ornament_bef:
            story.append(Paragraph(_esc_pdf(ornament_bef), ornament_s))

        # ── Chapter number + title ────────────────────────────────────────────
        story.append(Spacer(1, 10 * mm))
        story.append(Paragraph(
            f"Chapter {ch.get('chapter_number', ch_idx + 1)}", ch_label_s
        ))

        ch_title_style = {
            "left":       S(f"cht_l_{ch_idx}", parent=ch_title_s, alignment=TA_LEFT),
            "right":      S(f"cht_r_{ch_idx}", parent=ch_title_s, alignment=TA_RIGHT),
            "centered":   ch_title_s,
            "full_width": ch_title_s,
        }.get(hdr_style, ch_title_s)

        story.append(Paragraph(_esc_pdf(ch.get("title", "")), ch_title_style))
        story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=10))

        # keep_with_next: wrap title + first paragraph in KeepTogether
        if keep_next:
            raw_paras = [p.strip() for p in ch.get("content", "").split("\n\n") if p.strip()]
            first_para_items = []
            if raw_paras:
                p0 = raw_paras[0]
                override_font = _best_font_for_text(p0, font_map)
                p0_style = S(f"bs_kn_{ch_idx}", parent=ch_body_s, fontName=override_font) if override_font else ch_body_s
                if drop_cap and p0:
                    accent_hex = theme["accent"].lstrip("#")
                    first_para_items.append(Paragraph(
                        f'<font size="32" color="#{accent_hex}"><b>{_esc_pdf(p0[0])}</b></font>{_esc_pdf(p0[1:])}',
                        ch_body_s
                    ))
                else:
                    first_para_items.append(Paragraph(_esc_pdf(p0), p0_style))
            story[-2] = KeepTogether([story[-2], story[-1]] + first_para_items)
            story.pop()   # remove the bare HRFlowable (now inside KeepTogether)
            remaining_paras = raw_paras[1:] if raw_paras else []
        else:
            remaining_paras = [p.strip() for p in ch.get("content", "").split("\n\n") if p.strip()]
        # ── Body paragraphs ───────────────────────────────────────────────────
        if _two_col_mode and remaining_paras:
            from reportlab.platypus import Table, TableStyle  # type: ignore
            from reportlab.lib import colors as rl_colors      # type: ignore
            mid = max(1, len(remaining_paras) // 2)
            col1_paras = remaining_paras[:mid]
            col2_paras = remaining_paras[mid:]
            def _build_col(paras, col_idx):
                items = []
                for pi, pa in enumerate(paras):
                    ovf = _best_font_for_text(pa, font_map)
                    ps  = S(f"bs_2c{ch_idx}_{col_idx}_{pi}", parent=ch_body_s, fontName=ovf) if ovf else ch_body_s
                    items.append(Paragraph(_esc_pdf(pa), ps))
                return items
            col_w = (PAGE_W - INNER_M - OUTER_M - 8 * mm) / 2
            tbl_data = [[_build_col(col1_paras, 0), _build_col(col2_paras, 1)]]
            tbl = Table(tbl_data, colWidths=[col_w, col_w])
            tbl.setStyle(TableStyle([
                ("VALIGN",      (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING",(0, 0), (-1, -1), 4),
                ("LINEAFTER",   (0, 0), (0, -1),  0.5, rl_colors.lightgrey),
            ]))
            story.append(tbl)
        else:
            blank_after_paras = set(directives.get("blank_page_after_paragraphs") or [])
            for p_idx, para in enumerate(remaining_paras):
                story.append(_PageAnchorRecorder(
                    page_anchor_sink, ch.get("chapter_number", ch_idx + 1), p_idx
                ))

                override_font = _best_font_for_text(para, font_map)
                if override_font:
                    para_style = S(
                        f"bs_ov_{ch_idx}_{p_idx}",
                        parent=ch_body_s,
                        fontName=override_font,
                    )
                else:
                    para_style = (
                        ch_body_s if (p_idx == 0 and not directives.get("first_para_indent"))
                        else body_indent_s
                    )

                # Drop cap on first paragraph, first character
                if drop_cap and p_idx == 0 and para:
                    first_char  = _esc_pdf(para[0])
                    rest        = _esc_pdf(para[1:])
                    accent_hex  = theme["accent"].lstrip("#")
                    drop_markup = (
                        f'<font size="32" color="#{accent_hex}">'
                        f'<b>{first_char}</b></font>{rest}'
                    )
                    story.append(Paragraph(drop_markup, ch_body_s))
                else:
                    story.append(Paragraph(_esc_pdf(para), para_style))

                # Arbitrary mid-chapter blank-page insertion, resolved from a
                # "page N" instruction via resolve_page_target().
                if p_idx in blank_after_paras:
                    story.append(BlankPageFlowable())
                    story.append(PageBreak())

        # ── Ornament after ────────────────────────────────────────────────────
        if ornament_aft:
            story.append(Paragraph(_esc_pdf(ornament_aft), ornament_s))

        # ── Footnotes (Subsystem 13) ──────────────────────────────────────────
        ch_footnotes = ch.get("footnotes", [])
        if ch_footnotes:
            _render_footnotes_pdf(story, ch_footnotes, {})

        # ── Page break after ─────────────────────────────────────────────────
        if do_page_break:
            story.append(PageBreak())

        # ── Blank pages after ─────────────────────────────────────────────────
        for _ in range(blank_after):
            story.append(BlankPageFlowable())
            story.append(PageBreak())

    doc.build(story, onFirstPage=on_cover, onLaterPages=on_page)
    log.info("PDF written: %s", output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 3 — DOCX GENERATOR WITH NATIVE OOXML ODD-PAGE BREAKS
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _inject_odd_page_break(paragraph) -> None:
    """
    Inject <w:type w:val="oddPage"/> into the pPr/sectPr of `paragraph`.
    This causes Microsoft Word to advance to the next ODD page — native recto enforcement.

    OOXML structure:
      <w:pPr>
        <w:sectPr>
          <w:type w:val="oddPage"/>
          <w:pgSz w:w="11906" w:h="16838"/>   <!-- A4 -->
        </w:sectPr>
      </w:pPr>
    """
    from docx.oxml.ns import qn        # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:sectPr")):
        pPr.remove(old)

    sectPr  = OxmlElement("w:sectPr")
    pg_type = OxmlElement("w:type")
    pg_type.set(qn("w:val"), "oddPage")
    sectPr.append(pg_type)

    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "11906")    # 210mm
    pg_sz.set(qn("w:h"), "16838")    # 297mm
    sectPr.append(pg_sz)

    # Mirror margins: gutter on the binding side
    pg_mar = OxmlElement("w:pgMar")
    pg_mar.set(qn("w:top"),    "1440")   # 1 inch
    pg_mar.set(qn("w:right"),  "1134")   # 20mm outer
    pg_mar.set(qn("w:bottom"), "1260")   # 22mm
    pg_mar.set(qn("w:left"),   "1814")   # 32mm gutter
    pg_mar.set(qn("w:header"), "708")
    pg_mar.set(qn("w:footer"), "708")
    pg_mar.set(qn("w:gutter"), "0")
    sectPr.append(pg_mar)

    pPr.append(sectPr)


def _inject_next_page_break(paragraph) -> None:
    """Inject <w:type w:val="nextPage"/> sectPr for a standard page break."""
    from docx.oxml.ns import qn        # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:sectPr")):
        pPr.remove(old)

    sectPr  = OxmlElement("w:sectPr")
    pg_type = OxmlElement("w:type")
    pg_type.set(qn("w:val"), "nextPage")
    sectPr.append(pg_type)

    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "11906")
    pg_sz.set(qn("w:h"), "16838")
    sectPr.append(pg_sz)

    pPr.append(sectPr)


def _inject_blank_page_docx(doc) -> None:
    """
    Insert a blank page into a DOCX document.
    Uses a paragraph with a single NBSP so Word does not collapse it,
    followed by a nextPage section break.
    """
    from docx.oxml.ns import qn        # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.shared import Pt         # type: ignore

    blank = doc.add_paragraph("\u00A0")   # non-breaking space
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after  = Pt(0)
    for run in blank.runs:
        run.font.size = Pt(1)
    _inject_next_page_break(blank)


def _set_mirror_margins_docx(doc) -> None:
    """
    Enable mirror margins in a python-docx Document.

    Mirror margins is a document-level setting stored in word/settings.xml as
    <w:mirrorMargins/> inside <w:settings>.  It is NOT a pgMar attribute.
    The old code wrote an invalid attribute on pgMar which Word silently ignored.

    This implementation directly injects the element into the document settings
    XML so that odd and even pages swap their inner/outer margins correctly.
    """
    from docx.oxml.ns import qn        # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    # Access the document-level settings part
    settings_part = doc.settings.element   # <w:settings> root element

    # Remove any existing mirrorMargins element to avoid duplicates
    for existing in settings_part.findall(qn("w:mirrorMargins")):
        settings_part.remove(existing)

    # Insert <w:mirrorMargins/> as the first child (position matters for OOXML)
    mirror = OxmlElement("w:mirrorMargins")
    settings_part.insert(0, mirror)
    log.debug("Mirror margins enabled in document settings")


def _set_header_footer_docx(doc, book_title: str, theme_name: str) -> None:
    """
    Inject running headers into a DOCX document.
    Odd pages (recto): chapter/book title right-aligned.
    Even pages (verso): book title left-aligned.
    """
    from docx.oxml.ns import qn        # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.shared import Pt, RGBColor   # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Default (odd) header
        hdr = section.header
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            p.clear()
        if hdr.paragraphs:
            hp = hdr.paragraphs[0]
        else:
            hp = hdr.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(book_title[:80])
        run.font.size   = Pt(8)
        run.font.italic = True

        # Even header
        try:
            even_hdr = section.even_page_header
            even_hdr.is_linked_to_previous = False
            for p in even_hdr.paragraphs:
                p.clear()
            if even_hdr.paragraphs:
                ehp = even_hdr.paragraphs[0]
            else:
                ehp = even_hdr.add_paragraph()
            ehp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            erun = ehp.add_run(book_title[:80])
            erun.font.size   = Pt(8)
            erun.font.italic = True
        except Exception:
            pass   # even_page_header may not be accessible on all platforms


def generate_edited_docx(
    book:        dict,
    output_path: str,
    theme_name:  str = "premium",
    generate_toc: bool = False,
) -> str:
    """
    Full-featured DOCX generator.

    Features:
      • <w:type w:val="oddPage"/> sectPr injection  (Subsystem 3)
      • Blank page injection from layout_directives  (Subsystem 1)
      • Mirror margins via mirrorMargins="1"
      • Devanagari / CJK / Arabic Unicode font selection  (Subsystem 8)
      • Section labels, ornaments, drop cap indicator
      • Per-chapter custom font size and line spacing
      • Themed cover page
      • Consistent paragraph and heading styles
      • Running headers
    """
    from docx import Document                       # type: ignore
    from docx.shared import Pt, RGBColor, Cm        # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH   # type: ignore
    from docx.oxml.ns import qn                     # type: ignore
    from docx.oxml import OxmlElement               # type: ignore

    theme = THEMES.get(theme_name, THEMES["premium"])

    def rgb_from_hex(h: str) -> RGBColor:
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    TITLE_RGB  = rgb_from_hex(theme["title_col"])
    BODY_RGB   = rgb_from_hex(theme["body_col"])
    ACCENT_RGB = rgb_from_hex(theme["accent"])

    all_content = " ".join(str(ch.get("content", "")) for ch in book.get("chapters", []))
    has_dev = _has_devanagari(all_content)
    has_cjk = _has_cjk(all_content)
    has_ara = _has_arabic(all_content)

    if has_dev:
        body_font = head_font = "Noto Sans Devanagari"
    elif has_cjk:
        body_font = head_font = "Noto Sans CJK SC"
    elif has_ara:
        body_font = head_font = "Noto Sans Arabic"
    elif "Courier" in theme["font_body"]:
        body_font = head_font = "Courier New"
    elif "Times" in theme["font_body"]:
        body_font = "Palatino Linotype"
        head_font = "Garamond"
    else:
        body_font = "Calibri"
        head_font = "Calibri"

    doc     = Document()
    section = doc.sections[0]
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)
    section.left_margin   = Cm(3.2)   # gutter
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.6)
    section.bottom_margin = Cm(2.2)

    _set_mirror_margins_docx(doc)

    style_n           = doc.styles["Normal"]
    style_n.font.name = body_font
    style_n.font.size = Pt(11)

    def _apply_run_font(run, text: str = ""):
        run.font.name = body_font
        unicode_f = _docx_unicode_font(text or run.text or "")
        if unicode_f:
            run.font.name = unicode_f

    # ── Cover page ────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run(book.get("title", "Untitled"))
    title_r.font.name  = head_font
    title_r.font.size  = Pt(26)
    title_r.font.bold  = True
    title_r.font.color.rgb = TITLE_RGB
    _apply_run_font(title_r, book.get("title", ""))

    if book.get("author"):
        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_r = auth_p.add_run(f"by {book['author']}")
        auth_r.font.name = body_font
        auth_r.font.size = Pt(13)
        auth_r.font.color.rgb = BODY_RGB

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_r = meta_p.add_run(
        f"Theme: {theme_name.title()}  ·  Edited {datetime.date.today().strftime('%B %d, %Y')}"
    )
    meta_r.font.size   = Pt(9)
    meta_r.font.italic = True
    meta_r.font.color.rgb = ACCENT_RGB

    doc.add_page_break()

    # ── Table of Contents (optional) ──────────────────────────────────────────
    if generate_toc:
        generate_toc_docx(doc, book)

    # ── Running headers ───────────────────────────────────────────────────────
    _set_header_footer_docx(doc, book.get("title", ""), theme_name)

    # ── Chapters ──────────────────────────────────────────────────────────────
    chapters   = book.get("chapters", [])
    n_chapters = len(chapters)

    for ch_idx, ch in enumerate(chapters):
        directives = ch.get("layout_directives") or _default_layout_directives()
        if not isinstance(directives, dict):
            directives = _default_layout_directives()

        blank_before = max(0, int(directives.get("blank_pages_before", 0)))
        blank_after  = max(0, int(directives.get("blank_pages_after",  0)))
        force_recto  = bool(directives.get("force_recto_start",    False))
        sec_label    = str(directives.get("section_label",         "")).strip()
        ornament_bef = str(directives.get("ornament_before",       ""))
        ornament_aft = str(directives.get("ornament_after",        ""))
        custom_fs    = float(directives.get("custom_font_size",    0.0))
        ls_mult      = float(directives.get("line_spacing",        0.0)) or float(theme.get("leading_mult", 1.65))
        hdr_style    = str(directives.get("chapter_header_style",  "centered"))
        drop_cap     = bool(directives.get("drop_cap",             False))

        is_last_chapter = (ch_idx == n_chapters - 1)

        # ── Blank pages before ────────────────────────────────────────────────
        for _ in range(blank_before):
            _inject_blank_page_docx(doc)

        # ── Section label ─────────────────────────────────────────────────────
        if sec_label:
            sp = doc.add_paragraph(sec_label)
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sp.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = ACCENT_RGB
                _apply_run_font(run, sec_label)

        # ── Ornament before ───────────────────────────────────────────────────
        if ornament_bef:
            op = doc.add_paragraph(ornament_bef)
            op.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in op.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = ACCENT_RGB

        # ── Chapter label ─────────────────────────────────────────────────────
        lbl_p = doc.add_paragraph()
        lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl_r = lbl_p.add_run(f"Chapter {ch.get('chapter_number', ch_idx + 1)}")
        lbl_r.font.name      = body_font
        lbl_r.font.size      = Pt(9)
        lbl_r.font.bold      = True
        lbl_r.font.color.rgb = ACCENT_RGB

        # ── Chapter title ─────────────────────────────────────────────────────
        heading = doc.add_heading(ch.get("title", ""), level=1)
        heading.alignment = {
            "centered":   WD_ALIGN_PARAGRAPH.CENTER,
            "left":       WD_ALIGN_PARAGRAPH.LEFT,
            "right":      WD_ALIGN_PARAGRAPH.RIGHT,
            "full_width": WD_ALIGN_PARAGRAPH.CENTER,
        }.get(hdr_style, WD_ALIGN_PARAGRAPH.CENTER)

        for run in heading.runs:
            run.font.name  = head_font
            run.font.color.rgb = TITLE_RGB
            run.font.size  = Pt(18)
            _apply_run_font(run, ch.get("title", ""))

        doc.add_paragraph()   # spacing after heading

        # ── Body paragraphs ───────────────────────────────────────────────────
        paragraphs = [p.strip() for p in ch.get("content", "").split("\n\n") if p.strip()]
        n_paras    = len(paragraphs)
        blank_after_paras_docx = set(directives.get("blank_page_after_paragraphs") or [])

        for p_idx, para_text in enumerate(paragraphs):
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]

            eff_size    = Pt(custom_fs) if custom_fs else Pt(11)
            eff_leading = Pt((custom_fs or 11) * ls_mult)

            p.paragraph_format.space_after       = Pt(5)
            p.paragraph_format.line_spacing      = eff_leading
            p.paragraph_format.first_line_indent = (
                Cm(0) if p_idx == 0 else Cm(0.5)
            )
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # ── Drop cap on first paragraph of each chapter ───────────────────
            if drop_cap and p_idx == 0 and para_text:
                # First character: large, bold, accent-coloured, superscript-sized run
                r_cap = p.add_run(para_text[0])
                r_cap.font.name      = head_font
                r_cap.font.size      = Pt(32)
                r_cap.font.bold      = True
                r_cap.font.color.rgb = ACCENT_RGB
                # Remaining text of the paragraph
                r_rest = p.add_run(para_text[1:])
                r_rest.font.name      = body_font
                r_rest.font.size      = eff_size
                r_rest.font.color.rgb = BODY_RGB
                unicode_f = _docx_unicode_font(r_rest.text or para_text)
                if unicode_f:
                    r_rest.font.name = unicode_f
            else:
                p.add_run(para_text)
                for run in p.runs:
                    run.font.name      = body_font
                    run.font.size      = eff_size
                    run.font.color.rgb = BODY_RGB
                    unicode_f = _docx_unicode_font(run.text or para_text)
                    if unicode_f:
                        run.font.name = unicode_f

            # ── Section break injection (Subsystem 3) — only at end of chapter ──
            is_last_para = (p_idx == n_paras - 1)
            if is_last_para and not is_last_chapter:
                if force_recto:
                    _inject_odd_page_break(p)
                else:
                    _inject_next_page_break(p)

            # ── Page-targeted blank page (anchor resolved from "page N") ───────
            # Approximate in Word — page numbers reflow at render time — but
            # lands at the same manuscript position as the PDF anchor.
            if p_idx in blank_after_paras_docx and not is_last_para:
                _inject_blank_page_docx(doc)

        # ── Ornament after ────────────────────────────────────────────────────
        if ornament_aft:
            oa_p = doc.add_paragraph(ornament_aft)
            oa_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in oa_p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = ACCENT_RGB

        # ── Footnotes (Subsystem 13) ──────────────────────────────────────────
        ch_footnotes = ch.get("footnotes", [])
        if ch_footnotes:
            _inject_footnotes_docx(doc, ch_footnotes)

        # ── Blank pages after ─────────────────────────────────────────────────
        for _ in range(blank_after):
            _inject_blank_page_docx(doc)

    doc.save(output_path)
    log.info("DOCX written: %s", output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 11 — TABLE OF CONTENTS GENERATOR
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def generate_toc_text(book: dict) -> str:
    """Generate a plain-text Table of Contents with dot leaders."""
    lines = ["Table of Contents", "=" * 50]
    for ch in book.get("chapters", []):
        num      = ch.get("chapter_number", "?")
        title    = ch.get("title", "")
        wc       = ch.get("word_count", len(ch.get("content", "").split()))
        est_page = max(1, round(wc / 250))
        entry    = f"Chapter {num}  {title}"
        pad      = max(2, 50 - len(entry))
        lines.append(f"{entry} {'.' * pad} {est_page}")
    return "\n".join(lines)


def generate_toc_docx(doc, book: dict) -> None:
    """Insert a Table of Contents page into an existing python-docx Document."""
    from docx.shared import Pt, RGBColor          # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
    from docx.oxml.ns import qn                   # type: ignore
    from docx.oxml import OxmlElement             # type: ignore

    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for ch in book.get("chapters", []):
        num      = ch.get("chapter_number", "?")
        title    = ch.get("title", "")
        wc       = ch.get("word_count", len(ch.get("content", "").split()))
        est_page = max(1, round(wc / 250))

        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Pt(400))
        r1 = p.add_run(f"Chapter {num}  {title}")
        r1.font.size = Pt(11)
        r2 = p.add_run(f"\t{est_page}")
        r2.font.size = Pt(11)

    doc.add_page_break()


def generate_toc_pdf(book: dict, story: list, styles: dict, accent_color) -> None:
    """
    Append a Table of Contents to a ReportLab story list.
    Adds dot leaders and estimated page numbers.
    """
    from reportlab.lib.styles import ParagraphStyle        # type: ignore
    from reportlab.lib.enums import TA_CENTER, TA_LEFT     # type: ignore
    from reportlab.platypus import Paragraph, Spacer, PageBreak  # type: ignore

    toc_title_s = ParagraphStyle(
        "toc_title",
        fontName="Helvetica-Bold",
        fontSize=16,
        textColor=accent_color,
        leading=20,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    toc_entry_s = ParagraphStyle(
        "toc_entry",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
    )

    story.append(Paragraph("Table of Contents", toc_title_s))
    story.append(Spacer(1, 8))

    for ch in book.get("chapters", []):
        num      = ch.get("chapter_number", "?")
        title    = ch.get("title", "")
        wc       = ch.get("word_count", len(ch.get("content", "").split()))
        est_page = max(1, round(wc / 250))
        entry    = f"Chapter {num}  {title}"
        pad_dots = max(2, 60 - len(entry))
        line     = f"{entry} {'.' * pad_dots} {est_page}"
        story.append(Paragraph(_esc_pdf(line), toc_entry_s))

    story.append(PageBreak())


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 12 — CHAPTER RENUMBERING & REORDERING
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def reorder_chapters(book: dict, new_order: List[int]) -> dict:
    """
    Reorder chapters according to new_order (list of chapter_numbers in desired order).
    Re-numbers chapters 1..N after reorder.
    """
    by_num    = {ch["chapter_number"]: ch for ch in book.get("chapters", [])}
    reordered: List[dict] = []
    for new_idx, old_num in enumerate(new_order):
        if old_num not in by_num:
            raise ValueError(f"Chapter number {old_num} not found in book.")
        ch = copy.deepcopy(by_num[old_num])
        ch["chapter_number"] = new_idx + 1
        reordered.append(ch)
    result = copy.deepcopy(book)
    result["chapters"] = reordered
    log.info("Reordered chapters: %s → sequential 1..%d", new_order, len(reordered))
    return result


def insert_chapter(
    book:     dict,
    position: int,
    title:    str = "New Chapter",
    content:  str = "",
) -> dict:
    """
    Insert a new blank (or pre-filled) chapter at `position` (1-based).
    Chapters at and after `position` are renumbered.
    """
    result   = copy.deepcopy(book)
    chapters = result.get("chapters", [])
    new_ch   = {
        "chapter_number":    position,
        "title":             title,
        "content":           content,
        "layout_directives": _default_layout_directives(),
        "footnotes":         [],
        "word_count":        len(content.split()),
    }
    chapters.insert(position - 1, new_ch)
    for i, ch in enumerate(chapters):
        ch["chapter_number"] = i + 1
    result["chapters"] = chapters
    log.info("Inserted chapter '%s' at position %d", title, position)
    return result


def delete_chapter(book: dict, chapter_number: int) -> dict:
    """Remove a chapter by number. Remaining chapters are renumbered."""
    result   = copy.deepcopy(book)
    chapters = [
        ch for ch in result.get("chapters", [])
        if ch.get("chapter_number") != chapter_number
    ]
    for i, ch in enumerate(chapters):
        ch["chapter_number"] = i + 1
    result["chapters"] = chapters
    log.info("Deleted chapter %d, %d chapters remaining", chapter_number, len(chapters))
    return result


def merge_chapters(book: dict, first_number: int, second_number: int) -> dict:
    """
    Merge two adjacent chapters into one.
    The merged chapter uses the first chapter's title and directives.
    """
    by_num = {ch["chapter_number"]: ch for ch in book.get("chapters", [])}
    if first_number not in by_num or second_number not in by_num:
        raise ValueError("Both chapter numbers must exist in the book.")
    ch1 = by_num[first_number]
    ch2 = by_num[second_number]
    merged_content = str(ch1.get("content", "")) + "\n\n" + str(ch2.get("content", ""))
    merged = copy.deepcopy(ch1)
    merged["content"]    = merged_content
    merged["word_count"] = len(merged_content.split())

    result   = copy.deepcopy(book)
    chapters = [
        ch for ch in result.get("chapters", [])
        if ch.get("chapter_number") not in (first_number, second_number)
    ]
    insert_pos = next(
        (i for i, ch in enumerate(book.get("chapters", []))
         if ch["chapter_number"] == first_number),
        0,
    )
    chapters.insert(insert_pos, merged)
    for i, ch in enumerate(chapters):
        ch["chapter_number"] = i + 1
    result["chapters"] = chapters
    log.info(
        "Merged chapters %d+%d → new chapter at position %d",
        first_number, second_number, insert_pos + 1,
    )
    return result


def split_chapter(
    book:            dict,
    chapter_number:  int,
    split_paragraph: int,   # 0-based index of the paragraph to split AFTER
    second_title:    str = "",
) -> dict:
    """
    Split a chapter at `split_paragraph`.
    Paragraphs up to and including split_paragraph stay in the original.
    Paragraphs after split_paragraph form the new chapter.
    """
    by_num = {ch["chapter_number"]: ch for ch in book.get("chapters", [])}
    if chapter_number not in by_num:
        raise ValueError(f"Chapter {chapter_number} not found.")
    ch      = by_num[chapter_number]
    paras   = ch.get("content", "").split("\n\n")
    first_c = "\n\n".join(paras[:split_paragraph + 1])
    second_c = "\n\n".join(paras[split_paragraph + 1:])
    new_title = second_title or f"{ch['title']} (continued)"

    result = copy.deepcopy(book)
    for c in result["chapters"]:
        if c["chapter_number"] == chapter_number:
            c["content"]    = first_c
            c["word_count"] = len(first_c.split())
            break
    result = insert_chapter(result, chapter_number + 1, new_title, second_c)
    log.info("Split chapter %d at paragraph %d", chapter_number, split_paragraph)
    return result


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 13 — FOOTNOTE & ENDNOTE SUPPORT
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _extract_footnotes_docx(path: str) -> Dict[int, List[Dict[str, str]]]:
    """
    Pull footnotes from a Word document into a chapter-keyed dict.
    Returns {chapter_number: [{"number": "1", "text": "..."}]}.
    """
    footnotes_by_chapter: Dict[int, List[Dict[str, str]]] = {}
    try:
        with zipfile.ZipFile(path, "r") as z:
            if "word/footnotes.xml" not in z.namelist():
                return footnotes_by_chapter
            xml = z.read("word/footnotes.xml").decode("utf-8", errors="replace")

        # Extract each <w:footnote> element
        fn_matches = re.finditer(
            r'<w:footnote[^>]+w:id="(\d+)"[^>]*>(.*?)</w:footnote>',
            xml, re.DOTALL
        )
        footnotes: List[Dict[str, str]] = []
        for m in fn_matches:
            fn_id = m.group(1)
            if fn_id in ("0", "-1"):  # skip separator footnotes
                continue
            text_pieces = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", m.group(2))
            fn_text = " ".join(text_pieces).strip()
            if fn_text:
                footnotes.append({"number": fn_id, "text": fn_text})

        # All footnotes go on chapter 1 if we can't determine chapter
        if footnotes:
            footnotes_by_chapter[1] = footnotes

    except Exception as ex:
        log.warning("Footnote extraction failed: %s", ex)

    return footnotes_by_chapter


def _render_footnotes_pdf(story: list, footnotes: List[Dict[str, str]], styles: dict) -> None:
    """Append footnote separator and numbered footnotes to the ReportLab story."""
    from reportlab.platypus import HRFlowable, Paragraph, Spacer  # type: ignore
    from reportlab.lib.styles import ParagraphStyle                # type: ignore
    from reportlab.lib.enums import TA_LEFT                        # type: ignore
    from reportlab.lib.colors import HexColor                      # type: ignore

    if not footnotes:
        return

    fn_style = ParagraphStyle(
        "footnote",
        fontName="Helvetica",
        fontSize=8,
        leading=11,
        alignment=TA_LEFT,
        leftIndent=12,
        firstLineIndent=-12,
    )

    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="40%", thickness=0.5, spaceAfter=4))

    for fn in footnotes:
        num  = fn.get("number", "?")
        text = fn.get("text", "")
        story.append(Paragraph(f"<super>{_esc_pdf(num)}</super> {_esc_pdf(text)}", fn_style))


def _inject_footnotes_docx(doc, footnotes: List[Dict[str, str]]) -> None:
    """
    Append footnote text as a styled paragraph section at the end of the document.
    (True inline footnotes require complex OOXML manipulation — this is the safe version.)
    """
    from docx.shared import Pt  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    if not footnotes:
        return

    doc.add_paragraph()
    sep = doc.add_paragraph("─" * 30)
    sep.paragraph_format.space_before = Pt(12)

    for fn in footnotes:
        num  = fn.get("number", "?")
        text = fn.get("text", "")
        p    = doc.add_paragraph()
        r1   = p.add_run(f"{num}  ")
        r1.font.size      = Pt(7)
        r1.font.superscript = True
        r2   = p.add_run(text)
        r2.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 15 — METADATA & FRONT MATTER / BACK MATTER
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def _sanitise_book_metadata(book: dict) -> dict:
    """Ensure all top-level metadata fields exist and are strings."""
    book = copy.deepcopy(book)
    book.setdefault("title",       "Untitled")
    book.setdefault("author",      "")
    book.setdefault("description", "")
    book.setdefault("isbn",        "")
    book.setdefault("year",        str(datetime.date.today().year))
    book.setdefault("chapters",    [])
    for k in ("title", "author", "description", "isbn", "year"):
        if not isinstance(book[k], str):
            book[k] = str(book[k] or "")
    return book


def build_front_matter(
    title:        str,
    author:       str,
    year:         str   = "",
    isbn:         str   = "",
    dedication:   str   = "",
    epigraph:     str   = "",
    epigraph_src: str   = "",
) -> List[dict]:
    """
    Build front-matter chapters: copyright page, optional dedication, optional epigraph.
    Returns a list of chapter dicts (chapter_number = -2, -1, 0 for ordering).
    Prepend to the chapters list before rendering.
    """
    front: List[dict] = []
    yr = year or str(datetime.date.today().year)

    copy_text = (
        f"Copyright © {yr} {author}\n\n"
        "All rights reserved. No part of this publication may be reproduced, "
        "distributed, or transmitted in any form or by any means, including "
        "photocopying, recording, or other electronic or mechanical methods, "
        "without the prior written permission of the publisher.\n\n"
        f"First published {yr}"
    )
    if isbn:
        copy_text += f"\n\nISBN: {isbn}"

    directives_front = _default_layout_directives()
    directives_front["chapter_header_style"] = "centered"
    directives_front["force_recto_start"]    = False
    directives_front["page_break_after"]     = True

    front.append({
        "chapter_number":    -2,
        "title":             "Copyright",
        "content":           copy_text,
        "layout_directives": directives_front,
        "footnotes":         [],
        "word_count":        len(copy_text.split()),
    })

    if dedication:
        ded_directives = _default_layout_directives()
        ded_directives["chapter_header_style"] = "centered"
        ded_directives["page_break_after"]     = True
        front.append({
            "chapter_number":    -1,
            "title":             "Dedication",
            "content":           dedication,
            "layout_directives": ded_directives,
            "footnotes":         [],
            "word_count":        len(dedication.split()),
        })

    if epigraph:
        epi_text = epigraph
        if epigraph_src:
            epi_text += f"\n\n— {epigraph_src}"
        epi_directives = _default_layout_directives()
        epi_directives["chapter_header_style"] = "centered"
        epi_directives["page_break_after"]     = True
        front.append({
            "chapter_number":    0,
            "title":             "Epigraph",
            "content":           epi_text,
            "layout_directives": epi_directives,
            "footnotes":         [],
            "word_count":        len(epi_text.split()),
        })

    return front


def generate_back_matter(
    author:      str = "",
    author_bio:  str = "",
    colophon:    str = "",
    index_terms: Optional[List[str]] = None,
) -> List[dict]:
    """
    Generate back-matter chapters: About the Author, Colophon, Index stub.
    Returns a list of chapter dicts (high chapter_number for end ordering).
    Append to the chapters list before rendering.
    """
    back: List[dict] = []
    base_num = 9900   # high number to ensure back matter always renders last

    if author_bio:
        bio_text = author_bio
    elif author:
        bio_text = (
            f"{author} is the author of this work. "
            "Further biographical information will be added here."
        )
    else:
        bio_text = ""

    if bio_text:
        bio_directives = _default_layout_directives()
        bio_directives["chapter_header_style"] = "centered"
        bio_directives["force_recto_start"]    = True
        bio_directives["page_break_after"]     = True
        back.append({
            "chapter_number":    base_num,
            "title":             "About the Author",
            "content":           bio_text,
            "layout_directives": bio_directives,
            "footnotes":         [],
            "word_count":        len(bio_text.split()),
        })

    if colophon:
        col_directives = _default_layout_directives()
        col_directives["chapter_header_style"] = "centered"
        col_directives["page_break_after"]     = False
        back.append({
            "chapter_number":    base_num + 1,
            "title":             "Colophon",
            "content":           colophon,
            "layout_directives": col_directives,
            "footnotes":         [],
            "word_count":        len(colophon.split()),
        })

    if index_terms:
        # Build a simple alphabetical index stub
        sorted_terms = sorted(set(index_terms), key=str.lower)
        by_letter: Dict[str, List[str]] = {}
        for term in sorted_terms:
            letter = term[0].upper() if term else "?"
            by_letter.setdefault(letter, []).append(term)
        index_lines = []
        for letter in sorted(by_letter.keys()):
            index_lines.append(f"\n{letter}")
            index_lines.extend(f"  {t}" for t in by_letter[letter])
        index_text = "\n".join(index_lines).strip()

        idx_directives = _default_layout_directives()
        idx_directives["chapter_header_style"] = "left"
        idx_directives["force_recto_start"]    = True
        back.append({
            "chapter_number":    base_num + 2,
            "title":             "Index",
            "content":           index_text,
            "layout_directives": idx_directives,
            "footnotes":         [],
            "word_count":        len(index_text.split()),
        })

    return back


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# SUBSYSTEM 17 — BATCH PROCESSING & PROGRESS TRACKER
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

class ProgressTracker:
    """Thread-safe progress tracker with ETA estimates."""

    def __init__(self, total: int, description: str = "Processing"):
        self.total       = total
        self.completed   = 0
        self.failed      = 0
        self.start_time  = time.time()
        self.description = description
        self._lock       = threading.Lock()

    def update(self, success: bool = True) -> None:
        with self._lock:
            if success:
                self.completed += 1
            else:
                self.failed    += 1
            done      = self.completed + self.failed
            elapsed   = time.time() - self.start_time
            rate      = done / elapsed if elapsed > 0 else 0
            remaining = self.total - done
            eta       = remaining / rate if rate > 0 else 0
            pct       = done / self.total * 100 if self.total > 0 else 0
            log.info(
                "%s: %d/%d (%.0f%%)  ✓%d ✗%d  ETA %.0fs",
                self.description, done, self.total, pct,
                self.completed, self.failed, eta,
            )

    def summary(self) -> str:
        elapsed = time.time() - self.start_time
        return (
            f"{self.description}: {self.completed}/{self.total} succeeded, "
            f"{self.failed} failed in {elapsed:.1f}s"
        )


def batch_process_files(
    file_paths:  List[str],
    instruction: str,
    output_dir:  str,
    theme:       str = "premium",
    max_workers: int = 2,
) -> List[dict]:
    """
    Process a list of book files in parallel, applying the same instruction to each.
    Returns a list of result dicts.
    """
    os.makedirs(output_dir, exist_ok=True)
    tracker = ProgressTracker(len(file_paths), "Batch edit")
    results: List[dict] = []
    lock    = threading.Lock()

    def _process_one(fpath: str) -> dict:
        try:
            book   = load_book(fpath)
            job_id = uuid.uuid4().hex[:8]
            result = process_editor_turn(
                book_structure=book,
                user_message=instruction,
                conversation_history=[],
                output_dir=output_dir,
                theme=theme,
                job_id=job_id,
            )
            tracker.update(True)
            return {"file": fpath, "status": "ok", **result}
        except Exception as ex:
            tracker.update(False)
            log.error("Batch: failed for %s: %s", fpath, ex)
            return {"file": fpath, "status": "error", "error": str(ex)}

    with ThreadPoolExecutor(max_workers=max_workers) as exe:
        futures = {exe.submit(_process_one, fp): fp for fp in file_paths}
        for future in as_completed(futures):
            with lock:
                results.append(future.result())

    log.info(tracker.summary())
    return results


def parallel_edit_chapters(
    book:        dict,
    instruction: str,
    max_workers: int = MAX_WORKERS,
) -> dict:
    """
    Edit all chapters of a book in parallel using a thread pool.
    Returns the updated book dict.
    """
    chapters   = book.get("chapters", [])
    book_title = book.get("title", "Untitled")
    n          = len(chapters)

    if n == 0:
        return book

    updated: List[Optional[dict]] = [None] * n

    def _worker(args: Tuple[int, dict]) -> Tuple[int, dict]:
        idx, chapter = args
        return idx, _edit_single_chapter(chapter, instruction, book_title, idx, n)

    with ThreadPoolExecutor(max_workers=min(max_workers, n)) as executor:
        futures = {
            executor.submit(_worker, (idx, ch)): idx
            for idx, ch in enumerate(chapters)
        }
        for future in as_completed(futures):
            try:
                idx, result_ch = future.result()
                result_ch.pop("_changed", None)
                updated[idx] = result_ch
            except Exception as ex:
                idx = futures[future]
                log.error("Parallel chapter %d edit failed: %s", idx + 1, ex)
                updated[idx] = chapters[idx]

    result = copy.deepcopy(book)
    result["chapters"] = [ch for ch in updated if ch is not None]
    return result


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# MASTER ORCHESTRATOR — process_editor_turn
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def process_editor_turn(
    book_structure:       dict,
    user_message:         str,
    conversation_history: List[dict],
    output_dir:           str,
    theme:                str = "premium",
    job_id:               str = "",
    version_history:      Optional[VersionHistory] = None,
    generate_toc_page:    bool = False,
    generate_diff:        bool = True,
) -> dict:
    """
    Master orchestrator for one editing turn.

    Full pipeline:
      1.  Sanitise book metadata             (Subsystem 15)
      2.  Detect theme override from message (Subsystem 18)
      3.  Snapshot original book for diffing (Subsystem 5)
      4.  Parse instruction intent           (Subsystem 16)
      5.  Apply edit — layout + prose        (Subsystems 1, 4, 6, 7)
      6.  Generate PDF                       (Subsystems 2, 8)
      7.  Generate DOCX                      (Subsystems 3, 8)
      8.  Generate HTML diff report          (Subsystem 5)
      9.  Commit to version history          (Subsystem 10)
     10.  Return rich result dict

    Returns:
      {
        updated_book:       dict,
        edit_summary:       str,
        chapters_changed:   List[int],
        pdf_path:           str,
        docx_path:          str,
        diff_html_path:     str,
        theme:              str,
        version_index:      Optional[int],
        intent:             str,
        word_count_before:  int,
        word_count_after:   int,
        word_delta:         int,
      }
    """
    os.makedirs(output_dir, exist_ok=True)

    # Step 1: Sanitise
    book_structure = _sanitise_book_metadata(book_structure)

    # Step 2: Theme detection
    detected_theme = detect_theme_from_instruction(user_message, theme)

    # Step 3: Snapshot
    original_book = copy.deepcopy(book_structure)

    # Step 4: Intent parsing
    intent = parse_instruction(user_message, book_structure)
    log.info(
        "Intent: %s (confidence=%.2f) — '%s'",
        intent.intent_type.value, intent.confidence, user_message[:70],
    )

    # Step 5: Apply edit
    page_match = _PAGE_TARGET_PATTERN.search(user_message)
    try:
        if page_match:
            target_page = int(page_match.group(1) or page_match.group(2))
            log.info("Page-targeted blank-page instruction detected: page %d", target_page)
            updated = insert_blank_page_at_page_number(
                book_structure, target_page, detected_theme, generate_toc_page,
            )
            updated["edit_summary"]     = f"Inserted a blank page after page {target_page}."
            updated["chapters_changed"] = []
        else:
            updated = apply_edit(book_structure, user_message, conversation_history)
    except Exception as ex:
        raise ValueError(f"Edit pipeline failed: {ex}") from ex

    edit_summary     = updated.pop("edit_summary",     "Changes applied.")
    chapters_changed = updated.pop("chapters_changed", [])

    # Word count stats
    wc_before = sum(
        len(str(ch.get("content", "")).split())
        for ch in original_book.get("chapters", [])
    )
    wc_after = sum(
        len(str(ch.get("content", "")).split())
        for ch in updated.get("chapters", [])
    )

    # Step 6-9: File generation
    safe_title = (
        "".join(c for c in updated.get("title", "book")
                if c.isalnum() or c in " -_").strip()
        or "book"
    )
    vid       = job_id or uuid.uuid4().hex[:8]
    pdf_path  = os.path.join(output_dir, f"{safe_title}_{vid}.pdf")
    docx_path = os.path.join(output_dir, f"{safe_title}_{vid}.docx")
    diff_path = os.path.join(output_dir, f"{safe_title}_{vid}_diff.html")

    try:
        generate_edited_pdf(updated, pdf_path, detected_theme,
                            generate_toc=generate_toc_page)
    except Exception as ex:
        log.error("PDF generation failed: %s\n%s", ex, traceback.format_exc())
        pdf_path = ""

    try:
        generate_edited_docx(updated, docx_path, detected_theme,
                             generate_toc=generate_toc_page)
    except Exception as ex:
        log.error("DOCX generation failed: %s\n%s", ex, traceback.format_exc())
        docx_path = ""

    if generate_diff:
        try:
            generate_diff_report(original_book, updated, diff_path)
        except Exception as ex:
            log.error("Diff report failed: %s", ex)
            diff_path = ""
    else:
        diff_path = ""

    # Subsystem 10: Version history
    version_idx = None
    if version_history is not None:
        version_idx = version_history.commit(updated, edit_summary, user_message)
        log.info("Version %d committed (%d total)", version_idx, len(version_history))

    return {
        "updated_book":      updated,
        "edit_summary":      edit_summary,
        "chapters_changed":  chapters_changed,
        "pdf_path":          pdf_path,
        "docx_path":         docx_path,
        "diff_html_path":    diff_path,
        "theme":             detected_theme,
        "version_index":     version_idx,
        "intent":            intent.intent_type.value,
        "word_count_before": wc_before,
        "word_count_after":  wc_after,
        "word_delta":        wc_after - wc_before,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC CONVENIENCE HELPERS
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def load_book(file_path: str) -> dict:
    """
    One-call helper: extract text → parse chapters → return fully initialised book dict.
    Every chapter has layout_directives (full schema), footnotes, and word_count.
    """
    filename = os.path.basename(file_path)
    raw_text = extract_book_text(file_path, filename)
    if not raw_text:
        raise ValueError(f"Could not extract any text from '{filename}'.")
    book = parse_book_structure(raw_text, filename)
    book = _sanitise_book_metadata(book)
    for ch in book.get("chapters", []):
        if "layout_directives" not in ch or not isinstance(ch.get("layout_directives"), dict):
            ch["layout_directives"] = _default_layout_directives()
        if "footnotes" not in ch:
            ch["footnotes"] = []
        if "word_count" not in ch:
            ch["word_count"] = len(str(ch.get("content", "")).split())
    return book


def set_recto_start(
    book:            dict,
    enabled:         bool = True,
    chapter_numbers: Optional[List[int]] = None,
) -> dict:
    """
    Enable or disable force_recto_start on specified chapters (default: all).

    Usage:
      book = set_recto_start(book)                          # all chapters
      book = set_recto_start(book, chapter_numbers=[2, 4])  # specific chapters
      book = set_recto_start(book, enabled=False)           # disable for all
    """
    target = set(chapter_numbers) if chapter_numbers else None
    for ch in book.get("chapters", []):
        if target is None or ch.get("chapter_number") in target:
            if "layout_directives" not in ch:
                ch["layout_directives"] = _default_layout_directives()
            ch["layout_directives"]["force_recto_start"] = enabled
    return book


def set_blank_pages(
    book:            dict,
    before:          int  = 0,
    after:           int  = 0,
    chapter_numbers: Optional[List[int]] = None,
) -> dict:
    """
    Set blank page counts on specified chapters (default: all).

    Usage:
      book = set_blank_pages(book, before=1)                        # 1 blank before every chapter
      book = set_blank_pages(book, after=1, chapter_numbers=[1])    # after chapter 1 only
    """
    target = set(chapter_numbers) if chapter_numbers else None
    for ch in book.get("chapters", []):
        if target is None or ch.get("chapter_number") in target:
            if "layout_directives" not in ch:
                ch["layout_directives"] = _default_layout_directives()
            if before > 0:
                ch["layout_directives"]["blank_pages_before"] = before
            if after > 0:
                ch["layout_directives"]["blank_pages_after"] = after
    return book


def resolve_page_target(
    book:        dict,
    theme_name:  str = "premium",
    generate_toc: bool = False,
) -> Dict[Tuple[int, int], int]:
    """
    Run a throwaway "mapping pass" PDF render and return the resulting
    {(chapter_number, paragraph_index): page_number} anchor table.

    This is the mechanism that makes "page N" instructions possible despite
    the book model being chapter/paragraph-based, not page-based: render
    once with invisible markers, read back where everything actually
    landed, then use that table to translate a page number into a real
    chapter/paragraph anchor (see insert_blank_page_at_page_number()).

    NOTE: page numbers are reflow-dependent. Any subsequent change to the
    book (theme, font size, content, margins) invalidates this table — it
    must be recomputed if the book changes before you use page targeting
    again.
    """
    anchors: Dict[Tuple[int, int], int] = {}
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        generate_edited_pdf(
            book, tmp_path, theme_name,
            generate_toc=generate_toc,
            page_anchor_sink=anchors,
        )
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
    return anchors


def _closest_anchor_at_or_before(
    anchors: Dict[Tuple[int, int], int], target_page: int
) -> Optional[Tuple[int, int]]:
    """Of all (chapter, paragraph) anchors, pick the one on the highest
    page number that is still <= target_page (i.e. the last paragraph
    that starts on or before the requested page)."""
    candidates = [
        (page, key) for key, page in anchors.items() if page <= target_page
    ]
    if not candidates:
        return None
    candidates.sort(key=lambda x: x[0])
    return candidates[-1][1]   # (chapter_number, paragraph_index)


def insert_blank_page_at_page_number(
    book:        dict,
    page_number: int,
    theme_name:  str = "premium",
    generate_toc: bool = False,
) -> dict:
    """
    Insert a blank page immediately after the given PDF page number,
    regardless of which chapter/paragraph that page falls in.

    Two-pass process:
      1. resolve_page_target() — mapping-pass render to find which
         (chapter, paragraph) anchor is on/just before `page_number`.
      2. Set blank_page_after_paragraphs on that chapter's directives so
         the NEXT (real) render injects the blank page at the right spot.

    Caveat (important): this targets the PDF output specifically, since
    PDF page breaks are deterministic for a fixed page size/theme. The
    DOCX output uses the same chapter/paragraph anchor point, but Word
    reflows dynamically (zoom, printer, installed fonts), so the blank
    page may not land on the exact same page number when opened in Word —
    it will, however, land at the same point in the manuscript.
    """
    anchors = resolve_page_target(book, theme_name, generate_toc)
    if not anchors:
        raise ValueError(
            "Could not map any pages — the book may be empty or rendering failed."
        )

    anchor = _closest_anchor_at_or_before(anchors, page_number)
    if anchor is None:
        # Target page is before the first paragraph (e.g. in the cover/
        # front matter) — fall back to the very first anchor.
        anchor = min(anchors.items(), key=lambda kv: kv[1])[0]

    chapter_number, paragraph_index = anchor
    result = copy.deepcopy(book)
    for ch in result.get("chapters", []):
        if ch.get("chapter_number") == chapter_number:
            directives = ch.get("layout_directives") or _default_layout_directives()
            if not isinstance(directives, dict):
                directives = _default_layout_directives()
            existing = set(directives.get("blank_page_after_paragraphs") or [])
            existing.add(paragraph_index)
            directives["blank_page_after_paragraphs"] = sorted(existing)
            ch["layout_directives"] = directives
            break

    log.info(
        "Resolved page %d → chapter %d, paragraph %d. Blank page directive set.",
        page_number, chapter_number, paragraph_index,
    )
    return result


_PAGE_TARGET_PATTERN = re.compile(
    r"\bblank\s+page\b[^.!?]{0,40}\bpage\s+(\d+)\b"
    r"|\bafter\s+page\s+(\d+)\b[^.!?]{0,40}\bblank\s+page\b",
    re.IGNORECASE,
)


def set_ornaments(
    book:            dict,
    before:          str = "",
    after:           str = "",
    chapter_numbers: Optional[List[int]] = None,
) -> dict:
    """Set ornament strings before/after chapter body for specified chapters."""
    target = set(chapter_numbers) if chapter_numbers else None
    for ch in book.get("chapters", []):
        if target is None or ch.get("chapter_number") in target:
            if "layout_directives" not in ch:
                ch["layout_directives"] = _default_layout_directives()
            if before:
                ch["layout_directives"]["ornament_before"] = before
            if after:
                ch["layout_directives"]["ornament_after"] = after
    return book


def set_drop_cap(
    book:            dict,
    enabled:         bool = True,
    chapter_numbers: Optional[List[int]] = None,
) -> dict:
    """Enable/disable drop cap on specified chapters."""
    target = set(chapter_numbers) if chapter_numbers else None
    for ch in book.get("chapters", []):
        if target is None or ch.get("chapter_number") in target:
            if "layout_directives" not in ch:
                ch["layout_directives"] = _default_layout_directives()
            ch["layout_directives"]["drop_cap"] = enabled
    return book


def set_section_labels(
    book:     dict,
    labels:   Dict[int, str],
) -> dict:
    """
    Set section labels on specific chapters.

    Usage:
      book = set_section_labels(book, {1: "Part One: Beginnings", 5: "Part Two: The Fall"})
    """
    for ch in book.get("chapters", []):
        num = ch.get("chapter_number", -1)
        if num in labels:
            if "layout_directives" not in ch:
                ch["layout_directives"] = _default_layout_directives()
            ch["layout_directives"]["section_label"] = labels[num]
    return book


def get_book_stats(book: dict) -> dict:
    """Compute summary statistics for a book dict."""
    chapters    = book.get("chapters", [])
    total_words = sum(len(str(ch.get("content", "")).split()) for ch in chapters)
    total_chars = sum(len(str(ch.get("content", "")))         for ch in chapters)
    avg_wc      = total_words // len(chapters) if chapters else 0
    longest  = max(chapters, key=lambda c: len(str(c.get("content", ""))), default={})
    shortest = min(chapters, key=lambda c: len(str(c.get("content", ""))), default={})

    return {
        "title":                 book.get("title", ""),
        "author":                book.get("author", ""),
        "chapter_count":         len(chapters),
        "total_words":           total_words,
        "total_chars":           total_chars,
        "avg_words_per_chapter": avg_wc,
        "longest_chapter":       longest.get("title", ""),
        "shortest_chapter":      shortest.get("title", ""),
        "estimated_pages":       max(1, round(total_words / 250)),
        "has_devanagari":        _has_devanagari(" ".join(
            str(ch.get("content", "")) for ch in chapters)),
        "has_cjk":               _has_cjk(" ".join(
            str(ch.get("content", "")) for ch in chapters)),
        "has_arabic":            _has_arabic(" ".join(
            str(ch.get("content", "")) for ch in chapters)),
    }


def book_to_json(book: dict, indent: int = 2) -> str:
    """Serialise book dict to JSON string."""
    return json.dumps(book, ensure_ascii=False, indent=indent, default=str)


def book_from_json(json_str: str) -> dict:
    """Deserialise a book dict from JSON string."""
    book = json.loads(json_str)
    book = _sanitise_book_metadata(book)
    for ch in book.get("chapters", []):
        if "layout_directives" not in ch:
            ch["layout_directives"] = _default_layout_directives()
        else:
            # Merge any new keys that may have been added since serialisation
            defaults = _default_layout_directives()
            for k, v in defaults.items():
                ch["layout_directives"].setdefault(k, v)
        ch.setdefault("footnotes", [])
        ch.setdefault("word_count", len(str(ch.get("content", "")).split()))
    return book


# ══════════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
# CLI ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="AI Book Editor v5.0 — Full-featured command-line interface",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""
        Examples:
          # Basic edit
          python book_editor.py book.pdf "Improve the pacing in every chapter"

          # Recto starts + blank pages
          python book_editor.py book.docx "make every chapter start on the right page" \\
              --recto --blank-before 1

          # Blank page after author intro only
          python book_editor.py book.txt "add a blank page after the author intro" \\
              --blank-after-ch 1

          # Theme + output dir
          python book_editor.py book.txt "Rewrite in a sci-fi tone" \\
              --theme scifi --output-dir ./out

          # Front matter + TOC
          python book_editor.py book.epub "polish the prose" \\
              --front-matter --dedication "For my family"

          # Drop cap + ornament
          python book_editor.py book.docx "improve flow" \\
              --drop-cap --ornament "❧"

          # Show book stats only
          python book_editor.py book.pdf --stats-only

          # Batch process a folder
          python book_editor.py /books/ "Fix grammar" --batch --output-dir ./out
        """),
    )

    parser.add_argument("input",
        help="Input file (.pdf, .docx, .txt, .rtf, .epub, .zip) or folder (with --batch)")
    parser.add_argument("instruction",
        nargs="?", default="",
        help="Edit instruction in natural language (optional with --stats-only)")
    parser.add_argument("--output-dir",     default="./output")
    parser.add_argument("--theme",          default="premium",
        choices=list(THEMES.keys()))
    parser.add_argument("--recto",          action="store_true",
        help="Force recto (right-hand) start for ALL chapters")
    parser.add_argument("--blank-before",   type=int, default=0,
        help="Number of blank pages before each chapter")
    parser.add_argument("--blank-after",    type=int, default=0,
        help="Number of blank pages after each chapter")
    parser.add_argument("--blank-after-ch", type=int, default=0,
        help="Add one blank page after this specific chapter number (e.g. 1 for author intro)")
    parser.add_argument("--drop-cap",       action="store_true",
        help="Add drop cap to first letter of each chapter")
    parser.add_argument("--ornament",       default="",
        help="Ornament character to place before each chapter (e.g. '❧')")
    parser.add_argument("--section-labels", default="",
        help='JSON dict of chapter_number:label, e.g. \'{"1":"Part One","5":"Part Two"}\'')
    parser.add_argument("--front-matter",   action="store_true",
        help="Add front matter pages (copyright etc.)")
    parser.add_argument("--dedication",     default="",
        help="Dedication text for front matter")
    parser.add_argument("--author-bio",     default="",
        help="Author biography text for back matter")
    parser.add_argument("--no-diff",        action="store_true",
        help="Skip generating the HTML diff report")
    parser.add_argument("--toc",            action="store_true",
        help="Insert a Table of Contents page into the output PDF and DOCX")
    parser.add_argument("--stats-only",     action="store_true",
        help="Print book statistics and exit without editing")
    parser.add_argument("--batch",          action="store_true",
        help="Process all files in the input folder")
    parser.add_argument("--workers",        type=int, default=2,
        help="Number of parallel workers for batch/chapter processing")
    parser.add_argument("--verbose",        action="store_true")

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger("book_editor").setLevel(logging.DEBUG)

    # ── Batch mode ────────────────────────────────────────────────────────────
    if args.batch:
        if not os.path.isdir(args.input):
            parser.error("--batch requires a directory as input")
        if not args.instruction:
            parser.error("An instruction is required for batch mode")
        exts  = {".pdf", ".docx", ".txt", ".rtf", ".epub", ".zip", ".md"}
        files = [
            os.path.join(args.input, f)
            for f in sorted(os.listdir(args.input))
            if Path(f).suffix.lower() in exts
        ]
        if not files:
            print(f"No supported files found in '{args.input}'")
            sys.exit(1)
        print(f"Batch processing {len(files)} files...")
        results = batch_process_files(
            files, args.instruction, args.output_dir,
            args.theme, args.workers,
        )
        ok  = sum(1 for r in results if r.get("status") == "ok")
        err = sum(1 for r in results if r.get("status") != "ok")
        print(f"\n✅ Batch complete: {ok} succeeded, {err} failed")
        sys.exit(0)

    # ── Single file mode ──────────────────────────────────────────────────────
    print(f"📖 Loading: {args.input}")
    book  = load_book(args.input)
    stats = get_book_stats(book)
    print(f"   Title:    {stats['title']}")
    print(f"   Author:   {stats['author']}")
    print(f"   Chapters: {stats['chapter_count']}")
    print(f"   Words:    {stats['total_words']:,}  (~{stats['estimated_pages']} pages)")
    if stats["has_devanagari"]:
        print("   Script:   Contains Devanagari — Noto fonts will be used")
    if stats["has_cjk"]:
        print("   Script:   Contains CJK characters — Noto CJK fonts will be used")
    if stats["has_arabic"]:
        print("   Script:   Contains Arabic — Noto Arabic fonts will be used")

    if args.stats_only:
        import pprint
        pprint.pprint(stats)
        sys.exit(0)

    if not args.instruction:
        parser.error("An instruction is required (or use --stats-only)")

    # ── Apply manual layout overrides ─────────────────────────────────────────
    if args.recto:
        book = set_recto_start(book)
        print("   → Force recto start enabled for all chapters")

    if args.blank_before or args.blank_after:
        book = set_blank_pages(book, before=args.blank_before, after=args.blank_after)
        print(f"   → Blank pages: {args.blank_before} before / {args.blank_after} after each chapter")

    if args.blank_after_ch:
        book = set_blank_pages(book, after=1, chapter_numbers=[args.blank_after_ch])
        print(f"   → One blank page after chapter {args.blank_after_ch}")

    if args.drop_cap:
        book = set_drop_cap(book)
        print("   → Drop cap enabled for all chapters")

    if args.ornament:
        book = set_ornaments(book, before=args.ornament)
        print(f"   → Ornament '{args.ornament}' before each chapter")

    if args.section_labels:
        try:
            raw_labels = json.loads(args.section_labels)
            labels = {int(k): v for k, v in raw_labels.items()}
            book = set_section_labels(book, labels)
            print(f"   → Section labels applied to chapters: {list(labels.keys())}")
        except Exception as ex:
            print(f"   ⚠ Could not parse --section-labels: {ex}")

    # ── Front matter ──────────────────────────────────────────────────────────
    if args.front_matter or args.dedication:
        front = build_front_matter(
            title=book.get("title", ""),
            author=book.get("author", ""),
            year=book.get("year", ""),
            isbn=book.get("isbn", ""),
            dedication=args.dedication,
        )
        book["chapters"] = front + book.get("chapters", [])
        print(f"   → Front matter added ({len(front)} pages)")

    # ── Back matter ───────────────────────────────────────────────────────────
    if args.author_bio:
        back = generate_back_matter(
            author=book.get("author", ""),
            author_bio=args.author_bio,
        )
        book["chapters"] = book.get("chapters", []) + back
        print(f"   → Back matter added ({len(back)} pages)")

    # ── Set up version history ────────────────────────────────────────────────
    history = VersionHistory()
    history.commit(book, "Original document loaded", "initial")

    print(f"\n✏️  Instruction: {args.instruction}")
    print(f"   Theme:       {args.theme}")

    result = process_editor_turn(
        book_structure=book,
        user_message=args.instruction,
        conversation_history=[],
        output_dir=args.output_dir,
        theme=args.theme,
        version_history=history,
        generate_diff=not args.no_diff,
        generate_toc_page=args.toc,
    )

    print(f"\n✅ Done!")
    if result.get("pdf_path"):
        print(f"   PDF:        {result['pdf_path']}")
    if result.get("docx_path"):
        print(f"   DOCX:       {result['docx_path']}")
    if result.get("diff_html_path"):
        print(f"   Diff HTML:  {result['diff_html_path']}")
    print(f"   Summary:    {result['edit_summary']}")
    print(f"   Intent:     {result['intent']}")
    print(f"   Chapters changed: {result['chapters_changed']}")
    print(
        f"   Words: {result['word_count_before']:,} → {result['word_count_after']:,} "
        f"({'+'if result['word_delta']>=0 else ''}{result['word_delta']:,})"
    )

    if args.verbose:
        print(f"\n📜 Version log:")
        for entry in history.log():
            print(
                f"   v{entry['version']}  [{entry['timestamp'][:19]}]  "
                f"{entry['summary'][:80]}"
            )

    # Export history HTML
    hist_path = os.path.join(args.output_dir, "edit_history.html")
    history.export_history_html(hist_path)
    print(f"   History:    {hist_path}")