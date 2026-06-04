"""
docx_generator.py — Multilingual DOCX output using python-docx.

Font strategy
─────────────
Word/LibreOffice handle Unicode natively — they perform font substitution
automatically at render time.  However, we must:
  1. Set the font name to one that actually ships with the target OS OR
     embed fonts.  The safest cross-platform choice is "Noto Sans" for
     non-Latin scripts (works on Linux/Mac/Win if Noto is installed) and
     "Calibri" for Latin text.
  2. Set the <w:cs> (complex script) font alongside <w:ascii> so Word
     activates the correct shaping engine for Arabic, Hebrew, Devanagari, etc.
  3. Mark paragraphs that are right-to-left with the <w:bidi> property so
     Word renders them RTL.
  4. Ensure the document charset declaration is UTF-8 (python-docx does this
     by default, but we make it explicit).

ENTERPRISE UPGRADE:
- Interactive Clickable TOCs: Injects a native Word `w:fldChar` TOC Field allowing 
  users to Ctrl+Click navigate and right-click update dynamically.
"""

# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx.shared import Pt, RGBColor, Inches, Cm
# pyrefly: ignore [missing-import]
from docx.enum.text import WD_ALIGN_PARAGRAPH
# pyrefly: ignore [missing-import]
from docx.oxml.ns import qn
# pyrefly: ignore [missing-import]
from docx.oxml import OxmlElement
import datetime
import os
import unicodedata

# ── Script detection ───────────────────────────────────────────────────────────

_SCRIPT_RANGES = {
    "Arabic":     (0x0600, 0x06FF),
    "Hebrew":     (0x0590, 0x05FF),
    "Devanagari": (0x0900, 0x097F),
    "Bengali":    (0x0980, 0x09FF),
    "Tamil":      (0x0B80, 0x0BFF),
    "Telugu":     (0x0C00, 0x0C7F),
    "Kannada":    (0x0C80, 0x0CFF),
    "Malayalam":  (0x0D00, 0x0D7F),
    "Thai":       (0x0E00, 0x0E7F),
    "Hiragana":   (0x3040, 0x309F),
    "Katakana":   (0x30A0, 0x30FF),
    "CJK":        (0x4E00, 0x9FFF),
    "Hangul":     (0xAC00, 0xD7AF),
    "Cyrillic":   (0x0400, 0x04FF),
    "Greek":      (0x0370, 0x03FF),
}

_RTL_SCRIPTS = {"Arabic", "Hebrew"}

_SCRIPT_FONT = {
    "Arabic":     "Noto Sans Arabic",
    "Hebrew":     "Noto Sans Hebrew",
    "Devanagari": "Noto Sans Devanagari",
    "Bengali":    "Noto Sans Bengali",
    "Tamil":      "Noto Sans Tamil",
    "Telugu":     "Noto Sans Telugu",
    "Kannada":    "Noto Sans Kannada",
    "Malayalam":  "Noto Sans Malayalam",
    "Thai":       "Noto Sans Thai",
    "CJK":        "Noto Sans SC",
    "Hiragana":   "Noto Sans JP",
    "Katakana":   "Noto Sans JP",
    "Hangul":     "Noto Sans KR",
    "Cyrillic":   "Calibri",
    "Greek":      "Calibri",
    "Latin":      "Calibri",
}

def _detect_script(text: str) -> str:
    counts = {k: 0 for k in _SCRIPT_RANGES}
    for ch in text:
        cp = ord(ch)
        for name, (lo, hi) in _SCRIPT_RANGES.items():
            if lo <= cp <= hi:
                counts[name] += 1
                break
    best = max(counts, key=counts.get)
    return best if counts[best] > 0 else "Latin"

def _font_for(text: str) -> str:
    return _SCRIPT_FONT.get(_detect_script(text), "Calibri")

def _is_rtl(text: str) -> bool:
    return _detect_script(text) in _RTL_SCRIPTS

def _para_align(text: str):
    return WD_ALIGN_PARAGRAPH.RIGHT if _is_rtl(text) else WD_ALIGN_PARAGRAPH.JUSTIFY

# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_run_font(run, font_name: str):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.insert(0, el)
        el.set(qn("w:ascii"),    font_name)
        el.set(qn("w:hAnsi"),   font_name)
        el.set(qn("w:eastAsia"), font_name)
        el.set(qn("w:cs"),      font_name)

def _set_para_bidi(paragraph, rtl: bool):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if rtl:
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
        bidi.set(qn("w:val"), "1")
    else:
        if bidi is not None:
            pPr.remove(bidi)

def _add_run(paragraph, text: str, size_pt: float,
             bold: bool = False, italic: bool = False,
             color: RGBColor | None = None) -> None:
    font_name = _font_for(text)
    run = paragraph.add_run(text)
    _set_run_font(run, font_name)
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    _set_para_bidi(paragraph, _is_rtl(text))
    paragraph.alignment = _para_align(text)

def set_heading_color(paragraph, hex_color: str = "4F46E5"):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
        _set_run_font(run, _font_for(run.text))

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_native_toc(paragraph):
    """Injects a native Word Table of Contents field (Interactive/Clickable)."""
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # \o "1-3" limits heading levels, \h creates hyperlinks
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# ── Main generator ─────────────────────────────────────────────────────────────

def generate_docx(book_title: str, segments, output_dir: str = "output") -> str:
    os.makedirs(output_dir, exist_ok=True)
    safe_title = "".join(
        c for c in book_title if c.isalnum() or c in (" ", "-", "_")
    ).strip()
    filepath = os.path.join(output_dir, f"{safe_title}.docx")

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    settings = doc.settings.element
    compat = OxmlElement("w:compat")
    cs_el  = OxmlElement("w:compatSetting")
    cs_el.set(qn("w:name"),  "compatibilityMode")
    cs_el.set(qn("w:uri"),   "http://schemas.microsoft.com/office/word")
    cs_el.set(qn("w:val"),   "15")
    compat.append(cs_el)
    settings.append(compat)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = _font_for(book_title)
    style_normal.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title_para, book_title, 28, bold=True,
             color=RGBColor(0x1E, 0x29, 0x3B))

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(date_para,
             f"Generated on {datetime.date.today().strftime('%B %d, %Y')}",
             12, color=RGBColor(0x64, 0x74, 0x8B))

    doc.add_paragraph()
    agent_para = doc.add_paragraph()
    agent_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(agent_para, "Generated by AI Book Writing Agent", 10,
             italic=True, color=RGBColor(0x94, 0xA3, 0xB8))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    toc_heading = doc.add_heading("Table of Contents", level=1)
    set_heading_color(toc_heading, "1E293B")
    
    # Inject Native Interactive TOC
    toc_para = doc.add_paragraph()
    add_native_toc(toc_para)

    chapters: dict = {}
    for seg in segments:
        key = seg.chapter_number
        if key not in chapters:
            chapters[key] = {"title": seg.chapter_title, "subs": []}
        chapters[key]["subs"].append(seg.subheading)

    # Keeping static fallback list below native TOC so visual structure remains intact
    for ch_num in sorted(chapters.keys()):
        ch   = chapters[ch_num]
        para = doc.add_paragraph(style="List Bullet")
        _add_run(para, f"Chapter {ch_num}  —  {ch['title']}",
                 11, bold=True, color=RGBColor(0x4F, 0x46, 0xE5))
        for sub in ch["subs"]:
            sub_para = doc.add_paragraph(style="List Bullet 2")
            _add_run(sub_para, sub, 10, color=RGBColor(0x64, 0x74, 0x8B))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTERS
    # ══════════════════════════════════════════════════════════════════════════
    current_chapter = None

    for seg in segments:
        if seg.chapter_number != current_chapter:
            if current_chapter is not None:
                doc.add_page_break()
            current_chapter = seg.chapter_number

            label_para = doc.add_paragraph()
            _add_run(label_para, f"Chapter {seg.chapter_number}", 11,
                     bold=True, color=RGBColor(0x4F, 0x46, 0xE5))

            ch_heading = doc.add_heading(seg.chapter_title, level=1)
            set_heading_color(ch_heading, "1E293B")
            _set_para_bidi(ch_heading, _is_rtl(seg.chapter_title))
            doc.add_paragraph()

        sub_heading = doc.add_heading(seg.subheading, level=2)
        set_heading_color(sub_heading, "334155")
        _set_para_bidi(sub_heading, _is_rtl(seg.subheading))

        for para_text in seg.content.split("\n"):
            para_text = unicodedata.normalize("NFC", para_text.strip())
            if not para_text:
                continue
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.paragraph_format.space_after  = Pt(6)
            para.paragraph_format.line_spacing = Pt(16)
            _add_run(para, para_text, 11)

    doc.save(filepath)
    print(f"  ✅ DOCX saved: {filepath}")
    return filepath