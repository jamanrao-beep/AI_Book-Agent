"""
layout_designer.py  ·  v2.5 (Strict Dimensions, TOC Bypass & Intro Capture)
AI-powered internal book layout designer — with full book-type awareness
and proper Unicode (Devanagari, Hindi, multi-script) support.

Pipeline:
  1. Extract raw text from PDF / DOCX / ZIP
  2. Detect chapter boundaries (supports Hindi/Devanagari अध्याय headings)
  3. Build book-type defaults (novel, academic, religious, poetry, children, business)
  4. Ask GPT-4o to produce a complete typographic layout concept (JSON),
     seeded with type-aware defaults and any user overrides
  5. Apply hard user overrides on top of the AI concept  (user always wins)
  6. Render PDF  (ReportLab + registered Unicode/Noto fonts)
  7. Render DOCX (python-docx)
  8. Return paths + metadata to the caller
"""

from __future__ import annotations

import io
import json
import math
import os
import re
import shutil
import threading
import unicodedata
import urllib.request
import uuid
import zipfile
import traceback
from pathlib import Path
from typing import Callable, Optional

# pyrefly: ignore [missing-import]
from dotenv import load_dotenv
# pyrefly: ignore [missing-import]
from openai import OpenAI

load_dotenv()
_api_key = os.getenv("OPENAI_API_KEY")
if not _api_key:
    raise RuntimeError(
        "OPENAI_API_KEY is not set. Add it to your .env file or environment variables."
    )
client = OpenAI(api_key=_api_key)
MODEL = "gpt-4o"


# ─────────────────────────────────────────────────────────────────────────────
# Unicode font system
# Strategy (in priority order):
#   1. Local ./fonts/ directory bundled with the app
#   2. Common system paths (Ubuntu/Debian/Alpine)
#   3. Auto-download from Google Fonts CDN at first run → cached in ./fonts/
#
# This guarantees Hindi/Devanagari PDFs always work regardless of what is
# installed on the server.
# ─────────────────────────────────────────────────────────────────────────────

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_FONTS_DIR   = os.path.join(_SCRIPT_DIR, "fonts")   # local cache / bundled dir

# Google Fonts static CDN URLs for the two Noto fonts we need.
# These are stable direct-download links (not the CSS API endpoint).
_FONT_URLS: dict[str, str] = {
    "NotoSerifDevanagari-Regular.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSerifDevanagari/NotoSerifDevanagari-Regular.ttf"
    ),
    "NotoSerifDevanagari-Bold.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSerifDevanagari/NotoSerifDevanagari-Bold.ttf"
    ),
    "NotoSansDevanagari-Regular.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSansDevanagari/NotoSansDevanagari-Regular.ttf"
    ),
    "NotoSansDevanagari-Bold.ttf": (
        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/"
        "NotoSansDevanagari/NotoSansDevanagari-Bold.ttf"
    ),
}

# ReportLab name → filename
_NOTO_FONT_FILES: dict[str, str] = {
    "NotoSerifDevanagari":      "NotoSerifDevanagari-Regular.ttf",
    "NotoSerifDevanagari-Bold": "NotoSerifDevanagari-Bold.ttf",
    "NotoSansDevanagari":       "NotoSansDevanagari-Regular.ttf",
    "NotoSansDevanagari-Bold":  "NotoSansDevanagari-Bold.ttf",
}

_SYSTEM_FONT_DIRS = [
    "/usr/share/fonts/truetype/noto",   # Ubuntu/Debian: apt install fonts-noto-core
    "/usr/share/fonts/noto",
    "/usr/share/fonts/opentype/noto",
    "/usr/share/fonts/truetype/freefont",
    "/usr/share/fonts/freefont",
    "/usr/share/fonts",
    # Alpine / Docker slim images
    "/usr/share/fonts/noto-cjk",
    # macOS (local dev)
    "/Library/Fonts",
    "/System/Library/Fonts",
    # Windows (local dev)
    "C:/Windows/Fonts",
]

_REGISTERED_FONTS: set[str] = set()
_FONTS_REGISTERED  = False
_FONT_LOCK         = threading.Lock()


def _find_font_on_system(filename: str) -> Optional[str]:
    """Search known system font directories for a TTF file.
    Falls back to a recursive os.walk search under /usr/share/fonts."""
    # Check local ./fonts/ dir first
    local = os.path.join(_FONTS_DIR, filename)
    if os.path.isfile(local):
        return local
    for d in _SYSTEM_FONT_DIRS:
        p = os.path.join(d, filename)
        if os.path.isfile(p):
            return p
    # Last resort: recursive search under /usr/share/fonts (handles distro variations)
    base_search = "/usr/share/fonts"
    if os.path.isdir(base_search):
        for root, _dirs, files in os.walk(base_search):
            if filename in files:
                found = os.path.join(root, filename)
                print(f"  🔍  Found font via recursive search: {found}")
                return found
    return None


def _download_font(filename: str) -> Optional[str]:
    """
    Download a font file from GitHub/Google Fonts into _FONTS_DIR.
    Returns the local path on success, None on failure.
    """
    url = _FONT_URLS.get(filename)
    if not url:
        return None
    dest = os.path.join(_FONTS_DIR, filename)
    if os.path.isfile(dest):          # already downloaded in a previous run
        return dest
    try:
        os.makedirs(_FONTS_DIR, exist_ok=True)
        print(f"  ⬇️   Downloading font {filename} …")
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp, \
             open(dest, "wb") as f:
            f.write(resp.read())
        print(f"  ✅  Downloaded {filename} → {dest}")
        return dest
    except Exception as e:
        print(f"  ⚠️   Could not download {filename}: {e}")
        # Clean up partial file
        if os.path.isfile(dest):
            try:
                os.remove(dest)
            except Exception:
                pass
        return None


def _ensure_unicode_fonts() -> None:
    """
    Register Noto Devanagari TTFs with ReportLab (idempotent, thread-safe).
    Resolution order: local cache → system paths → auto-download from GitHub.
    """
    global _FONTS_REGISTERED, _REGISTERED_FONTS
    if _FONTS_REGISTERED:
        return
    with _FONT_LOCK:
        if _FONTS_REGISTERED:   # double-checked inside lock
            return
        try:
            from reportlab.pdfbase import pdfmetrics      # pyrefly: ignore [missing-import]
            from reportlab.pdfbase.ttfonts import TTFont  # pyrefly: ignore [missing-import]

            for rl_name, filename in _NOTO_FONT_FILES.items():
                # 1. Try system / local cache
                path = _find_font_on_system(filename)
                # 2. Auto-download if not found
                if not path:
                    path = _download_font(filename)
                if not path:
                    print(f"  ⚠️   Font unavailable: {rl_name} ({filename}) — skipping")
                    continue
                try:
                    pdfmetrics.registerFont(TTFont(rl_name, path))
                    _REGISTERED_FONTS.add(rl_name)
                    print(f"  ✅  Registered: {rl_name} from {path}")
                except Exception as e:
                    print(f"  ⚠️   registerFont failed for {rl_name}: {e}")

            if _REGISTERED_FONTS:
                print(f"  ✅  Unicode fonts ready: {sorted(_REGISTERED_FONTS)}")
            else:
                raise RuntimeError(
                    "CRITICAL: No Unicode/Devanagari fonts could be registered. "
                    "Fix: Run `apt-get install -y fonts-noto-core` on your server, "
                    "or place NotoSerifDevanagari-Regular.ttf (and Bold/Sans variants) "
                    "in a ./fonts/ folder next to layout_designer.py. "
                    "Without these fonts, Hindi text renders as square boxes."
                )
        except Exception as e:
            print(f"  ⚠️   _ensure_unicode_fonts failed: {e}\n{traceback.format_exc()}")
        finally:
            _FONTS_REGISTERED = True


def _has_non_latin(text: str) -> bool:
    """True if text contains Devanagari or other Indic scripts (U+0900+)."""
    return any(ord(c) >= 0x0900 for c in text
               if not unicodedata.category(c).startswith("Z"))


def _clean_extracted_text(text: str) -> str:
    """
    Remove null bytes, private-use Unicode characters, and other garbage that
    PDF extractors emit when a font uses a custom encoding map.
    Also normalises non-breaking spaces and zero-width characters.
    Returns the cleaned string.
    """
    # Strip null bytes and C0/C1 control characters (except newline/tab)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    # Strip Unicode private-use area blocks (U+E000–U+F8FF, U+F0000–U+FFFFF)
    text = re.sub(r"[\ue000-\uf8ff]", "", text)
    # Normalise non-breaking spaces and zero-width joiners/non-joiners to regular space
    text = text.replace("\u00a0", " ").replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
    # Collapse runs of whitespace-only lines (3+ blank lines → 2)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _text_looks_corrupt(text: str) -> bool:
    """
    Returns True when the extracted text is mostly garbage (null bytes,
    private-use codepoints, or extremely low printable-character density).
    This is the signal to discard PDF extraction and try another source.
    """
    if not text:
        return True
    sample = text[:2000]
    null_count = sample.count("\x00")
    pua_count  = sum(1 for c in sample if "\ue000" <= c <= "\uf8ff")
    printable  = sum(1 for c in sample if c.isprintable() and c not in "\x00")
    total      = max(len(sample), 1)
    # Corrupt if >5 % null/PUA bytes OR printable ratio below 60 %
    return (null_count + pua_count) / total > 0.05 or printable / total < 0.60


def _unicode_body_font(rl_name: str, has_unicode: bool) -> str:
    """
    Return the best available Unicode-capable font name for the requested style.
    Falls back through the registered set; never returns a Latin-only font name
    when Unicode content is present (unless nothing at all was registered).
    """
    if not has_unicode:
        return rl_name

    _PREF: dict[str, list[str]] = {
        "Times-Roman":       ["NotoSerifDevanagari", "NotoSansDevanagari"],
        "Times-Italic":      ["NotoSerifDevanagari", "NotoSansDevanagari"],
        "Helvetica":         ["NotoSansDevanagari",  "NotoSerifDevanagari"],
        "Helvetica-Oblique": ["NotoSansDevanagari",  "NotoSerifDevanagari"],
        "Courier":           ["NotoSansDevanagari",  "NotoSerifDevanagari"],
    }
    for candidate in _PREF.get(rl_name, ["NotoSerifDevanagari", "NotoSansDevanagari"]):
        if candidate in _REGISTERED_FONTS:
            return candidate

    print(f"  🚨  No Unicode font for '{rl_name}' — text may render as boxes.")
    return rl_name


# ─────────────────────────────────────────────────────────────────────────────
# Book-type default profiles
# ─────────────────────────────────────────────────────────────────────────────

BOOK_TYPE_PROFILES: dict[str, dict] = {
    "novel": {
        "_label": "Novel / Literary Fiction",
        "_description": (
            "A literary novel demands an intimate, reader-friendly interior. "
            "Use a warm cream page with a classic serif body font (Times-Roman), "
            "generous side margins, 1.5× line spacing, a chapter-opening drop cap, "
            "and a subtle ornamental divider.  Page numbers centred at the bottom.  "
            "Chapter headings should feel understated and elegant."
        ),
        "page_bg":             "#fffdf6",
        "text_color":          "#1c1a17",
        "chapter_title_color": "#2d2416",
        "accent_color":        "#8b6914",
        "body_font":           "Times-Roman",
        "body_font_size":      11.5,
        "line_spacing":        1.55,
        "first_para_indent_mm": 6,
        "margin_top_mm":       22,
        "margin_bottom_mm":    22,
        "margin_left_mm":      25,
        "margin_right_mm":     22,
        "chapter_font":        "Times-Roman",
        "chapter_font_size":   24,
        "chapter_prefix":      "Chapter",
        "show_drop_cap":       True,
        "ornament":            "—◆—",
        "header_text":         "",
        "show_page_numbers":   True,
    },

    "academic": {
        "_label": "Academic / Educational",
        "_description": (
            "An academic or educational book needs clear visual hierarchy and maximum "
            "readability.  White page background, sans-serif body font (Helvetica), "
            "structured numbered chapter headings, no drop cap, tight but comfortable "
            "1.4× spacing.  Running header with the book title; page numbers in the "
            "footer.  Accent colour should be a professional blue or teal.  "
            "No decorative ornaments."
        ),
        "page_bg":             "#ffffff",
        "text_color":          "#111827",
        "chapter_title_color": "#1e3a5f",
        "accent_color":        "#2563eb",
        "body_font":           "Helvetica",
        "body_font_size":      11,
        "line_spacing":        1.4,
        "first_para_indent_mm": 0,
        "margin_top_mm":       25,
        "margin_bottom_mm":    25,
        "margin_left_mm":      28,
        "margin_right_mm":     25,
        "chapter_font":        "Helvetica",
        "chapter_font_size":   20,
        "chapter_prefix":      "Chapter",
        "show_drop_cap":       False,
        "ornament":            "",
        "header_text":         "",
        "show_page_numbers":   True,
    },

    "religious": {
        "_label": "Religious / Spiritual",
        "_description": (
            "A religious or spiritual text calls for a reverent, traditional feel.  "
            "Warm ivory page, classic serif body (Times-Roman), generous margins, "
            "gold or saffron accent colour, ornate chapter-opening ornaments, "
            "decorative drop cap on every chapter.  Chapter prefix may be omitted "
            "or replaced with a verse reference.  Centred chapter titles with "
            "a double-rule accent below.  Comfortable 1.6× leading."
        ),
        "page_bg":             "#fef9f0",
        "text_color":          "#2d1f0a",
        "chapter_title_color": "#7c3d0a",
        "accent_color":        "#c8830a",
        "body_font":           "Times-Roman",
        "body_font_size":      11.5,
        "line_spacing":        1.6,
        "first_para_indent_mm": 5,
        "margin_top_mm":       24,
        "margin_bottom_mm":    24,
        "margin_left_mm":      28,
        "margin_right_mm":     28,
        "chapter_font":        "Times-Roman",
        "chapter_font_size":   22,
        "chapter_prefix":      "",
        "show_drop_cap":       True,
        "ornament":            "✦  ✦  ✦",
        "header_text":         "",
        "show_page_numbers":   True,
    },

    "poetry": {
        "_label": "Poetry / Shayari",
        "_description": (
            "A poetry collection must preserve the poet's line breaks and white space.  "
            "Soft off-white page, elegant italic serif (Times-Italic) body, "
            "very generous left and right margins to frame each poem, "
            "1.8× line spacing for breathing room.  No first-line indent — poetry "
            "is left-aligned.  Minimal ornamentation; a thin floral or asterism "
            "ornament between poems works well.  Chapter (poem) titles should be "
            "small-caps-style in a complementary serif."
        ),
        "page_bg":             "#fdfaf5",
        "text_color":          "#1e1523",
        "chapter_title_color": "#6b2d8b",
        "accent_color":        "#b45fc0",
        "body_font":           "Times-Italic",
        "body_font_size":      12,
        "line_spacing":        1.8,
        "first_para_indent_mm": 0,
        "margin_top_mm":       30,
        "margin_bottom_mm":    30,
        "margin_left_mm":      35,
        "margin_right_mm":     35,
        "chapter_font":        "Times-Roman",
        "chapter_font_size":   18,
        "chapter_prefix":      "",
        "show_drop_cap":       False,
        "ornament":            "❧",
        "header_text":         "",
        "show_page_numbers":   True,
    },

    "children": {
        "_label": "Children's Book",
        "_description": (
            "A children's book needs large, clear type and lots of white space for "
            "illustrations.  Pure white background, large sans-serif body font "
            "(Helvetica, 14pt+), very wide margins to leave room for artwork, "
            "double-spaced (2.0×) text, friendly short chapter titles in a bold "
            "round-looking font, bright cheerful accent colour (coral, teal, or "
            "sunshine yellow).  No drop cap — just friendly text.  Centred page "
            "numbers at the bottom.  No running header."
        ),
        "page_bg":             "#ffffff",
        "text_color":          "#1a1a2e",
        "chapter_title_color": "#e05a2b",
        "accent_color":        "#f4a535",
        "body_font":           "Helvetica",
        "body_font_size":      14,
        "line_spacing":        2.0,
        "first_para_indent_mm": 0,
        "margin_top_mm":       30,
        "margin_bottom_mm":    30,
        "margin_left_mm":      32,
        "margin_right_mm":     32,
        "chapter_font":        "Helvetica",
        "chapter_font_size":   22,
        "chapter_prefix":      "",
        "show_drop_cap":       False,
        "ornament":            "★",
        "header_text":         "",
        "show_page_numbers":   True,
    },

    "biography": {
        "_label": "Biography / Memoir",
        "_description": (
            "A biography or memoir needs an editorial, narrative feel that reads like "
            "quality journalism.  Off-white page, classic serif body (Times-Roman), "
            "generous margins, 1.5× line spacing, subtle drop caps on chapter opens.  "
            "Chapter headings are left-aligned and understated.  Running header carries "
            "the book title.  Page numbers at the bottom centre or bottom-right.  "
            "An optional thin rule under the chapter title gives a clean editorial look.  "
            "No heavy ornaments — dignity and restraint are key."
        ),
        "page_bg":             "#fdfcfa",
        "text_color":          "#1a1a1a",
        "chapter_title_color": "#1c2b3a",
        "accent_color":        "#5c6f7e",
        "body_font":           "Times-Roman",
        "body_font_size":      11.5,
        "line_spacing":        1.5,
        "first_para_indent_mm": 5,
        "margin_top_mm":       24,
        "margin_bottom_mm":    24,
        "margin_left_mm":      26,
        "margin_right_mm":     24,
        "chapter_font":        "Times-Roman",
        "chapter_font_size":   22,
        "chapter_prefix":      "Chapter",
        "show_drop_cap":       True,
        "ornament":            "—",
        "header_text":         "",
        "show_page_numbers":   True,
    },

    "business": {
        "_label": "Business / Self-help",
        "_description": (
            "A business or self-help book should feel modern, authoritative, and "
            "easy to skim.  Crisp white page, clean Helvetica body font (11pt), "
            "1.45× spacing, bold sans-serif chapter headings, a strong accent colour "
            "(deep navy, electric blue, or confident purple), no drop cap, "
            "a thin top accent rule under each chapter title.  "
            "Running header on even pages; page numbers bottom-right.  "
            "Tight margins for a modern 'trade paperback' feel."
        ),
        "page_bg":             "#ffffff",
        "text_color":          "#0f172a",
        "chapter_title_color": "#1e3a5f",
        "accent_color":        "#4f46e5",
        "body_font":           "Helvetica",
        "body_font_size":      11,
        "line_spacing":        1.45,
        "first_para_indent_mm": 0,
        "margin_top_mm":       22,
        "margin_bottom_mm":    22,
        "margin_left_mm":      24,
        "margin_right_mm":     22,
        "chapter_font":        "Helvetica",
        "chapter_font_size":   26,
        "chapter_prefix":      "Chapter",
        "show_drop_cap":       False,
        "ornament":            "—",
        "header_text":         "",
        "show_page_numbers":   True,
    },
}


def get_book_type_profile(book_type: Optional[str]) -> Optional[dict]:
    """Return the profile dict for a known book type, or None if unrecognised."""
    if not book_type:
        return None
    return BOOK_TYPE_PROFILES.get(book_type.lower().strip())


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _safe_title(raw: str, fallback: str = "book") -> str:
    cleaned = "".join(
        c for c in raw
        if unicodedata.category(c) not in ("Cc", "Cs") and c not in r'\/:*?"<>|'
    ).strip()
    return cleaned[:120] or fallback


def _hex_to_rgb(h: str) -> tuple[float, float, float]:
    h = h.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return tuple(int(h[i: i + 2], 16) / 255 for i in (0, 2, 4))  # type: ignore[return-value]


def _hex_to_docx_rgb(h: str):
    from docx.shared import RGBColor  # pyrefly: ignore [missing-import]
    h = h.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ─────────────────────────────────────────────────────────────────────────────
# Step 1 — Text extraction
# ─────────────────────────────────────────────────────────────────────────────

def _extract_from_pdf(path: str) -> str:
    """Extract text from PDF. Tries pypdf first, pdfplumber as fallback.
    Cleans null-bytes and private-use garbage from both extractors.
    Returns cleaned text, or raises RuntimeError if both fail."""
    text = ""
    # Primary: pypdf
    try:
        from pypdf import PdfReader  # pyrefly: ignore [missing-import]
        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        text = _clean_extracted_text("\n\n".join(pages))
    except Exception as e:
        print(f"  ⚠️  pypdf extraction failed: {e}\n{traceback.format_exc()}")

    # Detect custom-encoding corruption (null bytes / private-use characters)
    if _text_looks_corrupt(text):
        print("  ⚠️  pypdf produced corrupt/null-byte text — trying pdfplumber…")
        text = ""
        try:
            import pdfplumber  # pyrefly: ignore [missing-import]
            pages = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            text = _clean_extracted_text("\n\n".join(pages))
        except Exception as exc:
            print(f"  ⚠️  pdfplumber fallback failed: {exc}\n{traceback.format_exc()}")
            if not text:
                raise RuntimeError(
                    f"PDF text extraction failed — both pypdf and pdfplumber produced no usable text. "
                    f"This usually means the PDF uses a custom/private font encoding. "
                    f"Please provide a .docx version of the manuscript instead. ({exc})"
                ) from exc

    # Final corruption check after both attempts
    if _text_looks_corrupt(text):
        raise RuntimeError(
            "PDF text extraction produced only garbage (null bytes / private-use characters). "
            "The PDF likely uses a custom font encoding that cannot be decoded without the "
            "original font. Please upload a .docx version of the manuscript."
        )

    return text


def _extract_from_docx(path: str) -> str:
    try:
        from docx import Document  # pyrefly: ignore [missing-import]
        doc = Document(path)
        raw = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return _clean_extracted_text(raw)
    except Exception as exc:
        print(f"  ⚠️  DOCX extraction failed: {exc}\n{traceback.format_exc()}")
        raise RuntimeError(f"DOCX extraction failed: {exc}\n{traceback.format_exc()}") from exc


def _extract_from_zip(zip_path: str) -> str:
    texts: list[str] = []
    scratch = zip_path + "_scratch"
    os.makedirs(scratch, exist_ok=True)
    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            members = [
                m for m in zf.namelist()
                if os.path.splitext(m)[1].lower() in {".pdf", ".docx"}
                and not m.startswith("__MACOSX")
                and not os.path.basename(m).startswith(".")
            ]
            if not members:
                raise ValueError("No .pdf or .docx files found inside the zip.")
            for member in members:
                ext = os.path.splitext(member)[1].lower()
                tmp = os.path.join(scratch, f"{uuid.uuid4().hex}{ext}")
                with zf.open(member) as src, open(tmp, "wb") as dst:
                    shutil.copyfileobj(src, dst)
                texts.append(_extract_from_pdf(tmp) if ext == ".pdf" else _extract_from_docx(tmp))
    except Exception as e:
        print(f"  ⚠️  ZIP extraction failed: {e}\n{traceback.format_exc()}")
        raise
    finally:
        shutil.rmtree(scratch, ignore_errors=True)
    return "\n\n".join(texts)


def extract_text(file_path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    if ext == ".docx":
        return _extract_from_docx(file_path)
    if ext == ".zip":
        return _extract_from_zip(file_path)
    if ext in (".txt", ".md", ".text"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    if ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text  # pyrefly: ignore [missing-import]
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return rtf_to_text(f.read())
        except ImportError as e:
            print(f"  ⚠️  striprtf not installed: {e}\n{traceback.format_exc()}")
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception as e:
            print(f"  ⚠️  Error extracting RTF: {e}\n{traceback.format_exc()}")
            raise
    raise ValueError(f"Unsupported file type: {ext}")


# ─────────────────────────────────────────────────────────────────────────────
# Step 2 — Chapter detection
# BUG FIX: added Hindi/Devanagari chapter patterns (अध्याय, भाग, etc.)
# ─────────────────────────────────────────────────────────────────────────────

_CHAPTER_RE = re.compile(
    r"^(?:"
    # English: "Chapter 1", "Part II", "1. Title", "ALL CAPS TITLE"
    r"chapter\s+(?:\d+|[ivxlcdm]+)[^\n]*"
    r"|part\s+(?:\d+|[ivxlcdm]+)[^\n]*"
    r"|\d{1,3}[.\)]\s+[A-Z][^\n]{3,60}"
    r"|[A-Z][A-Z\s]{4,50}$"
    # Hindi/Devanagari: "अध्याय 1", "भाग 2", standalone Devanagari headings
    r"|अध्याय\s*[\d\u0966-\u096F]+"       # अध्याय + digits (ASCII or Devanagari)
    r"|भाग\s*[\d\u0966-\u096F]+"           # भाग (part)
    r"|प्रकरण\s*[\d\u0966-\u096F]+"        # प्रकरण (section/chapter)
    r"|खंड\s*[\d\u0966-\u096F]+"           # खंड (section)
    r"|सर्ग\s*[\d\u0966-\u096F]+"          # सर्ग (canto)
    r")",
    re.IGNORECASE | re.MULTILINE | re.UNICODE,
)

MAX_CHAPTERS = 10_000
MIN_CHAPTER_CHARS = 50


def parse_chapters(raw_text: str) -> list[dict]:
    try:
        lines = raw_text.split("\n")
        splits: list[int] = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped and _CHAPTER_RE.match(stripped) and len(stripped) < 200:
                splits.append(i)

        # --- FIX 2: TOC Filter (Bypass Table of Contents) ---
        bad_splits = set()
        for i in range(1, len(splits)):
            # If "chapters" are detected less than 12 lines apart, it's likely a TOC
            if splits[i] - splits[i-1] < 12:
                bad_splits.add(i)
                bad_splits.add(i-1)

        valid_splits = [s for i, s in enumerate(splits) if i not in bad_splits]
        splits = valid_splits
        # --------------------------------------------------

        if len(splits) < 2:
            # No chapter headings detected — split by word count.
            words = raw_text.split()
            total_words = len(words)
            target_sections = min(60, max(1, total_words // 500))
            chunk_size = max(500, math.ceil(total_words / target_sections))
            chapters = []
            for idx in range(0, total_words, chunk_size):
                chunk = " ".join(words[idx: idx + chunk_size])
                if len(chunk) >= MIN_CHAPTER_CHARS:
                    chapters.append({"title": f"Section {len(chapters) + 1}", "body": chunk})
            return chapters

        chapters: list[dict] = []

        # --- FIX 3: CAPTURE TOC & INTRO ---
        # Grabs all text BEFORE the first matched "Chapter 1"
        if splits and splits[0] > 0:
            intro_text = "\n".join(lines[0:splits[0]]).strip()
            if len(intro_text) > 15:
                chapters.append({
                    "title": "Front Matter & Introduction",
                    "body": intro_text
                })
        # ------------------------------------

        for k, start_line in enumerate(splits[:MAX_CHAPTERS]):
            end_line = splits[k + 1] if k + 1 < len(splits) else len(lines)
            heading = lines[start_line].strip()
            body = "\n".join(lines[start_line + 1: end_line]).strip()
            if len(body) < MIN_CHAPTER_CHARS:
                if chapters:
                    chapters[-1]["body"] += "\n\n" + heading + "\n" + body
                    continue
            chapters.append({"title": heading, "body": body})

        return chapters or [{"title": "Full Text", "body": raw_text}]
    except Exception as e:
        print(f"  ⚠️  Error in parse_chapters: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# Step 3 — AI layout concept  (book-type aware)
# ─────────────────────────────────────────────────────────────────────────────

_LAYOUT_SYSTEM_BASE = """You are a world-class book typographer and interior layout designer with 25 years of experience designing print-ready book interiors for major publishing houses.

Given a book title, its type/genre, a sample of the text, page dimensions, and optional design instructions, you create a complete, production-quality typographic layout specification.

You MUST respond ONLY with valid JSON — no markdown, no code fences, no commentary, nothing else.

The JSON must contain exactly these keys:
{
  "style_name":            "<short evocative name, e.g. 'Warm Classic Serif' or 'Modern Academic'>",
  "page_bg":               "<hex color for page background>",
  "text_color":            "<hex for main body text — must be readable on page_bg>",
  "chapter_title_color":   "<hex for chapter heading text>",
  "accent_color":          "<hex for rules, ornaments, running header, page-number color>",
  "body_font":             "<one of: Helvetica | Times-Roman | Courier | Helvetica-Oblique | Times-Italic>",
  "body_font_size":        <number 9–14 — appropriate for the book type>,
  "line_spacing":          <number 1.2–2.0 — the leading multiplier>,
  "first_para_indent_mm":  <number 0–10 — first-line indent; use 0 for poetry/children>,
  "margin_top_mm":         <number 15–45>,
  "margin_bottom_mm":      <number 15–45>,
  "margin_left_mm":        <number 15–45>,
  "margin_right_mm":       <number 15–45>,
  "chapter_font":          "<same allowable set as body_font>",
  "chapter_font_size":     <number 16–36>,
  "chapter_prefix":        "<e.g. 'Chapter' or 'Part' or '' to omit>",
  "show_drop_cap":         <true|false>,
  "ornament":              "<a short unicode ornament, e.g. '—◆—' or '✦  ✦  ✦' or '❧' or '' to skip>",
  "header_text":           "<running header text, usually the book title, or '' to omit>",
  "show_page_numbers":     <true|false>
}

Typography rules you must follow:
- NEVER choose a body_font_size below 9 or above 14.
- NEVER choose a chapter_font_size below 16 or above 36.
- NEVER choose line_spacing below 1.2 or above 2.0.
- All colour pairs must have sufficient contrast for print (WCAG AA on paper).
- For cream/ivory backgrounds, always use dark brown or near-black text, never grey.
- For dark backgrounds, always use near-white or light text.
- font choices must be from the five allowed values only.
"""


def _build_system_prompt(profile: Optional[dict]) -> str:
    """Append book-type guidance to the base system prompt if a profile exists."""
    if not profile:
        return _LAYOUT_SYSTEM_BASE
    return (
        _LAYOUT_SYSTEM_BASE
        + f"\n\n--- BOOK TYPE GUIDANCE ---\n"
        + f"This book is a {profile['_label']}.\n"
        + f"{profile['_description']}\n"
        + "Apply these genre conventions unless the user has explicitly overridden a specific value.\n"
        + "--- END GUIDANCE ---\n"
    )


def generate_layout_concept(
    book_title: str,
    sample_text: str,
    design_instructions: str = "",
    page_width_mm: float = 210,
    page_height_mm: float = 297,
    book_type: Optional[str] = None,
    profile_defaults: Optional[dict] = None,
) -> dict:
    """
    Call GPT-4o to produce a layout concept.
    profile_defaults (if supplied) are injected into the user message so the
    AI knows what field values are already 'strongly suggested'.
    """
    sample = (
        sample_text[:6_000]
        if len(sample_text) <= 6_000
        else sample_text[:5_000] + "\n…\n" + sample_text[-1_000:]
    )
    system_prompt = _build_system_prompt(
        BOOK_TYPE_PROFILES.get(book_type.lower().strip()) if book_type else None
    )

    user_msg = (
        f"Book title: {book_title}\n"
        f"Page size: {page_width_mm:.0f} × {page_height_mm:.0f} mm\n"
    )
    if book_type:
        user_msg += f"Book type: {book_type}\n"
    # BUG FIX: renamed local variable from `pd` to `_pd` to avoid shadowing
    _pd = profile_defaults or {}
    if _pd:
        subset = {k: v for k, v in _pd.items() if not k.startswith("_")}
        user_msg += f"Suggested defaults for this book type: {json.dumps(subset)}\n"
    if design_instructions:
        user_msg += f"Design instructions: {design_instructions}\n"
    user_msg += f"\nSample text (first 3,000 chars):\n{sample}"

    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_msg},
            ],
            temperature=0.7,
            max_tokens=1500,
        )
        raw = response.choices[0].message.content.strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        s = raw.find("{")
        e = raw.rfind("}") + 1
        if s == -1 or e == 0:
            raise ValueError(f"No JSON returned by the layout AI. Raw response: {raw[:300]}")
        try:
            concept = json.loads(raw[s:e])
        except json.JSONDecodeError as exc:
            raise ValueError(f"Layout AI returned invalid JSON: {exc}. Raw snippet: {raw[s:s+200]}\n{traceback.format_exc()}") from exc
    except Exception as e:
        print(f"  ⚠️  Error in generate_layout_concept API Call: {e}\n{traceback.format_exc()}")
        raise

    # ── Normalise & clamp ─────────────────────────────────────────────────────
    concept.setdefault("style_name",            "Custom Layout")
    concept.setdefault("page_bg",               _pd.get("page_bg", "#ffffff"))
    concept.setdefault("text_color",            _pd.get("text_color", "#1a1a1a"))
    concept.setdefault("chapter_title_color",   _pd.get("chapter_title_color", "#111111"))
    concept.setdefault("accent_color",          _pd.get("accent_color", "#555555"))
    concept.setdefault("body_font",             _pd.get("body_font", "Times-Roman"))
    concept.setdefault("chapter_font",          _pd.get("chapter_font", "Times-Roman"))
    concept.setdefault("chapter_prefix",        _pd.get("chapter_prefix", "Chapter"))
    concept.setdefault("show_drop_cap",         _pd.get("show_drop_cap", True))
    concept.setdefault("ornament",              _pd.get("ornament", "—◆—"))
    concept.setdefault("header_text",           book_title)
    concept.setdefault("show_page_numbers",     _pd.get("show_page_numbers", True))

    concept["body_font_size"]        = max(7,  min(20,  float(concept.get("body_font_size",  _pd.get("body_font_size",  11)))))
    concept["line_spacing"]          = max(1.0, min(3.0, float(concept.get("line_spacing",   _pd.get("line_spacing",   1.5)))))
    concept["first_para_indent_mm"]  = max(0,  min(20,  float(concept.get("first_para_indent_mm", _pd.get("first_para_indent_mm", 5)))))
    concept["chapter_font_size"]     = max(10, min(72,  float(concept.get("chapter_font_size", _pd.get("chapter_font_size", 22)))))
    for key, default in [
        ("margin_top_mm",    _pd.get("margin_top_mm",    20)),
        ("margin_bottom_mm", _pd.get("margin_bottom_mm", 20)),
        ("margin_left_mm",   _pd.get("margin_left_mm",   22)),
        ("margin_right_mm",  _pd.get("margin_right_mm",  22)),
    ]:
        concept[key] = max(5, min(100, float(concept.get(key, default))))

    _ALLOWED_FONTS = {"Helvetica", "Times-Roman", "Courier", "Helvetica-Oblique", "Times-Italic"}
    if concept["body_font"] not in _ALLOWED_FONTS:
        concept["body_font"] = _pd.get("body_font", "Times-Roman")
    if concept["chapter_font"] not in _ALLOWED_FONTS:
        concept["chapter_font"] = _pd.get("chapter_font", "Times-Roman")

    return concept


# ─────────────────────────────────────────────────────────────────────────────
# Step 4 — PDF typesetting with ReportLab
# BUG FIX: Unicode font substitution + safe Unicode drop-cap
# ─────────────────────────────────────────────────────────────────────────────

def render_layout_pdf(
    chapters: list[dict],
    concept: dict,
    output_path: str,
    page_width_mm: float,
    page_height_mm: float,
    book_title: str,
) -> str:
    _ensure_unicode_fonts()

    try:
        from reportlab.lib.units import mm                                       # pyrefly: ignore [missing-import]
        from reportlab.lib.colors import Color                                   # pyrefly: ignore [missing-import]
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak  # pyrefly: ignore [missing-import]
        from reportlab.lib.styles import ParagraphStyle                          # pyrefly: ignore [missing-import]
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY          # pyrefly: ignore [missing-import]
    except ImportError as e:
        print(f"  ⚠️  ReportLab import failed: {e}\n{traceback.format_exc()}")
        raise

    try:
        # ── Detect if the document contains non-Latin (e.g. Devanagari) text ─────
        all_text = book_title + " ".join(
            c.get("title", "") + " " + c.get("body", "") for c in chapters
        )
        has_unicode = _has_non_latin(all_text)

        # ── Resolve actual font names (Unicode-capable if needed) ─────────────────
        raw_body_font    = concept["body_font"]
        raw_chapter_font = concept["chapter_font"]
        body_font    = _unicode_body_font(raw_body_font, has_unicode)
        chapter_font = _unicode_body_font(raw_chapter_font, has_unicode)

        PW = page_width_mm * mm
        PH = page_height_mm * mm
        mt = concept["margin_top_mm"]    * mm
        mb = concept["margin_bottom_mm"] * mm
        ml = concept["margin_left_mm"]   * mm
        mr = concept["margin_right_mm"]  * mm

        bg_r,  bg_g,  bg_b  = _hex_to_rgb(concept["page_bg"])
        tx_r,  tx_g,  tx_b  = _hex_to_rgb(concept["text_color"])
        ch_r,  ch_g,  ch_b  = _hex_to_rgb(concept["chapter_title_color"])
        ac_r,  ac_g,  ac_b  = _hex_to_rgb(concept["accent_color"])

        body_size      = concept["body_font_size"]
        leading        = body_size * concept["line_spacing"]
        indent_pt      = concept["first_para_indent_mm"] * mm
        chapter_size   = concept["chapter_font_size"]
        # paragraph_spacing_mm: if set by user, use it directly; else derive from leading
        _para_sp_mm    = concept.get("paragraph_spacing_mm")
        para_space_after  = float(_para_sp_mm) * mm if _para_sp_mm else leading * 0.45
        para_space_before = float(_para_sp_mm) * mm * 0.35 if _para_sp_mm else leading * 0.15
        # color_mode: bw forces monochrome palette
        _color_mode    = concept.get("color_mode", "")
        if _color_mode == "bw":
            bg_r, bg_g, bg_b = 1.0, 1.0, 1.0
            tx_r, tx_g, tx_b = 0.0, 0.0, 0.0
            ch_r, ch_g, ch_b = 0.0, 0.0, 0.0
            ac_r, ac_g, ac_b = 0.2, 0.2, 0.2
        # mirror_margins: alternate left/right per page (handled in _on_page)
        _mirror        = concept.get("mirror_margins", False)
        # bleed: expand page dimensions outward
        _bleed_mm      = float(concept.get("bleed_mm", 0) or 0)
        bleed_pt       = _bleed_mm * mm
        if bleed_pt > 0:
            PW += bleed_pt * 2
            PH += bleed_pt * 2
            ml += bleed_pt
            mr += bleed_pt
            mt += bleed_pt
            mb += bleed_pt
        # page number start and style
        _pn_start      = int(concept.get("page_number_start", 1) or 1)
        _pn_roman      = concept.get("page_number_style", "") == "roman"
        # BUG FIX: only show drop cap for Latin scripts — Devanagari drop caps
        # require a Unicode-aware font that also supports large-size Devanagari,
        # and ReportLab's inline <font> tag does not re-shape multi-byte glyphs
        # correctly. Safe to disable for non-Latin.
        show_drop      = concept["show_drop_cap"] and not has_unicode
        # BUG FIX: strip ornaments that may not render in the chosen font family
        ornament       = concept.get("ornament", "")
        if has_unicode and ornament:
            # Keep only safe ASCII ornaments; strip complex Unicode symbols
            safe_ornament = "".join(c for c in ornament if ord(c) < 0x0300 or c in "—–•·")
            ornament = safe_ornament or ""
        header_text    = concept.get("header_text", book_title) or book_title
        show_pn        = concept["show_page_numbers"]
        chapter_prefix = concept.get("chapter_prefix", "Chapter")
        _section_breaks = concept.get("section_breaks", False)

        # Footer config — always-on, 3 slots: left / middle / right
        footer_left   = concept.get("footer_left_text") or ""
        footer_middle = concept.get("footer_middle_text") or ""
        footer_right_is_pagenum = concept.get("footer_right_pagenum", True)
        # _show_footer: still gates everything off when show_page_numbers is False AND
        # neither a custom left nor middle text is set (pure "no footer" case).
        _show_footer  = show_pn or bool(footer_left) or bool(footer_middle)

        def _roman(n: int) -> str:
            """Convert positive integer to lowercase roman numeral."""
            val = [(1000,'m'),(900,'cm'),(500,'d'),(400,'cd'),(100,'c'),(90,'xc'),
                   (50,'l'),(40,'xl'),(10,'x'),(9,'ix'),(5,'v'),(4,'iv'),(1,'i')]
            result = ''
            for (arabic, roman) in val:
                while n >= arabic:
                    result += roman; n -= arabic
            return result

        def _on_page(canvas, doc):
            canvas.saveState()
            # ── Page background ──────────────────────────────────────────────────
            canvas.setFillColorRGB(bg_r, bg_g, bg_b)
            canvas.rect(0, 0, PW, PH, fill=1, stroke=0)

            # ── Mirror margins: swap left/right on even pages ────────────────────
            _ml = mr if (_mirror and doc.page % 2 == 0) else ml
            _mr = ml if (_mirror and doc.page % 2 == 0) else mr

            # ── Running header (page 2+, centred at top) ──────────────────────
            if header_text and doc.page > 1:
                canvas.setFillColorRGB(ac_r, ac_g, ac_b)
                canvas.setFont(body_font, 8)
                canvas.drawCentredString(PW / 2, PH - mt * 0.55, header_text)
                canvas.setStrokeColorRGB(ac_r, ac_g, ac_b, alpha=0.35)
                canvas.setLineWidth(0.4)
                canvas.line(_ml, PH - mt * 0.65, PW - _mr, PH - mt * 0.65)

            # ── Footer: ALL pages ─────────────────────────────────────────────
            if _show_footer:
                footer_y = mb * 0.45
                canvas.setFont(body_font, 8)
                canvas.setFillColorRGB(ac_r, ac_g, ac_b)
                # Left slot
                if footer_left:
                    canvas.drawString(_ml, footer_y, footer_left)
                # Middle slot (centred)
                if footer_middle:
                    canvas.drawCentredString(PW / 2, footer_y, footer_middle)
                # Right slot: page number (when enabled) or custom text
                if footer_right_is_pagenum and show_pn:
                    real_page = doc.page + _pn_start - 1
                    pn_str = _roman(real_page) if _pn_roman else str(real_page)
                    canvas.drawRightString(PW - _mr, footer_y, pn_str)

            canvas.restoreState()

        # ── Mirror-margins: build the correct inner margin for each side ─────────
        # When mirror_margins is True:
        #   odd pages  (recto): inner = right side  → leftMargin  = mr+gutter, rightMargin = ml
        #   even pages (verso): inner = left side   → leftMargin  = ml,        rightMargin = mr+gutter
        # We use BaseDocTemplate with two PageTemplates so the text frame itself
        # shifts; SimpleDocTemplate cannot do this natively.
        # NOTE: _mirror was already set above before the _on_page closure; do NOT
        # re-assign it here to avoid a fragile double-assignment.
        # _gutter_pt is computed for completeness but gutter is already folded into ml.
        _gutter_pt = float(concept.get("gutter_mm", 0) or 0) * mm  # already folded into ml by design_layout()

        if _mirror:
            from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame  # pyrefly: ignore [missing-import]
            # odd (recto): binding on left → left margin is the INNER (larger) one
            inner_margin = ml   # ml already has gutter added by design_layout()
            outer_margin = mr
            frame_odd  = Frame(inner_margin, mb, PW - inner_margin - outer_margin, PH - mt - mb,
                               leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, id="odd")
            frame_even = Frame(outer_margin, mb, PW - inner_margin - outer_margin, PH - mt - mb,
                               leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, id="even")
            pt_odd  = PageTemplate(id="Odd",  frames=[frame_odd],  onPage=_on_page)
            pt_even = PageTemplate(id="Even", frames=[frame_even], onPage=_on_page)
            mirror_doc = BaseDocTemplate(
                output_path, pagesize=(PW, PH),
                leftMargin=inner_margin, rightMargin=outer_margin,
                topMargin=mt, bottomMargin=mb,
            )
            mirror_doc.addPageTemplates([pt_odd, pt_even])
        else:
            simple_doc = SimpleDocTemplate(
                output_path,
                pagesize=(PW, PH),
                leftMargin=ml, rightMargin=mr,
                topMargin=mt,  bottomMargin=mb,
            )

        # ── Heading design: apply to ch_style alignment and decoration ───────────
        _hd = concept.get("heading_design", "")
        _ch_align    = TA_LEFT
        _ch_caps     = False   # signal for ALL CAPS transform in text
        _ch_italic   = False   # signal for italic_elegant rendering
        _ch_smallcaps_letter_spacing = 0
        if _hd == "centered_decorative":
            _ch_align = TA_CENTER
        elif _hd == "left_bold_clean":
            _ch_align = TA_LEFT   # explicit; already default, but clear intent
        elif _hd == "allcaps_rule":
            _ch_caps  = True
            _ch_align = TA_LEFT
        elif _hd == "italic_elegant":
            _ch_align  = TA_CENTER
            _ch_italic = True
        elif _hd == "numbered":
            # chapter numbers handled in loop via real_chapter_num; left-aligned
            _ch_align = TA_LEFT
        elif _hd == "smallcaps_ornament":
            # Small-caps effect: use a smaller font size with letter-spacing,
            # centred, and ALL CAPS transform — distinct from centered_decorative
            _ch_align = TA_CENTER
            _ch_caps  = True
            _ch_smallcaps_letter_spacing = 2.5

        # Resolve chapter font name, applying italic for italic_elegant
        _ch_font_for_style = chapter_font
        if _ch_italic:
            # Map to italic variant if available
            _ITALIC_MAP = {
                "Times-Roman":  "Times-Italic",
                "Helvetica":    "Helvetica-Oblique",
            }
            _ch_font_for_style = _ITALIC_MAP.get(chapter_font, chapter_font)

        # For left_bold_clean / numbered: slightly reduce chapter_size for a clean look
        _ch_size_for_style = chapter_size
        if _hd in ("left_bold_clean", "numbered"):
            _ch_size_for_style = chapter_size * 0.92
        elif _hd == "smallcaps_ornament":
            _ch_size_for_style = chapter_size * 0.80   # smaller size mimics small-caps

        ch_style = ParagraphStyle(
            "ChapterTitle",
            fontName=_ch_font_for_style, fontSize=_ch_size_for_style,
            leading=_ch_size_for_style * 1.25,
            textColor=Color(ch_r, ch_g, ch_b),
            spaceAfter=_ch_size_for_style * 0.55, spaceBefore=_ch_size_for_style * 0.35,
            alignment=_ch_align,
            letterSpacing=_ch_smallcaps_letter_spacing,
        )
        prefix_style = ParagraphStyle(
            "ChapterPrefix",
            fontName=body_font, fontSize=body_size * 0.82,
            leading=body_size * 1.2,
            textColor=Color(ac_r, ac_g, ac_b),
            spaceBefore=0, spaceAfter=3, alignment=_ch_align, letterSpacing=1.8,
        )
        body_style = ParagraphStyle(
            "Body",
            fontName=body_font, fontSize=body_size, leading=leading,
            textColor=Color(tx_r, tx_g, tx_b),
            firstLineIndent=indent_pt,
            alignment=TA_JUSTIFY,
            spaceAfter=para_space_after,
            spaceBefore=para_space_before,
            wordWrap="LTR",
        )
        orn_style = ParagraphStyle(
            "Ornament",
            fontName=body_font, fontSize=body_size + 2,
            leading=(body_size + 2) * 1.5,
            textColor=Color(ac_r, ac_g, ac_b),
            alignment=TA_CENTER, spaceBefore=10, spaceAfter=10,
        )

        # ── Dual-font run builder for mixed Devanagari + Latin text ─────────────
        # When a paragraph contains both Hindi (Devanagari, U+0900–U+097F) and
        # English/Latin characters (or digits/punctuation), we need to tag the
        # Latin portions with a Latin-capable font so they render correctly.
        # Noto Devanagari fonts do include Latin glyphs, but their metrics and
        # hinting are optimised for Devanagari; using a proper Latin font (e.g.
        # Times-Roman) for the Latin runs gives much crisper output.
        #
        # Strategy: split the escaped text into runs of "Devanagari" vs "Latin",
        # wrap each Latin run in a <font name="..."> tag.
        #
        # We only apply this when has_unicode is True (the document has Devanagari)
        # AND a proper Latin fallback font is available.

        _LATIN_FALLBACK: dict[str, str] = {
            "NotoSerifDevanagari":      "Times-Roman",
            "NotoSerifDevanagari-Bold": "Times-Roman",
            "NotoSansDevanagari":       "Helvetica",
            "NotoSansDevanagari-Bold":  "Helvetica",
        }

        def _mixed_font_html(safe_escaped_text: str, deva_font: str) -> str:
            """
            Given HTML-escaped paragraph text and the Devanagari font name,
            return ReportLab XML markup with dual-font tags for Latin runs.

            Only called when has_unicode is True.  If no Latin fallback is
            registered, returns the text unchanged.
            """
            latin_font = _LATIN_FALLBACK.get(deva_font)
            if not latin_font:
                return safe_escaped_text

            # We work on the *unescaped* text to correctly classify codepoints,
            # then re-escape each segment individually.
            # NOTE: safe_escaped_text may contain <br/> tags — preserve them.
            # Split on <br/> first, process each fragment, then rejoin.
            fragments = re.split(r"(<br\s*/>)", safe_escaped_text)
            result_parts: list[str] = []

            for frag in fragments:
                if re.fullmatch(r"<br\s*/>", frag):
                    result_parts.append(frag)
                    continue
                if not frag:
                    continue

                # Un-escape HTML entities in this fragment so we can inspect codepoints
                frag_plain = (
                    frag.replace("&amp;", "&")
                        .replace("&lt;",  "<")
                        .replace("&gt;",  ">")
                        .replace("&quot;", '"')
                        .replace("&#39;",  "'")
                )

                def _is_latin_char(ch: str) -> bool:
                    """
                    Returns True for characters that benefit from a Latin font:
                    Basic Latin, Latin-1 Supplement, common punctuation,
                    digits, and ASCII symbols.  Devanagari and whitespace → False.
                    """
                    cp = ord(ch)
                    if ch.isspace():
                        return False   # spaces get the surrounding font context
                    if 0x0900 <= cp <= 0x097F:
                        return False   # Devanagari block → Devanagari font
                    if 0x0020 <= cp <= 0x024F:
                        return True    # Basic Latin + Latin Extended A/B
                    if 0x2000 <= cp <= 0x206F:
                        return True    # General Punctuation (—, ", ", …)
                    if 0x0030 <= cp <= 0x0039:
                        return True    # ASCII digits (redundant but explicit)
                    return False

                # Build runs: (is_latin, text_chunk)
                runs: list[tuple[bool, str]] = []
                if frag_plain:
                    cur_latin = _is_latin_char(frag_plain[0])
                    cur_buf   = frag_plain[0]
                    for ch in frag_plain[1:]:
                        ch_latin = _is_latin_char(ch)
                        if ch_latin == cur_latin or ch.isspace():
                            cur_buf += ch
                        else:
                            runs.append((cur_latin, cur_buf))
                            cur_latin = ch_latin
                            cur_buf   = ch
                    if cur_buf:
                        runs.append((cur_latin, cur_buf))

                # Build the HTML for this fragment
                frag_html = ""
                for is_latin, chunk in runs:
                    # Re-escape the chunk
                    esc = (
                        chunk.replace("&", "&amp;")
                             .replace("<", "&lt;")
                             .replace(">", "&gt;")
                    )
                    if is_latin and chunk.strip():
                        frag_html += f'<font name="{latin_font}">{esc}</font>'
                    else:
                        frag_html += esc

                result_parts.append(frag_html)

            return "".join(result_parts)

        # ── Story (flowable elements list) ────────────────────────────────────
        # This MUST be initialised here — after all styles and helpers are
        # defined, before any story.append() calls.
        story: list = []

        title_style = ParagraphStyle(
            "TitlePage",
            fontName=chapter_font,
            fontSize=min(36, chapter_size * 1.6),
            leading=min(36, chapter_size * 1.6) * 1.2,
            textColor=Color(ch_r, ch_g, ch_b),
            alignment=TA_CENTER, spaceAfter=20,
            # Devanagari/Indic scripts use LTR word-based wrapping just like Latin.
            # "CJK" mode breaks individual codepoints (wrong for Devanagari conjuncts).
            wordWrap="LTR",
        )
        # ── Front/back matter config ──────────────────────────────────────────────
        _front_matter = concept.get("front_matter") or []
        _back_matter  = concept.get("back_matter")  or []
        # Always generate a title page (it's the cover/splash); only skip if
        # front_matter is explicitly set AND does NOT include "title_page"
        _show_title_page = (not _front_matter) or ("title_page" in _front_matter)

        # ── Helper: small body-text paragraph for matter pages ───────────────────
        matter_style = ParagraphStyle(
            "Matter",
            fontName=body_font, fontSize=body_size,
            leading=leading,
            textColor=Color(tx_r, tx_g, tx_b),
            alignment=TA_CENTER, spaceAfter=leading * 0.6,
            wordWrap="LTR",
        )
        matter_small = ParagraphStyle(
            "MatterSmall",
            fontName=body_font, fontSize=max(7, body_size - 2),
            leading=leading * 0.85,
            textColor=Color(tx_r, tx_g, tx_b),
            alignment=TA_CENTER, spaceAfter=leading * 0.4,
            wordWrap="LTR",
        )

        def _safe(t: str) -> str:
            return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # ── Title page ────────────────────────────────────────────────────────────
        if _show_title_page:
            story.append(Spacer(1, PH * 0.28))
            story.append(Paragraph(_safe(book_title), title_style))
            if ornament:
                story.append(Spacer(1, 14))
                story.append(Paragraph(_safe(ornament), orn_style))
            story.append(PageBreak())

        # ── Copyright page ────────────────────────────────────────────────────────
        if "copyright_page" in _front_matter:
            import datetime as _dt
            year = _dt.datetime.now().year
            story.append(Spacer(1, PH * 0.40))
            story.append(Paragraph(_safe(f"Copyright © {year} {book_title}"), matter_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("All rights reserved. No part of this publication may be reproduced, "
                                         "distributed, or transmitted in any form or by any means without the "
                                         "prior written permission of the publisher."), matter_small))
            story.append(Spacer(1, leading * 0.5))
            story.append(Paragraph(_safe("First published edition."), matter_small))
            story.append(PageBreak())

        # ── Dedication ────────────────────────────────────────────────────────────
        if "dedication" in _front_matter:
            story.append(Spacer(1, PH * 0.38))
            story.append(Paragraph(_safe("For those who love stories."), matter_style))
            story.append(PageBreak())

        # ── Foreword ─────────────────────────────────────────────────────────────
        if "foreword" in _front_matter:
            story.append(Paragraph(_safe("Foreword"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("This foreword was generated as a placeholder. "
                                         "Please replace it with your own foreword text."), matter_style))
            story.append(PageBreak())

        # ── Preface ───────────────────────────────────────────────────────────────
        if "preface" in _front_matter:
            story.append(Paragraph(_safe("Preface"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("This preface was generated as a placeholder. "
                                         "Please replace it with your own preface text."), matter_style))
            story.append(PageBreak())

        # ── Acknowledgements ──────────────────────────────────────────────────────
        if "acknowledgement" in _front_matter:
            story.append(Paragraph(_safe("Acknowledgements"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("The author wishes to thank everyone who made this book possible."), matter_style))
            story.append(PageBreak())

        # ── Table of Contents ─────────────────────────────────────────────────────
        if "toc" in _front_matter:
            story.append(Paragraph(_safe("Contents"), ch_style))
            story.append(Spacer(1, leading * 0.5))
            toc_num = 0
            for ch in chapters:
                ch_tl = ch["title"].lower()
                if not ("introduction" in ch_tl or "front matter" in ch_tl):
                    toc_num += 1
                    story.append(Paragraph(
                        _safe(f"{toc_num}.  {ch['title']}"),
                        ParagraphStyle("TOCEntry", fontName=body_font, fontSize=body_size,
                                       leading=leading * 1.1, textColor=Color(tx_r, tx_g, tx_b),
                                       spaceAfter=leading * 0.25, wordWrap="LTR")
                    ))
            story.append(PageBreak())

        # ── Chapters ──────────────────────────────────────────────────────────────
        _chapter_start = concept.get("chapter_start", "")  # "right_page", "left_page", "any_page"
        real_chapter_num = 0
        _story_page_counter = [2]  # title page = 1, body starts at 2; track parity for chapter_start

        for chapter in chapters:
            ch_title_lower = chapter["title"].lower()
            already_has_chapter = ch_title_lower.startswith("chapter") or ch_title_lower.startswith("part")
            is_intro = "introduction" in ch_title_lower or "front matter" in ch_title_lower

            if not is_intro:
                real_chapter_num += 1

            # ── Chapter start: ensure right-hand (odd) or left-hand (even) page ──
            # We track physical page via _story_page_counter[0].
            # A PageBreak advances the page before we render headings.
            if not is_intro and _chapter_start in ("right_page", "left_page"):
                cur_page = _story_page_counter[0]
                target_odd = (_chapter_start == "right_page")
                # If already on the correct parity we don't need a blank
                if target_odd and cur_page % 2 == 0:
                    story.append(PageBreak())  # add blank to land on next odd page
                    _story_page_counter[0] += 1
                elif not target_odd and cur_page % 2 == 1:
                    story.append(PageBreak())  # add blank to land on next even page
                    _story_page_counter[0] += 1

            if chapter_prefix and not already_has_chapter and not is_intro:
                safe_prefix = chapter_prefix.replace("&", "&amp;")
                # For 'numbered' design, embed the number directly in the heading text
                # (no separate prefix line — the heading IS "1. Title" style)
                if _hd != "numbered":
                    story.append(Paragraph(f"{safe_prefix.upper()} {real_chapter_num}".strip(), prefix_style))

            # Apply heading caps/transforms
            if _hd == "numbered" and not already_has_chapter and not is_intro:
                _raw_ch_title = f"{real_chapter_num}. {chapter['title']}"
            else:
                _raw_ch_title = chapter["title"].upper() if _ch_caps else chapter["title"]
            safe_ch_title = _raw_ch_title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_ch_title, ch_style))

            # Post-heading decoration per design
            if _hd == "allcaps_rule":
                from reportlab.platypus import HRFlowable  # pyrefly: ignore [missing-import]
                story.append(HRFlowable(width="100%", thickness=1.2, color=Color(ac_r, ac_g, ac_b), spaceAfter=6))
            elif _hd == "left_bold_clean":
                from reportlab.platypus import HRFlowable  # pyrefly: ignore [missing-import]
                story.append(HRFlowable(width="40%", thickness=0.8, color=Color(ac_r, ac_g, ac_b), spaceAfter=6, hAlign="LEFT"))
            elif _hd == "numbered":
                from reportlab.platypus import HRFlowable  # pyrefly: ignore [missing-import]
                story.append(HRFlowable(width="15%", thickness=2, color=Color(ac_r, ac_g, ac_b), spaceAfter=8, hAlign="LEFT"))
            elif _hd == "smallcaps_ornament":
                # Extra ornament line directly under the small-caps title
                if ornament:
                    safe_orn2 = ornament.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe_orn2, orn_style))

            story.append(Spacer(1, 4))

            if ornament:
                safe_orn = ornament.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_orn, orn_style))
                story.append(Spacer(1, 6))

            raw_body = chapter.get("body", "").strip()
            # Detect explicit section dividers (---, ***, ~~~, ###) in source
            # Split on them when section_breaks is enabled
            _SEC_DIV_RE = re.compile(r"^\s*(?:---+|\*\*\*+|~~~+|###)\s*$", re.MULTILINE)
            if _section_breaks:
                # Split into sections; each section is separated by a visual break
                raw_sections = _SEC_DIV_RE.split(raw_body)
            else:
                raw_sections = [raw_body]

            all_para_items: list = []   # will be appended to story at end
            for sec_idx, section_text in enumerate(raw_sections):
                if sec_idx > 0 and _section_breaks:
                    # Insert section break ornament (asterism or custom ornament)
                    sec_orn = ornament if ornament else "* * *"
                    safe_sec_orn = sec_orn.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    all_para_items.append(Spacer(1, leading * 0.5))
                    all_para_items.append(Paragraph(safe_sec_orn, orn_style))
                    all_para_items.append(Spacer(1, leading * 0.5))

                paragraphs = [p.strip() for p in re.split(r"\n{2,}", section_text) if p.strip()]
                if not paragraphs and sec_idx == 0:
                    paragraphs = ["[No content]"]

                for p_idx, para_text in enumerate(paragraphs):
                    # FIX: Smart Line Break & Bullet Handling
                    lines = para_text.split('\n')
                    cleaned_lines = []
                    for line in lines:
                        line = line.strip()
                        if not line: continue
                        # If it's a bullet point, force a hard break
                        if line.startswith(('•', '-', '*')) or re.match(r'^\d+\.', line):
                            cleaned_lines.append('<br/>' + line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
                        else:
                            # Otherwise, join physical PDF lines with a space
                            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            if cleaned_lines and not cleaned_lines[-1].startswith('<br/>'):
                                cleaned_lines[-1] += " " + safe_line
                            else:
                                cleaned_lines.append(safe_line)

                    safe = "".join(cleaned_lines)

                    # Apply dual-font markup for mixed Hindi + Latin text
                    if has_unicode:
                        safe = _mixed_font_html(safe, body_font)

                    # Drop cap logic
                    if p_idx == 0 and show_drop and not is_intro and len(para_text) > 1:
                        first_char = para_text[0]   # original codepoint
                        # Only do drop cap if first char is a basic Latin letter
                        if first_char.isalpha() and ord(first_char) < 0x0250:
                            first_esc = first_char.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            rest_orig = safe[len(first_esc):]
                            drop_html = (
                                f'<font name="{_ch_font_for_style}" size="{int(body_size * 2.8)}">'
                                f"{first_esc}</font>{rest_orig}"
                            )
                            all_para_items.append(Paragraph(drop_html, body_style))
                        else:
                            all_para_items.append(Paragraph(safe, body_style))
                    else:
                        all_para_items.append(Paragraph(safe, body_style))

            # Flush collected flowables to story
            story.extend(all_para_items)

            story.append(PageBreak())
            _story_page_counter[0] += 1

        # ── Back matter pages ─────────────────────────────────────────────────────
        if "about_author" in _back_matter:
            story.append(Paragraph(_safe("About the Author"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("The author is a writer and storyteller. "
                                         "Please replace this placeholder with your biographical note."), matter_style))
            story.append(PageBreak())
        if "about_publisher" in _back_matter:
            story.append(Paragraph(_safe("About the Publisher"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("Published by [Publisher Name]. "
                                         "Please replace this with your publisher information."), matter_style))
            story.append(PageBreak())
        if "references" in _back_matter:
            story.append(Paragraph(_safe("References"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("[References list placeholder — replace with your citations.]"), matter_style))
            story.append(PageBreak())
        if "bibliography" in _back_matter:
            story.append(Paragraph(_safe("Bibliography"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("[Bibliography placeholder — replace with your bibliography.]"), matter_style))
            story.append(PageBreak())
        if "index" in _back_matter:
            story.append(Paragraph(_safe("Index"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("[Index placeholder — replace with your index entries.]"), matter_style))
            story.append(PageBreak())
        if "other_books" in _back_matter:
            story.append(Paragraph(_safe("Other Books by the Author"), ch_style))
            story.append(Spacer(1, leading))
            story.append(Paragraph(_safe("[List your other books here.]"), matter_style))
            story.append(PageBreak())

        if _mirror:
            from reportlab.platypus import NextPageTemplate  # pyrefly: ignore [missing-import]
            # Inject NextPageTemplate switches before each PageBreak so frames alternate
            mirrored_story: list = []
            page_counter = [1]
            for item in story:
                if isinstance(item, PageBreak):
                    page_counter[0] += 1
                    next_tpl = "Even" if page_counter[0] % 2 == 0 else "Odd"
                    mirrored_story.append(NextPageTemplate(next_tpl))
                mirrored_story.append(item)
            # mirror_doc was already created and templates added above
            mirror_doc.build(mirrored_story)
        else:
            # simple_doc was already created above
            simple_doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
        return output_path
    except Exception as e:
        print(f"  ⚠️  render_layout_pdf failed completely: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# Step 5 — DOCX typesetting
# BUG FIX: correct italic flag for Times-Italic / Helvetica-Oblique
# ─────────────────────────────────────────────────────────────────────────────

def render_layout_docx(
    chapters: list[dict],
    concept: dict,
    output_path: str,
    page_width_mm: float,
    page_height_mm: float,
    book_title: str,
) -> str:
    try:
        from docx import Document                                   # pyrefly: ignore [missing-import]
        from docx.shared import Pt, Cm, RGBColor                   # pyrefly: ignore [missing-import]
        from docx.enum.text import WD_ALIGN_PARAGRAPH              # pyrefly: ignore [missing-import]
        from docx.oxml.ns import qn                                # pyrefly: ignore [missing-import]
        from docx.oxml import OxmlElement                          # pyrefly: ignore [missing-import]

        body_size  = float(concept["body_font_size"])
        ch_size    = float(concept["chapter_font_size"])
        ls         = float(concept["line_spacing"])
        ornament   = concept.get("ornament", "")
        prefix     = concept.get("chapter_prefix", "Chapter")

        # Advanced fields
        _color_mode_d   = concept.get("color_mode", "")
        _mirror_d       = concept.get("mirror_margins", False)
        _para_sp_mm_d   = concept.get("paragraph_spacing_mm")
        _pn_start_d     = int(concept.get("page_number_start", 1) or 1)
        _pn_roman_d     = concept.get("page_number_style", "") == "roman"
        _chapter_start_d= concept.get("chapter_start", "")
        _hd_d           = concept.get("heading_design", "")
        _ch_caps_d      = _hd_d == "allcaps_rule"
        _ch_center_d    = _hd_d in ("centered_decorative", "italic_elegant", "smallcaps_ornament")
        _ch_numbered_d  = _hd_d == "numbered"
        _ch_italic_d    = _hd_d == "italic_elegant"
        _ch_smallcaps_d = _hd_d == "smallcaps_ornament"

        # Font/size adjustments per heading design
        _ch_size_d      = ch_size
        if _hd_d in ("left_bold_clean", "numbered"):
            _ch_size_d  = ch_size * 0.92
        elif _hd_d == "smallcaps_ornament":
            _ch_size_d  = ch_size * 0.80

        # color_mode bw: override palette
        if _color_mode_d == "bw":
            concept = dict(concept)  # shallow copy so we don't mutate the original
            concept["page_bg"]            = "#ffffff"
            concept["text_color"]         = "#000000"
            concept["chapter_title_color"]= "#000000"
            concept["accent_color"]       = "#333333"

        doc     = Document()
        section = doc.sections[0]

        # ── Bleed: expand page canvas on all sides ───────────────────────────────
        _bleed_mm_d   = float(concept.get("bleed_mm", 0) or 0)
        _bleed_cm_d   = _bleed_mm_d / 10.0
        section.page_width    = Cm((page_width_mm  + _bleed_mm_d * 2) / 10)
        section.page_height   = Cm((page_height_mm + _bleed_mm_d * 2) / 10)
        section.left_margin   = Cm((concept["margin_left_mm"]   + _bleed_mm_d) / 10)
        section.right_margin  = Cm((concept["margin_right_mm"]  + _bleed_mm_d) / 10)
        section.top_margin    = Cm((concept["margin_top_mm"]    + _bleed_mm_d) / 10)
        section.bottom_margin = Cm((concept["margin_bottom_mm"] + _bleed_mm_d) / 10)
        section.footer_distance = Cm(0.8)

        # ── Mirror margins XML flag ──────────────────────────────────────────────
        if _mirror_d:
            pgMar = section._sectPr.find(qn("w:pgMar"))
            if pgMar is None:
                pgMar = OxmlElement("w:pgMar"); section._sectPr.append(pgMar)
            mirror_el = section._sectPr.find(qn("w:mirrorMargins"))
            if mirror_el is None:
                mirror_el = OxmlElement("w:mirrorMargins"); section._sectPr.append(mirror_el)

        # ── Page number start and style via w:pgNumType ──────────────────────────
        if _pn_start_d != 1 or _pn_roman_d:
            pgNumType = section._sectPr.find(qn("w:pgNumType"))
            if pgNumType is None:
                pgNumType = OxmlElement("w:pgNumType")
                section._sectPr.append(pgNumType)
            pgNumType.set(qn("w:start"), str(_pn_start_d))
            if _pn_roman_d:
                pgNumType.set(qn("w:fmt"), "lowerRoman")
            else:
                pgNumType.set(qn("w:fmt"), "decimal")

        # ── Footer: 3-slot (left / centre / right) ──────────────────────────────
        show_pn_docx         = concept.get("show_page_numbers", True)
        footer_left_text_d   = concept.get("footer_left_text") or ""
        footer_middle_text_d = concept.get("footer_middle_text") or ""
        footer_right_pn_d    = concept.get("footer_right_pagenum", True)
        _render_footer = show_pn_docx or bool(footer_left_text_d) or bool(footer_middle_text_d)

        if _render_footer:
            from docx.oxml.ns import qn as _qn_footer   # pyrefly: ignore [missing-import]
            from docx.oxml import OxmlElement as _OxmlEl # pyrefly: ignore [missing-import]

            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]

            # Tab stops: centre at page midpoint (~4500 twips) + right-align (~9000 twips)
            pPr = fp._p.get_or_add_pPr()
            tabs_el = _OxmlEl("w:tabs")
            tab_ctr = _OxmlEl("w:tab")
            tab_ctr.set(_qn_footer("w:val"), "center")
            tab_ctr.set(_qn_footer("w:pos"), "4500")
            tabs_el.append(tab_ctr)
            tab_rt = _OxmlEl("w:tab")
            tab_rt.set(_qn_footer("w:val"),  "right")
            tab_rt.set(_qn_footer("w:pos"),  "9000")
            tabs_el.append(tab_rt)
            pPr.append(tabs_el)

            def _footer_run(text: str) -> None:
                r = fp.add_run(text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(8)
                try:
                    r.font.color.rgb = _hex_to_docx_rgb(concept.get("accent_color", "#555555"))
                except Exception:
                    pass

            # Left slot
            _footer_run(footer_left_text_d)
            # Centre slot
            fp.add_run("\t")
            if footer_middle_text_d:
                _footer_run(footer_middle_text_d)
            # Right slot: page number field
            fp.add_run("\t")
            if footer_right_pn_d and show_pn_docx:
                run_pg = fp.add_run()
                run_pg.font.name = "Times New Roman"
                run_pg.font.size = Pt(8)
                try:
                    run_pg.font.color.rgb = _hex_to_docx_rgb(concept.get("accent_color", "#555555"))
                except Exception:
                    pass
                fld_begin  = _OxmlEl("w:fldChar"); fld_begin.set(_qn_footer("w:fldCharType"),  "begin")
                instr      = _OxmlEl("w:instrText")
                instr.text = " PAGE \\* lowerRoman " if _pn_roman_d else " PAGE "
                instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                fld_sep    = _OxmlEl("w:fldChar"); fld_sep.set(_qn_footer("w:fldCharType"),    "separate")
                fld_text   = _OxmlEl("w:t");        fld_text.text = "1"
                fld_end    = _OxmlEl("w:fldChar"); fld_end.set(_qn_footer("w:fldCharType"),    "end")
                run_pg._r.append(fld_begin)
                rpr = _OxmlEl("w:rPr"); run_pg._r.insert(0, rpr)
                run_pg._r.append(instr)
                run_pg._r.append(fld_sep)
                run_pg._r.append(fld_text)
                run_pg._r.append(fld_end)

        # Detect whether this book has mixed Hindi + Latin content
        all_text_docx = book_title + " ".join(
            c.get("title", "") + " " + c.get("body", "") for c in chapters
        )
        has_unicode_docx = _has_non_latin(all_text_docx)

        # Map ReportLab font names → (Word font name, is_italic)
        _FONT_MAP = {
            "Times-Roman":       ("Times New Roman", False),
            "Times-Italic":      ("Times New Roman", True),
            "Helvetica":         ("Arial",            False),
            "Helvetica-Oblique": ("Arial",            True),
            "Courier":           ("Courier New",      False),
        }
        # Latin fallback fonts for body / chapter when document has Devanagari
        _DOCX_LATIN_FALLBACK: dict[str, str] = {
            "NotoSerifDevanagari":      "Times New Roman",
            "NotoSerifDevanagari-Bold": "Times New Roman",
            "NotoSansDevanagari":       "Arial",
            "NotoSansDevanagari-Bold":  "Arial",
        }

        def docx_font(rl_name: str) -> tuple[str, bool]:
            """Return (word_font_name, is_italic)."""
            return _FONT_MAP.get(rl_name, ("Times New Roman", False))

        body_fn, body_italic = docx_font(concept["body_font"])
        ch_fn,   ch_italic   = docx_font(concept["chapter_font"])

        def rgb(hex_str: str) -> RGBColor:
            return _hex_to_docx_rgb(hex_str)

        def _set_run_font(run, font_name: str, size_pt: float, bold: bool,
                          italic: bool, color_hex: str) -> None:
            """Apply font attributes to a single run."""
            run.font.name   = font_name
            run.font.size   = Pt(size_pt)
            run.font.bold   = bold
            run.font.italic = italic
            run.font.color.rgb = rgb(color_hex)
            # Also set the East-Asian / Complex-script font element so Word
            # uses the Devanagari font for Devanagari codepoints, rather than
            # falling through to the theme Latin font.
            from docx.oxml.ns import qn as _qn  # pyrefly: ignore [missing-import]
            rPr = run._r.get_or_add_rPr()
            for tag in ("w:rFonts",):
                existing = rPr.find(_qn(tag))
                if existing is None:
                    from docx.oxml import OxmlElement as _OxmlElement  # pyrefly: ignore [missing-import]
                    existing = _OxmlElement(tag)
                    rPr.insert(0, existing)
                existing.set(_qn("w:ascii"),       font_name)
                existing.set(_qn("w:hAnsi"),       font_name)
                existing.set(_qn("w:cs"),          font_name)   # complex-script
                existing.set(_qn("w:eastAsia"),    font_name)

        def add_para(text: str, font_name: str, size: float, bold: bool = False,
                     italic: bool = False, color: str = "#1a1a1a",
                     align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before: float = 0, space_after: float = 0,
                     line_space: float = 1.5) -> None:
            """
            Add a paragraph with proper dual-font support for mixed Hindi + English.

            When the document contains Devanagari text (has_unicode_docx is True),
            the paragraph is split into runs so that Latin characters (digits,
            English words, punctuation) use the appropriate Latin-script font
            instead of the Devanagari font — giving correct glyph metrics for both
            scripts within the same paragraph.
            """
            p  = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(space_before)
            pf.space_after  = Pt(space_after)
            pf.line_spacing = Pt(size * line_space)

            if not has_unicode_docx or not text.strip():
                # Simple single-run path (no Devanagari in this document)
                run = p.add_run(text)
                _set_run_font(run, font_name, size, bold, italic, color)
                return

            # Dual-run path: split text into Devanagari and Latin segments
            latin_fallback = _DOCX_LATIN_FALLBACK.get(font_name, font_name)

            def _is_latin_char_docx(ch: str) -> bool:
                cp = ord(ch)
                if ch.isspace():
                    return False
                if 0x0900 <= cp <= 0x097F:
                    return False   # Devanagari → Devanagari font
                if 0x0020 <= cp <= 0x024F:
                    return True    # Basic Latin + Latin Extended A/B
                if 0x2000 <= cp <= 0x206F:
                    return True    # General Punctuation
                return False

            # Build character-level runs: (is_latin, chunk)
            runs_list: list[tuple[bool, str]] = []
            if text:
                cur_latin = _is_latin_char_docx(text[0])
                cur_buf   = text[0]
                for ch in text[1:]:
                    ch_lat = _is_latin_char_docx(ch)
                    if ch.isspace() or ch_lat == cur_latin:
                        cur_buf += ch
                    else:
                        runs_list.append((cur_latin, cur_buf))
                        cur_latin = ch_lat
                        cur_buf   = ch
                if cur_buf:
                    runs_list.append((cur_latin, cur_buf))

            for is_latin, chunk in runs_list:
                chosen_font = latin_fallback if (is_latin and chunk.strip()) else font_name
                run = p.add_run(chunk)
                _set_run_font(run, chosen_font, size, bold, italic, color)

        def add_rule(color_hex: str, width_pct: int = 100, thickness: float = 1.0) -> None:
            """Add a paragraph with a bottom border that acts as a horizontal rule.
            width_pct is advisory via indentation (100% = full width, 40% = half etc).
            thickness maps to w:sz in eighths of a point."""
            p   = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            # Indent right side to simulate partial-width rule
            if width_pct < 100:
                ind = OxmlElement("w:ind")
                # Approximate: at 9000 twips body width, 40% ≈ 5400 twips right indent
                right_twips = int(9000 * (1 - width_pct / 100))
                ind.set(qn("w:right"), str(right_twips))
                pPr.append(ind)
            pBdr = OxmlElement("w:pBdr")
            bt   = OxmlElement("w:bottom")
            bt.set(qn("w:val"),   "single")
            bt.set(qn("w:sz"),    str(max(2, int(thickness * 8))))   # w:sz in 1/8 pt
            bt.set(qn("w:space"), "1")
            bt.set(qn("w:color"), color_hex.lstrip("#"))
            pBdr.append(bt)
            pPr.append(pBdr)

        # Title page
        _front_matter_d = concept.get("front_matter") or []
        _back_matter_d  = concept.get("back_matter")  or []
        _show_title_page_d = (not _front_matter_d) or ("title_page" in _front_matter_d)

        if _show_title_page_d:
            for _ in range(4):
                doc.add_paragraph()
            add_para(book_title, ch_fn, min(36, ch_size * 1.5), bold=True,
                     italic=ch_italic or _ch_italic_d,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
            if ornament:
                add_para(ornament, body_fn, body_size + 2, color=concept["accent_color"],
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
            doc.add_page_break()

        # Copyright page
        if "copyright_page" in _front_matter_d:
            import datetime as _dt_d
            year_d = _dt_d.datetime.now().year
            for _ in range(3):
                doc.add_paragraph()
            add_para(f"Copyright © {year_d} {book_title}", body_fn, body_size,
                     color=concept["text_color"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
            add_para("All rights reserved. No part of this publication may be reproduced, "
                     "distributed, or transmitted in any form or by any means without the "
                     "prior written permission of the publisher.",
                     body_fn, max(7, body_size - 2), color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_space=ls)
            add_para("First published edition.", body_fn, max(7, body_size - 2),
                     color=concept["text_color"], align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_page_break()

        if "dedication" in _front_matter_d:
            for _ in range(3):
                doc.add_paragraph()
            add_para("For those who love stories.", body_fn, body_size + 1,
                     italic=True, color=concept["text_color"], align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_page_break()

        if "foreword" in _front_matter_d:
            add_para("Foreword", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("This foreword was generated as a placeholder. "
                     "Please replace it with your own foreword text.",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()

        if "preface" in _front_matter_d:
            add_para("Preface", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("This preface was generated as a placeholder. "
                     "Please replace it with your own preface text.",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()

        if "acknowledgement" in _front_matter_d:
            add_para("Acknowledgements", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("The author wishes to thank everyone who made this book possible.",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()

        if "toc" in _front_matter_d:
            add_para("Contents", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            toc_num_d = 0
            for ch_toc in chapters:
                tl = ch_toc["title"].lower()
                if not ("introduction" in tl or "front matter" in tl):
                    toc_num_d += 1
                    add_para(f"{toc_num_d}.  {ch_toc['title']}", body_fn, body_size,
                             color=concept["text_color"], align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_after=round(body_size * ls * 0.2, 1))
            doc.add_page_break()

        # Chapters
        real_chapter_num = 0
        _docx_page_counter = [2]  # track for chapter_start parity

        for chapter in chapters:
            ch_title_lower = chapter["title"].lower()
            already_has_chapter = ch_title_lower.startswith("chapter") or ch_title_lower.startswith("part")
            is_intro = "introduction" in ch_title_lower or "front matter" in ch_title_lower

            if not is_intro:
                real_chapter_num += 1

            # ── Chapter start: insert blank page to land on correct parity ─────
            if not is_intro and _chapter_start_d in ("right_page", "left_page"):
                cur_p = _docx_page_counter[0]
                want_odd = (_chapter_start_d == "right_page")
                if want_odd and cur_p % 2 == 0:
                    doc.add_page_break()
                    _docx_page_counter[0] += 1
                elif not want_odd and cur_p % 2 == 1:
                    doc.add_page_break()
                    _docx_page_counter[0] += 1

            _ch_align_d = WD_ALIGN_PARAGRAPH.CENTER if _ch_center_d else WD_ALIGN_PARAGRAPH.LEFT

            if prefix and not already_has_chapter and not is_intro and not _ch_numbered_d:
                add_para(f"{prefix.upper()} {real_chapter_num}".strip(), body_fn, body_size * 0.82,
                         color=concept["accent_color"], space_before=6, space_after=2, align=_ch_align_d)

            # Apply heading design transforms
            if _ch_numbered_d and not already_has_chapter and not is_intro:
                _ch_title_text = f"{real_chapter_num}. {chapter['title']}"
            elif _ch_caps_d or _ch_smallcaps_d:
                _ch_title_text = chapter["title"].upper()
            else:
                _ch_title_text = chapter["title"]

            add_para(_ch_title_text, ch_fn, _ch_size_d, bold=True,
                     italic=ch_italic or _ch_italic_d,
                     color=concept["chapter_title_color"], space_after=8, align=_ch_align_d)

            # Post-heading decoration per design
            if _hd_d == "allcaps_rule":
                add_rule(concept["accent_color"], width_pct=100, thickness=1.2)
            elif _hd_d == "left_bold_clean":
                add_rule(concept["accent_color"], width_pct=40, thickness=0.8)
            elif _hd_d == "numbered":
                add_rule(concept["accent_color"], width_pct=15, thickness=2.0)
            elif _hd_d in ("centered_decorative", "italic_elegant", "smallcaps_ornament"):
                pass   # ornament paragraph handles decoration; no rule
            else:
                # Default (no heading_design set): thin full-width rule
                add_rule(concept["accent_color"], width_pct=100, thickness=0.5)

            if ornament:
                add_para(ornament, body_fn, body_size + 1, color=concept["accent_color"],
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
            elif _ch_smallcaps_d and ornament:
                # Extra ornament pass already handled above via general ornament block
                pass

            # Paragraph spacing: use explicit mm if set, else derive from line spacing
            _para_sp_after  = round(float(_para_sp_mm_d) * 2.835, 1) if _para_sp_mm_d else round(body_size * ls * 0.45, 1)
            _para_sp_before = round(float(_para_sp_mm_d) * 2.835 * 0.35, 1) if _para_sp_mm_d else round(body_size * ls * 0.10, 1)

            raw_body = chapter.get("body", "").strip()
            # Split on explicit section dividers when section_breaks enabled
            _SEC_DIV_RE_DOCX = re.compile(r"^\s*(?:---+|\*\*\*+|~~~+|###)\s*$", re.MULTILINE)
            _section_breaks_d = concept.get("section_breaks", False)
            if _section_breaks_d:
                raw_sections_d = _SEC_DIV_RE_DOCX.split(raw_body)
            else:
                raw_sections_d = [raw_body]

            for sec_idx_d, section_text_d in enumerate(raw_sections_d):
                if sec_idx_d > 0 and _section_breaks_d:
                    sec_orn_d = ornament if ornament else "* * *"
                    add_para(sec_orn_d, body_fn, body_size + 1, color=concept["accent_color"],
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

                paragraphs = [p.strip() for p in re.split(r"\n{2,}", section_text_d) if p.strip()]
                if not paragraphs and sec_idx_d == 0:
                    paragraphs = ["[No content]"]

                for para_text in paragraphs:
                    lines = para_text.split('\n')
                    cleaned_lines = []
                    for line in lines:
                        line = line.strip()
                        if not line: continue
                        if line.startswith(('•', '-', '*')) or re.match(r'^\d+\.', line):
                            cleaned_lines.append(line)
                        else:
                            if cleaned_lines:
                                cleaned_lines[-1] += " " + line
                            else:
                                cleaned_lines.append(line)

                    for sub_line in cleaned_lines:
                        align = WD_ALIGN_PARAGRAPH.LEFT if sub_line.startswith(('•', '-', '*')) else WD_ALIGN_PARAGRAPH.JUSTIFY
                        add_para(sub_line, body_fn, body_size, italic=body_italic,
                                 color=concept["text_color"],
                                 align=align,
                                 space_after=_para_sp_after,
                                 space_before=_para_sp_before,
                                 line_space=ls)

            doc.add_page_break()
            _docx_page_counter[0] += 1

        # ── DOCX back matter ──────────────────────────────────────────────────────
        if "about_author" in _back_matter_d:
            add_para("About the Author", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("The author is a writer and storyteller. "
                     "Please replace this placeholder with your biographical note.",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()
        if "about_publisher" in _back_matter_d:
            add_para("About the Publisher", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("Published by [Publisher Name]. "
                     "Please replace this with your publisher information.",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()
        if "references" in _back_matter_d:
            add_para("References", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("[References list placeholder — replace with your citations.]",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()
        if "bibliography" in _back_matter_d:
            add_para("Bibliography", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("[Bibliography placeholder — replace with your bibliography.]",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()
        if "index" in _back_matter_d:
            add_para("Index", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("[Index placeholder — replace with your index entries.]",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()
        if "other_books" in _back_matter_d:
            add_para("Other Books by the Author", ch_fn, ch_size * 0.85, bold=True,
                     color=concept["chapter_title_color"], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
            add_rule(concept["accent_color"], width_pct=100, thickness=0.5)
            add_para("[List your other books here.]",
                     body_fn, body_size, color=concept["text_color"],
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_space=ls)
            doc.add_page_break()
        doc.save(output_path)  # BUG FIX: was missing — DOCX was never written to disk
        return output_path
    except Exception as e:
        print(f"  ⚠️  render_layout_docx failed: {e}\n{traceback.format_exc()}")
        raise


# ─────────────────────────────────────────────────────────────────────────────
# Main orchestrator
# ─────────────────────────────────────────────────────────────────────────────

def design_layout(
    file_path: str,
    filename: str,
    output_dir: str,
    page_width_mm: float = 210.0,
    page_height_mm: float = 297.0,
    book_title: str = "",
    design_instructions: str = "",
    book_type: Optional[str] = None,
    visual_template: Optional[str] = None,
    progress_callback: Optional[Callable[[str, int, str], None]] = None,
    # ── Typography overrides (None = AI/profile decides) ─────────────────────
    body_font: Optional[str] = None,
    chapter_font: Optional[str] = None,
    body_font_size: Optional[float] = None,
    chapter_font_size: Optional[float] = None,
    line_spacing: Optional[float] = None,
    margin_top_mm: Optional[float] = None,
    margin_bottom_mm: Optional[float] = None,
    margin_left_mm: Optional[float] = None,
    margin_right_mm: Optional[float] = None,
    show_drop_cap: Optional[bool] = None,
    show_page_numbers: Optional[bool] = None,
    # ── Footer overrides ────────────────────────────────────────────────────
    footer_left_text: Optional[str] = None,
    footer_middle_text: Optional[str] = None,
    footer_right_pagenum: Optional[bool] = True,
    # ── Advanced layout overrides ───────────────────────────────────────────
    mirror_margins: Optional[bool] = None,
    gutter_mm: Optional[float] = None,
    paragraph_spacing_mm: Optional[float] = None,
    indent_mm: Optional[float] = None,
    color_mode: Optional[str] = None,
    bleed_mm: Optional[float] = None,
    chapter_start: Optional[str] = None,
    page_number_start: Optional[int] = None,
    page_number_style: Optional[str] = None,
    header_custom_text: Optional[str] = None,
    heading_design: Optional[str] = None,
    section_breaks: Optional[bool] = None,
    front_matter: Optional[list] = None,
    back_matter: Optional[list] = None,
) -> dict:
    """
    Full pipeline — book-type aware:
      1. Extract text
      2. Detect chapters (supports Hindi/Devanagari)
      3. Look up book-type profile (smart genre defaults)
      4. Build override hints (user values + profile)
      5. Ask GPT-4o for a layout concept (seeded with profile)
      6. Apply hard user overrides (user always wins)
      7. Render PDF (ReportLab + Unicode fonts)
      8. Render DOCX (python-docx)
    """

    def progress(stage: str, pct: int, message: str) -> None:
        if progress_callback:
            progress_callback(stage, pct, message)

    try:
        os.makedirs(output_dir, exist_ok=True)
        ext = os.path.splitext(filename)[1].lower()

        # --- CRITICAL FIX 1: REMOVED THE PDF NATIVE SIZE INHERITANCE ---
        # The backend now completely ignores the original PDF's dimensions.
        # It strictly uses the page_width_mm and page_height_mm sent by the frontend
        # (or defaults to 127.0 x 203.2 mm / 5x8 if none are sent).

        if not book_title:
            book_title = Path(filename).stem.replace("_", " ").replace("-", " ").title()
        book_title = _safe_title(book_title, fallback="My Book")

        # ── 1. Extract ────────────────────────────────────────────────────────────
        progress("extracting", 8, "Extracting text from your manuscript…")
        raw_text = extract_text(file_path, filename)
        if not raw_text.strip():
            raise ValueError("The uploaded file appears to contain no extractable text.")

        # ── 2. Parse chapters ─────────────────────────────────────────────────────
        progress("parsing", 20, "Detecting chapters and structure…")
        chapters = parse_chapters(raw_text)
        if not chapters:
            raise ValueError("Could not detect any chapters or sections in the manuscript.")

        # ── 3. Book-type profile ──────────────────────────────────────────────────
        profile = get_book_type_profile(book_type)
        if profile:
            progress("designing", 30, f"Applying {profile['_label']} design profile…")
        else:
            progress("designing", 30, "AI is designing your layout concept…")

        # ── 4. Build effective design instructions ────────────────────────────────
        override_hints: list[str] = []
        if body_font:            override_hints.append(f"body font MUST be {body_font}")
        if chapter_font:         override_hints.append(f"chapter heading font MUST be {chapter_font}")
        if body_font_size:       override_hints.append(f"body font size MUST be {body_font_size}pt")
        if chapter_font_size:    override_hints.append(f"chapter font size MUST be {chapter_font_size}pt")
        if line_spacing:         override_hints.append(f"line spacing MUST be {line_spacing}×")
        if margin_top_mm is not None:    override_hints.append(f"margin top MUST be {margin_top_mm}mm")
        if margin_bottom_mm is not None: override_hints.append(f"margin bottom MUST be {margin_bottom_mm}mm")
        if margin_left_mm is not None:   override_hints.append(f"margin left MUST be {margin_left_mm}mm")
        if margin_right_mm is not None:  override_hints.append(f"margin right MUST be {margin_right_mm}mm")
        if show_drop_cap is not None:
            override_hints.append("drop caps: " + ("ENABLED" if show_drop_cap else "DISABLED"))
        if show_page_numbers is not None:
            override_hints.append("page numbers: " + ("SHOWN" if show_page_numbers else "HIDDEN"))

        effective_instructions = design_instructions or ""

        _TEMPLATE_HINTS = {
            "classic_novel":    "Classic cream pages with serif fonts, generous margins, drop caps and ornamental chapter dividers — think vintage Penguin Classics.",
            "premium_hardcover":"Luxury dark background (#0f0f0f), cream/gold text, gold accent (#c8a200), wide margins — elegant premium edition.",
            "modern_minimal":   "Pure white page, clean Helvetica, minimal decoration, thin accent rule under chapter titles, airy spacing.",
            "sanskrit_style":   "Warm ivory page, saffron/gold accent (#c8830a), ornate ornament dividers, classic serif — traditional sacred text aesthetic.",
            "school_guide":     "White page, structured sans-serif layout, blue accent (#2563eb), numbered chapters, no drop cap — clear academic style.",
            "thriller_dark":    "Dark page (#111827) with near-white body text (#f3f4f6), red accent (#ef4444), high contrast, sharp Helvetica headings.",
            "retro_vintage":    "Warm sepia page (#f5ead0), brown text, antique brown accent, italic serif body, diagonal/decorative ornament.",
            "poetry_bloom":     "Soft blush page (#fff0f5), purple/magenta accent (#d63384), italic serif body, floral ornaments, wide margins.",
        }
        if visual_template:
            hint = _TEMPLATE_HINTS.get(visual_template, "")
            if hint:
                effective_instructions = hint + (("\n" + effective_instructions) if effective_instructions else "")

        if override_hints:
            hint_str = "; ".join(override_hints)
            effective_instructions = (
                (effective_instructions + "\n" if effective_instructions else "")
                + f"[HARD USER OVERRIDES — honour exactly: {hint_str}]"
            )

        # ── 5. AI concept (seeded with profile defaults) ──────────────────────────
        progress("designing", 40, "AI is crafting your personalised layout…")
        concept = generate_layout_concept(
            book_title=book_title,
            sample_text=raw_text,
            design_instructions=effective_instructions,
            page_width_mm=page_width_mm,
            page_height_mm=page_height_mm,
            book_type=book_type,
            profile_defaults=profile,
        )

        # ── 6. Apply hard user overrides (user always wins over AI + profile) ─────
        if body_font:
            concept["body_font"]          = body_font
        if chapter_font:
            concept["chapter_font"]       = chapter_font
        if body_font_size is not None:
            concept["body_font_size"]     = float(body_font_size)
        if chapter_font_size is not None:
            concept["chapter_font_size"]  = float(chapter_font_size)
        if line_spacing is not None:
            concept["line_spacing"]       = float(line_spacing)
        if margin_top_mm is not None:
            concept["margin_top_mm"]      = float(margin_top_mm)
        if margin_bottom_mm is not None:
            concept["margin_bottom_mm"]   = float(margin_bottom_mm)
        if margin_left_mm is not None:
            concept["margin_left_mm"]     = float(margin_left_mm)
        if margin_right_mm is not None:
            concept["margin_right_mm"]    = float(margin_right_mm)
        if show_drop_cap is not None:
            concept["show_drop_cap"]      = show_drop_cap
        if show_page_numbers is not None:
            concept["show_page_numbers"]  = show_page_numbers

        # ── Footer overrides ─────────────────────────────────────────────────────
        concept["footer_left_text"]     = footer_left_text if footer_left_text is not None else (book_title or "")
        concept["footer_middle_text"]   = footer_middle_text if footer_middle_text is not None else ""
        concept["footer_right_pagenum"] = footer_right_pagenum if footer_right_pagenum is not None else True

        # ── Advanced layout overrides ────────────────────────────────────────────
        if mirror_margins is not None:
            concept["mirror_margins"] = mirror_margins
        if gutter_mm is not None:
            concept["gutter_mm"]      = float(gutter_mm)
            # Gutter adds to the inner (binding) margin.
            # For single-sided: inner = left.  For mirror_margins: the PDF renderer
            # already alternates left/right per page; we add to left_mm (recto inner).
            concept["margin_left_mm"] = float(concept.get("margin_left_mm", 22)) + float(gutter_mm)
        if paragraph_spacing_mm is not None:
            concept["paragraph_spacing_mm"] = float(paragraph_spacing_mm)
        if indent_mm is not None:
            concept["first_para_indent_mm"] = float(indent_mm)
        if color_mode:
            concept["color_mode"] = color_mode
        if bleed_mm is not None:
            concept["bleed_mm"] = float(bleed_mm)
        if chapter_start:
            concept["chapter_start"] = chapter_start
        if page_number_start is not None:
            concept["page_number_start"] = int(page_number_start)
        if page_number_style:
            concept["page_number_style"] = page_number_style
        if header_custom_text:
            concept["header_text"] = header_custom_text
        if heading_design:
            concept["heading_design"] = heading_design
        if section_breaks is not None:
            concept["section_breaks"] = section_breaks
        if front_matter:
            concept["front_matter"] = front_matter
        if back_matter:
            concept["back_matter"] = back_matter

        concept["_book_type"]       = book_type or "auto"
        concept["_book_type_label"] = profile["_label"] if profile else "Auto (AI chosen)"

        job_id    = uuid.uuid4().hex
        safe_name = _safe_title(book_title, "book").replace(" ", "_")

        # ── 7. Render PDF ─────────────────────────────────────────────────────────
        progress("rendering", 58, "Typesetting PDF with your layout…")
        pdf_path = os.path.join(output_dir, f"layout_{safe_name}_{job_id}.pdf")
        render_layout_pdf(
            chapters=chapters, concept=concept, output_path=pdf_path,
            page_width_mm=page_width_mm, page_height_mm=page_height_mm, book_title=book_title,
        )

        # ── 8. Render DOCX ────────────────────────────────────────────────────────
        progress("rendering_docx", 80, "Generating DOCX version…")
        docx_path = os.path.join(output_dir, f"layout_{safe_name}_{job_id}.docx")
        render_layout_docx(
            chapters=chapters, concept=concept, output_path=docx_path,
            page_width_mm=page_width_mm, page_height_mm=page_height_mm, book_title=book_title,
        )

        progress("done", 100, "Layout design complete!")

        return {
            "title":            book_title,
            "style_name":       concept["style_name"],
            "concept":          concept,
            "chapter_count":    len(chapters),
            "chapter_titles":   [c["title"] for c in chapters],
            "pdf_path":         pdf_path,
            "docx_path":        docx_path,
            "job_id":           job_id,
            "book_type":        book_type or "auto",
            "book_type_label":  profile["_label"] if profile else "Auto (AI chosen)",
        }
    except Exception as e:
        print(f"  🚨 CRITICAL ERROR in design_layout: {e}\n{traceback.format_exc()}")
        if progress_callback:
            progress_callback("error", -1, f"Failed: {e}")
        raise