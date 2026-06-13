"""
translator.py
─────────────────────────────────────────────────────────────────────────────
Book Translator — accepts PDF / DOCX / ZIP, translates to any target language
using OpenAI GPT-4o, and exports both a PDF and a DOCX of the translated book.

MAXIMIZED ENTERPRISE EDITION:
- Multi-Agent Translation Swarm (Translator -> Critic -> Reconciler)
- Automated Terminology Glossary Extraction
- Exponential Backoff Rate-Limit Handling (Zero Drop-offs)
- Dynamic Unicode/Devanagari Font Engine for Flawless Global PDF/DOCX Exports
- Semantic Overlap Chunking (Flawless grammar continuity across chunk boundaries)
"""

from __future__ import annotations

import os
import re
import io
import uuid
import json
import zipfile
import shutil
import time
import traceback
import threading
import unicodedata
import urllib.request
from pathlib import Path
from typing import Callable, Optional

# ── Third-party ───────────────────────────────────────────────────────────────
# pyrefly: ignore [missing-import]
import openai                          
# pyrefly: ignore [missing-import]
import pdfplumber                      
# pyrefly: ignore [missing-import]
from docx import Document              
# pyrefly: ignore [missing-import]
from docx.shared import Pt, RGBColor, Inches
# pyrefly: ignore [missing-import]
from docx.enum.text import WD_ALIGN_PARAGRAPH
# pyrefly: ignore [missing-import]
from docx.oxml.ns import qn
# pyrefly: ignore [missing-import]
from docx.oxml import OxmlElement
# pyrefly: ignore [missing-import]
from reportlab.lib.pagesizes import A4
# pyrefly: ignore [missing-import]
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# pyrefly: ignore [missing-import]
from reportlab.lib.units import cm
# pyrefly: ignore [missing-import]
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
# pyrefly: ignore [missing-import]
from reportlab.lib.colors import HexColor, white
# pyrefly: ignore [missing-import]
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
# pyrefly: ignore [missing-import]
from reportlab.pdfbase import pdfmetrics
# pyrefly: ignore [missing-import]
from reportlab.pdfbase.ttfonts import TTFont

# ── Constants ─────────────────────────────────────────────────────────────────
_client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
_MODEL  = "gpt-4o"
_CHUNK_CHARS = 12_000
MAX_RETRIES = 8
RETRY_BASE_DELAY = 2.0

SUPPORTED_UPLOAD_EXTS = {".pdf", ".docx", ".zip"}

# ─────────────────────────────────────────────────────────────────────────────
# 0. UNIVERSAL FONT ENGINE (Fixes the Black Box Error for Hindi/Global text)
# ─────────────────────────────────────────────────────────────────────────────
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_FONTS_DIR  = os.path.join(_SCRIPT_DIR, "fonts")

_FONT_URLS: dict[str, str] = {
    "NotoSerifDevanagari-Regular.ttf": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifDevanagari/NotoSerifDevanagari-Regular.ttf",
    "NotoSerifDevanagari-Bold.ttf":    "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifDevanagari/NotoSerifDevanagari-Bold.ttf",
    "NotoSansDevanagari-Regular.ttf":  "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Regular.ttf",
    "NotoSansDevanagari-Bold.ttf":     "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Bold.ttf",
    "NotoNaskhArabic-Regular.ttf":     "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
    "NotoNaskhArabic-Bold.ttf":        "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Bold.ttf",
    "NotoSerifHebrew-Regular.ttf":     "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifHebrew/NotoSerifHebrew-Regular.ttf",
    "NotoSerifThai-Regular.ttf":       "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifThai/NotoSerifThai-Regular.ttf",
    "NotoSerifBengali-Regular.ttf":    "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifBengali/NotoSerifBengali-Regular.ttf",
    "NotoSerifTamil-Regular.ttf":      "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifTamil/NotoSerifTamil-Regular.ttf",
    "NotoSansGujarati-Regular.ttf":    "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansGujarati/NotoSansGujarati-Regular.ttf",
    "NotoSansGurmukhi-Regular.ttf":    "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansGurmukhi/NotoSansGurmukhi-Regular.ttf",
    "NotoSerifCyrillic-Regular.ttf":   "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSerifDisplay/NotoSerifDisplay-Regular.ttf",
}

_NOTO_FONT_FILES = {
    "NotoSerifDevanagari":      "NotoSerifDevanagari-Regular.ttf",
    "NotoSerifDevanagari-Bold": "NotoSerifDevanagari-Bold.ttf",
    "NotoSansDevanagari":       "NotoSansDevanagari-Regular.ttf",
    "NotoSansDevanagari-Bold":  "NotoSansDevanagari-Bold.ttf",
    "NotoNaskhArabic":          "NotoNaskhArabic-Regular.ttf",
    "NotoNaskhArabic-Bold":     "NotoNaskhArabic-Bold.ttf",
    "NotoSerifHebrew":          "NotoSerifHebrew-Regular.ttf",
    "NotoSerifThai":            "NotoSerifThai-Regular.ttf",
    "NotoSerifBengali":         "NotoSerifBengali-Regular.ttf",
    "NotoSerifTamil":           "NotoSerifTamil-Regular.ttf",
    "NotoSansGujarati":         "NotoSansGujarati-Regular.ttf",
    "NotoSansGurmukhi":         "NotoSansGurmukhi-Regular.ttf",
    "NotoSerifCyrillic":        "NotoSerifCyrillic-Regular.ttf",
}

# Script range → (rl_font_name, word_font_name)
_SCRIPT_FONT_MAP: list[tuple[int, int, str, str]] = [
    (0x0600, 0x06FF, "NotoNaskhArabic",     "Noto Naskh Arabic"),
    (0x0590, 0x05FF, "NotoSerifHebrew",     "Noto Serif Hebrew"),
    (0x0900, 0x097F, "NotoSerifDevanagari", "Noto Serif Devanagari"),
    (0x0980, 0x09FF, "NotoSerifBengali",    "Noto Serif Bengali"),
    (0x0A80, 0x0AFF, "NotoSansGujarati",    "Noto Sans Gujarati"),
    (0x0A00, 0x0A7F, "NotoSansGurmukhi",    "Noto Sans Gurmukhi"),
    (0x0B80, 0x0BFF, "NotoSerifTamil",      "Noto Serif Tamil"),
    (0x0E00, 0x0E7F, "NotoSerifThai",       "Noto Serif Thai"),
    (0x0400, 0x04FF, "NotoSerifCyrillic",   "Times New Roman"),
    # CJK: for PDF we have no reliable TTFont-compatible CJK font to download;
    # use the best registered Noto as a graceful fallback rather than crashing.
    # For DOCX, Word on Windows/Mac has SimSun/MS Gothic/Malgun built in.
    (0x4E00, 0x9FFF, "",  "SimSun"),       # Chinese Simplified
    (0x3040, 0x30FF, "",  "MS Gothic"),    # Japanese
    (0xAC00, 0xD7AF, "",  "Malgun Gothic"), # Korean
]

_SYSTEM_FONT_DIRS = [
    "/usr/share/fonts/truetype/noto",
    "/usr/share/fonts/noto",
    "/usr/share/fonts",
    "C:/Windows/Fonts",
    "/Library/Fonts"
]

_REGISTERED_FONTS: set[str] = set()
_FONTS_REGISTERED  = False
_FONT_LOCK         = threading.Lock()

def _find_font_on_system(filename: str) -> Optional[str]:
    local = os.path.join(_FONTS_DIR, filename)
    if os.path.isfile(local):
        return local
    for d in _SYSTEM_FONT_DIRS:
        p = os.path.join(d, filename)
        if os.path.isfile(p):
            return p
    return None

def _download_font(filename: str) -> Optional[str]:
    url = _FONT_URLS.get(filename)
    if not url:
        return None
    dest = os.path.join(_FONTS_DIR, filename)
    if os.path.isfile(dest):
        return dest
    try:
        os.makedirs(_FONTS_DIR, exist_ok=True)
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp, open(dest, "wb") as f:
            f.write(resp.read())
        return dest
    except Exception:
        return None

def _ensure_unicode_fonts() -> None:
    global _FONTS_REGISTERED, _REGISTERED_FONTS
    if _FONTS_REGISTERED:
        return
    with _FONT_LOCK:
        if _FONTS_REGISTERED:
            return
        try:
            for rl_name, filename in _NOTO_FONT_FILES.items():
                path = _find_font_on_system(filename) or _download_font(filename)
                if path:
                    try:
                        pdfmetrics.registerFont(TTFont(rl_name, path))
                        _REGISTERED_FONTS.add(rl_name)
                    except Exception:
                        pass
        except Exception:
            pass
        finally:
            _FONTS_REGISTERED = True

def _has_non_latin(text: str) -> bool:
    return any(ord(c) >= 0x0250 for c in text
               if not unicodedata.category(c).startswith("Z")
               and not unicodedata.category(c).startswith("P"))

def _detect_dominant_script(text: str) -> tuple[str, str]:
    """Return (pdf_rl_font, word_font) for the dominant non-Latin script.
    rl_font may be "" for CJK (no reliable TTFont-compatible download exists);
    callers must handle empty string by falling back to best available font."""
    counts: dict[int, int] = {}
    for ch in text:
        cp = ord(ch)
        if cp < 0x0250:
            continue
        for i, (lo, hi, _, _) in enumerate(_SCRIPT_FONT_MAP):
            if lo <= cp <= hi:
                counts[i] = counts.get(i, 0) + 1
    if not counts:
        return ("", "")
    idx = max(counts, key=counts.get)
    _, _, rl_name, word_name = _SCRIPT_FONT_MAP[idx]
    return (rl_name, word_name)

def _best_pdf_font(text: str) -> str:
    """Best ReportLab font for the dominant script. Never Helvetica for non-Latin."""
    _ensure_unicode_fonts()
    if not _has_non_latin(text):
        return "Helvetica"
    rl_name, _ = _detect_dominant_script(text)
    if rl_name and rl_name in _REGISTERED_FONTS:
        return rl_name
    for candidate in ["NotoSerifDevanagari", "NotoNaskhArabic", "NotoSerifThai",
                      "NotoSerifBengali", "NotoSerifTamil", "NotoSerifCyrillic",
                      "NotoSansDevanagari"]:
        if candidate in _REGISTERED_FONTS:
            return candidate
    return "Courier"

def _best_word_font(text: str) -> str:
    """Best Word font for the dominant script."""
    if not _has_non_latin(text):
        return "Calibri"
    _, word_name = _detect_dominant_script(text)
    return word_name if word_name else "Calibri"


# ─────────────────────────────────────────────────────────────────────────────
# 1. API RETRY ENGINE (Zero-Loss Rate Limit Handling)
# ─────────────────────────────────────────────────────────────────────────────
def _api_call_with_retry(fn, *args, **kwargs):
    """
    Wraps API calls in an exponential backoff loop to prevent mid-book failures
    caused by transient API errors or rate limits (429).
    """
    delay = RETRY_BASE_DELAY
    last_exc = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            return fn(*args, **kwargs)
        except Exception as e:
            last_exc = e
            err_str = str(e).lower()
            if any(k in err_str for k in ("rate limit", "429", "500", "502", "503", "timeout")):
                if attempt < MAX_RETRIES:
                    wait = delay * (2 ** (attempt - 1))
                    print(f"  ⏳  API error (attempt {attempt}/{MAX_RETRIES}), retrying in {wait:.1f}s...")
                    time.sleep(wait)
                    continue
            raise
    raise last_exc


# ─────────────────────────────────────────────────────────────────────────────
# 2. ADVANCED MULTI-AGENT PROMPTS
# ─────────────────────────────────────────────────────────────────────────────
GLOSSARY_EXTRACTOR_PROMPT = """You are an expert Terminologist and Book Localization Planner.
Review this sample text from a manuscript and build a bilingual translation glossary map.
Identify:
1) Recurring character names
2) Distinct locations
3) Crucial technical terms

Output strictly a valid JSON object:
{
  "glossary": {
    "Original Term/Name": "Strict Translated Equivalent"
  }
}"""

TRANSLATOR_SYSTEM_PROMPT = """You are an elite literary translator translating into the target language.

Strict Guidelines:
1. Preserve original emotional tone, style, author voice, and nuances.
2. Adapt idioms into natural equivalents in the target language.
3. You MUST rigidly follow the provided Translation Glossary Map for names and nouns.
4. Maintain paragraph breaks exactly as presented.
5. If a PREVIOUS CHUNK CONTEXT is provided, use it ONLY to understand the immediate preceding context for continuity. Do NOT translate it.
6. Output ONLY the translated text. No preamble, no "Here is the translation:", no commentary. Begin immediately with the first translated word."""

CRITIC_SYSTEM_PROMPT = """You are a ruthless Cultural Editor and Localization Expert. 
Compare the original book section with its newly translated version.

Flag:
1. Literal translations that sound awkward.
2. Missed idioms.
3. Loss of author style.

Output critique of issues found. If flawless, output ONLY the word: "PASSED"."""

RECONCILER_SYSTEM_PROMPT = """You are a Structural Reconciliation Specialist. 
Verify the polished translation against the original text structure.

Ensure:
1. EVERY original paragraph is present. 
2. Glossary terms were applied.

CRITICAL OUTPUT RULES:
- Output ONLY the final translated book content. Nothing else.
- Do NOT output any commentary, notes, explanations, or preamble whatsoever.
- Do NOT start with phrases like "Note:", "I have ensured", "Here is the", "Translation:".
- Begin IMMEDIATELY with the first word of the translated content."""


# ─────────────────────────────────────────────────────────────────────────────
# 3. TEXT EXTRACTION & STRUCTURING
# ─────────────────────────────────────────────────────────────────────────────
_CID_PATTERN = re.compile(r"\(cid:(\d+)\)")
_CID_MAP: dict[int, str] = {
    127: "\u2022", 183: "\u2022", 149: "\u2022", 164: "\u2022",
    150: "\u2013", 151: "\u2014",
    145: "\u2018", 146: "\u2019",
    147: "\u201c", 148: "\u201d",
    133: "\u2026", 160: " ", 173: "-",
}

def _clean_cid(text: str) -> str:
    """Replace PDF (cid:NNN) glyph refs with proper Unicode chars.
    Without this, GPT-4o translates literal '(cid:127)' as text."""
    return _CID_PATTERN.sub(lambda m: _CID_MAP.get(int(m.group(1)), "\u2022"), text)

# ── FIX TR-1: _CHAPTER_RE must be defined BEFORE _extract_docx which references it ──
_CHAPTER_RE = re.compile(
    r"^(?:"
    # Latin: Chapter 1, Part II, Section 3, Unit IV, Ch. 5
    r"(?:chapter|ch\.?|part|section|unit)\s+[\dIVXivx]+[:\.\s].*"
    r"|[\dIVX]+[\.]\s+[A-Z].{2,}"
    # Hindi/Devanagari: अध्याय, भाग, खंड + digit or Devanagari numeral
    r"|(?:अध्याय|भाग|खंड|पाठ|प्रकरण|सर्ग)\s*[\d\u0966-\u096F].*"
    # Arabic: الفصل, الجزء, القسم
    r"|(?:الفصل|الجزء|القسم)\s+.*"
    # Chinese: 第X章/节/部
    r"|第[一二三四五六七八九十百\d]+[章节部篇].*"
    # Russian: Глава, Часть
    r"|(?:Глава|Часть)\s+[\dIVXivx]+.*"
    # Injected by _extract_docx for DOCX heading styles
    r"|Chapter:\s+.+"
    r")",
    re.IGNORECASE | re.MULTILINE,
)

def _extract_pdf(path: str) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(_clean_cid(txt.strip()))
    return "\n\n".join(parts)

def _extract_docx(path: str) -> str:
    doc = Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower()
        is_heading = any(h in style_name for h in (
            "heading 1", "heading1", "title", "chapter", "h1",
        ))
        if is_heading and not _CHAPTER_RE.match(text):
            text = "Chapter: " + text
        lines.append(text)
    return _clean_cid("\n\n".join(lines))

def _extract_zip(zip_path: str, scratch_dir: str) -> tuple[str, str]:
    with zipfile.ZipFile(zip_path, "r") as zf:
        members = [
            m for m in zf.namelist()
            if os.path.splitext(m)[1].lower() in {".pdf", ".docx"}
            and not m.startswith("__MACOSX")
            and not os.path.basename(m).startswith(".")
        ]
        if not members:
            raise ValueError("No .pdf or .docx files found in zip.")

        member = members[0]
        ext = os.path.splitext(member)[1].lower()
        tmp = os.path.join(scratch_dir, f"zip_extracted_{uuid.uuid4().hex}{ext}")
        with zf.open(member) as src, open(tmp, "wb") as dst:
            shutil.copyfileobj(src, dst)

    base_name = os.path.basename(member)
    if ext == ".pdf":
        text = _extract_pdf(tmp)
    else:
        text = _extract_docx(tmp)

    os.remove(tmp)
    return text, base_name

def extract_book_text(file_path: str, filename: str, scratch_dir: str) -> tuple[str, str]:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path), filename
    if ext == ".docx":
        return _extract_docx(file_path), filename
    if ext == ".zip":
        return _extract_zip(file_path, scratch_dir)
    raise ValueError(f"Unsupported file type: {ext}")

def _parse_structure(text: str, title_hint: str = "") -> dict:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    title = title_hint or (lines[0] if lines else "Translated Book")
    splits = list(_CHAPTER_RE.finditer(text))
    chapters: list[dict] = []

    if not splits:
        chapters = [{"title": "Full Text", "body": text}]
    else:
        preamble = text[: splits[0].start()].strip()
        if preamble:
            chapters.append({"title": "Introduction", "body": preamble})
        for idx, m in enumerate(splits):
            ch_title = m.group(0).strip()
            start = m.end()
            end = splits[idx + 1].start() if idx + 1 < len(splits) else len(text)
            body = text[start:end].strip()
            # TR-7: always keep the chapter so it appears in the output;
            # an empty body will be caught in _translate_chapter and skipped gracefully
            chapters.append({"title": ch_title, "body": body or ""})

    return {"title": title, "chapters": chapters}


# ─────────────────────────────────────────────────────────────────────────────
# 4. SWARM AI PIPELINE
# ─────────────────────────────────────────────────────────────────────────────
def _detect_language(sample: str) -> str:
    resp = _api_call_with_retry(
        _client.chat.completions.create, 
        model=_MODEL, 
        max_tokens=20, 
        temperature=0,
        messages=[
            {"role": "system", "content": "Reply with ONLY the full English name of the language."},
            {"role": "user", "content": f"What language is this text written in?\n\n{sample[:600]}"}
        ]
    )
    return resp.choices[0].message.content.strip()

def _extract_automated_glossary(text_sample: str, target_lang: str) -> dict:
    try:
        resp = _api_call_with_retry(
            _client.chat.completions.create, 
            model=_MODEL, 
            temperature=0.2,
            max_tokens=2048,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": GLOSSARY_EXTRACTOR_PROMPT},
                {"role": "user", "content": f"TARGET LANGUAGE: {target_lang}\n\nMANUSCRIPT TEXT SAMPLE:\n{text_sample[:15000]}"}
            ]
        )
        return json.loads(resp.choices[0].message.content.strip()).get("glossary", {})
    except Exception as e: 
        print(f"    ⚠️ Glossary extraction failed, bypassing: {e}")
        return {}

def _translate_title(title: str, target_language: str, source_language: str) -> str:
    resp = _api_call_with_retry(
        _client.chat.completions.create, 
        model=_MODEL, 
        max_tokens=80, 
        temperature=0.2,
        messages=[
            {"role": "system", "content": f"Translate the following book title into {target_language}. Return ONLY the translated title."},
            {"role": "user", "content": title}
        ]
    )
    return resp.choices[0].message.content.strip()

def _translate_chunk_swarm(
    text_chunk: str, 
    target_language: str, 
    source_language: str, 
    context_hint: str,
    overlap_context: str,
    glossary_map: dict
) -> str:
    str_glossary = json.dumps(glossary_map or {}, ensure_ascii=False)

    # --- Phase 1: Translator ---
    user_content = (
        f"SOURCE: {source_language}\n"
        f"TARGET: {target_language}\n"
        f"GLOSSARY:\n{str_glossary}\n"
        f"CONTEXT: {context_hint}\n"
    )
    if overlap_context:
        user_content += f"\n[PREVIOUS CHUNK CONTEXT - DO NOT TRANSLATE THIS]:\n{overlap_context}\n\n"
        
    user_content += f"ORIGINAL TO TRANSLATE:\n{text_chunk}"
    
    resp1 = _api_call_with_retry(
        _client.chat.completions.create,
        model=_MODEL,
        messages=[
            {"role": "system", "content": TRANSLATOR_SYSTEM_PROMPT},
            {"role": "user", "content": user_content}
        ],
        temperature=0.3,
        max_tokens=4096,
    )
    initial_translation = resp1.choices[0].message.content.strip()

    # --- Phase 2: Critic & Refinement ---
    critic_user = (
        f"ORIGINAL:\n{text_chunk}\n\n"
        f"TRANSLATED:\n{initial_translation}\n\n"
        f"TARGET: {target_language}"
    )
    
    resp2 = _api_call_with_retry(
        _client.chat.completions.create,
        model=_MODEL,
        messages=[
            {"role": "system", "content": CRITIC_SYSTEM_PROMPT},
            {"role": "user", "content": critic_user}
        ],
        temperature=0.2,
        max_tokens=1024,  # critique is short; 1024 is plenty and keeps cost low
    )
    critic_feedback = resp2.choices[0].message.content.strip()

    refined_translation = initial_translation
    if "PASSED" not in critic_feedback.upper() and len(critic_feedback) > 10:
        refine_user = (
            f"ORIGINAL:\n{text_chunk}\n\n"
            f"INITIAL TRANSLATION:\n{initial_translation}\n\n"
            f"CRITIQUE:\n{critic_feedback}\n\n"
            f"Please output the refined translation."
        )
        
        resp3 = _api_call_with_retry(
            _client.chat.completions.create,
            model=_MODEL,
            messages=[
                {"role": "system", "content": TRANSLATOR_SYSTEM_PROMPT},
                {"role": "user", "content": refine_user}
            ],
            temperature=0.3,
            max_tokens=4096,
        )
        refined_translation = resp3.choices[0].message.content.strip()

    # --- Phase 3: Reconciler (skipped when critic says PASSED) ---
    if "PASSED" in critic_feedback.upper() and len(refined_translation) > 20:
        return refined_translation

    rec_user = (
        f"ORIGINAL:\n{text_chunk}\n\n"
        f"CURRENT DRAFT:\n{refined_translation}\n\n"
        f"GLOSSARY TARGETS:\n{str_glossary}\n\n"
        f"Ensure exact structural mapping. Output final text only."
    )
    
    resp4 = _api_call_with_retry(
        _client.chat.completions.create, 
        model=_MODEL, 
        messages=[
            {"role": "system", "content": RECONCILER_SYSTEM_PROMPT}, 
            {"role": "user", "content": rec_user}
        ], 
        temperature=0.1,
        max_tokens=4096,
    )
    
    return resp4.choices[0].message.content.strip()

def _translate_chapter(
    chapter: dict, 
    target_language: str, 
    source_language: str, 
    book_title: str, 
    glossary_map: dict
) -> dict:
    translated_title = _translate_title(chapter["title"], target_language, source_language)

    # TR-9: skip translation entirely for empty chapters
    if not chapter["body"].strip():
        return {"title": translated_title, "body": ""}

    paragraphs = chapter["body"].split("\n\n")
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    
    for para in paragraphs:
        if current_len + len(para) + 2 > _CHUNK_CHARS and current:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
            
        current.append(para)
        current_len += len(para) + 2
        
    if current: 
        chunks.append("\n\n".join(current))
    if not chunks: 
        chunks = [chapter["body"]]

    translated_parts = []
    rolling_memory = f'Book: "{book_title}"'
    previous_chunk_tail = "" # NEW: Holds the last ~200 words of the previous chunk
    
    for chunk in chunks:
        translated_text = _translate_chunk_swarm(
            text_chunk=chunk, 
            target_language=target_language, 
            source_language=source_language, 
            context_hint=rolling_memory, 
            overlap_context=previous_chunk_tail, # NEW: Semantic overlap injected here
            glossary_map=glossary_map
        )
        translated_parts.append(translated_text)
        
        # T-9: use the TRANSLATED tail as overlap, not the source chunk.
        # Source-language overlap confuses the model (it skips translating "previous context").
        t_words = translated_text.split()
        previous_chunk_tail = " ".join(t_words[-200:]) if len(t_words) > 200 else translated_text
        
        rolling_memory = f"Book: '{book_title}'. Last translated stylistic line: {translated_text[-250:]}"

    return {"title": translated_title, "body": "\n\n".join(translated_parts)}


# ─────────────────────────────────────────────────────────────────────────────
# 5. DYNAMIC OUTPUT GENERATORS (PDF/DOCX)
# ─────────────────────────────────────────────────────────────────────────────
def _xml_escape(text: str) -> str:
    """Full XML escape for ReportLab Paragraph content."""
    import unicodedata as _ud
    text = _ud.normalize("NFC", text)
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;"))


def _build_pdf(book: dict, target_language: str, output_path: str,
               source_language: str = "", author: str = "") -> None:
    """
    Render the translated book to a PDF using ReportLab with Dynamic Fonts.
    Includes a publisher-quality cover page and an auto-generated TOC page.

    FIX TR-2: _CHAPTER_RE was referenced before it was defined — now resolved
              by moving _CHAPTER_RE above the extraction functions.
    FIX TR-3: TOC page added after the cover, before chapter content.
    FIX TR-4: Cover page now includes author name block, translation credit line,
              language pair banner, and ISBN placeholder — proper publisher layout.
    """
    _ensure_unicode_fonts()

    all_text = book["title"] + " ".join([ch["body"] for ch in book["chapters"]])
    body_font = _best_pdf_font(all_text)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2.5*cm,
        leftMargin=2.5*cm,
        topMargin=2.8*cm,
        bottomMargin=2.5*cm,
        title=book["title"],
    )

    styles = getSampleStyleSheet()

    # ── Style definitions ────────────────────────────────────────────────────
    cover_title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontSize=30,
        leading=36,
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName=body_font,
        textColor=HexColor("#1E293B"),
    )
    cover_author_style = ParagraphStyle(
        "CoverAuthor",
        parent=styles["Normal"],
        fontSize=15,
        leading=20,
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName=body_font,
        textColor=HexColor("#334155"),
    )
    cover_label_style = ParagraphStyle(
        "CoverLabel",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        fontName=body_font,
        textColor=HexColor("#64748B"),
    )
    cover_credit_style = ParagraphStyle(
        "CoverCredit",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
        alignment=TA_CENTER,
        fontName=body_font,
        textColor=HexColor("#94A3B8"),
    )
    toc_title_style = ParagraphStyle(
        "TocTitle",
        parent=styles["Heading1"],
        fontSize=14,
        leading=18,
        spaceAfter=16,
        spaceBefore=0,
        alignment=TA_LEFT,
        fontName=body_font,
        textColor=HexColor("#1E293B"),
    )
    toc_entry_style = ParagraphStyle(
        "TocEntry",
        parent=styles["Normal"],
        fontSize=11,
        leading=18,
        spaceAfter=2,
        leftIndent=0,
        fontName=body_font,
        textColor=HexColor("#334155"),
    )
    ch_head_style = ParagraphStyle(
        "ChHead",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        spaceBefore=24,
        spaceAfter=10,
        fontName=body_font,
        textColor=HexColor("#1E293B"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=11,
        leading=17,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        fontName=body_font,
    )

    # ── Cover page ───────────────────────────────────────────────────────────
    # Language pair banner  e.g. "English → Hindi"
    if source_language and target_language:
        lang_banner = f"{source_language} → {target_language}"
    else:
        lang_banner = f"Translated into {target_language}"

    author_line = author.strip() if author.strip() else ""

    story: list = [
        Spacer(1, 5.5 * cm),
        Paragraph(_xml_escape(book["title"]), cover_title_style),
        Spacer(1, 0.4 * cm),
        HRFlowable(width="50%", thickness=1.5, color=HexColor("#334155")),
        Spacer(1, 0.4 * cm),
    ]

    if author_line:
        story.append(Paragraph(_xml_escape(author_line), cover_author_style))
        story.append(Spacer(1, 0.3 * cm))

    story += [
        Spacer(1, 5 * cm),
        HRFlowable(width="80%", thickness=0.5, color=HexColor("#CBD5E1")),
        Spacer(1, 0.3 * cm),
        Paragraph(_xml_escape(lang_banner), cover_label_style),
        Spacer(1, 0.15 * cm),
        Paragraph("Translated by Enterprise AI Swarm", cover_credit_style),
        Spacer(1, 0.15 * cm),
        Paragraph("ISBN: 000-0-000-00000-0", cover_credit_style),
        PageBreak(),
    ]

    # ── Table of Contents page ───────────────────────────────────────────────
    story.append(Paragraph("Contents", toc_title_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#CBD5E1")))
    story.append(Spacer(1, 0.3 * cm))

    for idx, ch in enumerate(book["chapters"], start=1):
        ch_title_raw = (ch.get("title") or "").strip() or f"Chapter {idx}"
        entry_text = f"{idx}.&nbsp;&nbsp;{_xml_escape(ch_title_raw)}"
        story.append(Paragraph(entry_text, toc_entry_style))

    story.append(PageBreak())

    # ── Chapter content ──────────────────────────────────────────────────────
    for ch in book["chapters"]:
        ch_title_safe = _xml_escape((ch.get("title") or "").strip())
        story.append(Paragraph(ch_title_safe or "&nbsp;", ch_head_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#CBD5E1")))
        story.append(Spacer(1, 0.3 * cm))

        for para in ch["body"].split("\n\n"):
            raw = para.strip()
            if not raw:
                continue
            story.append(Paragraph(_xml_escape(raw), body_style))

        story.append(PageBreak())

    doc.build(story)

def _build_docx(book: dict, target_language: str, output_path: str,
                source_language: str = "", author: str = "") -> None:
    """
    Render the translated book to a DOCX using python-docx with CS Font injections.
    Includes a publisher-quality cover and a Table of Contents page.

    FIX TR-3: TOC page added after cover, before chapter content.
    FIX TR-4: Cover now has author name, language pair banner, and ISBN placeholder.
    """
    import unicodedata as _ud

    doc = Document()

    all_text = book["title"] + " ".join([ch["body"] for ch in book["chapters"]])
    word_font = _best_word_font(all_text)

    def apply_font(run, font_name: str) -> None:
        run.font.name = font_name
        rPr = run._r.get_or_add_rPr()
        fonts = rPr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rPr.insert(0, fonts)
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        fonts.set(qn("w:cs"), font_name)

    def _add_centered(text: str, bold: bool = False, size_pt: float = 12,
                      color: Optional[RGBColor] = None, space_before_pt: float = 0) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if space_before_pt:
            p.paragraph_format.space_before = Pt(space_before_pt)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size_pt)
        if color:
            r.font.color.rgb = color
        apply_font(r, word_font)

    # ── Cover page ───────────────────────────────────────────────────────────
    lang_banner = (f"{source_language} → {target_language}"
                   if source_language and target_language
                   else f"Translated into {target_language}")

    # Push title down the page with blank paragraphs
    for _ in range(8):
        doc.add_paragraph()

    _add_centered(book["title"], bold=True, size_pt=26)

    if author.strip():
        _add_centered(author.strip(), bold=False, size_pt=14,
                      color=RGBColor(0x33, 0x41, 0x55))

    # Spacer gap then meta block
    for _ in range(6):
        doc.add_paragraph()

    _add_centered(lang_banner, size_pt=11, color=RGBColor(0x64, 0x74, 0x8B))
    _add_centered("Translated by Enterprise AI Swarm", size_pt=9,
                  color=RGBColor(0x94, 0xA3, 0xB8))
    _add_centered("ISBN: 000-0-000-00000-0", size_pt=9,
                  color=RGBColor(0x94, 0xA3, 0xB8))

    doc.add_page_break()

    # ── Table of Contents page ───────────────────────────────────────────────
    toc_heading = doc.add_heading("Contents", level=1)
    if toc_heading.runs:
        apply_font(toc_heading.runs[0], word_font)

    for idx, ch in enumerate(book["chapters"], start=1):
        ch_title_raw = _ud.normalize("NFC", (ch.get("title") or "").strip()) or f"Chapter {idx}"
        toc_p = doc.add_paragraph()
        toc_r = toc_p.add_run(f"{idx}.  {ch_title_raw}")
        toc_r.font.size = Pt(11)
        apply_font(toc_r, word_font)

    doc.add_page_break()

    # ── Chapter content ──────────────────────────────────────────────────────
    for ch in book["chapters"]:
        ch_title = _ud.normalize("NFC", (ch.get("title") or "").strip())
        h = doc.add_heading(ch_title or " ", level=1)
        # Guard against empty heading (zero runs → IndexError)
        if h.runs:
            apply_font(h.runs[0], word_font)
        else:
            run_h = h.add_run(ch_title or " ")
            apply_font(run_h, word_font)

        for para in ch["body"].split("\n\n"):
            cleaned = _ud.normalize("NFC", para.strip())
            if not cleaned:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for line_idx, line in enumerate(cleaned.split("\n")):
                if line_idx > 0:
                    p.add_run().add_break()
                run_p = p.add_run(line)
                run_p.font.size = Pt(11)
                apply_font(run_p, word_font)

        doc.add_page_break()

    doc.save(output_path)


# ─────────────────────────────────────────────────────────────────────────────
# 6. ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────
def translate_book(
    file_path: str, 
    filename: str, 
    output_dir: str, 
    target_language: str, 
    source_language: str = "", 
    progress_callback: Optional[Callable[[str, int, str], None]] = None
) -> dict:
    
    def _progress(stage: str, pct: int, msg: str) -> None:
        if progress_callback: 
            progress_callback(stage, pct, msg)

    os.makedirs(output_dir, exist_ok=True)
    scratch_dir = os.path.join(output_dir, f"translate_scratch_{uuid.uuid4().hex}")
    os.makedirs(scratch_dir, exist_ok=True)

    try:
        _progress("extracting", 5, "Extracting text from document…")
        raw_text, source_filename = extract_book_text(file_path, filename, scratch_dir)
        if not raw_text.strip(): 
            raise ValueError("Uploaded document is empty.")

        _progress("extracting", 12, "Detecting source language…")
        detected_src = source_language.strip() or _detect_language(raw_text[:1200])
        
        _progress("extracting", 18, "Building automated terminology glossary...")
        global_glossary = _extract_automated_glossary(raw_text, target_language)

        _progress("structuring", 22, "Analysing book structure…")
        title_hint = os.path.splitext(source_filename)[0].replace("_", " ").replace("-", " ").title()
        structure  = _parse_structure(raw_text, title_hint)
        n_chapters = len(structure["chapters"])
        
        translated_title = _translate_title(structure["title"], target_language, detected_src)
        translated_chapters = []

        for idx, chapter in enumerate(structure["chapters"]):
            pct = 30 + int((idx / max(n_chapters, 1)) * 55)  # TR-8: guard ZeroDivisionError
            _progress("translating", pct, f'Translating chapter {idx + 1}/{n_chapters}: "{chapter["title"][:60]}"…')
            translated_chapters.append(
                _translate_chapter(
                    chapter, 
                    target_language, 
                    detected_src, 
                    structure["title"], 
                    global_glossary
                )
            )

        translated_book = {
            "title": translated_title, 
            "chapters": translated_chapters
        }

        _progress("assembling", 87, "Building Global PDF…")
        job_id = uuid.uuid4().hex
        pdf_path  = os.path.join(output_dir, f"translated_{job_id}.pdf")
        docx_path = os.path.join(output_dir, f"translated_{job_id}.docx")

        # Extract author from first line of raw text if it looks like a byline
        # (heuristic: second non-empty line after title, short, no sentence punctuation)
        _author_hint = ""
        _raw_lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
        if len(_raw_lines) >= 2:
            candidate = _raw_lines[1]
            if len(candidate) < 80 and not any(c in candidate for c in ".!?:"):
                _author_hint = candidate

        _build_pdf(translated_book, target_language, pdf_path,
                   source_language=detected_src, author=_author_hint)

        _progress("assembling", 93, "Building Global DOCX…")
        _build_docx(translated_book, target_language, docx_path,
                    source_language=detected_src, author=_author_hint)

        total_words = sum(len(ch["body"].split()) for ch in translated_chapters)
        _progress("done", 100, "Enterprise Translation complete!")

        return {
            "title": translated_title, 
            "source_language": detected_src, 
            "target_language": target_language,
            "total_words": total_words, 
            "chapters": n_chapters, 
            "chapter_titles": [ch["title"] for ch in translated_chapters],
            "pdf_path": pdf_path, 
            "docx_path": docx_path,
        }

    finally:
        shutil.rmtree(scratch_dir, ignore_errors=True)